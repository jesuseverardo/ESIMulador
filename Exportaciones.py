import Interfaz as _base
from Interfaz import *
try:
    _ = _base._
except Exception:
    def _(text):
        return text

def exportar_proyecto_esim(nombre_archivo: str, componentes_scene: List[Any]) -> None:
    try:
        model = EsimulatorCore.construir_modelo(componentes_scene)
        comps_data = []
        for c in model.components:
            comp_entry = {
                "name": c.name,
                "type": c.type,
                "params": c.params,
                "pins": [
                    {"net_id": p.net_id, "index": p.index} for p in c.pins
                ]
            }
            comps_data.append(comp_entry)
        nets_data = []
        for n in model.nets.values():
            net_entry = {
                "id": n.id,
                "is_gnd": n.is_gnd,
                "pins": [
                    {"comp_name": p.comp_name, "index": p.index} for p in n.pins
                ]
            }
            nets_data.append(net_entry)
        data = {
            "components": comps_data,
            "nets": nets_data
        }
        bom: Dict[str, int] = {}
        for comp in model.components:
            bom[comp.type] = bom.get(comp.type, 0) + 1
        metadata = {
            "generator": "ESIMulador",
            "datetime": datetime.now().isoformat(),
            "component_count": len(model.components),
            "net_count": len(model.nets)
        }
        with zipfile.ZipFile(nombre_archivo, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('circuit.json', json.dumps(data, indent=4, ensure_ascii=False))
            zf.writestr('bom.json', json.dumps(bom, indent=4, ensure_ascii=False))
            zf.writestr('metadata.json', json.dumps(metadata, indent=4, ensure_ascii=False))
        estado_correcto(None, _("Proyecto exportado correctamente"))
        try:
            llamar_hooks("project_exported", nombre_archivo=nombre_archivo, metadata=metadata)
        except Exception:
            pass
    except Exception as exc:
        mostrar_error(None, "exportar proyecto", exc)
        try:
            logger.exception("Error al exportar proyecto .esim: %s", exc)
        except Exception:
            pass
        estado_error(None, _("La exportación del proyecto ha fallado"))

def exportar_imagen(self):
        from PyQt5.QtWidgets import QFileDialog, QInputDialog, QMessageBox
        from PyQt5.QtGui import QImage
        opciones = ["Con cuadriculas", "Sin cuadriculas"]
        opcion, ok = QInputDialog.getItem(
            self,
            "Exportar imagen",
            "Seleccione cómo exportar la imagen:",
            opciones,
            0,
            False
        )
        if not ok:
            return
        incluir_grid = opcion == "Con cuadriculas"
        rect = self.scene.itemsBoundingRect()
        imagen = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        imagen.fill(Qt.white)
        original_pen = None
        if not incluir_grid and hasattr(self.scene, 'grid_pen'):
            try:
                original_pen = QPen(self.scene.grid_pen)  
                transparent_pen = QPen(Qt.transparent)
                self.scene.grid_pen = transparent_pen
            except Exception:
                original_pen = None
        painter = QPainter(imagen)
        self.scene.render(painter, target=QRectF(imagen.rect()), source=rect)
        painter.end()
        if not incluir_grid and original_pen is not None:
            try:
                self.scene.grid_pen = original_pen
            except Exception:
                pass
        archivo, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar como imagen",
            "",
            "Imagen PNG (*.png);;Imagen JPEG (*.jpg)"
        )
        if archivo:
            imagen.save(archivo)
            try:
                QMessageBox.information(self, "Guardar imagen", f"Imagen guardada en:\n{archivo}")
            except Exception:
                try:
                    self.statusBar().showMessage(f"Imagen guardada en: {archivo}", 5000)
                except Exception:
                    pass

def exportar_pdf(self):
        from PyQt5.QtWidgets import QFileDialog, QInputDialog
        from PyQt5.QtGui import QImage, QPen
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        import tempfile

        opciones = ["Con cuadriculas", "Sin cuadriculas"]
        opcion, ok = QInputDialog.getItem(
            self,
            "Exportar imagen",
            "Seleccione cómo exportar la imagen:",
            opciones,
            0,
            False
        )
        if not ok:
            return
        incluir_grid = opcion == "Con cuadriculas"
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        original_pen_pdf = None
        if not incluir_grid and hasattr(self.scene, 'grid_pen'):
            try:
                original_pen_pdf = QPen(self.scene.grid_pen)
                self.scene.grid_pen = QPen(Qt.transparent)
            except Exception:
                original_pen_pdf = None
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        if not incluir_grid and original_pen_pdf is not None:
            try:
                self.scene.grid_pen = original_pen_pdf
            except Exception:
                pass

        archivo, _ = QFileDialog.getSaveFileName(self, "Guardar como PDF", "", "PDF (*.pdf)")
        if not archivo:
            return

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img_path = tmp.name
            img.save(img_path)

        c = canvas.Canvas(archivo, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(50, 750, "ESIMulador - Exportación PDF")
        c.drawImage(ImageReader(img_path), 50, 450, width=500, preserveAspectRatio=True)

        data = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            data.append([
                comp.nombre,
                comp.tipo,
                ", ".join(comp.parametros)
            ])

        table = Table(data, colWidths=[100, 120, 280])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))

        table.wrapOn(c, 480, 300)
        table.drawOn(c, 50, 100)

        c.save()
        try:
            QMessageBox.information(self, "Guardar PDF", f"PDF guardado en:\n{archivo}")
        except Exception:
            try:
                self.statusBar().showMessage(f"PDF guardado en: {archivo}", 5000)
            except Exception:
                pass

def exportar_resultados(self):
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        if not self.resultados_simulaciones:
            QMessageBox.information(self, "Exportar resultados", "No hay resultados para exportar.")
            return
        opciones = "Archivo de texto (*.txt);;Archivo CSV (*.csv)"
        ruta, filtro = QFileDialog.getSaveFileName(self, "Guardar resultados", "", opciones)
        if not ruta:
            return
        try:
            es_csv = ruta.lower().endswith(".csv") or "csv" in filtro.lower()
            from PyQt5.QtCore import Qt as _Qt
            results_iter: list[tuple[str, str]] = []
            try:
                if hasattr(self, 'lista_resultados') and self.lista_resultados is not None:
                    selected_items = self.lista_resultados.selectedItems()
                    if selected_items:
                        for _it in selected_items:
                            data = _it.data(_Qt.UserRole)
                            if data:
                                results_iter.append(data)
                    elif self.resultados_simulaciones:
                        results_iter.append(self.resultados_simulaciones[-1])
                elif self.resultados_simulaciones:
                    results_iter.append(self.resultados_simulaciones[-1])
            except Exception:
                results_iter = list(self.resultados_simulaciones)
            if not results_iter:
                results_iter = list(self.resultados_simulaciones)
            with open(ruta, "w", encoding="utf-8") as f:
                for tipo, mensaje in results_iter:
                    if es_csv:
                        import re
                        texto = re.sub(r'<br\\s*/?>', '\n', mensaje, flags=re.IGNORECASE)
                        texto = re.sub(r'<[^>]+>', '', texto)  
                        lineas = [l.strip() for l in texto.splitlines() if l.strip()]
                        for linea in lineas:
                            f.write(f"{tipo},{linea}\n")
                    else:
                        f.write(f"=== {tipo.upper()} ===\n")
                        texto = mensaje.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
                        import re
                        texto = re.sub(r'<[^>]+>', '', texto)
                        f.write(texto)
                        f.write("\n\n")
            QMessageBox.information(self, "Exportar resultados", f"Resultados exportados en: {ruta}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron exportar los resultados:\n{e}")

def exportar_logica_pdf(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.utils import ImageReader
        except Exception:
            title = traducir('Exportar lógica a PDF')
            msg = (
                "La librería reportlab no está instalada. Instálala para generar archivos PDF."
                if CURRENT_LANGUAGE == 'es' else
                "The reportlab library is not installed. Install it to generate PDF files."
            )
            QMessageBox.critical(self, title, msg)
            return
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a PDF'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a PDF'),
            "",
            "PDF (*.pdf)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".pdf"):
            ruta = f"{ruta}.pdf"
        try:
            c = canvas.Canvas(ruta, pagesize=letter)
            try:
                from PyQt5.QtGui import QImage, QPainter
                from PyQt5.QtCore import QRectF, Qt as _Qt
                import tempfile, os
                rect = self.scene.itemsBoundingRect()
                img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
                img.fill(_Qt.white)
                painter = QPainter(img)
                self.scene.render(painter, target=QRectF(img.rect()), source=rect)
                painter.end()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    img_path = tmp_img.name
                    img.save(img_path)
                c.setFont("Helvetica-Bold", 14)
                circ_label = "Circuito" if CURRENT_LANGUAGE == 'es' else "Circuit"
                c.drawString(50, 750, circ_label)
                try:
                    c.drawImage(ImageReader(img_path), 50, 450, width=500, preserveAspectRatio=True)
                except Exception:
                    pass
                c.showPage()
                try:
                    os.unlink(img_path)
                except Exception:
                    pass
            except Exception:
                c.showPage()
            c.setFont("Helvetica", 12)
            title_text = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
            c.drawString(50, 750, title_text)
            data = [columns]
            for r in rows:
                data.append([r.get(col, 0) for col in columns])
            col_width = 500 // len(columns) if columns else 500
            table = Table(data, colWidths=[col_width] * len(columns))
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d9d9d9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ]))
            w, h = table.wrapOn(c, 480, 600)
            table.drawOn(c, 50, 750 - 30 - h)
            y_position = 750 - 30 - h - 20
            try:
                functions_info = api.expresiones_booleanas_digitales()
            except Exception:
                functions_info = {}
            if functions_info:
                label_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                c.drawString(50, y_position, label_funcs)
                y_position -= 15
                for out_name, exprs in functions_info.items():
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                    canonical_expr = exprs.get('canonical', '')
                    simplified_expr = exprs.get('simplified', '')
                    line_text = f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}"
                    c.drawString(60, y_position, line_text)
                    y_position -= 12
            c.save()
            success_msg = (
                f"PDF guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"PDF saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a PDF'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a PDF\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to PDF\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a PDF'), error_msg)

def exportar_logica_word(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception:
            title = traducir('Exportar lógica a Word')
            msg = (
                "La librería python‑docx no está instalada. Instálala para generar archivos DOCX."
                if CURRENT_LANGUAGE == 'es' else
                "The python‑docx library is not installed. Install it to generate DOCX files."
            )
            QMessageBox.critical(self, title, msg)
            return
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a Word'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a Word'),
            "",
            "Word Document (*.docx)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".docx"):
            ruta = f"{ruta}.docx"
        try:
            doc = Document()
            heading_text = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
            doc.add_heading(heading_text, level=1)
            try:
                from PyQt5.QtGui import QImage, QPainter
                from PyQt5.QtCore import QRectF, Qt as _Qt
                import tempfile, os
                rect = self.scene.itemsBoundingRect()
                img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
                img.fill(_Qt.white)
                painter = QPainter(img)
                self.scene.render(painter, target=QRectF(img.rect()), source=rect)
                painter.end()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    circ_path = tmp_img.name
                    img.save(circ_path)
                circ_label = "Circuito" if CURRENT_LANGUAGE == 'es' else "Circuit"
                doc.add_heading(circ_label, level=2)
                try:
                    doc.add_picture(circ_path, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph("No se pudo insertar la imagen del circuito.")
                try:
                    os.unlink(circ_path)
                except Exception:
                    pass
                doc.add_page_break()
            except Exception:
                pass
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            table = doc.add_table(rows=len(rows_sorted) + 1, cols=len(columns))
            table.style = "Table Grid"
            for j, col in enumerate(columns):
                table.cell(0, j).text = str(col)
            for i, row in enumerate(rows_sorted):
                for j, col in enumerate(columns):
                    val = row.get(col, 0)
                    table.cell(i + 1, j).text = str(val)
            try:
                functions_info = api.expresiones_booleanas_digitales()
            except Exception:
                functions_info = {}
            if functions_info:
                heading_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                doc.add_heading(heading_funcs, level=2)
                for out_name, exprs in functions_info.items():
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                    canonical_expr = exprs.get('canonical', '')
                    simplified_expr = exprs.get('simplified', '')
                    doc.add_paragraph(
                        f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}"
                    )
            doc.save(ruta)
            success_msg = (
                f"Word guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"Word saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a Word'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a Word\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to Word\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a Word'), error_msg)

def exportar_logica_txt(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a TXT'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a TXT'),
            "",
            "Archivo de texto (*.txt)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".txt"):
            ruta = f"{ruta}.txt"
        try:
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            with open(ruta, 'w', encoding='utf-8') as f:
                header = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
                f.write(header + "\n")
                f.write('\t'.join(columns) + '\n')
                for row in rows_sorted:
                    f.write('\t'.join(str(row.get(col, 0)) for col in columns) + '\n')
                try:
                    functions_info = api.expresiones_booleanas_digitales()
                except Exception:
                    functions_info = {}
                if functions_info:
                    label_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                    f.write('\n' + label_funcs + '\n')
                    for out_name, exprs in functions_info.items():
                        if CURRENT_LANGUAGE == 'en':
                            canonical_label = 'Canonical'
                            simplified_label = 'Simplified'
                        else:
                            canonical_label = 'Canónica'
                            simplified_label = 'Simplificada'
                        canonical_expr = exprs.get('canonical', '')
                        simplified_expr = exprs.get('simplified', '')
                        f.write(f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}\n")
            success_msg = (
                f"TXT guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"TXT saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a TXT'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a TXT\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to TXT\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a TXT'), error_msg)

def exportar_logica_excel(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a Excel'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a Excel'),
            "",
            "Excel (*.xlsx)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".xlsx"):
            ruta = f"{ruta}.xlsx"
        try:
            try:
                import pandas as pd
            except Exception:
                title = traducir('Exportar lógica a Excel')
                msg = (
                    "La librería pandas no está instalada. Instálala para generar archivos Excel."
                    if CURRENT_LANGUAGE == 'es' else
                    "The pandas library is not installed. Install it to generate Excel files."
                )
                QMessageBox.critical(self, title, msg)
                return
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            df = pd.DataFrame(rows_sorted)
            try:
                df = df.reindex(columns=columns)
            except Exception:
                pass
            sheet_truth = 'TablaDeVerdad' if CURRENT_LANGUAGE == 'es' else 'TruthTable'
            sheet_funcs = 'FuncionesBooleanas' if CURRENT_LANGUAGE == 'es' else 'BooleanFunctions'
            engine_name = None
            try:
                import xlsxwriter  
                engine_name = 'xlsxwriter'
            except Exception:
                try:
                    import openpyxl  
                    engine_name = 'openpyxl'
                except Exception:
                    engine_name = None
            if engine_name is None:
                title = traducir('Exportar lógica a Excel')
                msg = (
                    "No se encontró un motor de Excel (necesitas instalar 'xlsxwriter' o 'openpyxl') para generar archivos .xlsx."
                    if CURRENT_LANGUAGE == 'es' else
                    "No Excel engine found (please install 'xlsxwriter' or 'openpyxl') to generate .xlsx files."
                )
                QMessageBox.critical(self, title, msg)
                return
            writer_obj = pd.ExcelWriter(ruta, engine=engine_name)
            with writer_obj as writer:
                df.to_excel(writer, index=False, sheet_name=sheet_truth)
                try:
                    functions_info = api.expresiones_booleanas_digitales()
                except Exception:
                    functions_info = {}
                if functions_info:
                    funcs_rows: list[dict[str, str]] = []
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                        output_label = 'Output'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                        output_label = 'Salida'
                    for out_name, exprs in functions_info.items():
                        funcs_rows.append({
                            output_label: out_name,
                            canonical_label: exprs.get('canonical', ''),
                            simplified_label: exprs.get('simplified', '')
                        })
                    df2 = pd.DataFrame(funcs_rows)
                    df2.to_excel(writer, index=False, sheet_name=sheet_funcs)
            success_msg = (
                f"Excel guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"Excel saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a Excel'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a Excel\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to Excel\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a Excel'), error_msg)

def exportar_mediciones_csv(self) -> None:
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    import datetime

    try:
        no_med = not getattr(self, 'historial_mediciones', None)
        no_graf = not getattr(self, 'ultima_grafica', None)
    except Exception:
        no_med = True
        no_graf = True

    if no_med and no_graf:
        try:
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is not None:
                voltajes, corrientes_dict, _ = self.analizar_circuito(elementos, nodo_gnd)
                nodos_set = set()
                for e in elementos:
                    for n in e["nodos"]:
                        nodos_set.add(n)
                nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
                nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
                volt_labels = []
                volt_values = []
                for elem in elementos:
                    try:
                        n1, n2 = elem.get("nodos", [None, None])
                        v1 = 0.0 if n1 == nodo_gnd else voltajes[nodo_idx.get(n1, 0)]
                        v2 = 0.0 if n2 == nodo_gnd else voltajes[nodo_idx.get(n2, 0)]
                        dv = abs(v1 - v2)
                        volt_labels.append(elem.get("nombre", ""))
                        volt_values.append(dv)
                    except Exception:
                        continue
                current_labels = []
                current_values = []
                for name, val in corrientes_dict.items():
                    current_labels.append(name)
                    current_values.append(val)
                try:
                    self.ultima_grafica = (volt_labels, volt_values, current_labels, current_values)
                except Exception:
                    self.ultima_grafica = None
        except Exception:
            pass
    has_hist = bool(getattr(self, 'historial_mediciones', None))
    has_graph = False
    try:
        if getattr(self, 'ultima_grafica', None):
            vlabels, vvals, clabels, cvals = self.ultima_grafica
            has_graph = bool(vlabels) or bool(clabels)
    except Exception:
        has_graph = False
    if not has_hist and not has_graph:
        QMessageBox.information(self, "Exportar CSV", "No hay datos de mediciones ni gráficas para exportar.")
        return

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    sugerido = f"mediciones_{timestamp}.csv"
    ruta, _ = QFileDialog.getSaveFileName(
        self,
        "Guardar datos de mediciones y gráficas",
        sugerido,
        "CSV (*.csv)"
    )
    if not ruta:
        return
    try:
        import csv
        with open(ruta, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if has_hist:
                writer.writerow(["Tipo de medición", "Valor"])
                try:
                    for tipo, valor in self.historial_mediciones:
                        writer.writerow([tipo, valor])
                except Exception:
                    pass
                writer.writerow([])
            try:
                if getattr(self, 'ultima_grafica', None):
                    volt_labels, volt_values, current_labels, current_values = self.ultima_grafica
                    if volt_labels:
                        writer.writerow(["Elemento", "Voltaje (V)"])
                        for lbl, val in zip(volt_labels, volt_values):
                            writer.writerow([lbl, f"{val}"])
                        writer.writerow([])
                    if current_labels:
                        writer.writerow(["Elemento", "Corriente (A)"])
                        for lbl, val in zip(current_labels, current_values):
                            writer.writerow([lbl, f"{val}"])
                        writer.writerow([])
            except Exception:
                pass
        try:
            import os
            preferences['last_folder'] = os.path.dirname(ruta)
            guardar_preferencias()
        except Exception:
            pass
        mostrar_informacion(self, "Exportar CSV", f"Datos exportados exitosamente a:\n{ruta}")
        try:
            self.agregar_resultado("exportar mediciones", f"Mediciones exportadas a: {ruta}")
        except Exception:
            pass
        try:
            llamar_hooks("measurement_exported", ruta=ruta)
        except Exception:
            pass
    except Exception as exc:
        try:
            mostrar_error(self, "exportar mediciones CSV", exc)
        except Exception:
            pass
        mostrar_error_detallado(self, "Exportar CSV", f"No se pudo exportar a CSV:\n{exc}")

def exportar_reporte_pdf_completo(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    from PyQt5.QtGui import QImage, QPainter
    from PyQt5.QtCore import QRectF, Qt
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    import datetime
    import tempfile
    import os
    archivo, _ = QFileDialog.getSaveFileName(self, "Guardar reporte completo", "", "PDF (*.pdf)")
    if not archivo:
        return
    if not archivo.lower().endswith('.pdf'):
        archivo = f"{archivo}.pdf"
    try:
        c = canvas.Canvas(archivo, pagesize=letter)
        c.setFont("Helvetica-Bold", 20)
        c.drawCentredString(300, 700, "Reporte de Simulación")
        c.setFont("Helvetica", 12)
        fecha = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        c.drawCentredString(300, 670, f"Generado el {fecha}")
        try:
            c.drawImage(ImageReader(ICON_PATH), 200, 480, width=200, preserveAspectRatio=True)
        except Exception:
            pass
        c.showPage()
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
            img_path = tmp_img.name
            img.save(img_path)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Circuito")
        try:
            c.drawImage(ImageReader(img_path), 50, 400, width=500, preserveAspectRatio=True)
        except Exception:
            pass
        c.showPage()
        datos_bom, err_bom = self._calcular_bom()
        if err_bom:
            raise Exception(err_bom)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Lista de materiales (BOM)")
        if datos_bom:
            ncols = len(datos_bom[0])
            if ncols == 9:
                col_widths = [80, 70, 180, 50, 60, 50, 60, 60, 40]
            elif ncols == 6:
                col_widths = [90, 80, 180, 60, 60, 80]
            else:
                col_widths = [500.0 / ncols] * ncols
            table_bom = Table(datos_bom, colWidths=col_widths)
            table_bom.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            w, h = table_bom.wrapOn(c, 500, 600)
            table_bom.drawOn(c, 50, 720 - h)
        c.showPage()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Gráficas de voltaje y corriente")
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is None:
                raise Exception("No se encuentra nodo GND para graficar")
            voltajes, corrientes_fuente, _ = self.analizar_circuito(elementos, nodo_gnd)
            nodos_set_g = set()
            for e in elementos:
                for n in e["nodos"]:
                    nodos_set_g.add(n)
            nodos_sorted_g = sorted(n for n in nodos_set_g if n != nodo_gnd)
            nodo_idx_g = {n: i for i, n in enumerate(nodos_sorted_g)}
            volt_dict_g = {nodo_gnd: 0.0}
            for n, idx in nodo_idx_g.items():
                volt_dict_g[n] = voltajes[idx]
            volt_labels = []
            volt_vals = []
            curr_labels = []
            curr_vals = []
            lista_fuentes_g = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
            for elem in elementos:
                try:
                    nombre = elem.get("nombre", "")
                    tipo = elem.get("tipo", "")
                    n1, n2 = elem.get("nodos", [None, None])
                    v1 = volt_dict_g.get(n1, 0.0)
                    v2 = volt_dict_g.get(n2, 0.0)
                    dv = abs(v1 - v2)
                    volt_labels.append(nombre)
                    volt_vals.append(dv)
                    i_val = None
                    if tipo == "Resistencia":
                        try:
                            r_val = float(elem.get("valor"))
                            i_val = dv / r_val if r_val != 0 else None
                        except Exception:
                            i_val = None
                    elif tipo in ["Fuente Voltaje", "Batería"]:
                        try:
                            idxf = lista_fuentes_g.index(elem)
                            i_src = corrientes_fuente[idxf]
                            i_val = abs(i_src)
                        except Exception:
                            i_val = None
                    elif tipo == "Fuente Corriente":
                        try:
                            i_val = abs(elem.get("valor", 0.0))
                        except Exception:
                            i_val = None
                    if i_val is not None:
                        curr_labels.append(nombre)
                        curr_vals.append(i_val)
                except Exception:
                    continue
            if volt_labels or curr_labels:
                fig, axes = plt.subplots(2, 1, figsize=(8, 6))
                if volt_labels:
                    xv = list(range(len(volt_labels)))
                    axes[0].plot(xv, volt_vals, marker='o', linestyle='-', color='tab:blue')
                    axes[0].set_xticks(xv)
                    axes[0].set_xticklabels(volt_labels, rotation=45, ha='right', fontsize=6)
                    axes[0].set_ylabel('Voltaje (V)')
                    axes[0].set_title('Voltajes por elemento')
                    axes[0].grid(True, linestyle='--', alpha=0.5)
                if curr_labels:
                    xi = list(range(len(curr_labels)))
                    axes[1].plot(xi, curr_vals, marker='o', linestyle='-', color='tab:green')
                    axes[1].set_xticks(xi)
                    axes[1].set_xticklabels(curr_labels, rotation=45, ha='right', fontsize=6)
                    axes[1].set_ylabel('Corriente (A)')
                    axes[1].set_title('Corrientes por elemento')
                    axes[1].grid(True, linestyle='--', alpha=0.5)
                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_plot:
                    plot_path = tmp_plot.name
                fig.savefig(plot_path)
                plt.close(fig)
                try:
                    c.drawImage(ImageReader(plot_path), 50, 400, width=500, height=300, preserveAspectRatio=True)
                except Exception:
                    pass
            else:
                c.setFont("Helvetica", 12)
                c.drawString(50, 720, "No se pudieron generar gráficas debido a la falta de datos.")
        except Exception as gexc:
            c.setFont("Helvetica", 12)
            c.drawString(50, 720, f"No se pudieron generar gráficas: {gexc}")
        c.showPage()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Parámetros de simulación")
        datos_param = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            try:
                datos_param.append([
                    getattr(comp, "nombre", ""),
                    getattr(comp, "tipo", ""),
                    ", ".join(getattr(comp, "parametros", []))
                ])
            except Exception:
                continue
        if datos_param:
            ncols_p = len(datos_param[0])
            col_widths_p = [160, 100, 240] if ncols_p == 3 else [500.0 / ncols_p] * ncols_p
            table_params = Table(datos_param, colWidths=col_widths_p)
            table_params.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            w_p, h_p = table_params.wrapOn(c, 500, 600)
            table_params.drawOn(c, 50, 720 - h_p)
        c.save()
        try:
            if os.path.exists(img_path):
                os.unlink(img_path)
        except Exception:
            pass
        try:
            if 'plot_path' in locals() and os.path.exists(plot_path):
                os.unlink(plot_path)
        except Exception:
            pass
        QMessageBox.information(self, "Reporte PDF", f"Reporte guardado en:\n{archivo}")
    except Exception as exc:
        QMessageBox.critical(self, "Reporte PDF", f"No se pudo exportar el reporte completo:\n{exc}")

def exportar_reporte_word_completo(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    from PyQt5.QtGui import QImage, QPainter
    from PyQt5.QtCore import QRectF, Qt
    import datetime
    import tempfile
    import os

    archivo, _ = QFileDialog.getSaveFileName(
        self, "Guardar reporte completo", "", "Word Document (*.docx)"
    )
    if not archivo:
        return
    if not archivo.lower().endswith(".docx"):
        archivo = f"{archivo}.docx"

    try:
        try:
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except Exception:
            QMessageBox.critical(
                self,
                "Reporte Word",
                "No se pudo exportar el reporte completo:\n"
                "La librería python-docx no está instalada. Instálala para generar archivos DOCX.",
            )
            return

        doc = Document()
        titulo = doc.add_heading("Reporte de Simulación", level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fecha = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        par_fecha = doc.add_paragraph(f"Generado el {fecha}")
        par_fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if ICON_PATH and os.path.isfile(ICON_PATH):
            try:
                doc.add_picture(ICON_PATH, width=Inches(3))
            except Exception:
                pass
        doc.add_page_break()

        doc.add_heading("Circuito", level=2)
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_circ:
            circ_path = tmp_circ.name
            img.save(circ_path)
        try:
            doc.add_picture(circ_path, width=Inches(5.5))
        except Exception:
            doc.add_paragraph("No se pudo insertar la imagen del circuito.")
        try:
            if os.path.exists(circ_path):
                os.unlink(circ_path)
        except Exception:
            pass
        doc.add_page_break()

        doc.add_heading("Lista de materiales (BOM)", level=2)
        datos_bom, err_bom = self._calcular_bom()
        if err_bom:
            raise Exception(err_bom)
        if datos_bom:
            nrows = len(datos_bom)
            ncols = len(datos_bom[0])
            table = doc.add_table(rows=nrows, cols=ncols)
            table.style = "Table Grid"
            for i, row in enumerate(datos_bom):
                for j, val in enumerate(row):
                    cell = table.cell(i, j)
                    try:
                        cell.text = str(val)
                    except Exception:
                        cell.text = ""
                    for paragraph in cell.paragraphs:
                        paragraph.style = doc.styles['Normal']
            doc.add_page_break()
        else:
            doc.add_paragraph("No se pudo generar la lista de materiales.")

        doc.add_heading("Gráficas de voltaje y corriente", level=2)
        plot_temp_path = None
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is None:
                raise Exception("No se encuentra nodo GND para graficar")
            voltajes, corrientes_fuente, _ = self.analizar_circuito(elementos, nodo_gnd)
            nodos_set_g = set()
            for e in elementos:
                for n in e["nodos"]:
                    nodos_set_g.add(n)
            nodos_sorted_g = sorted(n for n in nodos_set_g if n != nodo_gnd)
            nodo_idx_g = {n: i for i, n in enumerate(nodos_sorted_g)}
            volt_dict_g = {nodo_gnd: 0.0}
            for n, idx in nodo_idx_g.items():
                volt_dict_g[n] = voltajes[idx]
            volt_labels = []
            volt_vals = []
            curr_labels = []
            curr_vals = []
            lista_fuentes_g = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
            for elem in elementos:
                try:
                    nombre = elem.get("nombre", "")
                    tipo = elem.get("tipo", "")
                    n1, n2 = elem.get("nodos", [None, None])
                    v1 = volt_dict_g.get(n1, 0.0)
                    v2 = volt_dict_g.get(n2, 0.0)
                    dv = abs(v1 - v2)
                    volt_labels.append(nombre)
                    volt_vals.append(dv)
                    i_val = None
                    if tipo == "Resistencia":
                        try:
                            r_val = float(elem.get("valor"))
                            i_val = dv / r_val if r_val != 0 else None
                        except Exception:
                            i_val = None
                    elif tipo in ["Fuente Voltaje", "Batería"]:
                        try:
                            idxf = lista_fuentes_g.index(elem)
                            i_src = corrientes_fuente[idxf]
                            i_val = abs(i_src)
                        except Exception:
                            i_val = None
                    elif tipo == "Fuente Corriente":
                        try:
                            i_val = abs(elem.get("valor", 0.0))
                        except Exception:
                            i_val = None
                    if i_val is not None:
                        curr_labels.append(nombre)
                        curr_vals.append(i_val)
                except Exception:
                    continue
            if volt_labels or curr_labels:
                fig, axes = plt.subplots(2, 1, figsize=(8, 6))
                if volt_labels:
                    xv = list(range(len(volt_labels)))
                    axes[0].plot(
                        xv,
                        volt_vals,
                        marker="o",
                        linestyle="-", color="tab:blue"
                    )
                    axes[0].set_xticks(xv)
                    axes[0].set_xticklabels(
                        volt_labels, rotation=45, ha="right", fontsize=6
                    )
                    axes[0].set_ylabel("Voltaje (V)")
                    axes[0].set_title("Voltajes por elemento")
                    axes[0].grid(True, linestyle="--", alpha=0.5)
                if curr_labels:
                    xi = list(range(len(curr_labels)))
                    axes[1].plot(
                        xi,
                        curr_vals,
                        marker="o",
                        linestyle="-", color="tab:green"
                    )
                    axes[1].set_xticks(xi)
                    axes[1].set_xticklabels(
                        curr_labels, rotation=45, ha="right", fontsize=6
                    )
                    axes[1].set_ylabel("Corriente (A)")
                    axes[1].set_title("Corrientes por elemento")
                    axes[1].grid(True, linestyle="--", alpha=0.5)
                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_plot:
                    plot_temp_path = tmp_plot.name
                fig.savefig(plot_temp_path)
                plt.close(fig)
                doc.add_picture(plot_temp_path, width=Inches(5.5))
            else:
                doc.add_paragraph(
                    "No se pudieron generar gráficas debido a la falta de datos."
                )
        except Exception as gexc:
            doc.add_paragraph(f"No se pudieron generar gráficas: {gexc}")
        try:
            if plot_temp_path and os.path.exists(plot_temp_path):
                os.unlink(plot_temp_path)
        except Exception:
            pass
        doc.add_page_break()

        doc.add_heading("Parámetros de simulación", level=2)
        datos_param = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            try:
                datos_param.append(
                    [
                        getattr(comp, "nombre", ""),
                        getattr(comp, "tipo", ""),
                        ", ".join(getattr(comp, "parametros", [])),
                    ]
                )
            except Exception:
                continue
        if len(datos_param) > 1:
            nrows_p = len(datos_param)
            ncols_p = len(datos_param[0])
            table_p = doc.add_table(rows=nrows_p, cols=ncols_p)
            table_p.style = "Table Grid"
            for i, row in enumerate(datos_param):
                for j, val in enumerate(row):
                    cell = table_p.cell(i, j)
                    cell.text = str(val)
                    for paragraph in cell.paragraphs:
                        paragraph.style = doc.styles['Normal']
        else:
            doc.add_paragraph("No hay parámetros registrados.")

        doc.save(archivo)
        QMessageBox.information(self, "Reporte Word", f"Reporte guardado en:\n{archivo}")
    except Exception as exc:
        QMessageBox.critical(self, "Reporte Word", f"No se pudo exportar el reporte completo:\n{exc}")

def exportar_netlist_spice(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    gnd_id = self._obtener_nodo_gnd()
    if gnd_id is None:
        QMessageBox.warning(self, "Exportar SPICE", "Debes colocar un GND en el circuito para exportar.")
        return
    nodemap = {gnd_id: "0"}
    next_num = 1

    def id_spice(node_id):
        nonlocal next_num
        if node_id in nodemap:
            return nodemap[node_id]
        nodemap[node_id] = str(next_num)
        next_num += 1
        return nodemap[node_id]

    lines = ["* ESIMulador netlist"]
    for comp in self.componentes:
        if not getattr(comp, "terminales", []):
            continue
        ctype = getattr(comp, "tipo", "")
        if ctype == "GND":
            continue
        if ctype in ["Resistencia", "Amperímetro", "Voltímetro", "Ohmímetro"] or "(eq." in comp.nombre:
            pref = "R"
        elif ctype == "Capacitor":
            pref = "C"
        elif ctype == "Bobina":
            pref = "L"
        elif ctype in ["Fuente Voltaje", "Batería", "Generador de funciones"]:
            pref = "V"
        elif ctype == "Fuente Corriente":
            pref = "I"
        else:
            pref = "R"
        nombre_spice = comp.nombre.replace(" ", "_")
        nodos = []
        for term in comp.terminales:
            try:
                nodos.append(id_spice(getattr(term, "node_group", None)))
            except Exception:
                nodos.append("0")
            if len(nodos) == 2:
                break
        if len(nodos) < 2:
            continue
        valor = "1"
        try:
            if getattr(comp, "parametros", []):
                raw = comp.parametros[0]
                if ":" in raw:
                    val = raw.split(":", 1)[1].strip()
                else:
                    val = raw.strip()
                try:
                    valor_num = interpretar_valor_prefijo(val)
                    valor = f"{valor_num}"
                except Exception:
                    valor = val
        except Exception:
            valor = "1"
        lines.append(f"{pref}{nombre_spice} {nodos[0]} {nodos[1]} {valor}")
    ruta, _ = QFileDialog.getSaveFileName(self, "Guardar netlist SPICE", "circuito.cir", "Netlist (*.cir *.sp *.spice *.txt)")
    if not ruta:
        return
    try:
        with open(ruta, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
        try:
            logger.info(f"Netlist SPICE exportado a {ruta}")
        except Exception:
            pass
    except Exception as e:
        QMessageBox.critical(self, "Exportar SPICE", f"No se pudo guardar el netlist:\n{e}")

def exportar_netlist_modelos(self) -> None:
    from PyQt5.QtWidgets import QInputDialog, QFileDialog, QMessageBox
    try:
        tipos = sorted({getattr(c, 'tipo', '') for c in self.componentes if getattr(c, 'tipo', '') and getattr(c, 'tipo', '') != 'GND'})
    except Exception:
        tipos = []
    pref_def = {
        'Resistencia': 'R',
        'Resistor': 'R',
        'Capacitor': 'C',
        'Bobina': 'L',
        'Inductor': 'L',
        'Fuente Voltaje': 'V',
        'Batería': 'V',
        'Generador de funciones': 'V',
        'Fuente Corriente': 'I',
        'Amperímetro': 'I',
        'Voltímetro': 'V',
        'Ohmímetro': 'R',
    }
    modelo_map = {}
    for tipo_item in tipos:
        valor_def = pref_def.get(tipo_item, tipo_item[:1].upper())
        try:
            valor, ok = QInputDialog.getText(self, _( "Modelo SPICE"), f"Ingrese modelo o prefijo para {tipo_item}:", text=valor_def)
        except Exception:
            valor, ok = (valor_def, True)
        if not ok or not valor:
            valor = valor_def
        modelo_map[tipo_item] = valor
    ruta, _ = QFileDialog.getSaveFileName(
        self,
        _( "Guardar netlist SPICE/Qucs"),
        "netlist.cir",
        "Netlist (*.cir *.net);;Todos los archivos (*)",
    )
    if not ruta:
        return
    try:
        netlist_lines = ["* Netlist generado por ESIMulador"]
        try:
            gnd_id = self._obtener_nodo_gnd()
        except Exception:
            gnd_id = None
        nodemap = {}
        next_num = 1
        def id_spice(node_id):
            nonlocal next_num
            if node_id is None:
                return "0"
            if node_id in nodemap:
                return nodemap[node_id]
            nodemap[node_id] = str(next_num)
            next_num += 1
            return nodemap[node_id]
        if gnd_id is not None:
            nodemap[gnd_id] = "0"
        for comp in getattr(self, 'componentes', []):
            if not getattr(comp, 'terminales', []):
                continue
            tipo_comp = getattr(comp, 'tipo', '')
            if tipo_comp == 'GND':
                continue
            pref = modelo_map.get(tipo_comp, tipo_comp[:1].upper())
            nombre_comp = getattr(comp, 'nombre', '').replace(' ', '_')
            nodos = []
            for term in getattr(comp, 'terminales', []):
                try:
                    nodos.append(id_spice(getattr(term, 'node_group', None)))
                except Exception:
                    nodos.append('0')
                if len(nodos) == 2:
                    break
            if len(nodos) < 2:
                continue
            valor = '1'
            try:
                params_list = getattr(comp, 'parametros', [])
                if params_list:
                    raw = params_list[0]
                    if ':' in raw:
                        val = raw.split(':', 1)[1].strip()
                    else:
                        val = raw.strip()
                    try:
                        valor_num = interpretar_valor_prefijo(val)
                        valor = f"{valor_num}"
                    except Exception:
                        valor = val
            except Exception:
                valor = '1'
            netlist_lines.append(f"{pref}{nombre_comp} {nodos[0]} {nodos[1]} {valor}")
        with open(ruta, 'w', encoding='utf-8') as f:
            f.write('\n'.join(netlist_lines) + '\n')
        QMessageBox.information(self, _( "Información"), _( "Proyecto exportado correctamente"))
    except Exception:
        try:
            QMessageBox.critical(self, _( "Error"), _( "La exportación del proyecto ha fallado"))
        except Exception:
            pass

def ejecutar_netlist_spice(self) -> None:
    from PyQt5.QtWidgets import QInputDialog, QFileDialog, QMessageBox
    import subprocess
    motores = ["ngspice", "qucs"]
    try:
        motor, ok = QInputDialog.getItem(self, _( "Seleccionar motor SPICE"), _( "Motor:"), motores, 0, False)
    except Exception:
        motor, ok = (motores[0], True)
    if not ok or not motor:
        return
    ruta, _ = QFileDialog.getOpenFileName(
        self,
        _( "Seleccionar netlist SPICE/Qucs"),
        "",
        "Netlist (*.cir *.net);;Todos los archivos (*)",
    )
    if not ruta:
        return
    try:
        if motor == "ngspice":
            resultado = subprocess.run([motor, "-b", ruta], capture_output=True, text=True)
        else:
            resultado = subprocess.run([motor, ruta], capture_output=True, text=True)
        salida_total = (resultado.stdout or '') + (resultado.stderr or '')
        try:
            self.resultados_simulaciones = getattr(self, 'resultados_simulaciones', [])
        except Exception:
            self.resultados_simulaciones = []
        self.resultados_simulaciones.append(("SPICE", salida_total))
        try:
            if hasattr(self, 'lista_resultados'):
                from PyQt5.QtCore import Qt
                resumen_out = salida_total.replace('\n', ' ')
                if len(resumen_out) > 60:
                    resumen_out = resumen_out[:57] + '...'
                item_res = QListWidgetItem(f"SPICE: {resumen_out}")
                item_res.setData(Qt.UserRole, ("SPICE", salida_total))
                self.lista_resultados.addItem(item_res)
        except Exception:
            pass
        QMessageBox.information(self, _( "Información"), _( "Proyecto exportado correctamente"))
    except Exception as exc:
        try:
            QMessageBox.critical(self, _( "Error"), f"No se pudo ejecutar {motor}: {exc}")
        except Exception:
            pass


def aplicar_a(clase=None):
    clase = clase or _base.CircuitSimulator
    for nombre in ['exportar_imagen','exportar_pdf','exportar_resultados','exportar_logica_pdf','exportar_logica_word','exportar_logica_txt','exportar_logica_excel','exportar_reporte_pdf_completo','exportar_reporte_word_completo','exportar_mediciones_csv','exportar_netlist_spice','exportar_netlist_modelos','ejecutar_netlist_spice','exportar_proyecto_extendido']:
        if nombre in globals():
            setattr(clase, nombre, globals()[nombre])
    if 'exportar_proyecto_esim' in globals():
        _base.exportar_proyecto_esim = exportar_proyecto_esim
    return clase
