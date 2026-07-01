from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
from PyQt5.QtCore import Qt
from typing import Any, Optional, Dict, List, Tuple

CURRENT_LANGUAGE: str = 'es'

TRANSLATIONS = {
    'Archivo': {'es': 'Archivo', 'en': 'File'},
    'Simulación': {'es': 'Simulación', 'en': 'Simulation'},
    'Edición': {'es': 'Edición', 'en': 'Edit'},
    'Herramientas': {'es': 'Herramientas', 'en': 'Tools'},
    'Plantillas': {'es': 'Plantillas', 'en': 'Templates'},
    'Insertar': {'es': 'Insertar', 'en': 'Insert'},
    'Ver': {'es': 'Ver', 'en': 'View'},
    'Ayuda': {'es': 'Ayuda', 'en': 'Help'},
    'Plugins': {'es': 'Plugins', 'en': 'Plugins'},
    'Idiomas': {'es': 'Idiomas', 'en': 'Languages'},
    'Español': {'es': 'Español', 'en': 'Spanish'},
    'Inglés': {'es': 'Inglés', 'en': 'English'},
    'Reestablecer': {'es': 'Reestablecer', 'en': 'Undo'},
    'Restablecer': {'es': 'Restablecer', 'en': 'Reset'},
    'solo números': {'es': 'solo números', 'en': 'numbers only'},
    'Entrada de parámetro': {'es': 'Entrada de parámetro', 'en': 'Parameter input'},
    'Tolerancia': {'es': 'Tolerancia', 'en': 'Tolerance'},
    'Selecciona la tolerancia (%):': {'es': 'Selecciona la tolerancia (%):', 'en': 'Select the tolerance (%):'},
    'Potencia nominal': {'es': 'Potencia nominal', 'en': 'Nominal power'},
    'Selecciona la potencia nominal:': {'es': 'Selecciona la potencia nominal:', 'en': 'Select the nominal power:'},
    'Valor del Potenciómetro': {'es': 'Valor del Potenciómetro', 'en': 'Potentiometer value'},
    'Selecciona valor comercial:': {'es': 'Selecciona valor comercial:', 'en': 'Select commercial value:'},
    'Porcentaje del cursor': {'es': 'Porcentaje del cursor', 'en': 'Cursor percentage'},
    'Ingresa el % del cursor (0-100):': {'es': 'Ingresa el % del cursor (0-100):', 'en': 'Enter the cursor percentage (0-100):'},
    'Voltaje del generador': {'es': 'Voltaje del generador', 'en': 'Generator voltage'},
    'Voltaje': {'es': 'Voltaje', 'en': 'Voltage'},
    'Voltaje (solo números):': {'es': 'Voltaje (solo números):', 'en': 'Voltage (numbers only):'},
    'Frecuencia del generador': {'es': 'Frecuencia del generador', 'en': 'Generator frequency'},
    'Frecuencia': {'es': 'Frecuencia', 'en': 'Frequency'},
    'Frecuencia (solo números):': {'es': 'Frecuencia (solo números):', 'en': 'Frequency (numbers only):'},
    'Offset DC': {'es': 'Offset DC', 'en': 'DC Offset'},
    'Offset DC (solo números):': {'es': 'Offset DC (solo números):', 'en': 'DC Offset (numbers only):'},
    'Fase': {'es': 'Fase', 'en': 'Phase'},
    'Fase (grados, solo números):': {'es': 'Fase (grados, solo números):', 'en': 'Phase (degrees, numbers only):'},
    'Forma de onda': {'es': 'Forma de onda', 'en': 'Waveform'},
    'Selecciona la forma de onda:': {'es': 'Selecciona la forma de onda:', 'en': 'Select the waveform:'},
    'Tensión de la batería': {'es': 'Tensión de la batería', 'en': 'Battery voltage'},
    'Selecciona tensión:': {'es': 'Selecciona tensión:', 'en': 'Select voltage:'},
    'Color del LED': {'es': 'Color del LED', 'en': 'LED color'},
    'Selecciona color:': {'es': 'Selecciona color:', 'en': 'Select color:'},
    'Modelo Transistor': {'es': 'Modelo Transistor', 'en': 'Transistor model'},
    'Modelo Op-Amp': {'es': 'Modelo Op-Amp', 'en': 'Op-Amp model'},
    'Selecciona modelo:': {'es': 'Selecciona modelo:', 'en': 'Select model:'},
    'Ciclo de trabajo': {'es': 'Ciclo de trabajo', 'en': 'Duty cycle'},
    'Ciclo de trabajo (%) (0-100):': {'es': 'Ciclo de trabajo (%) (0-100):', 'en': 'Duty cycle (%) (0-100):'},
    'Rehacer': {'es': 'Rehacer', 'en': 'Redo'},
    'Cortar': {'es': 'Cortar', 'en': 'Cut'},
    'Copiar': {'es': 'Copiar', 'en': 'Copy'},
    'Pegar': {'es': 'Pegar', 'en': 'Paste'},
    'Eliminar': {'es': 'Eliminar', 'en': 'Delete'},
    'Seleccionar': {'es': 'Seleccionar', 'en': 'Select'},
    'Seleccionar todo': {'es': 'Seleccionar todo', 'en': 'Select All'},
    'Rotar': {'es': 'Rotar', 'en': 'Rotate'},
    'Tierra': {'es': 'Tierra', 'en': 'Ground'},
    'Etiqueta Cable': {'es': 'Etiqueta Cable', 'en': 'Cable Label'},
    'Conexión': {'es': 'Conexión', 'en': 'Connection'},
    'Ampliar': {'es': 'Ampliar', 'en': 'Zoom In'},
    'Reducir': {'es': 'Reducir', 'en': 'Zoom Out'},
    'Ver Todo': {'es': 'Ver Todo', 'en': 'View All'},
    'Vista 1:1': {'es': 'Vista 1:1', 'en': '1:1 View'},
    'Gráfica Bode (AC)': {'es': 'Gráfica Bode (AC)', 'en': 'Bode Plot (AC)'},
    'Gráfica LGR': {'es': 'Gráfica LGR', 'en': 'LGR Plot'},
    'Calculadora de subred VSLM': {'es': 'Calculadora de subred VSLM', 'en': 'VLSM Subnet Calculator'},
    'Calculadora de resistencias': {'es': 'Calculadora de resistencias', 'en': 'Resistance Calculator'},
    'VNA View': {'es': 'VNA View', 'en': 'VNA View'},
    'Exportar Proyecto Extendido': {'es': 'Exportar Proyecto Extendido', 'en': 'Extended Project Export'},
    'Atajos de Teclado': {'es': 'Atajos de Teclado', 'en': 'Keyboard Shortcuts'},
    'Agregar Componente': {'es': 'Agregar Componente', 'en': 'Add Component'},
    'Conectar': {'es': 'Conectar', 'en': 'Connect'},
    'Simular Voltajes': {'es': 'Simular Voltajes', 'en': 'Simulate Voltages'},
    'Simular Corrientes': {'es': 'Simular Corrientes', 'en': 'Simulate Currents'},
    'Simular todo': {'es': 'Simular todo', 'en': 'Simulate All'},
    'Inspector de propiedades': {'es': 'Inspector de propiedades', 'en': 'Properties Inspector'},
    'Salir': {'es': 'Salir', 'en': 'Exit'},
    '¿Estás seguro que quieres salir?': {'es': '¿Estás seguro que quieres salir?', 'en': 'Are you sure you want to exit?'},
    'Sí': {'es': 'Sí', 'en': 'Yes'},
    'No': {'es': 'No', 'en': 'No'},
    'Modo Claro': {'es': 'Modo Claro', 'en': 'Light Mode'},
    'Modo Oscuro': {'es': 'Modo Oscuro', 'en': 'Dark Mode'},
    'Modo conexión activado': {'es': 'Modo conexión activado', 'en': 'Connection mode activated'},
    'Pestaña': {'es': 'Pestaña', 'en': 'Tab'},
    'Buscar componente...': {'es': 'Buscar componente...', 'en': 'Search component...'},
    'Atajos de teclado disponibles:': {'es': 'Atajos de teclado disponibles:', 'en': 'Available keyboard shortcuts:'},
    'Cortar componentes': {'es': 'Cortar componentes', 'en': 'Cut components'},
    'Copiar componentes': {'es': 'Copiar componentes', 'en': 'Copy components'},
    'Pegar componentes': {'es': 'Pegar componentes', 'en': 'Paste components'},
    'Eliminar selección': {'es': 'Eliminar selección', 'en': 'Delete selection'},
    'Deshacer': {'es': 'Deshacer', 'en': 'Undo'},
    'Rehacer': {'es': 'Rehacer', 'en': 'Redo'},
    'Seleccionar todo': {'es': 'Seleccionar todo', 'en': 'Select all'},
    'Acercar vista (zoom in)': {'es': 'Acercar vista (zoom in)', 'en': 'Zoom in'},
    'Alejar vista (zoom out)': {'es': 'Alejar vista (zoom out)', 'en': 'Zoom out'},
    'Ver todo': {'es': 'Ver todo', 'en': 'View all'},
    'Restablecer escala 1:1': {'es': 'Restablecer escala 1:1', 'en': 'Reset scale 1:1'},
    'Cargar circuito': {'es': 'Cargar circuito', 'en': 'Load circuit'},
    'Nueva pestaña': {'es': 'Nueva pestaña', 'en': 'New Tab'},
    'Guardar Circuito': {'es': 'Guardar Circuito', 'en': 'Save Circuit'},
    'Cargar': {'es': 'Cargar', 'en': 'Load'},
    'PDF': {'es': 'PDF', 'en': 'PDF'},
    'Txt': {'es': 'Txt', 'en': 'Txt'},
    'PNG': {'es': 'PNG', 'en': 'PNG'},
    'Lista de Materiales (BOM)': {'es': 'Lista de Materiales (BOM)', 'en': 'Bill of Materials (BOM)'},
    'Exportar datos y mediciones': {'es': 'Exportar datos y mediciones', 'en': 'Export data and measurements'},
    'Reporte PDF completo': {'es': 'Reporte PDF completo', 'en': 'Full PDF Report'},
    'Reporte Word completo': {'es': 'Reporte Word completo', 'en': 'Full Word Report'},
    'Abrir Proyecto Extendido': {'es': 'Abrir Proyecto Extendido', 'en': 'Open Extended Project'},
    'Recientes': {'es': 'Recientes', 'en': 'Recent'},
    'Exportar netlist SPICE/Qucs': {'es': 'Exportar netlist SPICE/Qucs', 'en': 'Export SPICE/Qucs netlist'},
    'Ejecutar SPICE/Qucs': {'es': 'Ejecutar SPICE/Qucs', 'en': 'Run SPICE/Qucs'},
    'Simular Voltaje': {'es': 'Simular Voltaje', 'en': 'Simulate Voltage'},
    'Simular Corriente': {'es': 'Simular Corriente', 'en': 'Simulate Current'},
    'Simular todo': {'es': 'Simular todo', 'en': 'Simulate All'},
    'Simular lógica': {'es': 'Simular lógica', 'en': 'Simulate Logic'},
    'Barrido Monte Carlo': {'es': 'Barrido Monte Carlo', 'en': 'Monte Carlo Sweep'},
    'Modo Laboratorio': {'es': 'Modo Laboratorio', 'en': 'Laboratory Mode'},
    'Guardar Medición': {'es': 'Guardar Medición', 'en': 'Save Measurement'},
    'Ver mediciones': {'es': 'Ver mediciones', 'en': 'View measurements'},
    'Medir Voltaje': {'es': 'Medir Voltaje', 'en': 'Measure Voltage'},
    'Medir Corriente': {'es': 'Medir Corriente', 'en': 'Measure Current'},
    'Medir Resistencia': {'es': 'Medir Resistencia', 'en': 'Measure Resistance'},
    'Mostrar Gráficas': {'es': 'Mostrar Gráficas', 'en': 'Show Graphs'},
    'Osciloscopio': {'es': 'Osciloscopio', 'en': 'Oscilloscope'},
    'Mediciones activas': {'es': 'Mediciones activas', 'en': 'Active Measurements'},
    'Consola de scripts': {'es': 'Consola de scripts', 'en': 'Script Console'},
    'Cantidad de bandas:': {'es': 'Cantidad de bandas:', 'en': 'Number of bands:'},
    '4 bandas': {'es': '4 bandas', 'en': '4 bands'},
    '5 bandas': {'es': '5 bandas', 'en': '5 bands'},
    '6 bandas': {'es': '6 bandas', 'en': '6 bands'},
    '1.ª banda de color': {'es': '1.ª banda de color', 'en': '1st color band'},
    '2.ª banda de color': {'es': '2.ª banda de color', 'en': '2nd color band'},
    '3.ª banda de color': {'es': '3.ª banda de color', 'en': '3rd color band'},
    'Multiplicador': {'es': 'Multiplicador', 'en': 'Multiplier'},
    'Tolerancia': {'es': 'Tolerancia', 'en': 'Tolerance'},
    'ppm': {'es': 'ppm', 'en': 'ppm'},
    'Calcular': {'es': 'Calcular', 'en': 'Calculate'},
    'Limpiar': {'es': 'Limpiar', 'en': 'Clear'},
    'Valor nominal': {'es': 'Valor nominal', 'en': 'Nominal value'},
    'Confiabilidad': {'es': 'Confiabilidad', 'en': 'Reliability'},
    'Bode de función de transferencia': {'es': 'Bode de función de transferencia', 'en': 'Bode of transfer function'},
    'Ingresa H(s) (usa s para la variable compleja):': {'es': 'Ingresa H(s) (usa s para la variable compleja):', 'en': 'Enter H(s) (use s for the complex variable):'},
    'Gráfica LGR limpia': {'es': 'Gráfica LGR limpia', 'en': 'Clean LGR plot'},
    'Ingresa la función de transferencia G(s):': {'es': 'Ingresa la función de transferencia G(s):', 'en': 'Enter the transfer function G(s):'},
    'LGR de función de transferencia': {'es': 'LGR de función de transferencia', 'en': 'LGR of transfer function'},
    'Dirección IP': {'es': 'Dirección IP', 'en': 'IP Address'},
    'Prefijo de red': {'es': 'Prefijo de red', 'en': 'Network prefix'},
    'Número de subredes': {'es': 'Número de subredes', 'en': 'Number of subnets'},
    'Subred': {'es': 'Subred', 'en': 'Subnet'},
    'Número de Hosts': {'es': 'Número de Hosts', 'en': 'Number of hosts'},
    'Nº de Hosts': {'es': 'Nº de Hosts', 'en': 'No. of hosts'},
    'IP de red': {'es': 'IP de red', 'en': 'Network IP'},
    'Prefijo': {'es': 'Prefijo', 'en': 'Prefix'},
    'Máscara': {'es': 'Máscara', 'en': 'Mask'},
    'Primer Host': {'es': 'Primer Host', 'en': 'First Host'},
    'Último Host': {'es': 'Último Host', 'en': 'Last Host'},
    'Broadcast': {'es': 'Broadcast', 'en': 'Broadcast'},
    'Resistencia': {'es': 'Resistencia', 'en': 'Resistor'},
    'Potenciómetro': {'es': 'Potenciómetro', 'en': 'Potentiometer'},
    'Push Button': {'es': 'Push Button', 'en': 'Push Button'},
    'Push Button Cerrado': {'es': 'Push Button Cerrado', 'en': 'Push Button Closed'},
    'Push Button Abierto': {'es': 'Push Button Abierto', 'en': 'Push Button Open'},
    'Fuente Voltaje': {'es': 'Fuente Voltaje', 'en': 'Voltage Source'},
    'Fuente Corriente': {'es': 'Fuente Corriente', 'en': 'Current Source'},
    'Transformador': {'es': 'Transformador', 'en': 'Transformer'},
    'Generador de funciones': {'es': 'Generador de funciones', 'en': 'Function generator'},
    'Capacitor': {'es': 'Capacitor', 'en': 'Capacitor'},
    'Bobina': {'es': 'Bobina', 'en': 'Inductor'},
    'Batería': {'es': 'Batería', 'en': 'Battery'},
    'Motor': {'es': 'Motor', 'en': 'Motor'},
    'GND': {'es': 'GND', 'en': 'GND'},
    'LED': {'es': 'LED', 'en': 'LED'},
    'Diodo Zener': {'es': 'Diodo Zener', 'en': 'Zener Diode'},
    'Mosfet': {'es': 'Mosfet', 'en': 'MOSFET'},
    'OpAmp': {'es': 'OpAmp', 'en': 'OpAmp'},
    'Transistor NPN': {'es': 'Transistor NPN', 'en': 'NPN Transistor'},
    'Transistor PNP': {'es': 'Transistor PNP', 'en': 'PNP Transistor'},
    'AND': {'es': 'AND', 'en': 'AND'},
    'OR': {'es': 'OR', 'en': 'OR'},
    'NAND': {'es': 'NAND', 'en': 'NAND'},
    'NOR': {'es': 'NOR', 'en': 'NOR'},
    'XOR': {'es': 'XOR', 'en': 'XOR'},
    'NOT': {'es': 'NOT', 'en': 'NOT'},
    'Amperímetro': {'es': 'Amperímetro', 'en': 'Ammeter'},
    'Voltímetro': {'es': 'Voltímetro', 'en': 'Voltmeter'},
    'Ohmímetro': {'es': 'Ohmímetro', 'en': 'Ohmmeter'},
    'Resultados y Mediciones': {'es': 'Resultados y Mediciones', 'en': 'Results and Measurements'},
    'Frecuencia (Hz)': {'es': 'Frecuencia (Hz)', 'en': 'Frequency (Hz)'},
    'Fase (°)': {'es': 'Fase (°)', 'en': 'Phase (°)'},
    'Valor fuera de rango': {'es': 'Valor fuera de rango', 'en': 'Value out of range'},
    'Por favor ingresa un valor entre 0 y 100.': {
        'es': 'Por favor ingresa un valor entre 0 y 100.',
        'en': 'Please enter a value between 0 and 100.'
    },
    'Valor inválido': {'es': 'Valor inválido', 'en': 'Invalid value'},
    'El valor no puede ser negativo.': {
        'es': 'El valor no puede ser negativo.',
        'en': 'The value cannot be negative.'
    },
    'Por favor ingresa un número válido para el voltaje.': {
        'es': 'Por favor ingresa un número válido para el voltaje.',
        'en': 'Please enter a valid number for the voltage.'
    },
    'El voltaje no puede ser negativo.': {
        'es': 'El voltaje no puede ser negativo.',
        'en': 'The voltage cannot be negative.'
    },
    'Por favor ingresa un número válido para la frecuencia.': {
        'es': 'Por favor ingresa un número válido para la frecuencia.',
        'en': 'Please enter a valid number for the frequency.'
    },
    'La frecuencia no puede ser negativa.': {
        'es': 'La frecuencia no puede ser negativa.',
        'en': 'The frequency cannot be negative.'
    },
    'Por favor ingresa un número válido para el offset.': {
        'es': 'Por favor ingresa un número válido para el offset.',
        'en': 'Please enter a valid number for the offset.'
    },
    'Por favor ingresa un número válido para la fase.': {
        'es': 'Por favor ingresa un número válido para la fase.',
        'en': 'Please enter a valid number for the phase.'
    },
    'Por favor ingresa un número válido para el ciclo de trabajo.': {
        'es': 'Por favor ingresa un número válido para el ciclo de trabajo.',
        'en': 'Please enter a valid number for the duty cycle.'
    },
    'El ciclo de trabajo debe estar entre 0 y 100.': {
        'es': 'El ciclo de trabajo debe estar entre 0 y 100.',
        'en': 'The duty cycle must be between 0 and 100.'
    },
    'Por favor ingresa un número válido (ej. 4.7, 100, 0.01).': {
        'es': 'Por favor ingresa un número válido (ej. 4.7, 100, 0.01).',
        'en': 'Please enter a valid number (e.g., 4.7, 100, 0.01).'
    },
    'Valor': {'es': 'Valor', 'en': 'Value'},
    'Cursor': {'es': 'Cursor', 'en': 'Cursor'},
    'Unidad': {'es': 'Unidad', 'en': 'Unit'},
    'Selecciona unidad para': {'es': 'Selecciona unidad para', 'en': 'Select unit for'},
    'Selecciona unidad para el voltaje:': {
        'es': 'Selecciona unidad para el voltaje:',
        'en': 'Select unit for the voltage:'
    },
    'Selecciona unidad para la frecuencia:': {
        'es': 'Selecciona unidad para la frecuencia:',
        'en': 'Select unit for the frequency:'
    },
    'Selecciona unidad para el offset:': {
        'es': 'Selecciona unidad para el offset:',
        'en': 'Select unit for the offset:'
    },
    'Unidad de voltaje': {'es': 'Unidad de voltaje', 'en': 'Voltage unit'},
    'Unidad de frecuencia': {'es': 'Unidad de frecuencia', 'en': 'Frequency unit'},
    'Unidad de offset': {'es': 'Unidad de offset', 'en': 'Offset unit'},
    'Voltaje del generador': {'es': 'Voltaje del generador', 'en': 'Generator voltage'},
    'Voltaje (solo números):': {'es': 'Voltaje (solo números):', 'en': 'Voltage (numbers only):'},
    'Frecuencia del generador': {'es': 'Frecuencia del generador', 'en': 'Generator frequency'},
    'Frecuencia (solo números):': {'es': 'Frecuencia (solo números):', 'en': 'Frequency (numbers only):'},
    'Offset DC': {'es': 'Offset DC', 'en': 'DC Offset'},
    'Offset DC (solo números):': {'es': 'Offset DC (solo números):', 'en': 'DC Offset (numbers only):'},
    'Fase': {'es': 'Fase', 'en': 'Phase'},
    'Fase (grados, solo números):': {'es': 'Fase (grados, solo números):', 'en': 'Phase (degrees, numbers only):'},
    'Forma de onda': {'es': 'Forma de onda', 'en': 'Waveform'},
    'Selecciona la forma de onda:': {
        'es': 'Selecciona la forma de onda:',
        'en': 'Select the waveform:'
    },
    'Ciclo de trabajo': {'es': 'Ciclo de trabajo', 'en': 'Duty cycle'},
    'Ciclo de trabajo (%) (0-100):': {
        'es': 'Ciclo de trabajo (%) (0-100):',
        'en': 'Duty cycle (%) (0-100):'
    },
    'Senoidal': {'es': 'Senoidal', 'en': 'Sine'},
    'Cuadrada': {'es': 'Cuadrada', 'en': 'Square'},
    'Triangular': {'es': 'Triangular', 'en': 'Triangular'},
    'Diente de sierra': {'es': 'Diente de sierra', 'en': 'Sawtooth'},
    'Forma': {'es': 'Forma', 'en': 'Waveform'},
    'Modelo MOSFET': {'es': 'Modelo MOSFET', 'en': 'MOSFET model'},
    'Tipo': {'es': 'Tipo', 'en': 'Type'},
    'Canal:': {'es': 'Canal:', 'en': 'Channel:'},
    'Canal N': {'es': 'Canal N', 'en': 'N-channel'},
    'Canal P': {'es': 'Canal P', 'en': 'P-channel'},
    'Tensión umbral (V):': {'es': 'Tensión umbral (V):', 'en': 'Threshold voltage (V):'},
    'Resistencia de conducción (Ω):': {
        'es': 'Resistencia de conducción (Ω):',
        'en': 'On resistance (Ω):'
    },
    'Ganancia de corriente (β):': {
        'es': 'Ganancia de corriente (β):',
        'en': 'Current gain (β):'
    },
    'Voltaje base-emisor (V):': {
        'es': 'Voltaje base-emisor (V):',
        'en': 'Base-emitter voltage (V):'
    },
    'hfe (factor de amplificación):': {
        'es': 'hfe (factor de amplificación):',
        'en': 'hfe (amplification factor):'
    },
    'Color': {'es': 'Color', 'en': 'Color'},
    'Offset': {'es': 'Offset', 'en': 'Offset'},
    'Voltaje Zener (V):': {'es': 'Voltaje Zener (V):', 'en': 'Zener voltage (V):'},
    'Corriente máxima (mA):': {'es': 'Corriente máxima (mA):', 'en': 'Maximum current (mA):'},
    'Tensión': {'es': 'Tensión', 'en': 'Voltage'},
    'Potencia': {'es': 'Potencia', 'en': 'Power'},
    'Exportar Resultados': {'es': 'Exportar Resultados', 'en': 'Export Results'},
    'Exportar Imagen': {'es': 'Exportar Imagen', 'en': 'Export Image'},
    'Exportar PDF': {'es': 'Exportar PDF', 'en': 'Export PDF'},
    'Limpiar Todo': {'es': 'Limpiar Todo', 'en': 'Clear All'},
    'Texto': {'es': 'Texto', 'en': 'Text'},
    'Exportar CSV (Med./Gráf.)': {'es': 'Exportar CSV (Med./Gráf.)', 'en': 'Export CSV (Meas./Graphs)'},
    'Unidad': {'es': 'Unidad', 'en': 'Unit'},
    'Unidad de voltaje': {'es': 'Unidad de voltaje', 'en': 'Voltage unit'},
    'Unidad de frecuencia': {'es': 'Unidad de frecuencia', 'en': 'Frequency unit'},
    'Unidad de offset': {'es': 'Unidad de offset', 'en': 'Offset unit'},
    'Selecciona unidad para': {'es': 'Selecciona unidad para', 'en': 'Select unit for'},
    'Selecciona unidad para el voltaje:': {'es': 'Selecciona unidad para el voltaje:', 'en': 'Select unit for the voltage:'},
    'Selecciona unidad para la frecuencia:': {'es': 'Selecciona unidad para la frecuencia:', 'en': 'Select unit for the frequency:'},
    'Selecciona unidad para el offset:': {'es': 'Selecciona unidad para el offset:', 'en': 'Select unit for the offset:'},
    'Selecciona unidad para Resistencia:': {'es': 'Selecciona unidad para Resistencia:', 'en': 'Select unit for Resistance:'},
    'Valor inválido': {'es': 'Valor inválido', 'en': 'Invalid value'},
    'Cargar Circuito': {'es': 'Cargar Circuito', 'en': 'Load Circuit'},
    'Mostrar Conexiones': {'es': 'Mostrar Conexiones', 'en': 'Show Connections'},
    'Inspector de Errores': {'es': 'Inspector de Errores', 'en': 'Error Inspector'},
    'Zoom +': {'es': 'Zoom +', 'en': 'Zoom +'},
    'Zoom -': {'es': 'Zoom -', 'en': 'Zoom -'},
    'Nodos explícitos': {'es': 'Nodos explícitos', 'en': 'Explicit nodes'},
    'Mostrar nodos': {'es': 'Mostrar nodos', 'en': 'Show nodes'},
    'Ocultar nodos': {'es': 'Ocultar nodos', 'en': 'Hide nodes'},
    'Etiquetas de red': {'es': 'Etiquetas de red', 'en': 'Net labels'},
    'Resaltado de red': {'es': 'Resaltado de red', 'en': 'Net highlighting'},
    'Alinear/Distribuir': {'es': 'Alinear/Distribuir', 'en': 'Align/Distribute'},
    'Alinear izquierda': {'es': 'Alinear izquierda', 'en': 'Align left'},
    'Alinear derecha': {'es': 'Alinear derecha', 'en': 'Align right'},
    'Alinear centro': {'es': 'Alinear centro', 'en': 'Align center'},
    'Distribuir horizontal': {'es': 'Distribuir horizontal', 'en': 'Distribute horizontally'},
    'Distribuir vertical': {'es': 'Distribuir vertical', 'en': 'Distribute vertically'},
    'Duplicar': {'es': 'Duplicar', 'en': 'Duplicate'},
    'Bloquear elementos': {'es': 'Bloquear elementos', 'en': 'Lock elements'},
    'Solo seleccionar cables': {'es': 'Solo seleccionar cables', 'en': 'Select only wires'},
    'Solo seleccionar componentes': {'es': 'Solo seleccionar componentes', 'en': 'Select only components'},
    'Ruteo inteligente': {'es': 'Ruteo inteligente', 'en': 'Smart routing'},
    'Puentes': {'es': 'Puentes', 'en': 'Jump-over'},
    'Función en desarrollo': {'es': 'Función en desarrollo', 'en': 'Feature under development'},
    'Chequeo eléctrico': {'es': 'Chequeo eléctrico', 'en': 'Electrical check'},
    'sin problemas.': {'es': 'sin problemas.', 'en': 'no problems.'},
    'Exportar como Imagen': {'es': 'Exportar como Imagen', 'en': 'Export Image'},
    'Acercar vista': {'es': 'Acercar vista', 'en': 'Zoom in'},
    'Alejar vista': {'es': 'Alejar vista', 'en': 'Zoom out'},
    'Mostrar y exportar la lista de materiales (BOM)': {
        'es': 'Mostrar y exportar la lista de materiales (BOM)',
        'en': 'Show and export the bill of materials (BOM)'
    },
    'Exporta datos de mediciones y gráficas a un archivo CSV con timestamp': {
        'es': 'Exporta datos de mediciones y gráficas a un archivo CSV con timestamp',
        'en': 'Export measurement and graph data to a timestamped CSV file'
    },
    'Exportar lógica': {'es': 'Exportar lógica', 'en': 'Export Logic'},
    'Exportar lógica a PDF': {'es': 'Exportar lógica a PDF', 'en': 'Export logic to PDF'},
    'Exportar lógica a Word': {'es': 'Exportar lógica a Word', 'en': 'Export logic to Word'},
    'Exportar lógica a TXT': {'es': 'Exportar lógica a TXT', 'en': 'Export logic to TXT'},
    'Exportar lógica a Excel': {'es': 'Exportar lógica a Excel', 'en': 'Export logic to Excel'},
    'Agregar etiqueta de texto al circuito': {
        'es': 'Agregar etiqueta de texto al circuito',
        'en': 'Add a text label to the circuit'
    },
    'Comparar curva simulada y medida en tiempo real': {
        'es': 'Comparar curva simulada y medida en tiempo real',
        'en': 'Compare simulated and real‑time measured curves'
    },
    'Guardar medición actual en la base de datos': {
        'es': 'Guardar medición actual en la base de datos',
        'en': 'Save the current measurement to the database'
    },
    'Mostrar historial de mediciones guardadas': {
        'es': 'Mostrar historial de mediciones guardadas',
        'en': 'Show history of saved measurements'
    },
    'Exporta el proyecto con layout, capas y netlist SPICE/Qucs': {
        'es': 'Exporta el proyecto con layout, capas y netlist SPICE/Qucs',
        'en': 'Export the project with layout, layers and SPICE/Qucs netlist'
    },
    'Exportar como imagen': {'es': 'Exportar como imagen', 'en': 'Export image'},
    'Eliminar': {'es': 'Eliminar', 'en': 'Delete'},
    'Cargar': {'es': 'Cargar', 'en': 'Load'},
    'Guardar': {'es': 'Guardar', 'en': 'Save'},
    'Restablecer': {'es': 'Restablecer', 'en': 'Reset'},

    'Fuente': {'es': 'Fuente', 'en': 'Font'},
    'Tamaño': {'es': 'Tamaño', 'en': 'Size'},
    'Negrita': {'es': 'Negrita', 'en': 'Bold'},
    'Cursiva': {'es': 'Cursiva', 'en': 'Italic'},
    'Subrayado': {'es': 'Subrayado', 'en': 'Underline'},
    'Aceptar': {'es': 'Aceptar', 'en': 'OK'},
    'Cancelar': {'es': 'Cancelar', 'en': 'Cancel'},
    'Nombre': {'es': 'Nombre', 'en': 'Name'},
    'Color etiqueta': {'es': 'Color etiqueta', 'en': 'Label color'},
    'Ver todo': {'es': 'Ver todo', 'en': 'View all'},
    'Establecer escala 1:1': {'es': 'Establecer escala 1:1', 'en': 'Set 1:1 scale'},
    'Simulando voltajes...': {'es': 'Simulando voltajes...', 'en': 'Simulating voltages...'},
    'Simulando corrientes...': {'es': 'Simulando corrientes...', 'en': 'Simulating currents...'},
    'Simulando circuito...': {'es': 'Simulando circuito...', 'en': 'Simulating circuit...'},
    'Simulación de voltajes completada': {'es': 'Simulación de voltajes completada', 'en': 'Voltage simulation completed'},
    'Simulación de corrientes completada': {'es': 'Simulación de corrientes completada', 'en': 'Current simulation completed'},
    'Simulación completa finalizada': {'es': 'Simulación completa finalizada', 'en': 'Complete simulation finished'},
    'Simulación inválida': {'es': 'Simulación inválida', 'en': 'Invalid simulation'},
    'No hay suficientes componentes conectados para simular.': {
        'es': 'No hay suficientes componentes conectados para simular.',
        'en': 'Not enough components connected to simulate.'
    },
    'Error de simulación': {'es': 'Error de simulación', 'en': 'Simulation error'},
    'No se encontró GND.\n\nAgrega un nodo de tierra (GND) para poder simular el circuito.': {
        'es': 'No se encontró GND.\n\nAgrega un nodo de tierra (GND) para poder simular el circuito.',
        'en': 'GND not found.\n\nAdd a ground node (GND) to simulate the circuit.'
    },
    'No se puede resolver el circuito: el sistema es singular. Verifica las conexiones.': {
        'es': 'No se puede resolver el circuito: el sistema es singular. Verifica las conexiones.',
        'en': 'The circuit cannot be solved: the system is singular. Check the connections.'
    },
    'Voltajes por componente:': {'es': 'Voltajes por componente:', 'en': 'Voltages per component:'},
    'Corrientes por componente:': {'es': 'Corrientes por componente:', 'en': 'Currents per component:'},
    'Resultados de Simulación': {'es': 'Resultados de Simulación', 'en': 'Simulation Results'},
    'Resultados de Corrientes': {'es': 'Resultados de Corrientes', 'en': 'Current Results'},
    'Acción restablecida (Ctrl+Z)': {
        'es': 'Acción restablecida (Ctrl+Z)',
        'en': 'Action restored (Ctrl+Z)'
    },
    'Circuito limpiado completamente.': {
        'es': 'Circuito limpiado completamente.',
        'en': 'Circuit completely cleared.'
    },
    'Circuito guardado': {'es': 'Circuito guardado', 'en': 'Circuit saved'},
    'Circuito cargado': {'es': 'Circuito cargado', 'en': 'Circuit loaded'},
}

TOOLTIPS_TRANSLATIONS = {
    'Haz clic aquí para iniciar una conexión entre terminales': {
        'es': 'Haz clic aquí para iniciar una conexión entre terminales',
        'en': 'Click here to start a connection between terminals'
    },
    'Ejecuta la simulación de voltajes en el circuito': {
        'es': 'Ejecuta la simulación de voltajes en el circuito',
        'en': 'Run the voltage simulation on the circuit'
    },
    'Ejecuta la simulación de corrientes en el circuito': {
        'es': 'Ejecuta la simulación de corrientes en el circuito',
        'en': 'Run the current simulation on the circuit'
    },
    'Simula voltajes y corrientes en una sola operación': {
        'es': 'Simula voltajes y corrientes en una sola operación',
        'en': 'Simulate voltages and currents in a single operation'
    },
    'Simula lógica digital y muestra la tabla de verdad del circuito': {
        'es': 'Simula lógica digital y muestra la tabla de verdad del circuito',
        'en': 'Simulates digital logic and displays the circuit truth table'
    },
    'Mostrar los atajos de teclado disponibles': {
        'es': 'Mostrar los atajos de teclado disponibles',
        'en': 'Show the available keyboard shortcuts'
    },
    'Agregar un nodo de tierra (GND) al circuito': {
        'es': 'Agregar un nodo de tierra (GND) al circuito',
        'en': 'Add a ground node (GND) to the circuit'
    },
    'Permite colocar una etiqueta de texto en el circuito': {
        'es': 'Permite colocar una etiqueta de texto en el circuito',
        'en': 'Allows placing a text label in the circuit'
    },
    'Iniciar una conexión entre terminales': {
        'es': 'Iniciar una conexión entre terminales',
        'en': 'Start a connection between terminals'
    },
    'Acercar la vista (zoom in)': {
        'es': 'Acercar la vista (zoom in)',
        'en': 'Zoom in'
    },
    'Alejar la vista (zoom out)': {
        'es': 'Alejar la vista (zoom out)',
        'en': 'Zoom out'
    },
    'Ajustar la vista para mostrar todos los elementos': {
        'es': 'Ajustar la vista para mostrar todos los elementos',
        'en': 'Adjust the view to show all elements'
    },
    'Abrir módulo VNA para mediciones de red': {
        'es': 'Abrir módulo VNA para mediciones de red',
        'en': 'Open VNA module for network measurements'
    },
    'Restablecer la vista a escala 1:1': {
        'es': 'Restablecer la vista a escala 1:1',
        'en': 'Reset the view to 1:1 scale'
    },
    'Ejecuta la inspección de errores de conexión': {
        'es': 'Ejecuta la inspección de errores de conexión',
        'en': 'Run the connection error inspection'
    },
    'Acercar vista': {
        'es': 'Acercar vista',
        'en': 'Zoom in'
    },
    'Alejar vista': {
        'es': 'Alejar vista',
        'en': 'Zoom out'
    },
    'Exporta datos de mediciones y gráficas a un archivo CSV con timestamp': {
        'es': 'Exporta datos de mediciones y gráficas a un archivo CSV con timestamp',
        'en': 'Export measurement and graph data to a timestamped CSV file'
    },
    'Mostrar y exportar la lista de materiales (BOM)': {
        'es': 'Mostrar y exportar la lista de materiales (BOM)',
        'en': 'Show and export the bill of materials (BOM)'
    },
    'Agregar etiqueta de texto al circuito': {
        'es': 'Agregar etiqueta de texto al circuito',
        'en': 'Add a text label to the circuit'
    },
    'Abrir módulo VNA para mediciones de red': {
        'es': 'Abrir módulo VNA para mediciones de red',
        'en': 'Open the VNA module for network measurements'
    },
    'Exporta el proyecto con layout, capas y netlist SPICE/Qucs': {
        'es': 'Exporta el proyecto con layout, capas y netlist SPICE/Qucs',
        'en': 'Export the project with layout, layers and SPICE/Qucs netlist'
    },
    'Comparar curva simulada y medida en tiempo real': {
        'es': 'Comparar curva simulada y medida en tiempo real',
        'en': 'Compare simulated and real‑time measured curves'
    },
    'Guardar medición actual en la base de datos': {
        'es': 'Guardar medición actual en la base de datos',
        'en': 'Save the current measurement to the database'
    },
    'Mostrar historial de mediciones guardadas': {
        'es': 'Mostrar historial de mediciones guardadas',
        'en': 'Show history of saved measurements'
    },
    'Exporta layout, capas, parámetros de simulación y netlist SPICE/Qucs': {
        'es': 'Exporta layout, capas, parámetros de simulación y netlist SPICE/Qucs',
        'en': 'Export layout, layers, simulation parameters and SPICE/Qucs netlist'
    },
    'Abre un proyecto extendido existente': {
        'es': 'Abre un proyecto extendido existente',
        'en': 'Open an existing extended project'
    },
    'Simula la distribución de voltajes en el circuito': {
        'es': 'Simula la distribución de voltajes en el circuito',
        'en': 'Simulate the distribution of voltages in the circuit'
    },
    'Simula la distribución de corrientes en el circuito': {
        'es': 'Simula la distribución de corrientes en el circuito',
        'en': 'Simulate the distribution of currents in the circuit'
    },
    'Simula tanto voltajes como corrientes en una sola operación': {
        'es': 'Simula tanto voltajes como corrientes en una sola operación',
        'en': 'Simulate both voltages and currents in a single operation'
    },
    'Realiza un análisis Monte Carlo de las tolerancias de los componentes': {
        'es': 'Realiza un análisis Monte Carlo de las tolerancias de los componentes',
        'en': 'Perform a Monte Carlo analysis of component tolerances'
    },
    'Genera un netlist con prefijos/modelos seleccionables': {
        'es': 'Genera un netlist con prefijos/modelos seleccionables',
        'en': 'Generate a netlist with selectable prefixes/models'
    },
    'Ejecuta el netlist en un motor SPICE/Qucs externo': {
        'es': 'Ejecuta el netlist en un motor SPICE/Qucs externo',
        'en': 'Run the netlist in an external SPICE/Qucs engine'
    }
    ,
    'Guardar Circuito': {
        'es': 'Guardar Circuito',
        'en': 'Save Circuit'
    },
    'Reestablecer': {
        'es': 'Reestablecer',
        'en': 'Undo'
    },
    'Rehacer': {
        'es': 'Rehacer',
        'en': 'Redo'
    },
    'Cargar': {
        'es': 'Cargar',
        'en': 'Load'
    },
    'Exportar Imagen': {
        'es': 'Exportar Imagen',
        'en': 'Export Image'
    },
    'Exportar PDF': {
        'es': 'Exportar PDF',
        'en': 'Export PDF'
    },
    'Reporte Word completo': {
        'es': 'Reporte Word completo',
        'en': 'Full Word Report'
    },
    'Exportar Resultados': {
        'es': 'Exportar Resultados',
        'en': 'Export Results'
    },
    'Eliminar': {
        'es': 'Eliminar',
        'en': 'Delete'
    },
    'Limpiar Todo': {
        'es': 'Limpiar Todo',
        'en': 'Clear All'
    },
    'Texto': {
        'es': 'Texto',
        'en': 'Text'
    },
    'Exportar CSV (Med./Gráf.)': {
        'es': 'Exportar CSV (Med./Gráf.)',
        'en': 'Export CSV (Meas./Graphs)'
    },
    'Lista de Materiales (BOM)': {
        'es': 'Lista de Materiales (BOM)',
        'en': 'Bill of Materials (BOM)'
    },
    'Acercar vista': {
        'es': 'Acercar vista',
        'en': 'Zoom in'
    },
    'Alejar vista': {
        'es': 'Alejar vista',
        'en': 'Zoom out'
    },
    'Inspector de Errores': {
        'es': 'Inspector de Errores',
        'en': 'Error Inspector'
    },
    'Agregar etiqueta de texto al circuito': {
        'es': 'Agregar etiqueta de texto al circuito',
        'en': 'Add a text label to the circuit'
    },
    'Atajos de Teclado': {
        'es': 'Atajos de Teclado',
        'en': 'Keyboard Shortcuts'
    }
}

def traducir(text: str) -> str:
    try:
        mapping = TRANSLATIONS.get(text)
        if mapping is None:
            for key, val in TRANSLATIONS.items():
                try:
                    if val.get('es') == text or val.get('en') == text:
                        mapping = val
                        break
                except Exception:
                    continue
        if mapping is None:
            return text
        return mapping.get(CURRENT_LANGUAGE, text)
    except Exception:
        return text

def traducir_descripcion_emergente(text: str) -> str:
    try:
        mapping = TOOLTIPS_TRANSLATIONS.get(text)
        if mapping is None:
            for key, val in TOOLTIPS_TRANSLATIONS.items():
                try:
                    if val.get('es') == text or val.get('en') == text:
                        mapping = val
                        break
                except Exception:
                    continue
        if mapping is None:
            return text
        return mapping.get(CURRENT_LANGUAGE, text)
    except Exception:
        return text

def traducir_mensaje_problema(msg: str) -> str:
    try:
        if CURRENT_LANGUAGE != 'en' or not isinstance(msg, str):
            return msg
        translations = {
            "Nodo flotante": "Floating node",
            "solo un terminal conectado": "only one terminal connected",
            "Hay": "There are",
            "fuentes de voltaje": "voltage sources",
            "en paralelo entre nodos": "in parallel between nodes",
            "Si sus valores son distintos, el sistema será singular": "If their values are different, the system will be singular",
            "Se encontraron": "Found",
            "subredes eléctricas desconectadas": "disconnected electrical subnetworks",
            "Algunos componentes no interactúan entre sí": "Some components do not interact with each other",
            "Fuentes de voltaje enfrentadas entre nodos": "Opposing voltage sources between nodes",
            "No se encontró un nodo de tierra (GND).": "No ground node (GND) was found.",
            "Agrega al menos un 'GND' conectado al circuito.": "Add at least one 'GND' connected to the circuit.",
            "Se detectaron múltiples nodos de tierra (GND).": "Multiple ground nodes (GND) were detected.",
            "Verifica que sólo exista una referencia común.": "Ensure that only a single common reference exists.",
            "No fue posible completar el chequeo eléctrico": "Unable to complete the electrical check",
        }
        translated = msg
        for es_text, en_text in translations.items():
            translated = translated.replace(es_text, en_text)
        return translated
    except Exception:
        return msg


def formatear_con_prefijo(valor, unidad="V"):
    abs_val = abs(valor)
    if abs_val == 0:
        return f"0.0000 {unidad}"
    elif abs_val < 1e-9:
        return f"{valor * 1e12:.4g} p{unidad}"
    elif abs_val < 1e-6:
        return f"{valor * 1e9:.4g} n{unidad}"
    elif abs_val < 1e-3:
        return f"{valor * 1e6:.4g} µ{unidad}"
    elif abs_val < 1:
        return f"{valor * 1e3:.4g} m{unidad}"
    elif abs_val < 1e3:
        return f"{valor:.4g} {unidad}"
    elif abs_val < 1e6:
        return f"{valor / 1e3:.4g} k{unidad}"
    elif abs_val < 1e9:
        return f"{valor / 1e6:.4g} M{unidad}"
    else:
        return f"{valor / 1e9:.4g} G{unidad}"

import os
from pathlib import Path
import sys
from pathlib import Path
import subprocess  

import sqlite3
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas 
from matplotlib.figure import Figure  
PROJECT_DIR = Path(__file__).resolve().parent
IMAGE_DIR = PROJECT_DIR / "icons"
if __name__ == "__main__":
    os.chdir(PROJECT_DIR)
ICON_PATH = str(IMAGE_DIR / "Logotipo.png")
BOM_ICON_PATH = str(IMAGE_DIR / "BOM.png")

import logging

try:
    import config  
except Exception:
    config = None

import json
try:
    import yaml  
except Exception:
    yaml = None
try:
    import toml  
except Exception:
    toml = None
try:
    import simulation_core 
except Exception:
    simulation_core = None

if simulation_core is None:
    import numpy as _np
    import math as _math

    class _SimulationCoreFallback:

        @staticmethod
        def _extraer_valor_de_parametros(param_list):
            for param in param_list or []:
                if ':' in param:
                    try:
                        _, val_str = param.split(':', 1)
                        return interpretar_valor_prefijo(val_str.strip())
                    except Exception:
                        continue
            return None

        @staticmethod
        def calcular_bode(componentes, tinp, tinm, toutp, toutm,
                          fmin=1e-3, fmax=1e5, npts=401):
            
            nodo_ids = set()
            nodo_gnd = None
            for comp in componentes:
                ctype = getattr(comp, 'tipo', None)
                if ctype == 'GND':
                    try:
                        if comp.terminales:
                            nodo_gnd = getattr(comp.terminales[0], 'node_group', None)
                    except Exception:
                        pass
                for t in getattr(comp, 'terminales', []):
                    node = getattr(t, 'node_group', None)
                    if node is not None:
                        nodo_ids.add(node)
            if nodo_gnd is None:
                nodo_gnd = 0 if 0 in nodo_ids else (min(nodo_ids) if nodo_ids else 0)
            nodos_sorted = sorted(n for n in nodo_ids if n != nodo_gnd)
            N = len(nodos_sorted)
            nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
            def _nodo_de(term):
                try:
                    return getattr(term, 'node_group', None)
                except Exception:
                    return None
            n_inp_p = _nodo_de(tinp)
            n_inp_m = _nodo_de(tinm)
            n_out_p = _nodo_de(toutp)
            n_out_m = _nodo_de(toutm)
            if n_inp_p is None or n_inp_m is None or n_out_p is None or n_out_m is None:
                return _np.array([]), _np.array([]), _np.array([])
            resistors = []
            capacitors = []
            inductors = []
            current_srcs = []
            fuentes_v = []
            op_amps = []
            for comp in componentes:
                ctype = getattr(comp, 'tipo', None)
                nodes = [getattr(t, 'node_group', None) for t in getattr(comp, 'terminales', [])]
                val = _SimulationCoreFallback._extraer_valor_de_parametros(getattr(comp, 'parametros', None))
                if val is None:
                    val = getattr(comp, 'valor', None)
                    try:
                        val = float(val) if val is not None else None
                    except Exception:
                        val = None
                if ctype == 'Resistencia':
                    if len(nodes) >= 2 and val not in (None, 0):
                        resistors.append((nodes[0], nodes[1], val))
                elif ctype == 'Fuente Corriente':
                    if len(nodes) >= 2 and val is not None:
                        current_srcs.append((nodes[0], nodes[1], val))
                elif ctype in ('Fuente Voltaje', 'Batería', 'Generador de funciones'):
                    if len(nodes) >= 2:
                        fuentes_v.append((comp, nodes, val if val is not None else 0.0))
                elif ctype == 'Capacitor':
                    if len(nodes) >= 2 and val is not None:
                        capacitors.append((nodes[0], nodes[1], val))
                elif ctype in ('Bobina', 'Inductor'):
                    if len(nodes) >= 2 and val is not None:
                        inductors.append((nodes[0], nodes[1], val))
                elif ctype == 'OpAmp':
                    if len(nodes) >= 3:
                        op_amps.append((comp, nodes))
            M = len(fuentes_v)
            O = len(op_amps)
            B_total = _np.zeros((N, M + O), dtype=complex)
            E_total = _np.zeros(M + O, dtype=complex)
            for idx_f, (comp, nodes, v_val) in enumerate(fuentes_v):
                n1, n2 = nodes[0], nodes[1]
                if n1 != nodo_gnd:
                    B_total[nodo_idx[n1], idx_f] = 1.0
                if n2 != nodo_gnd:
                    B_total[nodo_idx[n2], idx_f] = -1.0
                E_total[idx_f] = v_val if v_val is not None else 0.0
            for j, (comp, nodes) in enumerate(op_amps):
                n_out = nodes[2]
                if n_out != nodo_gnd:
                    B_total[nodo_idx[n_out], M + j] = 1.0
            f = _np.logspace(_math.log10(fmin), _math.log10(fmax), int(npts))
            w = 2.0 * _math.pi * f
            H = _np.zeros_like(f, dtype=complex)
            dim = N + M + O
            for k, omega in enumerate(w):
                G_mat = _np.zeros((N, N), dtype=complex)
                I_vec = _np.zeros(N, dtype=complex)
                for n1, n2, r_val in resistors:
                    if r_val is None or r_val == 0:
                        continue
                    g = 1.0 / r_val
                    if n1 != nodo_gnd:
                        i1 = nodo_idx[n1]; G_mat[i1, i1] += g
                    if n2 != nodo_gnd:
                        i2 = nodo_idx[n2]; G_mat[i2, i2] += g
                    if n1 != nodo_gnd and n2 != nodo_gnd:
                        i1 = nodo_idx[n1]; i2 = nodo_idx[n2]
                        G_mat[i1, i2] -= g
                        G_mat[i2, i1] -= g
                for n1, n2, c_val in capacitors:
                    if c_val is None or c_val == 0:
                        continue
                    y = 1j * omega * c_val
                    if n1 != nodo_gnd:
                        i1 = nodo_idx[n1]; G_mat[i1, i1] += y
                    if n2 != nodo_gnd:
                        i2 = nodo_idx[n2]; G_mat[i2, i2] += y
                    if n1 != nodo_gnd and n2 != nodo_gnd:
                        i1 = nodo_idx[n1]; i2 = nodo_idx[n2]
                        G_mat[i1, i2] -= y
                        G_mat[i2, i1] -= y
                for n1, n2, l_val in inductors:
                    if l_val is None:
                        continue
                    if omega == 0 or l_val == 0:
                        y = 0.0
                    else:
                        y = 1.0 / (1j * omega * l_val)
                    if n1 != nodo_gnd:
                        i1 = nodo_idx[n1]; G_mat[i1, i1] += y
                    if n2 != nodo_gnd:
                        i2 = nodo_idx[n2]; G_mat[i2, i2] += y
                    if n1 != nodo_gnd and n2 != nodo_gnd:
                        i1 = nodo_idx[n1]; i2 = nodo_idx[n2]
                        G_mat[i1, i2] -= y
                        G_mat[i2, i1] -= y
                for n1, n2, i_val in current_srcs:
                    if i_val is None:
                        continue
                    if n1 != nodo_gnd:
                        I_vec[nodo_idx[n1]] -= i_val
                    if n2 != nodo_gnd:
                        I_vec[nodo_idx[n2]] += i_val
                A_mat = _np.zeros((dim, dim), dtype=complex)
                A_mat[:N, :N] = G_mat
                if M + O > 0:
                    A_mat[:N, N:] = B_total
                    A_mat[N:, :N] = B_total.T
                z_vec = _np.zeros(dim, dtype=complex)
                z_vec[:N] = I_vec
                if M > 0:
                    z_vec[N:N + M] = E_total[:M]
                mu = 1.0e8  
                for j, (comp, nodes) in enumerate(op_amps):
                    row = N + M + j
                    A_mat[row, :] = 0.0
                    n_plus, n_minus, n_out = nodes[0], nodes[1], nodes[2]
                    if n_out != nodo_gnd and n_out in nodo_idx:
                        A_mat[row, nodo_idx[n_out]] = 1.0
                    if n_plus != nodo_gnd and n_plus in nodo_idx:
                        A_mat[row, nodo_idx[n_plus]] -= mu
                    if n_minus != nodo_gnd and n_minus in nodo_idx:
                        A_mat[row, nodo_idx[n_minus]] += mu
                try:
                    x = _np.linalg.solve(A_mat, z_vec)
                except Exception:
                    H[k] = 0.0
                    continue
                V_nodes = x[:N]
                def _voltaje(node):
                    if node == nodo_gnd or node is None or node not in nodo_idx:
                        return 0.0
                    return V_nodes[nodo_idx[node]]
                v_in = _voltaje(n_inp_p) - _voltaje(n_inp_m)
                v_out = _voltaje(n_out_p) - _voltaje(n_out_m)
                if abs(v_in) > 0:
                    H[k] = v_out / v_in
                else:
                    H[k] = 0.0
            mag_db = 20.0 * _np.log10(_np.maximum(_np.abs(H), 1e-20))
            phase_deg = _np.angle(H, deg=True)
            return f, mag_db, phase_deg

    simulation_core = _SimulationCoreFallback()

from dataclasses import dataclass, field
from copy import deepcopy

def _obtener_valor_componente(comp) -> Optional[float]:
    try:
        plist = getattr(comp, 'parametros', None)
    except Exception:
        plist = comp.get('parametros') if isinstance(comp, dict) else None
    if plist:
        try:
            for param in plist:
                if isinstance(param, str) and ':' in param:
                    _, val_str = param.split(':', 1)
                    try:
                        return interpretar_valor_prefijo(val_str.strip())
                    except Exception:
                        continue
        except Exception:
            pass
    try:
        val = getattr(comp, 'valor')
    except Exception:
        val = comp.get('valor') if isinstance(comp, dict) else None
    if isinstance(val, str):
        try:
            return interpretar_valor_prefijo(val)
        except Exception:
            pass
    try:
        return float(val) if val is not None else None
    except Exception:
        return None

def _obtener_terminales(comp) -> list:
    try:
        terms = getattr(comp, 'terminales')
    except Exception:
        terms = comp.get('terminales') if isinstance(comp, dict) else None
    if terms is None:
        try:
            nodos = comp.get('nodos')
        except Exception:
            nodos = None
        if nodos:
            class SimpleTerm:
                def __init__(self, node_group):
                    self.node_group = node_group
            return [SimpleTerm(n) for n in nodos]
    return terms or []

class SimulationAPI:
   
    def __init__(self, componentes: list):
        self.componentes = componentes

    def _recopilar_nodos(self):
        nodo_ids = set()
        nodo_gnd = None
        for comp in self.componentes:
            ctype = getattr(comp, 'tipo', None) or (comp.get('tipo') if isinstance(comp, dict) else None)
            terms = _obtener_terminales(comp)
            if ctype == 'GND':
                if terms:
                    try:
                        nodo_gnd = getattr(terms[0], 'node_group', None)
                    except Exception:
                        nodo_gnd = terms[0].get('node_group') if isinstance(terms[0], dict) else None
            for t in terms:
                try:
                    node = getattr(t, 'node_group')
                except Exception:
                    node = t.get('node_group') if isinstance(t, dict) else None
                if node is not None:
                    nodo_ids.add(node)
        if nodo_gnd is None:
            if 0 in nodo_ids:
                nodo_gnd = 0
            elif nodo_ids:
                nodo_gnd = min(nodo_ids)
            else:
                nodo_gnd = 0
        nodos_sorted = sorted(n for n in nodo_ids if n != nodo_gnd)
        nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
        return nodo_gnd, nodo_idx

    def _construir_mna(self, omega: float = 0.0):
        import numpy as _np
        nodo_gnd, nodo_idx = self._recopilar_nodos()
        N = len(nodo_idx)
        resistors = []
        capacitors = []
        inductors = []
        current_srcs = []
        voltage_srcs = []
        op_amps = []
        for comp in self.componentes:
            ctype = getattr(comp, 'tipo', None) or (comp.get('tipo') if isinstance(comp, dict) else None)
            terms = _obtener_terminales(comp)
            nodes = []
            for t in terms:
                try:
                    node = getattr(t, 'node_group')
                except Exception:
                    node = t.get('node_group') if isinstance(t, dict) else None
                nodes.append(node)
            val = _obtener_valor_componente(comp)
            if ctype == 'Resistencia' or ctype == 'Resistor' or ctype == 'Res':  
                if len(nodes) >= 2 and val not in (None, 0):
                    resistors.append((nodes[0], nodes[1], val))
            elif ctype == 'Capacitor' or ctype == 'C':
                if len(nodes) >= 2 and val is not None:
                    capacitors.append((nodes[0], nodes[1], val))
            elif ctype == 'Bobina' or ctype == 'Inductor' or ctype == 'L':
                if len(nodes) >= 2 and val is not None:
                    inductors.append((nodes[0], nodes[1], val))
            elif ctype == 'Fuente Corriente' or ctype == 'CurrentSource':
                if len(nodes) >= 2 and val is not None:
                    current_srcs.append((nodes[0], nodes[1], val))
            elif ctype in ('Fuente Voltaje', 'Batería', 'FuenteCorrienteControlada', 'VoltageSource', 'Generador de funciones'):
                if len(nodes) >= 2:
                    voltage_srcs.append((comp, nodes, val if val is not None else 0.0))
            elif ctype == 'OpAmp':
                if len(nodes) >= 3:
                    op_amps.append((comp, nodes))
        M = len(voltage_srcs)
        O = len(op_amps)
        B_total = _np.zeros((N, M + O), dtype=complex)
        E_total = _np.zeros(M, dtype=complex)
        for idx_f, (comp, nodes, v_val) in enumerate(voltage_srcs):
            n1, n2 = nodes[0], nodes[1]
            if n1 != nodo_gnd:
                B_total[nodo_idx[n1], idx_f] = 1.0
            if n2 != nodo_gnd:
                B_total[nodo_idx[n2], idx_f] = -1.0
            E_total[idx_f] = v_val
        for j, (comp, nodes) in enumerate(op_amps):
            n_out = nodes[2]
            if n_out != nodo_gnd and n_out in nodo_idx:
                B_total[nodo_idx[n_out], M + j] = 1.0
        def agregar_admitancia(mat, i1, i2, y):
            if y == 0:
                return
            if i1 != -1:
                mat[i1, i1] += y
            if i2 != -1:
                mat[i2, i2] += y
            if i1 != -1 and i2 != -1:
                mat[i1, i2] -= y
                mat[i2, i1] -= y
        G_mat = _np.zeros((N, N), dtype=complex)
        I_vec = _np.zeros(N, dtype=complex)
        for n1, n2, r_val in resistors:
            if r_val in (None, 0):
                continue
            g = 1.0 / r_val
            i1 = nodo_idx.get(n1, -1) if n1 != nodo_gnd else -1
            i2 = nodo_idx.get(n2, -1) if n2 != nodo_gnd else -1
            agregar_admitancia(G_mat, i1, i2, g)
        for n1, n2, c_val in capacitors:
            if c_val is None or c_val == 0:
                continue
            y = 1j * omega * c_val
            i1 = nodo_idx.get(n1, -1) if n1 != nodo_gnd else -1
            i2 = nodo_idx.get(n2, -1) if n2 != nodo_gnd else -1
            agregar_admitancia(G_mat, i1, i2, y)
        for n1, n2, l_val in inductors:
            if l_val is None:
                continue
            if omega == 0 or l_val == 0:
                y = 0.0
            else:
                y = 1.0 / (1j * omega * l_val)
            i1 = nodo_idx.get(n1, -1) if n1 != nodo_gnd else -1
            i2 = nodo_idx.get(n2, -1) if n2 != nodo_gnd else -1
            agregar_admitancia(G_mat, i1, i2, y)
        for n1, n2, i_val in current_srcs:
            if i_val is None:
                continue
            if n1 != nodo_gnd and n1 in nodo_idx:
                I_vec[nodo_idx[n1]] -= i_val
            if n2 != nodo_gnd and n2 in nodo_idx:
                I_vec[nodo_idx[n2]] += i_val
        dim = N + M + O
        A_mat = _np.zeros((dim, dim), dtype=complex)
        A_mat[:N, :N] = G_mat
        if M + O > 0:
            A_mat[:N, N:] = B_total
            A_mat[N:, :N] = B_total.T
        z_vec = _np.zeros(dim, dtype=complex)
        z_vec[:N] = I_vec
        if M > 0:
            z_vec[N:N + M] = E_total[:M]
        mu = 1.0e8
        for j, (comp, nodes) in enumerate(op_amps):
            row = N + M + j
            A_mat[row, :] = 0.0
            n_plus, n_minus, n_out = nodes[0], nodes[1], nodes[2]
            if n_out != nodo_gnd and n_out in nodo_idx:
                A_mat[row, nodo_idx[n_out]] = 1.0
            if n_plus != nodo_gnd and n_plus in nodo_idx:
                A_mat[row, nodo_idx[n_plus]] -= mu
            if n_minus != nodo_gnd and n_minus in nodo_idx:
                A_mat[row, nodo_idx[n_minus]] += mu
        return A_mat, z_vec, nodo_idx, nodo_gnd, M, O

    def resolver_dc(self):
        import numpy as _np
        A_mat, z_vec, nodo_idx, nodo_gnd, M, O = self._construir_mna(omega=0.0)
        try:
            x = _np.linalg.solve(A_mat, z_vec)
        except Exception:
            return {'voltages': {}, 'currents': [], 'branch_details': []}

    def tabla_verdad_digital(self) -> List[Dict[str, Dict[str, int]]]:
        model = ModelBuilder.desde_escena(self.componentes)
        LOGIC_TYPES = {"AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"}
        gates: List[_LogicGate] = []
        net_to_gate: Dict[int, _LogicGate] = {}
        for comp in model.components:
            if comp.type not in LOGIC_TYPES:
                continue
            if not comp.pins:
                continue
            input_nets = [p.net_id for p in comp.pins[:-1]]
            output_net = comp.pins[-1].net_id
            gate = _LogicGate(comp.type, input_nets, output_net)
            gates.append(gate)
            net_to_gate[output_net] = gate
        if not gates:
            return []
        used_nets = set()
        for g in gates:
            used_nets.update(g.input_nets)
            used_nets.add(g.output_net)
        input_nets: List[int] = []
        for net in used_nets:
            if net not in net_to_gate:
                net_obj = model.nets.get(net)
                if net_obj is not None and not net_obj.is_gnd:
                    input_nets.append(net)
        def red_es_entrada_de_alguna_compuerta(net_id: int) -> bool:
            for g in gates:
                if net_id in g.input_nets:
                    return True
            return False
        output_nets: List[int] = []
        for g in gates:
            if not red_es_entrada_de_alguna_compuerta(g.output_net):
                output_nets.append(g.output_net)
        input_nets = list(dict.fromkeys(sorted(input_nets)))
        output_nets = list(dict.fromkeys(sorted(output_nets)))
        input_names: Dict[int, str] = {}
        for idx, net in enumerate(input_nets):
            input_names[net] = chr(ord('A') + idx)
        output_names: Dict[int, str] = {}
        for idx, net in enumerate(output_nets):
            output_names[net] = chr(ord('Z') - idx)
        def evaluar_red(net_id: int, assignment: Dict[int, int], memo: Dict[int, int]) -> int:
            if net_id in memo:
                return memo[net_id]
            if net_id in net_to_gate:
                gate = net_to_gate[net_id]
                vals = [evaluar_red(n, assignment, memo) for n in gate.input_nets]
                value = gate.calcular(vals)
                memo[net_id] = value
                return value
            return assignment.get(net_id, 0)
        table: List[Dict[str, Dict[str, int]]] = []
        num_inputs = len(input_nets)
        total_combinations = 1 << num_inputs
        for combo in range(total_combinations):
            assignment: Dict[int, int] = {}
            for bit_idx, net in enumerate(input_nets):
                assignment[net] = (combo >> bit_idx) & 1
            memo: Dict[int, int] = {}
            outputs_values: Dict[int, int] = {}
            for net in output_nets:
                outputs_values[net] = evaluar_red(net, assignment, memo)
            row = {
                'inputs': {input_names[net]: assignment[net] for net in input_nets},
                'outputs': {output_names[net]: outputs_values[net] for net in output_nets}
            }
            table.append(row)
        return table

    def tabla_verdad_digital_extendida(self) -> tuple[list[dict[str, int]], list[str]]:
        model = ModelBuilder.desde_escena(self.componentes)
        LOGIC_TYPES = {"AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"}
        gates: list[_LogicGate] = []
        net_to_gate: dict[int, _LogicGate] = {}
        for comp in model.components:
            if comp.type not in LOGIC_TYPES:
                continue
            if not comp.pins:
                continue
            input_nets = [p.net_id for p in comp.pins[:-1]]
            output_net = comp.pins[-1].net_id
            gate = _LogicGate(comp.type, input_nets, output_net)
            gates.append(gate)
            net_to_gate[output_net] = gate
        if not gates:
            self.last_net_name_map = {}
            return [], []
        used_nets: set[int] = set()
        for g in gates:
            used_nets.update(g.input_nets)
            used_nets.add(g.output_net)
        input_candidate_nets: set[int] = set()
        for net in used_nets:
            if net not in net_to_gate:
                net_obj = model.nets.get(net)
                if net_obj is not None and not net_obj.is_gnd:
                    input_candidate_nets.add(net)
        def red_es_entrada(net_id: int) -> bool:
            for g in gates:
                if net_id in g.input_nets:
                    return True
            return False
        input_nets: list[int] = []
        seen_inputs: set[int] = set()
        for comp in model.components:
            if comp.type not in {"AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"}:
                continue
            for pin in comp.pins[:-1]:
                net = pin.net_id
                if net not in net_to_gate and net not in seen_inputs and net in input_candidate_nets:
                    seen_inputs.add(net)
                    input_nets.append(net)
        for net in input_candidate_nets:
            if net not in seen_inputs:
                input_nets.append(net)
        output_nets: list[int] = []
        for g in gates:
            if not red_es_entrada(g.output_net):
                output_nets.append(g.output_net)
        intermediate_nets: list[int] = []
        for g in gates:
            net = g.output_net
            if net in output_nets:
                continue
            intermediate_nets.append(net)
        input_names: dict[int, str] = {}
        for idx, net in enumerate(input_nets):
            input_names[net] = chr(ord('A') + idx)
        intermediate_names: dict[int, str] = {}
        next_letter_code = ord('A') + len(input_nets)
        seen_intermediate: set[int] = set()
        for comp in model.components:
            if comp.type not in {"AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"}:
                continue
            if not comp.pins:
                continue
            net = comp.pins[-1].net_id
            if net in intermediate_nets and net not in seen_intermediate:
                intermediate_names[net] = chr(next_letter_code)
                next_letter_code += 1
                seen_intermediate.add(net)
        output_names: dict[int, str] = {}
        seen_outputs: set[int] = set()
        ordered_outputs: list[int] = []
        for comp in model.components:
            if comp.type not in {"AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"}:
                continue
            if not comp.pins:
                continue
            net = comp.pins[-1].net_id
            if net in output_nets and net not in seen_outputs:
                seen_outputs.add(net)
                ordered_outputs.append(net)
        for net in output_nets:
            if net not in seen_outputs:
                ordered_outputs.append(net)
        output_nets = ordered_outputs
        for idx, net in enumerate(output_nets):
            output_names[net] = chr(ord('Z') - idx)
        net_name_map: dict[int, str] = {}
        net_name_map.update(input_names)
        net_name_map.update(intermediate_names)
        net_name_map.update(output_names)
        self.last_net_name_map = net_name_map
        currently_eval: set[int] = set()
        def evaluar_red(net_id: int, assignment: dict[int, int], memo: dict[int, int]) -> int:
            if net_id in memo:
                return memo[net_id]
            if net_id in currently_eval:
                memo[net_id] = 0
                return 0
            if net_id in net_to_gate:
                gate = net_to_gate[net_id]
                currently_eval.add(net_id)
                vals = [evaluar_red(n, assignment, memo) for n in gate.input_nets]
                currently_eval.discard(net_id)
                value = gate.calcular(vals)
                memo[net_id] = value
                return value
            return assignment.get(net_id, 0)
        num_inputs = len(input_nets)
        total_combinations = 1 << num_inputs
        rows: list[dict[str, int]] = []
        column_order: list[str] = []
        for net in input_nets:
            column_order.append(input_names[net])
        for net in intermediate_nets:
            if net in intermediate_names:
                column_order.append(intermediate_names[net])
        for net in output_nets:
            column_order.append(output_names[net])
        for combo in range(total_combinations):
            assignment: dict[int, int] = {}
            for bit_idx, net in enumerate(input_nets):
                assignment[net] = (combo >> bit_idx) & 1
            memo: dict[int, int] = {}
            row: dict[str, int] = {}
            for net in input_nets:
                name = input_names[net]
                row[name] = assignment[net]
            for g in gates:
                net = g.output_net
                if net in intermediate_names:
                    val = evaluar_red(net, assignment, memo)
                    name = intermediate_names[net]
                    row[name] = val
            for net in output_nets:
                val = evaluar_red(net, assignment, memo)
                name = output_names[net]
                row[name] = val
            rows.append(row)
        return rows, column_order

    def funciones_booleanas_digitales(self) -> Dict[str, str]:
        
        try:
            table = self.tabla_verdad_digital()
        except Exception:
            return {}
        if not table:
            return {}
        first_row = table[0]
        input_vars = list(first_row.get('inputs', {}).keys())
        output_vars = list(first_row.get('outputs', {}).keys())
        functions: Dict[str, str] = {}
        for out in output_vars:
            terms: List[str] = []
            all_zero = True
            for row in table:
                out_val = row['outputs'].get(out, 0)
                if out_val == 1:
                    all_zero = False
                    term_parts: List[str] = []
                    for var in input_vars:
                        val = row['inputs'].get(var, 0)
                        if val:
                            term_parts.append(var)
                        else:
                            term_parts.append(f"{var}'")
                    term = ''.join(term_parts) if term_parts else '1'
                    terms.append(term)
            if not terms:
                if all_zero:
                    expr = '0'
                else:
                    expr = '0'
            else:
                if len(terms) == (1 << len(input_vars)):
                    expr = '1'
                else:
                    expr = ' + '.join(terms)
            functions[out] = expr

    def expresiones_booleanas_digitales(self) -> Dict[str, Dict[str, str]]:
        try:
            table = self.tabla_verdad_digital()
        except Exception:
            return {}
        if not table:
            return {}
        first_row = table[0]
        input_vars = list(first_row.get('inputs', {}).keys())
        output_vars = list(first_row.get('outputs', {}).keys())
        try:
            import sympy
        except Exception:
            sympy = None
        result: Dict[str, Dict[str, str]] = {}
        for out in output_vars:
            terms: List[str] = []
            minterms: List[List[int]] = []
            all_zero = True
            for row in table:
                out_val = row['outputs'].get(out, 0)
                if out_val == 1:
                    all_zero = False
                    term_parts: List[str] = []
                    mt: List[int] = []
                    for var in input_vars:
                        val = row['inputs'].get(var, 0)
                        mt.append(val)
                        if val:
                            term_parts.append(var)
                        else:
                            term_parts.append(f"{var}'")
                    term = ''.join(term_parts) if term_parts else '1'
                    terms.append(term)
                    minterms.append(mt)
            if not terms:
                if all_zero:
                    canonical_expr = '0'
                else:
                    canonical_expr = '0'
            else:
                if len(terms) == (1 << len(input_vars)):
                    canonical_expr = '1'
                else:
                    canonical_expr = ' + '.join(terms)
            simplified_expr: str
            if sympy is None or len(input_vars) == 0:
                simplified_expr = canonical_expr
            else:
                try:
                    symbols = sympy.symbols(' '.join(input_vars))
                    if not isinstance(symbols, (tuple, list)):
                        symbols = (symbols,)
                    if not minterms:
                        simplified_expr = '0'
                    elif len(minterms) == (1 << len(input_vars)):
                        simplified_expr = '1'
                    else:
                        sop_expr = sympy.SOPform(list(symbols), minterms)
                        try:
                            simple_expr = sympy.simplify_logic(sop_expr, form='dnf')
                        except Exception:
                            simple_expr = sop_expr
                        def expresion_a_texto(expr) -> str:
                            if expr == sympy.true:
                                return '1'
                            if expr == sympy.false:
                                return '0'
                            if isinstance(expr, sympy.Symbol):
                                return str(expr)
                            from sympy.logic.boolalg import Not, And, Or
                            if expr.func == Not:
                                return f"{expresion_a_texto(expr.args[0])}'"
                            if expr.func == And:
                                parts: List[str] = []
                                for arg in expr.args:
                                    if arg.func == Or:
                                        parts.append(f"({expresion_a_texto(arg)})")
                                    else:
                                        parts.append(expresion_a_texto(arg))
                                return ''.join(parts)
                            if expr.func == Or:
                                parts = [expresion_a_texto(arg) for arg in expr.args]
                                return ' + '.join(parts)
                            return str(expr)
                        simplified_expr = expresion_a_texto(simple_expr)
                except Exception:
                    simplified_expr = canonical_expr
            result[out] = {'canonical': canonical_expr, 'simplified': simplified_expr}
        return result
        return functions

class _LogicGate:
    def __init__(self, gate_type: str, input_nets: List[int], output_net: int) -> None:
        self.gate_type: str = (gate_type or '').upper()
        self.input_nets: List[int] = list(input_nets or [])
        self.output_net: int = output_net

    def calcular(self, inputs: List[int]) -> int:
        normalized = [1 if val else 0 for val in inputs]
        if not normalized:
            normalized = [0]
        t = self.gate_type
        if t == "AND":
            return 1 if all(normalized) else 0
        elif t == "OR":
            return 1 if any(normalized) else 0
        elif t == "NAND":
            return 0 if all(normalized) else 1
        elif t == "NOR":
            return 0 if any(normalized) else 1
        elif t == "XOR":
            return 1 if (sum(normalized) % 2 == 1) else 0
        elif t == "XNOR":
            return 1 if (sum(normalized) % 2 == 0) else 0
        elif t == "NOT":
            return 0 if normalized[0] else 1
        else:
            return 0

        N = len(nodo_idx)
        V_nodes = x[:N]
        voltages = {}
        for node_id, idx in nodo_idx.items():
            voltages[node_id] = V_nodes[idx].real  
        voltages[nodo_gnd] = 0.0
        currents = list(x[N:N + M + O])
        branch_details = []
        for comp in self.componentes:
            ctype = getattr(comp, 'tipo', None) or (comp.get('tipo') if isinstance(comp, dict) else None)
            if ctype not in ('Resistencia', 'Resistor', 'Res'):
                continue
            terms = _obtener_terminales(comp)
            if len(terms) < 2:
                continue
            n1 = getattr(terms[0], 'node_group', None) if not isinstance(terms[0], dict) else terms[0].get('node_group')
            n2 = getattr(terms[1], 'node_group', None) if not isinstance(terms[1], dict) else terms[1].get('node_group')
            val = _obtener_valor_componente(comp)
            if val in (None, 0):
                continue
            v1 = voltages.get(n1, 0.0)
            v2 = voltages.get(n2, 0.0)
            vdrop = v1 - v2
            i = vdrop / val
            p = vdrop * i  
            branch_details.append({
                'tipo': ctype,
                'nodos': (n1, n2),
                'valor': val,
                'voltaje': vdrop,
                'corriente': i,
                'potencia': p,
            })
        return {'voltages': voltages, 'currents': currents, 'branch_details': branch_details}

    def bode(self, tinp, tinm, toutp, toutm, fmin=1e-3, fmax=1e5, npts=401):
        return simulation_core.calcular_bode(self.componentes, tinp, tinm, toutp, toutm, fmin, fmax, npts)

TEMPLATES = {
    'divisor_resistivo': {
        'descripcion': 'Divisor resistivo simple: dos resistencias en serie conectadas a una fuente de voltaje.',
        'componentes': [
            {'tipo': 'Fuente Voltaje', 'valor': 10.0, 'nodos': (1, 0)},  
            {'tipo': 'Resistencia', 'valor': 1000.0, 'nodos': (1, 2)},   
            {'tipo': 'Resistencia', 'valor': 1000.0, 'nodos': (2, 0)},   
        ],
        'parametros_sugeridos': {
            'frecuencia_bode': {'fmin': 10, 'fmax': 10000, 'npts': 100},
        },
        'medidas': {
            'voltaje_nodo2_medido': None,
            'corriente_fuente_medido': None,
        },
    },
    'filtro_rc': {
        'descripcion': 'Filtro RC paso bajo: una resistencia en serie seguida de un capacitor a tierra.',
        'componentes': [
            {'tipo': 'Fuente Voltaje', 'valor': 1.0, 'nodos': (1, 0)},    
            {'tipo': 'Resistencia', 'valor': 1000.0, 'nodos': (1, 2)},    
            {'tipo': 'Capacitor', 'valor': 1e-6, 'nodos': (2, 0)},        
        ],
        'parametros_sugeridos': {
            'frecuencia_bode': {'fmin': 1, 'fmax': 100000, 'npts': 1000},
        },
        'medidas': {
            'frecuencia_corte_medido': None,
            'ganancia_db_1khz_medido': None,
        },
    },
}

def obtener_plantilla(name: str) -> dict:
    tpl = TEMPLATES.get(name)
    if tpl is None:
        raise KeyError(f"Template '{name}' not found")
    return deepcopy(tpl)


def _casi_iguales(a: float, b: float, tol: float = 1e-2) -> bool:
    return abs(a - b) <= tol * max(1.0, abs(b))

def ejecutar_pruebas_internas(verbose: bool = True) -> bool:
    success = True
    messages = []
    tpl = obtener_plantilla('divisor_resistivo')
    api = SimulationAPI(tpl['componentes'])
    dc = api.resolver_dc()
    v = dc['voltages'].get(2, None)
    if v is None or not _casi_iguales(v, 5.0, 1e-2):
        success = False
        messages.append(f"Resistive divider: expected 5.0 V at node 2, got {v}")
    tpl2 = obtener_plantilla('filtro_rc')
    api2 = SimulationAPI(tpl2['componentes'])
    class FakeTerm:
        def __init__(self, node_group):
            self.node_group = node_group
    tinp = FakeTerm(1)
    tinm = FakeTerm(0)
    toutp = FakeTerm(2)
    toutm = FakeTerm(0)
    f, mag_db, phase_deg = api2.bode(tinp, tinm, toutp, toutm, 10, 200000, 200)
    import numpy as _np
    theoretical_fc = 1.0 / (2.0 * 3.141592653589793 * 1000.0 * 1e-6)
    idx = _np.argmin(_np.abs(_np.array(f) - theoretical_fc))
    db_at_fc = mag_db[idx]
    if not _casi_iguales(db_at_fc, -3.0, 0.5):
        success = False
        messages.append(f"RC filter: expected ≈ -3 dB at fc, got {db_at_fc:.2f} dB at {f[idx]:.2f} Hz")
    if verbose:
        if success:
            print("All internal tests passed.")
        else:
            print("Some internal tests failed:")
            for msg in messages:
                print("  -", msg)
    return success

def ejecutar_cli(args=None):
    import argparse
    import json
    parser = argparse.ArgumentParser(description="Circuit simulator CLI", prog="Prueba5.py")
    parser.add_argument('command', nargs='?', help='Subcommand to run (list-templates, show-template, simulate, run-tests)')
    parser.add_argument('argument', nargs='?', help='Argument for the subcommand (template name or file name)')
    parser.add_argument('--no-verbose', action='store_true', help='Silence test output when running tests')
    parsed = parser.parse_args(args)
    cmd = parsed.command
    if cmd is None:
        parser.print_help()
        return 0
    if cmd == 'list-templates':
        print("Available templates:")
        for name, tpl in TEMPLATES.items():
            print(f"  - {name}: {tpl['descripcion']}")
    elif cmd == 'show-template':
        if not parsed.argument:
            print("Please specify a template name.")
            return 1
        name = parsed.argument
        try:
            tpl = obtener_plantilla(name)
        except KeyError:
            print(f"Template '{name}' not found.")
            return 1
        print(f"Template '{name}':")
        print(json.dumps(tpl, indent=4, ensure_ascii=False))
    elif cmd == 'simulate':
        if not parsed.argument:
            print("Please specify a JSON file containing the circuit description.")
            return 1
        fname = parsed.argument
        try:
            with open(fname, 'r', encoding='utf-8') as f:
                comp_list = json.load(f)
        except Exception as exc:
            print(f"Error loading '{fname}': {exc}")
            return 1
        comps = []
        for comp in comp_list:
            if 'terminales' not in comp and 'nodos' in comp:
                comp = dict(comp)  
                comp['terminales'] = [{'node_group': n} for n in comp['nodos']]
            comps.append(comp)
        api = SimulationAPI(comps)
        dc = api.resolver_dc()
        print("DC voltages (V):")
        for node_id, v in sorted(dc['voltages'].items()):
            print(f"  Node {node_id}: {v:.6g}")
        print("\nResistor details:")
        for detail in dc['branch_details']:
            n1, n2 = detail['nodos']
            print(f"  {detail['tipo']} between {n1}-{n2}: V={detail['voltaje']:.6g} V, I={detail['corriente']:.6g} A, P={detail['potencia']:.6g} W")
        return 0
    elif cmd == 'run-tests':
        ok = ejecutar_pruebas_internas(verbose=not parsed.no_verbose)
        return 0 if ok else 1
    else:
        print(f"Unknown command '{cmd}'.")
        parser.print_help()
        return 1
import zipfile
from datetime import datetime
PREF_FILE = Path(__file__).resolve().parent / 'preferences.json'

def cargar_preferencias() -> dict:
    defaults = {
        'modo_oscuro': False,
        'last_folder': str(Path(__file__).resolve().parent),
        'window_size': [1100, 600],
        'window_pos': [100, 100],
        'default_values': {},
        'language': 'es',
        'recent_files': []
    }
    try:
        if PREF_FILE.exists():
            with open(PREF_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for key, val in defaults.items():
                if key not in data:
                    data[key] = val
            return data
    except Exception:
        pass
    return defaults

def guardar_preferencias(prefs: Optional[dict] = None) -> None:
    data = prefs if prefs is not None else preferences
    try:
        with open(PREF_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=4)
    except Exception:
        pass

preferences = cargar_preferencias()

LANGUAGES = {
    'es': {
        "Información": "Información",
        "Error": "Error",
        "Advertencia": "Advertencia",
        "Proyecto exportado correctamente": "Proyecto exportado correctamente",
        "La exportación del proyecto ha fallado": "La exportación del proyecto ha fallado",
        "Valor inválido": "Valor inválido",
        "Por favor ingresa un número válido (ej. 4.7, 100, 0.01).": "Por favor ingresa un número válido (ej. 4.7, 100, 0.01).",
        "El valor no puede ser negativo.": "El valor no puede ser negativo.",
        "Valor fuera de rango": "Valor fuera de rango",
        "Por favor ingresa un valor entre 0 y 100.": "Por favor ingresa un valor entre 0 y 100."
    },
    'en': {
        "Información": "Information",
        "Error": "Error",
        "Advertencia": "Warning",
        "Proyecto exportado correctamente": "Project exported successfully",
        "La exportación del proyecto ha fallado": "Project export failed",
        "Valor inválido": "Invalid value",
        "Por favor ingresa un número válido (ej. 4.7, 100, 0.01).": "Please enter a valid number (e.g. 4.7, 100, 0.01).",
        "El valor no puede ser negativo.": "The value cannot be negative.",
        "Valor fuera de rango": "Out of range",
        "Por favor ingresa un valor entre 0 y 100.": "Please enter a value between 0 and 100.",
        "Gráfica LGR limpia": "Clean LGR plot",
        "Ingresa la función de transferencia G(s):": "Enter the transfer function G(s):",
        "LGR de función de transferencia": "LGR of transfer function",
        "Bode de función de transferencia": "Bode of transfer function",
        "Ingresa H(s) (usa s para la variable compleja):": "Enter H(s) (use s for the complex variable):",
        "Rango de frecuencias": "Frequency range",
        "fmin, fmax, puntos (Hz):": "fmin, fmax, points (Hz):"
    }
}

def _(text: str) -> str:
    lang = preferences.get('language', 'es')
    try:
        global CURRENT_LANGUAGE
        prev_lang = CURRENT_LANGUAGE
        CURRENT_LANGUAGE = lang
        translated = traducir(text)
    finally:
        CURRENT_LANGUAGE = prev_lang
    if translated != text:
        return translated
    return LANGUAGES.get(lang, {}).get(text, text)

def mostrar_mensaje(parent: Optional[Any], level: str, message: str, detailed: Optional[str] = None) -> None:
    
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    msg = QMessageBox(parent)
    if level == 'info':
        msg.setIcon(QMessageBox.Information)
        title = _("Información")
    elif level == 'warning':
        msg.setIcon(QMessageBox.Warning)
        title = _("Advertencia")
    else:
        msg.setIcon(QMessageBox.Critical)
        title = _("Error")
    msg.setWindowTitle(title)
    text_translated = _(message)
    if level == 'warning':
        msg.setText(f"<span style='color:orange'>{text_translated}</span>")
        msg.setTextFormat(Qt.RichText)
    elif level == 'error':
        msg.setText(f"<span style='color:red'>{text_translated}</span>")
        msg.setTextFormat(Qt.RichText)
    else:
        msg.setText(text_translated)
    if detailed:
        msg.setDetailedText(str(detailed))
    msg.exec_()

def mostrar_informacion(parent: Optional[Any], message: str, detailed: Optional[str] = None) -> None:
    mostrar_mensaje(parent, 'info', message, detailed)

def mostrar_advertencia(parent: Optional[Any], message: str, detailed: Optional[str] = None) -> None:
    mostrar_mensaje(parent, 'warning', message, detailed)

def mostrar_error_detallado(parent: Optional[Any], message: str, detailed: Optional[str] = None) -> None:
    mostrar_mensaje(parent, 'error', message, detailed)

def mostrar_error(parent: Optional[Any], accion: str, exc: Exception) -> None:
    try:
        logger.exception("Error al %s: %s", accion, exc)
    except Exception:
        pass
    err_msg = f"{accion}: {exc}"
    try:
        mostrar_error_detallado(parent, err_msg)
    except Exception:
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        try:
            QMessageBox.critical(parent, _("Error"), err_msg)
        except Exception:
            pass

LOG_FORMAT = "%(asctime)s | %(levelname)s | %(name)s | %(message)s"
logging.basicConfig(level=logging.INFO, format=LOG_FORMAT)
logger = logging.getLogger("ESIMulador")
try:
    import pytest  
    logger.info("PyTest detectado: puedes escribir tests de prefijos/unidades y MNA.")
except Exception:
    pass

from typing import Callable, List, Dict, Any  

plugins_hooks: Dict[str, List[Callable]] = {}

loaded_plugins: List[Any] = []

plugin_enabled: Dict[str, bool] = preferences.get('enabled_plugins', {}) if isinstance(preferences.get('enabled_plugins', {}), dict) else {}

plugin_hook_registry: Dict[str, List[tuple]] = {}

_current_plugin_name: Optional[str] = None

COMPONENT_SYNONYMS = {
    "resistencia": ["r", "res", "resistor", "resistance"],
    "opamp": ["opamp", "op amp", "operational amplifier", "amplificador operacional"],
    "bjt": ["bjt", "transistor", "npn", "pnp"],
    "condensador": ["c", "cap", "capacitor", "condensador"],
    "inductor": ["l", "ind", "coil", "inductor"],
}

def cargar_idiomas_externos() -> None:
    base_dir = Path(__file__).resolve().parent
    for lang_code in ["es", "en"]:
        archivo = base_dir / f"lang_{lang_code}.json"
        if archivo.exists():
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                lang_dict = LANGUAGES.setdefault(lang_code, {})
                lang_dict.update(data)
            except Exception:
                continue

try:
    cargar_idiomas_externos()
except Exception:
    pass

DB_PATH = Path(__file__).resolve().parent / 'mediciones.db'
CUSTOM_COMPONENTS_FILE = Path(__file__).resolve().parent / 'custom_components.json'

SCRIPT_EXAMPLES = {
    "Listar componentes": (
        "for c in componentes:\n"
        "    print(f'{c.nombre}: {c.tipo}')\n"
    ),
    "Contar tipos de componentes": (
        "from collections import Counter\n"
        "tipos = Counter(c.tipo for c in componentes)\n"
        "for tipo, cuenta in tipos.items():\n"
        "    print(f'{tipo}: {cuenta}')\n"
    ),
    "Mostrar preferencias": (
        "import json\n"
        "print(json.dumps(preferences, indent=2, ensure_ascii=False))\n"
    ),
    "Exportar netlist SPICE": (
        "if hasattr(sim, 'exportar_netlist_modelos'):\n"
        "    print(sim.exportar_netlist_modelos())\n"
        "else:\n"
        "    print('API de exportación no disponible en esta versión.')\n"
    ),
    "Bode sencillo": (
        "try:\n"
        "    import matplotlib.pyplot as plt\n"
        "    # Selecciona el primer componente y sus terminales como ejemplo.\n"
        "    if componentes:\n"
        "        comp = componentes[0]\n"
        "        tinp = comp.terminales[0] if comp.terminales else None\n"
        "        # Usa el mismo terminal para la salida por simplicidad.\n"
        "        toutp = tinp\n"
        "        # Ejecuta un barrido de Bode usando simulation_core.\n"
        "        f, mag_db, phase_deg = simulation_core.calcular_bode(\n"
        "            componentes, tinp, None, toutp, None\n"
        "        )\n"
        "        plt.semilogx(f, mag_db)\n"
        "        plt.xlabel('Frecuencia (Hz)')\n"
        "        plt.ylabel('Magnitud (dB)')\n"
        "        plt.title('Bode de magnitud')\n"
        "        plt.grid(True)\n"
        "        plt.show()\n"
        "    else:\n"
        "        print('No hay componentes para analizar.')\n"
    ),
}

def inicializar_bd() -> None:
    try:
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute(
            """
            CREATE TABLE IF NOT EXISTS mediciones (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                fecha TEXT,
                equipo TEXT,
                circuito TEXT,
                parametros TEXT
            )
            """
        )
        conn.commit()
        conn.close()
    except Exception as exc:
        try:
            logger.exception(f"Error inicializando base de datos de mediciones: {exc}")
        except Exception:
            pass

def cargar_componentes_personalizados() -> None:
    if CUSTOM_COMPONENTS_FILE.exists():
        try:
            with open(CUSTOM_COMPONENTS_FILE, 'r', encoding='utf-8') as f:
                data = json.load(f)
            for name in data.keys():
                COMPONENT_SYNONYMS.setdefault(name.lower(), [name.lower()])
        except Exception as exc:
            try:
                logger.exception(f"Error cargando componentes personalizados: {exc}")
            except Exception:
                pass

try:
    inicializar_bd()
    cargar_componentes_personalizados()
except Exception:
    pass


def registrar_hook(hook_name: str, func: Callable) -> None:
    plugin_name = getattr(func, "_plugin_name", None) or _current_plugin_name
    if plugin_name:
        try:
            setattr(func, "_plugin_name", plugin_name)
        except Exception:
            pass
        lst = plugin_hook_registry.setdefault(plugin_name, [])
        lst.append((hook_name, func))
    hooks = plugins_hooks.setdefault(hook_name, [])
    hooks.append(func)

def llamar_hooks(hook_name: str, *args, **kwargs) -> None:
    for func in plugins_hooks.get(hook_name, []):
        try:
            func(*args, **kwargs)
        except Exception as exc:
            try:
                logger.exception(f"Plugin hook '{hook_name}' failed: {exc}")
            except Exception:
                pass

def cargar_plugins() -> None:
    import importlib
    plugin_dir = Path(__file__).resolve().parent / "plugins"
    if not plugin_dir.exists():
        return
    sys.path.insert(0, str(plugin_dir))
    for file in plugin_dir.glob("*.py"):
        if file.name.startswith("_"):
            continue
        mod_name = file.stem
        try:
            if mod_name not in plugin_enabled:
                plugin_enabled[mod_name] = True
        except Exception:
            pass
        try:
            if not plugin_enabled.get(mod_name, True):
                continue
        except Exception:
            pass
        try:
            module = importlib.import_module(mod_name)
            try:
                loaded_plugins.append(module)
            except Exception:
                pass
            if hasattr(module, "init_plugin"):
                global _current_plugin_name
                _current_plugin_name = mod_name
                try:
                    module.init_plugin(registrar_hook)
                except Exception as exc:
                    logger.exception(f"Error inicializando plugin {mod_name}: {exc}")
                finally:
                    _current_plugin_name = None
        except Exception as exc:
            logger.exception(f"Error cargando plugin {mod_name}: {exc}")

try:
    cargar_plugins()
except Exception:
    pass

class EsimulatorCore:
    @staticmethod
    def construir_modelo(componentes_scene: List[Any]) -> "CircuitModel":
        return ModelBuilder.desde_escena(componentes_scene)

    @staticmethod
    def ejecutar_comprobaciones(model: "CircuitModel") -> List["Issue"]:
        checker = EarlyElectricalChecks(model)
        return checker.ejecutar_todo()

from PyQt5.QtWidgets import (
    QListWidget,
    QListWidgetItem,
    QLineEdit,
    QApplication,
    QWidget,
    QMainWindow,
    QPushButton,
    QGraphicsScene,
    QGraphicsView,
    QGraphicsEllipseItem,
    QGraphicsItemGroup,
    QMessageBox,
    QGraphicsItem,
    QInputDialog,
    QGraphicsTextItem,
    QAction,
    QDialog,
    QTableWidget,
    QTableWidgetItem,
    QHeaderView,
    QCheckBox,
    QVBoxLayout,
    QPushButton,
    QComboBox,
    QFormLayout,
    QLabel,
    QHBoxLayout,
    QMenu,
    QDockWidget,
    QGraphicsColorizeEffect,
    QTextEdit,
    QPlainTextEdit,
    QFileDialog,
    QSpinBox,
)
from PyQt5.QtGui import QPainter, QIcon, QPen
from PyQt5.QtCore import Qt, QPointF, QLineF, QRectF, QSize, QPropertyAnimation
from PyQt5.QtSvg import QGraphicsSvgItem

from PyQt5.QtWidgets import QLabel, QVBoxLayout, QWidget, QFormLayout, QLineEdit

from dataclasses import dataclass, field
from typing import Dict, List, Tuple, Optional, Any

import ipaddress

@dataclass
class Pin:
    comp_name: str
    comp_type: str
    index: int
    net_id: int
    def clave(self) -> Tuple[str, int]:
        return (self.comp_name, self.index)

@dataclass
class Net:
    id: int
    pins: List[Pin] = field(default_factory=list)
    is_gnd: bool = False

@dataclass
class Component:
    name: str
    type: str
    pins: List[Pin] = field(default_factory=list)
    params: Dict[str, Any] = field(default_factory=dict)

@dataclass
class CircuitModel:
    components: List[Component]
    nets: Dict[int, Net]
    def encontrar_componente(self, name: str) -> Optional[Component]:
        for c in self.components:
            if c.name == name:
                return c
        return None

class ModelBuilder:
    @staticmethod
    def desde_escena(componentes_scene: List[Any]) -> CircuitModel:
        nets: Dict[int, Net] = {}
        comps: List[Component] = []
        for comp in componentes_scene:
            nombre = getattr(comp, "nombre", "Comp")
            tipo   = getattr(comp, "tipo", "Desconocido")
            params_list = list(getattr(comp, "parametros", []))
            params: Dict[str, str] = {}
            for p in params_list:
                if ":" in p:
                    k, v = p.split(":", 1)
                    params[k.strip()] = v.strip()
            pins: List[Pin] = []
            for idx, term in enumerate(getattr(comp, "terminales", [])):
                net_id = getattr(term, "node_group", None)
                if net_id is None:
                    net_id = -100000 - len(pins)
                pin = Pin(nombre, tipo, idx, net_id)
                pins.append(pin)
                if net_id not in nets:
                    nets[net_id] = Net(id=net_id, pins=[], is_gnd=False)
            component = Component(nombre, tipo, pins, params)
            comps.append(component)
            for p in pins:
                nets[p.net_id].pins.append(p)
        for comp in comps:
            if comp.type == "GND" and comp.pins:
                nid = comp.pins[0].net_id
                if nid in nets:
                    nets[nid].is_gnd = True
        return CircuitModel(components=comps, nets=nets)

@dataclass
class Issue:
    level: str  
    code: str   
    msg: str    
    details: str = ""

class EarlyElectricalChecks:
    VOLTAGE_NAMES = {"Fuente Voltaje", "Batería", "Generador de funciones"}
    CURRENT_NAMES = {"Fuente Corriente"}
    def __init__(self, model: CircuitModel):
        self.m = model
    def ejecutar_todo(self) -> List[Issue]:
        issues: List[Issue] = []
        issues += self._comprobar_presencia_gnd()
        issues += self._comprobar_nodos_flotantes()
        issues += self._comprobar_cortos_en_componentes()
        issues += self._comprobar_fuentes_voltaje_en_conflicto()
        issues += self._comprobar_fuente_en_cortocircuito()
        issues += self._comprobar_gnd_duplicado()
        issues += self._comprobar_subredes_desconectadas()
        issues += self._comprobar_fuentes_voltaje_opuestas()
        return issues

    def _comprobar_gnd_duplicado(self) -> List[Issue]:
        gnd_nets = [net.id for net in self.m.nets.values() if net.is_gnd]
        if not gnd_nets or len(set(gnd_nets)) <= 1:
            return []
        detalles = f"IDs de tierras detectadas: {', '.join(str(n) for n in sorted(set(gnd_nets)))}"
        return [Issue("warning", "MULTI_GND",
                      "Se detectaron múltiples nodos de tierra (GND). Verifica que sólo exista una referencia común.",
                      detalles)]

    def _comprobar_subredes_desconectadas(self) -> List[Issue]:
        adjacency: Dict[int, set] = {}
        for comp in self.m.components:
            nets = [p.net_id for p in comp.pins]
            for n in nets:
                adjacency.setdefault(n, set())
            for i in range(len(nets)):
                for j in range(i + 1, len(nets)):
                    n1, n2 = nets[i], nets[j]
                    adjacency[n1].add(n2)
                    adjacency[n2].add(n1)
        if not adjacency:
            return []
        visited = set()
        subnets = 0
        for net_id in adjacency.keys():
            if net_id in visited:
                continue
            subnets += 1
            stack = [net_id]
            while stack:
                curr = stack.pop()
                if curr in visited:
                    continue
                visited.add(curr)
                for neigh in adjacency.get(curr, set()):
                    if neigh not in visited:
                        stack.append(neigh)
        if subnets > 1:
            return [Issue("warning", "DISCONNECTED",
                          f"Se encontraron {subnets} subredes eléctricas desconectadas. Algunos componentes no interactúan entre sí.")]
        return []

    def _comprobar_fuentes_voltaje_opuestas(self) -> List[Issue]:
        out: List[Issue] = []
        pair_map: Dict[Tuple[int, int], List[Tuple[int, int, str]]] = {}
        for c in self.m.components:
            if c.type in self.VOLTAGE_NAMES and len(c.pins) >= 2:
                a = c.pins[0].net_id
                b = c.pins[1].net_id
                key = tuple(sorted((a, b)))
                pair_map.setdefault(key, []).append((a, b, c.name))
        for key, lst in pair_map.items():
            for i in range(len(lst)):
                for j in range(i + 1, len(lst)):
                    a1, b1, name1 = lst[i]
                    a2, b2, name2 = lst[j]
                    if a1 == b2 and b1 == a2:
                        out.append(Issue("warning", "OPPOSING_VSRC",
                                         f"Fuentes de voltaje enfrentadas entre nodos {key}: {name1} y {name2}.",))
                        break
        return out
    def _comprobar_presencia_gnd(self) -> List[Issue]:
        if any(n.is_gnd for n in self.m.nets.values()):
            return []
        return [Issue("error", "NO_GND",
                      "No se encontró un nodo de tierra (GND). "
                      "Agrega al menos un 'GND' conectado al circuito.")]
    def _comprobar_nodos_flotantes(self) -> List[Issue]:
        out: List[Issue] = []
        for net in self.m.nets.values():
            if net.is_gnd:
                continue
            if len(net.pins) <= 1:
                comps = ", ".join({p.comp_name for p in net.pins}) or "(sin componente)"
                out.append(Issue("warning", "FLOAT_NODE",
                                 f"Nodo flotante id={net.id}: solo un terminal conectado ({comps})."))
        return out
    def _comprobar_cortos_en_componentes(self) -> List[Issue]:
        out: List[Issue] = []
        for c in self.m.components:
            if c.type in {"Push Button Cerrado"}:
                continue
            nets_ids = [p.net_id for p in c.pins]
            if len(nets_ids) >= 2 and len(set(nets_ids)) == 1:
                out.append(Issue("error", "INTRA_SHORT",
                                 f"Corto interno en '{c.name}' ({c.type}): sus terminales "
                                 f"están en el mismo nodo {nets_ids[0]}."))
        return out
    def _comprobar_fuentes_voltaje_en_conflicto(self) -> List[Issue]:
        out: List[Issue] = []
        pair_map: Dict[Tuple[int, int], List[str]] = {}
        for c in self.m.components:
            if c.type in self.VOLTAGE_NAMES and len(c.pins) >= 2:
                a, b = c.pins[0].net_id, c.pins[1].net_id
                key = tuple(sorted((a, b)))
                pair_map.setdefault(key, []).append(c.name)
        for key, lst in pair_map.items():
            if len(lst) >= 2:
                out.append(Issue("warning", "PAR_VSRC",
                                 f"Hay {len(lst)} fuentes de voltaje en paralelo entre nodos {key}: {', '.join(lst)}. "
                                 "Si sus valores son distintos, el sistema será singular."))
        return out
    def _comprobar_fuente_en_cortocircuito(self) -> List[Issue]:
        out: List[Issue] = []
        for c in self.m.components:
            if c.type in self.VOLTAGE_NAMES and len(c.pins) >= 2:
                a, b = c.pins[0].net_id, c.pins[1].net_id
                if a == b:
                    out.append(Issue("error", "VSRC_SHORT",
                                     f"La fuente '{c.name}' está cortocircuitada (ambos pines en el nodo {a})."))
        return out

def problemas_a_html(issues: List[Issue]) -> str:
    try:
        if not issues:
            header = traducir('Chequeo eléctrico')
            ok_msg = traducir('sin problemas.')
            return f"<b>{header}:</b> {ok_msg}"
        lines: List[str] = []
        level_emoji = {"error": "⛔", "warning": "⚠️", "info": "ℹ️"}
        for i in issues:
            pref = level_emoji.get(i.level, "•")
            msg_translated = traducir_mensaje_problema(i.msg)
            lines.append(f"{pref} <b>{i.code}</b>: {msg_translated}")
            if i.details:
                detail_translated = traducir_mensaje_problema(i.details)
                lines.append(f"<i>{detail_translated}</i>")
        return "<br>".join(lines)
    except Exception:
        basic_lines: List[str] = []
        level_emoji = {"error": "⛔", "warning": "⚠️", "info": "ℹ️"}
        for i in issues:
            pref = level_emoji.get(i.level, "•")
            basic_lines.append(f"{pref} <b>{i.code}</b>: {i.msg}")
            if i.details:
                basic_lines.append(f"<i>{i.details}</i>")
        return "<br>".join(basic_lines)

def estado_correcto(parent: Any, message: str) -> None:
    try:
        if hasattr(parent, "statusBar"):
            parent.statusBar().showMessage(_(message), 5000)
    except Exception:
        pass
    try:
        mostrar_informacion(parent, message)
    except Exception:
        pass

def estado_error(parent: Any, message: str) -> None:
    try:
        if hasattr(parent, "statusBar"):
            parent.statusBar().showMessage(_(message), 5000)
    except Exception:
        pass
    try:
        mostrar_error_detallado(parent, message)
    except Exception:
        pass

def exportar_proyecto_esim(nombre_archivo: str, componentes_scene: List[Any]) -> None:
    try:
        model = EsimulatorCore.construir_modelo(componentes_scene)
        comps_data = []
        for c in model.components:
            comp_entry = {
                "name": c.name,
                "type": c.type,
                "params": c.params,
                "pins": [
                    {"net_id": p.net_id, "index": p.index} for p in c.pins
                ]
            }
            comps_data.append(comp_entry)
        nets_data = []
        for n in model.nets.values():
            net_entry = {
                "id": n.id,
                "is_gnd": n.is_gnd,
                "pins": [
                    {"comp_name": p.comp_name, "index": p.index} for p in n.pins
                ]
            }
            nets_data.append(net_entry)
        data = {
            "components": comps_data,
            "nets": nets_data
        }
        bom: Dict[str, int] = {}
        for comp in model.components:
            bom[comp.type] = bom.get(comp.type, 0) + 1
        metadata = {
            "generator": "ESIMulador",
            "datetime": datetime.now().isoformat(),
            "component_count": len(model.components),
            "net_count": len(model.nets)
        }
        with zipfile.ZipFile(nombre_archivo, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
            zf.writestr('circuit.json', json.dumps(data, indent=4, ensure_ascii=False))
            zf.writestr('bom.json', json.dumps(bom, indent=4, ensure_ascii=False))
            zf.writestr('metadata.json', json.dumps(metadata, indent=4, ensure_ascii=False))
        estado_correcto(None, _("Proyecto exportado correctamente"))
        try:
            llamar_hooks("project_exported", nombre_archivo=nombre_archivo, metadata=metadata)
        except Exception:
            pass
    except Exception as exc:
        mostrar_error(None, "exportar proyecto", exc)
        try:
            logger.exception("Error al exportar proyecto .esim: %s", exc)
        except Exception:
            pass
        estado_error(None, _("La exportación del proyecto ha fallado"))

def test_formatear_con_prefijo_simple():
    assert formatear_con_prefijo(0, 'V').startswith("0.0000")
    assert 'mV' in formatear_valor_con_prefijo(1e-3)
    assert 'kV' in formatear_valor_con_prefijo(1e4)

def test_interpretar_valor_prefijo():
    assert abs(interpretar_valor_prefijo("1 k") - 1000.0) < 1e-6
    assert abs(interpretar_valor_prefijo("2.2 m") - 2.2e-3) < 1e-9
    assert abs(interpretar_valor_prefijo("5 µ") - 5e-6) < 1e-9

def test_exportar_proyecto_esim_crea_zip(tmp_path):
    class Dummy:
        pass
    r1 = Dummy()
    r1.nombre = "R1"
    r1.tipo = "Resistencia"
    r1.parametros = ["R=1k"]
    t1 = Dummy(); t1.node_group = 1
    t2 = Dummy(); t2.node_group = 2
    r1.terminales = [t1, t2]
    gnd = Dummy()
    gnd.nombre = "GND1"
    gnd.tipo = "GND"
    gnd.parametros = []
    g = Dummy(); g.node_group = 0
    gnd.terminales = [g]
    escena = [r1, gnd]
    salida = tmp_path / 'test.esim'
    exportar_proyecto_esim(str(salida), escena)
    assert salida.exists()
    with zipfile.ZipFile(str(salida), 'r') as zf:
        nombres = set(zf.namelist())
        assert 'circuit.json' in nombres
        assert 'bom.json' in nombres
        assert 'metadata.json' in nombres

def interpretar_valor_prefijo(cadena):
    prefijos = {
        'p': 1e-12,
        'n': 1e-9,
        'μ': 1e-6,
        'u': 1e-6,
        'm': 1e-3,
        '': 1,
        'k': 1e3,
        'K': 1e3,
        'M': 1e6,
        'G': 1e9
    }

    import re
    match = re.match(r"([0-9.,]+)\s*([pnumkKMGμ]*)", cadena.strip())
    if not match:
        raise ValueError("Formato no reconocido")

    numero, prefijo = match.groups()
    numero = float(numero.replace(",", "."))
    factor = prefijos.get(prefijo, 1)
    return numero * factor


def formatear_valor_con_prefijo(valor):
    unidades = [
        ("GV", 1e9), ("MV", 1e6), ("kV", 1e3), ("V", 1),
        ("mV", 1e-3), ("µV", 1e-6), ("nV", 1e-9), ("pV", 1e-12)
    ]
    for sufijo, factor in unidades:
        if abs(valor) >= factor:
            return f"{valor / factor:.4f} {sufijo}"
    return f"{valor:.4f} V"

def formatear_resistencia_con_prefijo(valor):
    unidades = [("GΩ", 1e9), ("MΩ", 1e6), ("kΩ", 1e3), ("Ω", 1),
                ("mΩ", 1e-3), ("µΩ", 1e-6), ("nΩ", 1e-9), ("pΩ", 1e-12)]
    for sufijo, factor in unidades:
        if abs(valor) >= factor:
            return f"{valor / factor:.4f} {sufijo}"
    return f"{valor:.4f} Ω"

class ResistenciasCalculatorDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle(traducir('Calculadora de resistencias'))
        self.setModal(True)
        if CURRENT_LANGUAGE == 'en':
            self.digit_colors = [
                "Black", "Brown", "Red", "Orange", "Yellow",
                "Green", "Blue", "Violet", "Gray", "White"
            ]
            self.digit_map = {color: i for i, color in enumerate(self.digit_colors)}
            self.multiplier_colors = [
                "Black", "Brown", "Red", "Orange", "Yellow",
                "Green", "Blue", "Violet", "Gray", "White",
                "Gold", "Silver"
            ]
            self.multiplier_map = {
                "Black": 1,
                "Brown": 10,
                "Red": 100,
                "Orange": 1_000,
                "Yellow": 10_000,
                "Green": 100_000,
                "Blue": 1_000_000,
                "Violet": 10_000_000,
                "Gray": 100_000_000,
                "White": 1_000_000_000,
                "Gold": 0.1,
                "Silver": 0.01
            }
            self.tolerance_colors = [
                "Brown", "Red", "Green", "Blue", "Violet",
                "Gray", "Gold", "Silver"
            ]
            self.tolerance_map = {
                "Brown": 0.01,     
                "Red": 0.02,       
                "Green": 0.005,    
                "Blue": 0.0025,    
                "Violet": 0.001,   
                "Gray": 0.0005,    
                "Gold": 0.05,      
                "Silver": 0.10     
            }
            self.ppm_colors = [
                "Brown", "Red", "Orange", "Yellow", "Green",
                "Blue", "Violet", "Gray"
            ]
            self.ppm_map = {
                "Brown": 100,
                "Red": 50,
                "Orange": 15,
                "Yellow": 25,
                "Green": 20,
                "Blue": 10,
                "Violet": 5,
                "Gray": 1
            }
        else:
            self.digit_colors = [
                "Negro", "Marrón", "Rojo", "Naranja", "Amarillo",
                "Verde", "Azul", "Violeta", "Gris", "Blanco"
            ]
            self.digit_map = {color: i for i, color in enumerate(self.digit_colors)}
            self.multiplier_colors = [
                "Negro", "Marrón", "Rojo", "Naranja", "Amarillo",
                "Verde", "Azul", "Violeta", "Gris", "Blanco",
                "Dorado", "Plateado"
            ]
            self.multiplier_map = {
                "Negro": 1,
                "Marrón": 10,
                "Rojo": 100,
                "Naranja": 1_000,
                "Amarillo": 10_000,
                "Verde": 100_000,
                "Azul": 1_000_000,
                "Violeta": 10_000_000,
                "Gris": 100_000_000,
                "Blanco": 1_000_000_000,
                "Dorado": 0.1,
                "Plateado": 0.01
            }
            self.tolerance_colors = [
                "Marrón", "Rojo", "Verde", "Azul", "Violeta",
                "Gris", "Dorado", "Plateado"
            ]
            self.tolerance_map = {
                "Marrón": 0.01,
                "Rojo": 0.02,
                "Verde": 0.005,
                "Azul": 0.0025,
                "Violeta": 0.001,
                "Gris": 0.0005,
                "Dorado": 0.05,
                "Plateado": 0.10
            }
            self.ppm_colors = [
                "Marrón", "Rojo", "Naranja", "Amarillo", "Verde",
                "Azul", "Violeta", "Gris"
            ]
            self.ppm_map = {
                "Marrón": 100,
                "Rojo": 50,
                "Naranja": 15,
                "Amarillo": 25,
                "Verde": 20,
                "Azul": 10,
                "Violeta": 5,
                "Gris": 1
            }
        main_layout = QVBoxLayout(self)

        bands_layout = QHBoxLayout()
        bands_label = QLabel(traducir('Cantidad de bandas:'))
        self.bands_combo = QComboBox()
        self.bands_combo.addItems([traducir('4 bandas'), traducir('5 bandas'), traducir('6 bandas')])
        self.bands_combo.currentIndexChanged.connect(self._actualizar_campos)
        bands_layout.addWidget(bands_label)
        bands_layout.addWidget(self.bands_combo)
        bands_layout.addStretch()
        main_layout.addLayout(bands_layout)

        self.fields_layout = QFormLayout()
        main_layout.addLayout(self.fields_layout)

        self.result_label = QLabel("")
        main_layout.addWidget(self.result_label)

        buttons_layout = QHBoxLayout()
        self.calc_button = QPushButton(traducir('Calcular'))
        self.calc_button.clicked.connect(self._calcular)
        self.clear_button = QPushButton(traducir('Limpiar'))
        self.clear_button.clicked.connect(self._limpiar_seleccion)
        buttons_layout.addWidget(self.calc_button)
        buttons_layout.addWidget(self.clear_button)
        buttons_layout.addStretch()
        main_layout.addLayout(buttons_layout)

        self._actualizar_campos()

    def _limpiar_disposicion(self, layout):
        while layout.count():
            item = layout.takeAt(0)
            widget = item.widget()
            if widget is not None:
                widget.deleteLater()

    def _actualizar_campos(self):
        self._limpiar_disposicion(self.fields_layout)
        self.band_combos = []
        nbands = 4 + self.bands_combo.currentIndex()  
        num_digit_bands = 2 if nbands == 4 else 3
        for i in range(num_digit_bands):
            combo = QComboBox()
            combo.addItems(self.digit_colors)
            self.band_combos.append(combo)
            label_text = traducir(f"{i+1}.ª banda de color")
            self.fields_layout.addRow(label_text, combo)
        self.multiplier_combo = QComboBox()
        self.multiplier_combo.addItems(self.multiplier_colors)
        self.fields_layout.addRow(traducir('Multiplicador'), self.multiplier_combo)
        self.tolerance_combo = QComboBox()
        self.tolerance_combo.addItems(self.tolerance_colors)
        self.fields_layout.addRow(traducir('Tolerancia'), self.tolerance_combo)
        if nbands == 6:
            self.ppm_combo = QComboBox()
            self.ppm_combo.addItems(self.ppm_colors)
            self.fields_layout.addRow(traducir('ppm'), self.ppm_combo)
        else:
            self.ppm_combo = None
        self.result_label.setText("")

    def _limpiar_seleccion(self):
        for combo in getattr(self, 'band_combos', []):
            combo.setCurrentIndex(0)
        if hasattr(self, 'multiplier_combo') and self.multiplier_combo:
            self.multiplier_combo.setCurrentIndex(0)
        if hasattr(self, 'tolerance_combo') and self.tolerance_combo:
            self.tolerance_combo.setCurrentIndex(0)
        if hasattr(self, 'ppm_combo') and self.ppm_combo:
            self.ppm_combo.setCurrentIndex(0)
        self.result_label.setText("")

    def _calcular(self):
        try:
            nbands = 4 + self.bands_combo.currentIndex()
            digits = [self.digit_map[combo.currentText()] for combo in self.band_combos]
            if nbands == 4:
                base = digits[0] * 10 + digits[1]
            else:
                base = digits[0] * 100 + digits[1] * 10 + digits[2]
            multiplier_value = self.multiplier_map[self.multiplier_combo.currentText()]
            resistance_value = base * multiplier_value
            tolerance_value = self.tolerance_map[self.tolerance_combo.currentText()]
            valor_str = formatear_resistencia_con_prefijo(resistance_value)
            if CURRENT_LANGUAGE == 'en':
                result_text = f"{traducir('Valor nominal')}: {valor_str}\n{traducir('Tolerancia')}: \u00b1{tolerance_value*100:.2f}%"
                if nbands == 6 and self.ppm_combo is not None:
                    ppm_value = self.ppm_map[self.ppm_combo.currentText()]
                    result_text += f"\n{traducir('Confiabilidad')}: {ppm_value} ppm/K"
            else:
                result_text = f"Valor nominal: {valor_str}\nTolerancia: \u00b1{tolerance_value*100:.2f}%"
                if nbands == 6 and self.ppm_combo is not None:
                    ppm_value = self.ppm_map[self.ppm_combo.currentText()]
                    result_text += f"\nConfiabilidad: {ppm_value}\u00a0ppm/K"
            self.result_label.setText(result_text)
        except Exception as exc:
            try:
                mostrar_error(self, "calcular resistencia", exc)
            except Exception:
                pass
            self.result_label.setText(f"Error al calcular: {exc}")

class VLSMCalculatorDialog(QDialog):
    def __init__(self, parent: Optional[QWidget] = None) -> None:
        super().__init__(parent)
        self.setWindowTitle(traducir('Calculadora de subred VSLM'))
        self.setModal(True)

        form_layout = QFormLayout()
        self.ip_edit = QLineEdit()
        self.ip_edit.setPlaceholderText("192.0.0.0")
        form_layout.addRow(traducir('Dirección IP'), self.ip_edit)

        self.prefix_spin = QSpinBox()
        self.prefix_spin.setRange(1, 32)
        self.prefix_spin.setValue(24)
        form_layout.addRow(traducir('Prefijo de red'), self.prefix_spin)

        self.subnets_spin = QSpinBox()
        self.subnets_spin.setRange(1, 64)
        self.subnets_spin.setValue(1)
        form_layout.addRow(traducir('Número de subredes'), self.subnets_spin)

        self.hosts_table = QTableWidget()
        self.hosts_table.setColumnCount(2)
        self.hosts_table.setHorizontalHeaderLabels([traducir('Subred'), traducir('Número de Hosts')])
        self.hosts_table.horizontalHeader().setStretchLastSection(True)
        self.hosts_table.verticalHeader().setVisible(False)

        self.calc_button = QPushButton(traducir('Calcular'))
        self.calc_button.clicked.connect(self.calcular)

        self.results_table = QTableWidget()
        self.results_table.setColumnCount(8)
        self.results_table.setHorizontalHeaderLabels([
            traducir('Subred'), traducir('Nº de Hosts'), traducir('IP de red'), traducir('Prefijo'),
            traducir('Máscara'), traducir('Primer Host'), traducir('Último Host'), traducir('Broadcast')
        ])
        self.results_table.horizontalHeader().setStretchLastSection(True)
        self.results_table.verticalHeader().setVisible(False)

        layout = QVBoxLayout(self)
        layout.addLayout(form_layout)
        layout.addWidget(self.hosts_table)
        layout.addWidget(self.calc_button)
        layout.addWidget(self.results_table)

        self.subnets_spin.valueChanged.connect(self.actualizar_tabla_hosts)
        self.actualizar_tabla_hosts()

    def actualizar_tabla_hosts(self) -> None:
        count = self.subnets_spin.value()
        self.hosts_table.setRowCount(count)
        for row in range(count):
            base_label = traducir('Subred')
            item0 = QTableWidgetItem(f"{base_label} {row + 1}")
            item0.setFlags(item0.flags() & ~Qt.ItemIsEditable)
            self.hosts_table.setItem(row, 0, item0)
            if not self.hosts_table.item(row, 1):
                self.hosts_table.setItem(row, 1, QTableWidgetItem(""))

    def sanear_ip(self, ip_str: str) -> str:
        parts = [p for p in ip_str.strip().split('.') if p != '']
        if len(parts) > 4:
            parts = parts[:4]
        return '.'.join(parts)

    def calcular(self) -> None:
        self.results_table.clearContents()
        self.results_table.setRowCount(0)
        ip_raw = self.ip_edit.text().strip()
        ip_clean = self.sanear_ip(ip_raw)
        prefix = self.prefix_spin.value()
        try:
            network = ipaddress.ip_network(f"{ip_clean}/{prefix}", strict=False)
        except Exception as exc:
            try:
                network = ipaddress.ip_network(f"0.0.0.0/{prefix}", strict=False)
            except Exception:
                QMessageBox.critical(self, "Error", f"Dirección IP o prefijo inválido:\n{exc}")
                return
        host_counts: List[int] = []
        for row in range(self.hosts_table.rowCount()):
            item = self.hosts_table.item(row, 1)
            txt = item.text().strip() if item else ''
            if not txt:
                host_counts.append(0)
            else:
                try:
                    val = int(txt)
                    host_counts.append(max(val, 0))
                except Exception:
                    host_counts.append(0)
        try:
            results = self.calcular_vlsm(network, host_counts)
        except Exception as exc:
            QMessageBox.critical(self, "Error", str(exc))
            return
        sorted_results = sorted(results, key=lambda r: r.get("hosts", 0), reverse=True)
        for idx, res in enumerate(sorted_results):
            res["subnet"] = f"{traducir('Subred')} {idx + 1}"
        self.results_table.setRowCount(len(sorted_results))
        self.results_table.setColumnCount(8)
        self.results_table.setHorizontalHeaderLabels([
            traducir('Subred'), traducir('Nº de Hosts'), traducir('IP de red'), traducir('Prefijo'),
            traducir('Máscara'), traducir('Primer Host'), traducir('Último Host'), traducir('Broadcast')
        ])
        for row, res in enumerate(sorted_results):
            self.results_table.setItem(row, 0, QTableWidgetItem(res["subnet"]))
            self.results_table.setItem(row, 1, QTableWidgetItem(str(res["hosts"])))
            self.results_table.setItem(row, 2, QTableWidgetItem(res["network"]))
            self.results_table.setItem(row, 3, QTableWidgetItem(f"/{res['prefix']}"))
            self.results_table.setItem(row, 4, QTableWidgetItem(res["netmask"]))
            self.results_table.setItem(row, 5, QTableWidgetItem(res["first_host"]))
            self.results_table.setItem(row, 6, QTableWidgetItem(res["last_host"]))
            self.results_table.setItem(row, 7, QTableWidgetItem(res["broadcast"]))
        try:
            self.results_table.resizeColumnsToContents()
        except Exception:
            pass

    def calcular_vlsm(self, base_network: ipaddress.IPv4Network, host_counts: List[int]) -> List[Dict[str, Any]]:
        indexed = []
        for idx, cnt in enumerate(host_counts):
            try:
                cnt_int = int(cnt)
            except Exception:
                cnt_int = 0
            if cnt_int < 0:
                cnt_int = 0
            indexed.append((idx, cnt_int))
        sorted_hosts = sorted(indexed, key=lambda x: -x[1])
        allocations: Dict[int, Dict[str, Any]] = {}
        current_addr = base_network.network_address
        for orig_idx, hosts_needed in sorted_hosts:
            needed = hosts_needed + 2 if hosts_needed > 0 else 2
            bits = (needed - 1).bit_length()
            prefix = 32 - bits
            if prefix < base_network.prefixlen:
                prefix = base_network.prefixlen
            try:
                subnet = ipaddress.ip_network(f"{current_addr}/{prefix}", strict=False)
            except Exception as exc:
                raise ValueError(f"No se puede crear la subred: {exc}")
            if not (subnet.network_address >= base_network.network_address and subnet.broadcast_address <= base_network.broadcast_address):
                raise ValueError("Las subredes exceden el rango de la red base.")
            total_addrs = subnet.num_addresses
            if total_addrs <= 2:
                first_host = str(subnet.network_address)
                last_host = str(subnet.broadcast_address)
            else:
                first_host = str(subnet.network_address + 1)
                last_host = str(subnet.broadcast_address - 1)
            allocations[orig_idx] = {
                "subnet": f"Subred {orig_idx + 1}",
                "hosts": hosts_needed,
                "network": str(subnet.network_address),
                "prefix": subnet.prefixlen,
                "netmask": str(subnet.netmask),
                "first_host": first_host,
                "last_host": last_host,
                "broadcast": str(subnet.broadcast_address),
            }
            try:
                current_addr = subnet.broadcast_address + 1
            except Exception:
                current_addr = subnet.broadcast_address
        return [allocations[i] for i in range(len(host_counts))]

class ItemComponente(QWidget):
    def __init__(self, icono: QIcon, texto: str, parent=None):
        super().__init__(parent)
        self.setFocusPolicy(Qt.NoFocus)
        self.setStyleSheet("""
            QWidget { outline: none; }
        """)
        layout = QVBoxLayout(self)
        layout.setContentsMargins(4, 4, 4, 4)
        layout.setSpacing(4)

        self.icono = QLabel(self)
        pixmap = icono.pixmap(36, 36).scaled(36, 36, Qt.KeepAspectRatio, Qt.SmoothTransformation)
        self.icono.setPixmap(pixmap)
        self.icono.setAlignment(Qt.AlignCenter)

        self.etiqueta = QLabel(texto, self)
        self.etiqueta.setTextInteractionFlags(Qt.NoTextInteraction)
        self.etiqueta.setStyleSheet("""
            font-size: 9pt;
            border: none;
            background: transparent;
        """)
        self.etiqueta.setWordWrap(True)
        self.etiqueta.setAlignment(Qt.AlignCenter)
        self.icono.setFocusPolicy(Qt.NoFocus)
        self.etiqueta.setFocusPolicy(Qt.NoFocus)
        self.setFocusPolicy(Qt.NoFocus)
        self.etiqueta.setStyleSheet("font-size: 9pt;")

        layout.addWidget(self.icono)
        layout.addWidget(self.etiqueta)

class TerminalItem(QGraphicsEllipseItem):
    def __init__(self, parent=None):
        super().__init__(-4, -4, 8, 8, parent)
        self.setBrush(Qt.black)
        self.setZValue(1)
        self.setFlag(QGraphicsItem.ItemSendsScenePositionChanges)
        self.conexiones = []
        self.is_logic_pin: bool = False
        self.is_logic_output: bool = False
        self.logic_value: int = 0

class Conexion:
    def __init__(self, scene, t1, t2):
        self.scene = scene
        self.t1 = t1
        self.t2 = t2
        self.is_logic: bool = bool(getattr(t1, 'is_logic_pin', False) and getattr(t2, 'is_logic_pin', False))
        self.lines = []
        self.crear_lineas()
        self.t1.conexiones.append(self)
        self.t2.conexiones.append(self)

    def ajustar_a_cuadricula(self, point, grid=10):
        x = round(point.x() / grid) * grid
        y = round(point.y() / grid) * grid
        return QPointF(x, y)

    def crear_lineas(self):
        if getattr(self, 'is_logic', False):
            pen = QPen(Qt.blue, 2)
        else:
            pen = QPen(Qt.black, 2)
        for _ in range(3):
            line = self.scene.addLine(0, 0, 0, 0, pen)
            line.setFlags(QGraphicsItem.ItemIsSelectable)
            self.lines.append(line)
        self.actualizar_posicion()

    def actualizar_posicion(self):
        p1 = self.ajustar_a_cuadricula(self.t1.scenePos())
        p2 = self.ajustar_a_cuadricula(self.t2.scenePos())
        dx = p2.x() - p1.x()
        dy = p2.y() - p1.y()
        grid = 20
        margin = grid

        horizontal_preferred = abs(dx) >= abs(dy)

        def segmento_interseca_rectangulo(s: QPointF, e: QPointF, rect: QRectF) -> bool:
            if s.x() == e.x():
                x = s.x()
                y1 = min(s.y(), e.y())
                y2 = max(s.y(), e.y())
                if x >= rect.left() and x <= rect.right() and y2 >= rect.top() and y1 <= rect.bottom():
                    return True
                return False
            else:
                y = s.y()
                x1 = min(s.x(), e.x())
                x2 = max(s.x(), e.x())
                if y >= rect.top() and y <= rect.bottom() and x2 >= rect.left() and x1 <= rect.right():
                    return True
                return False

        comp1 = self.t1.parentItem()
        comp2 = self.t2.parentItem()
        try:
            rect1 = comp1.mapRectToScene(comp1.boundingRect()).boundingRect()
            rect2 = comp2.mapRectToScene(comp2.boundingRect()).boundingRect()
        except Exception:
            rect1 = QRectF()
            rect2 = QRectF()
        union_rect = rect1.united(rect2)

        if self.scene is not None:
            for item in self.scene.items():
                if hasattr(item, 'terminales') and item not in [comp1, comp2]:
                    try:
                        rect = item.mapRectToScene(item.boundingRect()).boundingRect()
                    except Exception:
                        continue
                    union_rect = union_rect.united(rect)

        union_rect = union_rect.adjusted(-margin, -margin, margin, margin)

        def construir_segmentos_horizontales(mid_y: float):
            return [
                (p1, QPointF(p1.x(), mid_y)),
                (QPointF(p1.x(), mid_y), QPointF(p2.x(), mid_y)),
                (QPointF(p2.x(), mid_y), p2)
            ]

        def construir_segmentos_verticales(mid_x: float):
            return [
                (p1, QPointF(mid_x, p1.y())),
                (QPointF(mid_x, p1.y()), QPointF(mid_x, p2.y())),
                (QPointF(mid_x, p2.y()), p2)
            ]

        final_segs: List[Tuple[QPointF, QPointF]] = []

        if horizontal_preferred:
            mid_y = round(((p1.y() + p2.y()) / 2) / grid) * grid
            candidate_segs = construir_segmentos_horizontales(mid_y)
            intersects = any(segmento_interseca_rectangulo(s, e, union_rect) for s, e in candidate_segs)
            if not intersects:
                final_segs = candidate_segs
            else:
                above_y = round((union_rect.top() - grid) / grid) * grid
                below_y = round((union_rect.bottom() + grid) / grid) * grid
                len_above = abs(p1.y() - above_y) + abs(p2.y() - above_y) + abs(dx)
                len_below = abs(p1.y() - below_y) + abs(p2.y() - below_y) + abs(dx)
                chosen_y = above_y if len_above <= len_below else below_y
                final_segs = construir_segmentos_horizontales(chosen_y)
        else:
            mid_x = round(((p1.x() + p2.x()) / 2) / grid) * grid
            candidate_segs = construir_segmentos_verticales(mid_x)
            intersects = any(segmento_interseca_rectangulo(s, e, union_rect) for s, e in candidate_segs)
            if not intersects:
                final_segs = candidate_segs
            else:
                left_x = round((union_rect.left() - grid) / grid) * grid
                right_x = round((union_rect.right() + grid) / grid) * grid
                len_left = abs(p1.x() - left_x) + abs(p2.x() - left_x) + abs(dy)
                len_right = abs(p1.x() - right_x) + abs(p2.x() - right_x) + abs(dy)
                chosen_x = left_x if len_left <= len_right else right_x
                final_segs = construir_segmentos_verticales(chosen_x)

        for i in range(3):
            if i < len(final_segs):
                s, e = final_segs[i]
                self.lines[i].setLine(QLineF(s, e))
            else:
                pt = p1
                self.lines[i].setLine(QLineF(pt, pt))

    def eliminar(self):
        for line in self.lines:
            self.scene.removeItem(line)
        self.t1.conexiones.remove(self)
        self.t2.conexiones.remove(self)

class SceneConCuadricula(QGraphicsScene):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.grid_size = 20
        self.grid_pen = QPen(Qt.lightGray)
        self.grid_pen.setStyle(Qt.DotLine)

    def drawBackground(self, painter, rect):
        left = int(rect.left()) - (int(rect.left()) % self.grid_size)
        top = int(rect.top()) - (int(rect.top()) % self.grid_size)

        lines = []
        x = left
        while x < rect.right():
            lines.append(QLineF(x, rect.top(), x, rect.bottom()))
            x += self.grid_size

        y = top
        while y < rect.bottom():
            lines.append(QLineF(rect.left(), y, rect.right(), y))
            y += self.grid_size

        painter.setPen(self.grid_pen)
        painter.drawLines(lines)

class ComponenteInteractivo(QGraphicsItem):
    def hoverEnterEvent(self, event):
        translated_params = []
        for param in self.parametros:
            if ":" in param:
                key, val = param.split(":", 1)
                key_translated = traducir(key.strip())
                translated_params.append(f"{key_translated}:{val}")
            else:
                translated_params.append(param)
        if len(translated_params) > 1:
            tooltip = f"<b>{self.nombre}:</b><br>" + "<br>".join(translated_params)
        else:
            tooltip = f"<b>{self.nombre}:</b> " + " ".join(translated_params)
        self.setToolTip(tooltip)
        super().hoverEnterEvent(event)

    def hoverLeaveEvent(self, event):
        self.setToolTip("")
        super().hoverLeaveEvent(event)

    def itemChange(self, change, value):
        if change == QGraphicsItem.ItemPositionHasChanged:
            for terminal in self.terminales:
                for conn in terminal.conexiones:
                    conn.actualizar_posicion()
            try:
                self.actualizar_guias()
            except Exception:
                pass
        return super().itemChange(change, value)

    def actualizar_guias(self) -> None:
        scene = self.scene()
        if scene is None:
            return
        views = scene.views()
        if not views:
            return
        ventana = None
        try:
            ventana = views[0].window()
        except Exception:
            ventana = None
        if ventana is None or not hasattr(ventana, 'guideline_items'):
            return
        for linea in list(ventana.guideline_items):
            try:
                scene.removeItem(linea)
            except Exception:
                pass
        ventana.guideline_items = []
        rect = scene.sceneRect()
        try:
            punto = self.scenePos()
        except Exception:
            punto = self.pos()
        x = punto.x()
        y = punto.y()
        pen = QPen(Qt.blue)
        pen.setStyle(Qt.DashLine)
        linea_v = scene.addLine(x, rect.top(), x, rect.bottom(), pen)
        linea_h = scene.addLine(rect.left(), y, rect.right(), y, pen)
        ventana.guideline_items = [linea_v, linea_h]
    LOGIC_GATES = ["AND", "OR", "NAND", "NOR", "XOR", "XNOR", "NOT"]

    def __init__(self, pixmap, parametros, nombre, tipo, terminales=2, vertical=False, invertir=False):
        super().__init__()
        self.setAcceptHoverEvents(True)
        self.terminales = []
        self.vertical = vertical
        self.invertir = invertir
        self.setFlags(QGraphicsItem.ItemIsMovable | QGraphicsItem.ItemIsSelectable |
                      QGraphicsItem.ItemSendsGeometryChanges)
        self.tipo = tipo
        self.parametros = parametros
        self.nombre = nombre
        if not os.path.exists(pixmap):
            print(f"[ERROR] El archivo SVG no existe: {pixmap}")
        self.svg_item = QGraphicsSvgItem(pixmap, self)

        self.svg_item.setAcceptedMouseButtons(Qt.NoButton)
        self.setAcceptedMouseButtons(Qt.LeftButton | Qt.RightButton)

        self.tooltip_texto = f"{nombre}:" + "".join(self.parametros)
        self.svg_item.setScale(1.2)  

        self.nombre_label = QGraphicsTextItem(nombre, self)
        fuente = self.nombre_label.font()
        fuente.setPointSizeF(8)  
        self.nombre_label.setFont(fuente)

        self.nombre_label.setDefaultTextColor(Qt.blue)
        
        if self.tipo == "GND":
           self.svg_item.setOpacity(0.4)
           self.nombre_label.setDefaultTextColor(Qt.darkRed)
 
        self.nombre_label.setVisible(True)

        self.nombre_label.setZValue(2)
        self.nombre_label.setParentItem(self)

        bounding = self.svg_item.boundingRect()
        ancho = bounding.width()
        alto = bounding.height()
        cx = ancho / 2
        cy = alto / 2

        alto_escalado = bounding.height() * self.svg_item.scale()
        if self.tipo in ["LED", "Fuente Corriente", "Fuente Voltaje", "Diodo Zener", "Batería"]:
            x = ancho + 2
            y = (alto - self.nombre_label.boundingRect().height()) / 2
        else:
            x = (ancho - self.nombre_label.boundingRect().width()) / 2
            y = alto + 5 if not self.vertical else -25
        self.nombre_label.setPos(x, y) 

        if tipo == "GND":
            posiciones = [QPointF(cx, 0)]  
        elif tipo in ["Push Button", "Push Button Cerrado"]:
            posiciones = [QPointF(10, alto - 5), QPointF(ancho - 10, alto - 5)]
        elif tipo in ["Transistor NPN", "Transistor PNP", "Mosfet"]:
            posiciones = [
                QPointF(0, cy),
                QPointF(ancho - 2, 0),
                QPointF(ancho - 2, alto)
            ]
        elif tipo == "OpAmp":
            posiciones = [
                QPointF(0, alto * 0.35),  
                QPointF(0, alto * 0.65),  
                QPointF(ancho, cy)        
            ]
        elif tipo == "Potenciómetro":
            posiciones = [
                QPointF(0, cy),           
                QPointF(ancho, cy),       
                QPointF(cx, alto)         
            ]

        elif tipo == "Generador de funciones":
            posiciones = [QPointF(ancho * 0.18, alto), QPointF(ancho * 0.82, alto)]
        elif tipo in ["AND", "OR", "NAND", "NOR", "XOR"]:
            posiciones = [
                QPointF(0, alto * 0.25),
                QPointF(0, alto * 0.75),
                QPointF(ancho, cy)
            ]
        elif tipo == "NOT":
            posiciones = [
                QPointF(0, cy),
                QPointF(ancho, cy)
            ]
        elif self.vertical ^ self.invertir:
            posiciones = {
    2: [QPointF(0, cy), QPointF(ancho, cy)],
    3: [QPointF(0, cy), QPointF(cx, 0), QPointF(ancho, cy)]
}[terminales]
        else:
             posiciones = {
        2: [QPointF(cx, 0), QPointF(cx, alto)],
        3: [QPointF(cx, 0), QPointF(cx, alto / 2), QPointF(cx, alto)]
    }.get(terminales, [])

        for pos in posiciones:
            terminal = TerminalItem(self)
            terminal.setPos(pos)
            self.terminales.append(terminal)

        try:
            if self.tipo in self.LOGIC_GATES:
                self._inicializar_compuerta_logica()
        except Exception:
            pass

        if self.tipo in self.LOGIC_GATES:
            try:
                num_inputs = max(1, len(getattr(self, 'terminales', [])) - 1)
                self.logic_input_names = [chr(ord('A') + i) for i in range(num_inputs)]
                self.logic_output_name = 'Y'
            except Exception:
                self.logic_input_names = []
                self.logic_output_name = 'Y'
            try:
                self.actualizar_posiciones_etiquetas_logicas()
            except Exception:
                pass

    def actualizar_escala_terminales(self):
        componente_escala = self.scale() or 0.1
        factor_ampliacion = 1.0  
        for terminal in self.terminales:
            terminal.setScale((1.0 / componente_escala) * factor_ampliacion)

    def reposicionar_terminales(self):
        bounding = self.svg_item.boundingRect()
        ancho = bounding.width() * self.svg_item.scale()
        alto = bounding.height() * self.svg_item.scale()
        cx = ancho / 2
        cy = alto / 2

        alto_escalado = bounding.height() * self.svg_item.scale()
        if self.tipo in ["LED", "Fuente Corriente", "Fuente Voltaje", "Diodo Zener", "Batería"]:
            self.nombre_label.setPos(ancho + 5, (alto - self.nombre_label.boundingRect().height()) / 2)
        elif not self.vertical:
            self.nombre_label.setPos((ancho - self.nombre_label.boundingRect().width()) / 2, alto + 5)
        else:
            self.nombre_label.setPos((ancho - self.nombre_label.boundingRect().width()) / 2, -35)

        tipo = self.tipo
        terminales = len(self.terminales)

        if tipo in ["Amperímetro", "Voltímetro", "Ohmímetro"]:
            posiciones = [
                QPointF(ancho * 0.29, alto),
                QPointF(ancho * 0.71, alto)
            ]
            for i, pos in enumerate(posiciones):
                if i < len(self.terminales):
                    self.terminales[i].setPos(pos)
                    self.terminales[i].setVisible(True)
            self.nombre_label.setPos((ancho - self.nombre_label.boundingRect().width()) / 2, alto + 5)
            self.actualizar_escala_terminales()
            return

        if tipo == "GND":
            posiciones = [QPointF(cx, 0)]  
        elif tipo in ["Push Button", "Push Button Cerrado"]:
            posiciones = [QPointF(10, alto - 5), QPointF(ancho - 10, alto - 5)]
        elif tipo in ["Transistor NPN", "Transistor PNP", "Mosfet"]:
            posiciones = [
                QPointF(0, cy),
                QPointF(ancho - 2, 0),
                QPointF(ancho - 2, alto)
            ]
        elif tipo == "OpAmp":
            posiciones = [
                QPointF(0, alto * 0.35),
                QPointF(0, alto * 0.65),
                QPointF(ancho, cy)
            ]
        elif tipo == "Potenciómetro":
            posiciones = [
                QPointF(0, cy),
                QPointF(ancho, cy),
                QPointF(cx, alto)
            ]
        elif tipo == "Generador de funciones":
             posiciones = [QPointF(ancho * 0.18, alto), QPointF(ancho * 0.82, alto)]
        elif tipo == "Motor":
            posiciones = [QPointF(0, cy), QPointF(ancho, cy)]  
        elif tipo in ["AND", "OR", "NAND", "NOR", "XOR"]:
            posiciones = [
                QPointF(0, alto * 0.25),
                QPointF(0, alto * 0.75),
                QPointF(ancho, alto * 0.5)
            ]
        elif tipo == "NOT":
            posiciones = [
                QPointF(0, alto * 0.5),
                QPointF(ancho, alto * 0.5)
            ]
        elif tipo in ["Resistencia", "Capacitor", "Bobina"]:
            posiciones = [QPointF(0, cy), QPointF(ancho, cy)]
        elif tipo == "Transformador":
            posiciones = [
                QPointF(0, alto * 0.1),  
                QPointF(0, alto * 0.9),  
                QPointF(ancho, alto * 0.1),  
                QPointF(ancho, alto * 0.9)   
            ]
            for pos in posiciones:
                terminal = TerminalItem(self)
                terminal.setPos(pos)
                self.terminales.append(terminal)
                terminal.setVisible(True)  
            self.actualizar_escala_terminales() 
        elif terminales == 4:
            posiciones = [
                QPointF(0, alto * 0.3), QPointF(0, alto * 0.7),
                QPointF(ancho, alto * 0.3), QPointF(ancho, alto * 0.7)
            ]
        else:
            posiciones = {
                2: [QPointF(cx, 0), QPointF(cx, alto)],
                3: [QPointF(cx, 0), QPointF(cx, alto / 2), QPointF(cx, alto)]
            }.get(terminales, [])

        for i, pos in enumerate(posiciones):
            if i < len(self.terminales):
                self.terminales[i].setPos(pos)
                self.terminales[i].setVisible(True)  
        self.actualizar_escala_terminales()

        try:
            if hasattr(self, "logic_inputs"):
                self.actualizar_posiciones_etiquetas_logicas()
        except Exception:
            pass

    def _inicializar_compuerta_logica(self):
        try:
            num_inputs = max(1, len(self.terminales) - 1)
            self.logic_inputs = [0 for _ in range(num_inputs)]
            self.logic_output = 0
            self.logic_input_labels = []
            font = self.nombre_label.font() if hasattr(self, 'nombre_label') else None
            if font is not None:
                font.setPointSizeF(7)
            for idx in range(num_inputs):
                lbl = QGraphicsTextItem(self)
                lbl.logic_input_index = idx
                try:
                    lbl.setAcceptedMouseButtons(Qt.NoButton)
                except Exception:
                    pass
                if font is not None:
                    lbl.setFont(font)
                lbl.setDefaultTextColor(Qt.darkGreen)
                lbl.setZValue(3)
                try:
                    if hasattr(self, 'logic_input_names') and idx < len(self.logic_input_names):
                        var_name = self.logic_input_names[idx]
                    else:
                        var_name = chr(65 + idx)
                except Exception:
                    var_name = chr(65 + idx)
                lbl.setPlainText(f"{var_name}=0")
                self.logic_input_labels.append(lbl)
            self.output_indicator = QGraphicsEllipseItem(-4, -4, 8, 8, self)
            self.output_indicator.setBrush(Qt.red)
            self.output_indicator.setZValue(2)
            self.output_indicator.output_indicator_flag = True
            try:
                self.output_indicator.setAcceptedMouseButtons(Qt.NoButton)
            except Exception:
                pass
            self.output_label = QGraphicsTextItem(self)
            self.output_label.output_label_flag = True
            try:
                self.output_label.setAcceptedMouseButtons(Qt.NoButton)
            except Exception:
                pass
            if font is not None:
                self.output_label.setFont(font)
            self.output_label.setDefaultTextColor(Qt.darkGreen)
            self.output_label.setZValue(3)
            try:
                if hasattr(self, 'logic_output_name'):
                    out_var = self.logic_output_name
                else:
                    out_var = 'Y'
            except Exception:
                out_var = 'Y'
            self.output_label.setPlainText(f"{out_var}=0")
            for idx, term in enumerate(self.terminales):
                term.is_logic_pin = True
                term.is_logic_output = (idx == len(self.terminales) - 1)
                term.logic_value = 0
            self.actualizar_estado_logico()
            self.actualizar_posiciones_etiquetas_logicas()
        except Exception:
            pass

    def _calcular_salida(self, inputs):
        try:
            normalized = [1 if val else 0 for val in inputs]
            if not normalized:
                normalized = [0]
            t = self.tipo
            if t == "AND":
                return 1 if all(normalized) else 0
            elif t == "OR":
                return 1 if any(normalized) else 0
            elif t == "NAND":
                return 0 if all(normalized) else 1
            elif t == "NOR":
                return 0 if any(normalized) else 1
            elif t == "XOR":
                return 1 if (sum(normalized) % 2 == 1) else 0
            elif t == "XNOR":
                return 1 if (sum(normalized) % 2 == 0) else 0
            elif t == "NOT":
                return 0 if normalized[0] else 1
            else:
                return 0
        except Exception:
            return 0

    def actualizar_estado_logico(self):
        def _actualizar_compuerta(gate: 'ComponenteInteractivo', visited: Optional[set] = None) -> None:
            if visited is None:
                visited = set()
            if gate in visited:
                return
            visited.add(gate)
            try:
                gate.logic_output = gate._calcular_salida(gate.logic_inputs)
                if hasattr(gate, "output_indicator"):
                    gate.output_indicator.setBrush(Qt.green if gate.logic_output else Qt.red)
                if hasattr(gate, "output_label"):
                    gate.output_label.setPlainText(f"Y={gate.logic_output}")
                if gate.terminales:
                    out_term = gate.terminales[-1]
                    out_term.logic_value = gate.logic_output
                    for conn in list(out_term.conexiones):
                        try:
                            if not getattr(conn, 'is_logic', False):
                                continue
                            other = conn.t1 if conn.t2 is out_term else conn.t2
                            if getattr(other, 'is_logic_output', False):
                                continue
                            other.logic_value = gate.logic_output
                            dest_gate = other.parentItem()
                            if isinstance(dest_gate, ComponenteInteractivo) and hasattr(dest_gate, 'logic_inputs'):
                                try:
                                    idx = dest_gate.terminales.index(other)
                                except ValueError:
                                    idx = -1
                                if idx >= 0 and idx < len(dest_gate.logic_inputs):
                                    dest_gate.logic_inputs[idx] = other.logic_value
                                    _actualizar_compuerta(dest_gate, visited)
                        except Exception:
                            continue
            except Exception:
                pass
        _actualizar_compuerta(self)

    def actualizar_posiciones_etiquetas_logicas(self):
        try:
            num_inputs = len(self.logic_inputs) if hasattr(self, 'logic_inputs') else 0
            for idx in range(num_inputs):
                if idx >= len(self.terminales):
                    continue
                term = self.terminales[idx]
                if idx < len(self.logic_input_labels):
                    lbl = self.logic_input_labels[idx]
                else:
                    continue
                try:
                    if hasattr(self, 'logic_input_names') and idx < len(self.logic_input_names):
                        var_name = self.logic_input_names[idx]
                    else:
                        var_name = chr(ord('A') + idx)
                except Exception:
                    var_name = chr(ord('A') + idx)
                lbl.setPlainText(f"{var_name}={self.logic_inputs[idx]}")
                pos = term.pos()
                lbl.setPos(pos.x() - 38, pos.y() - 10)
            if self.terminales:
                out_term = self.terminales[-1]
                if hasattr(self, "output_indicator"):
                    self.output_indicator.setPos(out_term.pos())
                if hasattr(self, "output_label"):
                    try:
                        out_var = self.logic_output_name if hasattr(self, 'logic_output_name') else 'Y'
                    except Exception:
                        out_var = 'Y'
                    self.output_label.setPlainText(f"{out_var}={self.logic_output}")
                    self.output_label.setPos(out_term.pos().x() + 10, out_term.pos().y() - 10)
        except Exception:
            pass

    def establecer_nombres_logicos(self, input_names: List[str], output_name: str) -> None:
        try:
            if not hasattr(self, 'logic_inputs'):
                return
            self.logic_input_names = list(input_names) if input_names else []
            self.logic_output_name = output_name or 'Y'
            self.actualizar_posiciones_etiquetas_logicas()
        except Exception:
            pass

    def boundingRect(self):
        return self.svg_item.mapToParent(self.svg_item.boundingRect()).boundingRect()

    def paint(self, painter, option, widget):
        pass

    def mouseDoubleClickEvent(self, event):
        if event.button() == Qt.LeftButton and self.tipo in self.LOGIC_GATES:
            try:
                dlg = TruthTableDialog(self)
                dlg.exec_()
            except Exception:
                pass
            return
        if event.button() == Qt.LeftButton:
            print(f"Se hizo doble clic en el componente {self.nombre}")
            views = self.scene().views()
            if not views:
                super().mouseDoubleClickEvent(event)
                return

            parent = None
            for view in views:
                for cand in [getattr(view, "window", lambda: None)(), getattr(view, "parent", lambda: None)()]:
                    if cand and hasattr(cand, "pedir_valores_componente"):
                        parent = cand
                        break
                    try:
                        tw = getattr(cand, "tab_widget", None)
                        if tw is not None:
                            cw = tw.currentWidget()
                            if cw and hasattr(cw, "pedir_valores_componente"):
                                parent = cw
                                break
                    except Exception:
                        pass
                if parent:
                    break

            if not parent:
                super().mouseDoubleClickEvent(event)
                return

            nuevas_entradas = parent.pedir_valores_componente(self.tipo)
            if nuevas_entradas:
                self.parametros = nuevas_entradas

                translated_params = []
                for param in self.parametros:
                    if ":" in param:
                        key, val = param.split(":", 1)
                        key_translated = traducir(key.strip())
                        translated_params.append(f"{key_translated}:{val}")
                    else:
                        translated_params.append(param)
                if len(translated_params) > 1:
                    nuevo_tooltip = f"<b>{self.nombre}:</b><br>" + "<br>".join(translated_params)
                else:
                    nuevo_tooltip = f"<b>{self.nombre}:</b> " + " ".join(translated_params)
                self.svg_item.setToolTip(nuevo_tooltip)
                self.update()
            else:
                 print("No se ingresaron los valores.")
        else: 
            super().mouseDoubleClickEvent(event)

    def mousePressEvent(self, event):
        if event.button() == Qt.RightButton:
            self.setRotation((self.rotation() + 90) % 360)
            try:
                scene = self.scene()
                views = scene.views() if scene else []
                ventana = views[0].window() if views else None
                if ventana and hasattr(ventana, 'registrar_historial'):
                    ventana.registrar_historial()
            except Exception:
                pass
        else:
            super().mousePressEvent(event)

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.LeftButton:
            grid = 20
            pos = self.pos()
            snapped = QPointF(
                round(pos.x() / grid) * grid,
                round(pos.y() / grid) * grid
            )
            self.setPos(snapped)
        try:
            scene = self.scene()
            views = scene.views() if scene else []
            ventana = views[0].window() if views else None
            if ventana and hasattr(ventana, 'guideline_items'):
                for linea in list(ventana.guideline_items):
                    try:
                        scene.removeItem(linea)
                    except Exception:
                        pass
                ventana.guideline_items = []
        except Exception:
            pass
        try:
            scene = self.scene()
            views = scene.views() if scene else []
            ventana = views[0].window() if views else None
            if ventana and hasattr(ventana, 'registrar_historial'):
                ventana.registrar_historial()
        except Exception:
            pass
        super().mouseReleaseEvent(event)

class TruthTableDialog(QDialog):

    def __init__(self, gate: ComponenteInteractivo):
        super().__init__(gate)
        self.gate = gate
        self.setWindowTitle(f"Tabla de verdad - {gate.tipo}")
        layout = QVBoxLayout(self)
        num_inputs = len(getattr(gate, 'logic_inputs', [])) or 1
        row_count = 2 ** num_inputs
        if hasattr(gate, 'logic_input_names') and gate.logic_input_names:
            column_labels = list(gate.logic_input_names[:num_inputs])
        else:
            column_labels = [chr(ord('A') + i) for i in range(num_inputs)]
        if hasattr(gate, 'logic_output_name'):
            column_labels.append(gate.logic_output_name)
        else:
            column_labels.append("Y")
        self.table = QTableWidget(row_count, num_inputs + 1, self)
        self.table.setHorizontalHeaderLabels(column_labels)
        self.table.verticalHeader().setVisible(False)
        combinations = []
        for i in range(row_count):
            combo = []
            for bit_idx in range(num_inputs):
                combo.append((i >> bit_idx) & 1)
            combinations.append(combo)
        for row, combo in enumerate(combinations):
            for col, val in enumerate(combo):
                item = QTableWidgetItem(str(val))
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.table.setItem(row, col, item)
            out = gate._calcular_salida(combo)
            out_item = QTableWidgetItem(str(out))
            out_item.setFlags(out_item.flags() & ~Qt.ItemIsEditable)
            self.table.setItem(row, num_inputs, out_item)
        self.table.resizeColumnsToContents()
        self.table.resizeRowsToContents()
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.SingleSelection)
        self.table.cellClicked.connect(self.al_hacer_clic_en_celda)
        layout.addWidget(self.table)
        self.resaltar_estado_actual()

    def resaltar_estado_actual(self):
        try:
            num_inputs = len(getattr(self.gate, 'logic_inputs', [])) or 1
            idx = 0
            for i, val in enumerate(self.gate.logic_inputs):
                idx |= (val << (num_inputs - 1 - i))
            self.table.selectRow(idx)
            for row in range(self.table.rowCount()):
                color = Qt.yellow if row == idx else Qt.white
                for col in range(self.table.columnCount()):
                    item = self.table.item(row, col)
                    if item is not None:
                        item.setBackground(color)
        except Exception:
            pass

    def al_hacer_clic_en_celda(self, row: int, column: int):
        try:
            num_inputs = len(getattr(self.gate, 'logic_inputs', [])) or 1
            new_inputs = []
            for col in range(num_inputs):
                item = self.table.item(row, col)
                new_inputs.append(int(item.text()) if item else 0)
            self.gate.logic_inputs = new_inputs
            self.gate.actualizar_estado_logico()
            self.gate.actualizar_posiciones_etiquetas_logicas()
            self.resaltar_estado_actual()
        except Exception:
            pass

class DigitalTruthTableDialog(QDialog):

    def __init__(self, rows: list[dict[str, int]], columns: list[str], parent: Optional[QWidget] = None) -> None:
        super().__init__(parent)
        self.setWindowTitle("Tabla de verdad")
        layout = QVBoxLayout(self)
        row_count = len(rows)
        column_count = len(columns)
        self.table = QTableWidget(row_count, column_count, self)
        self.table.setHorizontalHeaderLabels(columns)
        self.table.verticalHeader().setVisible(False)
        for r_idx, row in enumerate(rows):
            for c_idx, col_name in enumerate(columns):
                val = row.get(col_name, 0)
                item = QTableWidgetItem(str(val))
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
                self.table.setItem(r_idx, c_idx, item)
        try:
            self.table.resizeColumnsToContents()
            self.table.resizeRowsToContents()
        except Exception:
            pass
        self.table.setSelectionBehavior(QTableWidget.SelectRows)
        self.table.setSelectionMode(QTableWidget.SingleSelection)
        layout.addWidget(self.table)

class VistaConZoom(QGraphicsView):
    def wheelEvent(self, event):
        try:
            from PyQt5.QtCore import Qt  
            if event.modifiers() & Qt.ControlModifier:
                zoom_in_factor = 1.15
                zoom_out_factor = 1 / zoom_in_factor
                zoom_factor = zoom_in_factor if event.angleDelta().y() > 0 else zoom_out_factor
                self.scale(zoom_factor, zoom_factor)
                return
        except Exception:
            pass
        super().wheelEvent(event)

class CircuitSimulator(QMainWindow):
    def cambiar_idioma(self, lang: str) -> None:
        global CURRENT_LANGUAGE
        CURRENT_LANGUAGE = lang
        try:
            preferences['language'] = lang
            guardar_preferencias()
        except Exception:
            pass
        try:
            self.aplicar_idioma()
        except Exception:
            pass
        try:
            parent = self
            while parent is not None and not hasattr(parent, 'tab_widget'):
                parent = parent.parent()
            if parent is not None and hasattr(parent, 'tab_widget'):
                for i in range(parent.tab_widget.count()):
                    parent.tab_widget.setTabText(i, f"{traducir('Pestaña')} {i+1}")
        except Exception:
            pass

    def aplicar_idioma(self) -> None:
        try:
            if hasattr(self, 'menu_archivo'):
                self.menu_archivo.setTitle(traducir('Archivo'))
            if hasattr(self, 'menu_exportar_logica'):
                self.menu_exportar_logica.setTitle(traducir('Exportar lógica'))
            if hasattr(self, 'menu_simulacion'):
                self.menu_simulacion.setTitle(traducir('Simulación'))
            if hasattr(self, 'menu_edicion') and self.menu_edicion is not None:
                self.menu_edicion.setTitle(traducir('Edición'))
            if hasattr(self, 'menu_herramientas'):
                self.menu_herramientas.setTitle(traducir('Herramientas'))
            if hasattr(self, 'menu_plantillas'):
                self.menu_plantillas.setTitle(traducir('Plantillas'))
            if hasattr(self, 'menu_insertar'):
                self.menu_insertar.setTitle(traducir('Insertar'))
            if hasattr(self, 'menu_ver'):
                self.menu_ver.setTitle(traducir('Ver'))
            if hasattr(self, 'menu_ayuda'):
                self.menu_ayuda.setTitle(traducir('Ayuda'))
            if hasattr(self, 'menu_plugins'):
                self.menu_plugins.setTitle(traducir('Plugins'))
            if hasattr(self, 'menu_idiomas'):
                self.menu_idiomas.setTitle(traducir('Idiomas'))
        except Exception:
            pass

        try:
            if hasattr(self, 'accion_modo_lab'):
                self.accion_modo_lab.setText(traducir('Modo Laboratorio'))
                self.accion_modo_lab.setToolTip(traducir_descripcion_emergente('Comparar curva simulada y medida en tiempo real'))
            if hasattr(self, 'accion_guardar_medicion'):
                self.accion_guardar_medicion.setText(traducir('Guardar Medición'))
                self.accion_guardar_medicion.setToolTip(traducir_descripcion_emergente('Guardar medición actual en la base de datos'))
            if hasattr(self, 'accion_ver_mediciones'):
                self.accion_ver_mediciones.setText(traducir('Ver mediciones'))
                self.accion_ver_mediciones.setToolTip(traducir_descripcion_emergente('Mostrar historial de mediciones guardadas'))
        except Exception:
            pass

        try:
            if hasattr(self, 'accion_restablecer_menu'):
                self.accion_restablecer_menu.setText(traducir('Reestablecer'))
            if hasattr(self, 'accion_rehacer_menu'):
                self.accion_rehacer_menu.setText(traducir('Rehacer'))
            if hasattr(self, 'accion_cortar'):
                self.accion_cortar.setText(traducir('Cortar'))
            if hasattr(self, 'accion_copiar'):
                self.accion_copiar.setText(traducir('Copiar'))
            if hasattr(self, 'accion_pegar'):
                self.accion_pegar.setText(traducir('Pegar'))
            if hasattr(self, 'accion_eliminar_menu'):
                self.accion_eliminar_menu.setText(traducir('Eliminar'))
            if hasattr(self, 'accion_seleccionar'):
                self.accion_seleccionar.setText(traducir('Seleccionar'))
            if hasattr(self, 'accion_select_all'):
                self.accion_select_all.setText(traducir('Seleccionar todo'))
            if hasattr(self, 'accion_rotar'):
                self.accion_rotar.setText(traducir('Rotar'))

            if hasattr(self, 'accion_tierra'):
                self.accion_tierra.setText(traducir('Tierra'))
                self.accion_tierra.setToolTip(traducir_descripcion_emergente(self.accion_tierra.toolTip()))
            if hasattr(self, 'accion_etiqueta'):
                self.accion_etiqueta.setText(traducir('Etiqueta Cable'))
                self.accion_etiqueta.setToolTip(traducir_descripcion_emergente(self.accion_etiqueta.toolTip()))
            if hasattr(self, 'accion_conexion'):
                self.accion_conexion.setText(traducir('Conexión'))
                self.accion_conexion.setToolTip(traducir_descripcion_emergente(self.accion_conexion.toolTip()))
            if hasattr(self, 'accion_ampliar'):
                self.accion_ampliar.setText(traducir('Ampliar'))
                self.accion_ampliar.setToolTip(traducir_descripcion_emergente(self.accion_ampliar.toolTip()))
            if hasattr(self, 'accion_reducir'):
                self.accion_reducir.setText(traducir('Reducir'))
                self.accion_reducir.setToolTip(traducir_descripcion_emergente(self.accion_reducir.toolTip()))
            if hasattr(self, 'accion_ver_todo'):
                self.accion_ver_todo.setText(traducir('Ver todo'))
                self.accion_ver_todo.setToolTip(traducir_descripcion_emergente(self.accion_ver_todo.toolTip()))
            if hasattr(self, 'accion_vista_11'):
                self.accion_vista_11.setText(traducir('Establecer escala 1:1'))
                self.accion_vista_11.setToolTip(traducir_descripcion_emergente(self.accion_vista_11.toolTip()))

            if hasattr(self, 'accion_duplicar'):
                self.accion_duplicar.setText(traducir('Duplicar'))
            if hasattr(self, 'menu_alinear_distribuir'):
                self.menu_alinear_distribuir.setTitle(traducir('Alinear/Distribuir'))
            if hasattr(self, 'accion_alinear_izquierda'):
                self.accion_alinear_izquierda.setText(traducir('Alinear izquierda'))
            if hasattr(self, 'accion_alinear_derecha'):
                self.accion_alinear_derecha.setText(traducir('Alinear derecha'))
            if hasattr(self, 'accion_alinear_centro'):
                self.accion_alinear_centro.setText(traducir('Alinear centro'))
            if hasattr(self, 'accion_distribuir_horizontal'):
                self.accion_distribuir_horizontal.setText(traducir('Distribuir horizontal'))
            if hasattr(self, 'accion_distribuir_vertical'):
                self.accion_distribuir_vertical.setText(traducir('Distribuir vertical'))
            if hasattr(self, 'accion_bloquear'):
                self.accion_bloquear.setText(traducir('Bloquear elementos'))
            if hasattr(self, 'accion_solo_cables'):
                self.accion_solo_cables.setText(traducir('Solo seleccionar cables'))
            if hasattr(self, 'accion_solo_componentes'):
                self.accion_solo_componentes.setText(traducir('Solo seleccionar componentes'))
            if hasattr(self, 'action_bode_top'):
                self.action_bode_top.setText(traducir('Gráfica Bode (AC)'))
            if hasattr(self, 'action_lgr_top'):
                self.action_lgr_top.setText(traducir('Gráfica LGR'))
            if hasattr(self, 'action_vlsm_top'):
                self.action_vlsm_top.setText(traducir('Calculadora de subred VSLM'))
            if hasattr(self, 'accion_calculadora'):
                self.accion_calculadora.setText(traducir('Calculadora de resistencias'))
            if hasattr(self, 'accion_vna'):
                self.accion_vna.setText(traducir('VNA View'))
                self.accion_vna.setToolTip(traducir_descripcion_emergente(self.accion_vna.toolTip()))
            if hasattr(self, 'accion_export_ext'):
                self.accion_export_ext.setText(traducir('Exportar Proyecto Extendido'))
                self.accion_export_ext.setToolTip(traducir_descripcion_emergente(self.accion_export_ext.toolTip()))

            if hasattr(self, 'buscador'):
                try:
                    self.buscador.setPlaceholderText(traducir('Buscar componente...'))
                except Exception:
                    pass

            try:
                if hasattr(self, 'asegurar_menu_ver'):
                    self.asegurar_menu_ver()
            except Exception:
                pass

            try:
                if hasattr(self, "accion_ampliar"):
                    toolbar.addAction(self.accion_ampliar)
                if hasattr(self, "accion_reducir"):
                    toolbar.addAction(self.accion_reducir)
                if hasattr(self, "accion_ver_todo"):
                    toolbar.addAction(self.accion_ver_todo)
                if hasattr(self, "accion_vista_11"):
                    toolbar.addAction(self.accion_vista_11)
            except Exception:
                pass

            try:
                if hasattr(self, 'accion_guardar_toolbar'):
                    self.accion_guardar_toolbar.setText(traducir('Guardar Circuito'))
                    self.accion_guardar_toolbar.setToolTip(traducir_descripcion_emergente('Guardar Circuito'))
                
                if hasattr(self, 'accion_restablecer_toolbar'):
                    self.accion_restablecer_toolbar.setText(traducir('Reestablecer'))
                    self.accion_restablecer_toolbar.setToolTip(traducir_descripcion_emergente('Reestablecer'))
                
                if hasattr(self, 'accion_rehacer_toolbar'):
                    self.accion_rehacer_toolbar.setText(traducir('Rehacer'))
                    self.accion_rehacer_toolbar.setToolTip(traducir_descripcion_emergente('Rehacer'))
                
                if hasattr(self, 'accion_cargar_toolbar'):
                    self.accion_cargar_toolbar.setText(traducir('Cargar'))
                    self.accion_cargar_toolbar.setToolTip(traducir_descripcion_emergente('Cargar'))
                
                if hasattr(self, 'accion_exportar_imagen_toolbar'):
                    self.accion_exportar_imagen_toolbar.setText(traducir('Exportar Imagen'))
                    self.accion_exportar_imagen_toolbar.setToolTip(traducir_descripcion_emergente('Exportar Imagen'))
                
                if hasattr(self, 'accion_exportar_pdf_toolbar'):
                    self.accion_exportar_pdf_toolbar.setText(traducir('Exportar PDF'))
                    self.accion_exportar_pdf_toolbar.setToolTip(traducir_descripcion_emergente('Exportar PDF'))
                
                if hasattr(self, 'accion_reporte_word_toolbar'):
                    self.accion_reporte_word_toolbar.setText(traducir('Reporte Word completo'))
                    self.accion_reporte_word_toolbar.setToolTip(traducir_descripcion_emergente('Reporte Word completo'))
                
                if hasattr(self, 'accion_exportar_resultados'):
                    self.accion_exportar_resultados.setText(traducir('Exportar Resultados'))
                    self.accion_exportar_resultados.setToolTip(traducir_descripcion_emergente('Exportar Resultados'))
                
                if hasattr(self, 'accion_eliminar_toolbar'):
                    self.accion_eliminar_toolbar.setText(traducir('Eliminar'))
                    self.accion_eliminar_toolbar.setToolTip(traducir_descripcion_emergente('Eliminar'))
                
                if hasattr(self, 'accion_limpiar_toolbar'):
                    self.accion_limpiar_toolbar.setText(traducir('Limpiar Todo'))
                    self.accion_limpiar_toolbar.setToolTip(traducir_descripcion_emergente('Limpiar Todo'))
                
                if hasattr(self, 'accion_texto_toolbar'):
                    self.accion_texto_toolbar.setText(traducir('Texto'))
                    self.accion_texto_toolbar.setToolTip(traducir_descripcion_emergente('Agregar etiqueta de texto al circuito'))
                
                if hasattr(self, 'accion_exportar_csv_toolbar'):
                    self.accion_exportar_csv_toolbar.setText(traducir('Exportar CSV (Med./Gráf.)'))
                    self.accion_exportar_csv_toolbar.setToolTip(
                        traducir_descripcion_emergente('Exporta datos de mediciones y gráficas a un archivo CSV con timestamp'))
                
                if hasattr(self, 'accion_bom_toolbar'):
                    self.accion_bom_toolbar.setText(traducir('Lista de Materiales (BOM)'))
                    self.accion_bom_toolbar.setToolTip(traducir_descripcion_emergente('Mostrar y exportar la lista de materiales (BOM)'))
                
                if hasattr(self, 'accion_zoom_in_toolbar'):
                    self.accion_zoom_in_toolbar.setText(traducir('Zoom +'))
                    self.accion_zoom_in_toolbar.setToolTip(traducir_descripcion_emergente('Acercar vista'))
                if hasattr(self, 'accion_zoom_out_toolbar'):
                    self.accion_zoom_out_toolbar.setText(traducir('Zoom -'))
                    self.accion_zoom_out_toolbar.setToolTip(traducir_descripcion_emergente('Alejar vista'))
                
                if hasattr(self, 'accion_inspector_toolbar'):
                    self.accion_inspector_toolbar.setText(traducir('Inspector de Errores'))
                    self.accion_inspector_toolbar.setToolTip(traducir_descripcion_emergente('Ejecuta la inspección de errores de conexión'))
                if hasattr(self, 'accion_conectar_toolbar'):
                    self.accion_conectar_toolbar.setText(traducir('Conectar'))
                    self.accion_conectar_toolbar.setToolTip(
                        traducir_descripcion_emergente('Haz clic aquí para iniciar una conexión entre terminales')
                    )
                if hasattr(self, 'accion_mostrar_conexiones_toolbar'):
                    self.accion_mostrar_conexiones_toolbar.setText(traducir('Mostrar Conexiones'))
                    self.accion_mostrar_conexiones_toolbar.setToolTip(traducir('Mostrar Conexiones'))
            except Exception:
                pass

            if hasattr(self, 'accion_ayuda_atajos'):
                self.accion_ayuda_atajos.setText(traducir('Atajos de Teclado'))
                self.accion_ayuda_atajos.setToolTip(traducir_descripcion_emergente(self.accion_ayuda_atajos.toolTip()))
            if hasattr(self, 'menu_idiomas'):
                if hasattr(self, 'action_espanol'):
                    self.action_espanol.setText(traducir('Español'))
                if hasattr(self, 'action_ingles'):
                    self.action_ingles.setText(traducir('Inglés'))
        except Exception:
            pass

        try:
            file_actions = [
                getattr(self, attr) for attr in [
                    'accion_nueva_pestana_menu', 'accion_guardar_menu', 'accion_cargar_menu',
                    'accion_pdf_menu', 'accion_txt_menu', 'accion_png_menu',
                    'accion_bom_menu', 'accion_exportar_datos_menu', 'accion_reporte_completo',
                    'accion_reporte_word',
                    'accion_export_ext', 'accion_abrir_ext',
                    'accion_exportar_logica_pdf_menu', 'accion_exportar_logica_word_menu',
                    'accion_exportar_logica_txt_menu', 'accion_exportar_logica_excel_menu'
                ] if hasattr(self, attr)
            ]
            for act in file_actions:
                try:
                    key = act.objectName() or act.text()
                    act.setText(traducir(key))
                except Exception:
                    pass
            if hasattr(self, 'menu_recientes') and self.menu_recientes is not None:
                self.menu_recientes.setTitle(traducir('Recientes'))
        except Exception:
            pass

        try:
            sim_actions = [
                getattr(self, attr) for attr in [
                    'accion_simular_voltaje', 'accion_simular_corriente', 'accion_simular_todo',
                    'accion_monte_carlo', 'accion_export_netlist', 'accion_ejecutar_spice'
                ] if hasattr(self, attr)
            ]
            for act in sim_actions:
                try:
                    key = act.objectName() or act.text()
                    act.setText(traducir(key))
                except Exception:
                    pass
        except Exception:
            pass

        try:
            if hasattr(self, 'menu_herramientas'):
                self._reconstruir_menu_herramientas()
        except Exception:
            pass

        try:
            for tipo, widget in getattr(self, 'item_widgets', {}).items():
                try:
                    translated = traducir(tipo)
                    widget.etiqueta.setText(translated)
                except Exception:
                    pass
        except Exception:
            pass

        try:
            if hasattr(self, 'boton_agregar'):
                self.boton_agregar.setText(traducir('Agregar Componente'))
            if hasattr(self, 'boton_conectar'):
                self.boton_conectar.setText(traducir('Conectar'))
                self.boton_conectar.setToolTip(traducir_descripcion_emergente(self.boton_conectar.toolTip()))
            if hasattr(self, 'boton_simular'):
                self.boton_simular.setText(traducir('Simular Voltajes'))
                self.boton_simular.setToolTip(traducir_descripcion_emergente(self.boton_simular.toolTip()))
            if hasattr(self, 'boton_corrientes'):
                self.boton_corrientes.setText(traducir('Simular Corrientes'))
                self.boton_corrientes.setToolTip(traducir_descripcion_emergente(self.boton_corrientes.toolTip()))
            if hasattr(self, 'boton_simular_todo'):
                self.boton_simular_todo.setText(traducir('Simular todo'))
                self.boton_simular_todo.setToolTip(traducir_descripcion_emergente(self.boton_simular_todo.toolTip()))
            if hasattr(self, 'boton_simular_logica'):
                self.boton_simular_logica.setText(traducir('Simular lógica'))
                original_tooltip = getattr(self.boton_simular_logica, 'toolTip', lambda: '')()
                if original_tooltip:
                    self.boton_simular_logica.setToolTip(traducir_descripcion_emergente(original_tooltip))
                else:
                    self.boton_simular_logica.setToolTip(
                        traducir_descripcion_emergente('Simula lógica digital y muestra la tabla de verdad del circuito')
                    )
            if hasattr(self, 'inspector_title'):
                self.inspector_title.setText(traducir('Inspector de propiedades'))
        except Exception:
            pass

        try:
            if hasattr(self, 'accion_tema'):
                if getattr(self, 'modo_oscuro', False):
                    self.accion_tema.setText(traducir('Modo Claro'))
                else:
                    self.accion_tema.setText(traducir('Modo Oscuro'))
        except Exception:
            pass

        try:
            if hasattr(self, 'dock_resultados'):
                self.dock_resultados.setWindowTitle(traducir('Resultados y Mediciones'))
        except Exception:
            pass

    def ejecutar_comprobaciones_previas(self) -> list:
        try:
            model = EsimulatorCore.construir_modelo(self.componentes)
            issues = EsimulatorCore.ejecutar_comprobaciones(model)
            try:
                self.resaltar_problemas(issues)
            except Exception:
                pass
            return issues
        except Exception as e:
            try:
                self.resaltar_problemas([])
            except Exception:
                pass
            return [Issue("warning", "PRECHECK_FAIL",
                          f"No fue posible completar el chequeo eléctrico: {e}")]

    def resaltar_problemas(self, issues):
        try:
            for c in getattr(self, "componentes", []):
                try:
                    if hasattr(c, "svg_item"):
                        c.svg_item.setGraphicsEffect(None)
                except Exception:
                    pass
        except Exception:
            pass
        if not issues:
            return
        import re
        nombres = set()
        for i in issues:
            try:
                matches = re.findall(r"'([^']+)'", i.msg)
                for m in matches:
                    nombres.add(m.strip())
                m2 = re.search(r"\(([^()]+)\)", i.msg)
                if m2:
                    nombres.add(m2.group(1).strip())
            except Exception:
                continue

    def inspector_errores(self):
        try:
            issues = self.ejecutar_comprobaciones_previas()
        except Exception as exc:
            issues = [Issue("warning", "PRECHECK_FAIL", f"No fue posible realizar el chequeo: {exc}")]
        try:
            from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
            from PyQt5.QtCore import Qt
        except Exception:
            return
        has_errors = any(i.level == "error" for i in issues)
        has_warnings = any(i.level == "warning" for i in issues)
        if has_errors:
            html = problemas_a_html(issues)
            msg = QMessageBox(self)
            msg.setWindowTitle(traducir('Inspector de Errores'))
            msg.setTextFormat(Qt.RichText)
            msg.setIcon(QMessageBox.Critical)
            msg.setText(html)
            msg.exec_()
        elif has_warnings:
            html = problemas_a_html([i for i in issues if i.level == "warning"])
            msg = QMessageBox(self)
            msg.setWindowTitle(traducir('Inspector de Errores'))
            msg.setTextFormat(Qt.RichText)
            msg.setIcon(QMessageBox.Warning)
            msg.setText(html)
            msg.exec_()
        else:
            msg = QMessageBox(self)
            msg.setWindowTitle(traducir('Inspector de Errores'))
            msg.setIcon(QMessageBox.Information)
            msg.setText(f"<b>{traducir('Chequeo eléctrico')}:</b> {traducir('sin problemas.')}")
            msg.exec_()
        for c in getattr(self, "componentes", []):
            try:
                if getattr(c, "nombre", None) in nombres:
                    effect = QGraphicsColorizeEffect()
                    effect.setColor(Qt.red)
                    effect.setStrength(1.0)
                    try:
                        c.svg_item.setGraphicsEffect(effect)
                    except Exception:
                        continue
                    try:
                        anim = QPropertyAnimation(effect, b"strength")
                        anim.setStartValue(0.0)
                        anim.setEndValue(1.0)
                        anim.setDuration(500)
                        anim.setLoopCount(6)
                        anim.start()
                        c._blink_animation = anim
                    except Exception:
                        pass
            except Exception:
                continue

    def _mostrar_detalle_resultado(self, item):
        try:
            data = item.data(Qt.UserRole)
            if not data:
                return
            tipo, mensaje = data
            from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
            msg = QMessageBox(self)
            msg.setWindowTitle(f"Detalle de {tipo}")
            msg.setTextFormat(Qt.RichText)
            msg.setIcon(QMessageBox.Information)
            msg.setText(mensaje)
            msg.exec_()
        except Exception:
            pass

    def _inicializar_cache_iconos(self) -> None:
        if not hasattr(self, "_icon_cache"):
            self._icon_cache: Dict[str, QIcon] = {}

    def obtener_icono(self, filename: str) -> QIcon:
        self._inicializar_cache_iconos()
        if filename in self._icon_cache:
            return self._icon_cache[filename]
        p = IMAGE_DIR / Path(filename).name
        icon = QIcon(str(p)) if p.exists() else QIcon()
        self._icon_cache[filename] = icon
        return icon

    def ampliar(self) -> None:
        try:
            factor = 1.15
            if hasattr(self, 'view') and self.view is not None:
                self.view.scale(factor, factor)
        except Exception:
            pass

    def reducir(self) -> None:
        try:
            factor = 1.15
            if hasattr(self, 'view') and self.view is not None:
                self.view.scale(1.0 / factor, 1.0 / factor)
        except Exception:
            pass

    def ver_todo(self) -> None:
        try:
            from PyQt5.QtCore import Qt
            if hasattr(self, 'view') and hasattr(self, 'scene') and self.view is not None and self.scene is not None:
                bounding_rect = self.scene.itemsBoundingRect()
                if not bounding_rect.isNull():
                    self.view.fitInView(bounding_rect, Qt.KeepAspectRatio)
        except Exception:
            pass

    def vista_uno_a_uno(self) -> None:
        try:
            if hasattr(self, 'view') and self.view is not None:
                self.view.resetTransform()
        except Exception:
            pass

    def abrir_calculadora_resistencias(self) -> None:
        try:
            dlg = ResistenciasCalculatorDialog(self)
            dlg.exec_()
        except Exception as exc:
            try:
                mostrar_error(self, "abrir calculadora de resistencias", exc)
            except Exception:
                pass

    def abrir_calculadora_subred_vslm(self) -> None:
        try:
            dlg = VLSMCalculatorDialog(self)
            dlg.exec_()
        except Exception as exc:
            try:
                QMessageBox.critical(
                    self,
                    "Calculadora de subred VSLM",
                    f"No se pudo abrir la calculadora: {exc}",
                )
            except Exception:
                pass
            try:
                QMessageBox.critical(self, "Calculadora de resistencias",
                                     f"No se pudo abrir la calculadora: {exc}")
            except Exception:
                pass

    def abrir_vna(self) -> None:
        from pathlib import Path
        import sys
        import subprocess
        try:
            script_path = Path(__file__).resolve().parent / 'VNA3.1.51.py'
            if not script_path.exists():
                raise FileNotFoundError(f"No se encontró {script_path.name} en {script_path.parent}")
            subprocess.Popen([sys.executable, str(script_path)], cwd=str(script_path.parent))
        except Exception as exc:
            try:
                QMessageBox.critical(self, "VNA View", f"No se pudo abrir el módulo VNA: {exc}")
            except Exception:
                pass

    def exportar_proyecto_extendido(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar proyecto extendido",
            "proyecto_extendido.esimx",
            "Proyecto ESIM extendido (*.esimx *.zip);;Todos los archivos (*)",
        )
        if not ruta:
            return
        try:
            model = EsimulatorCore.construir_modelo(self.componentes)
            comps_data: list[dict[str, Any]] = []
            for c in model.components:
                comps_data.append(
                    {
                        "name": c.name,
                        "type": c.type,
                        "params": c.params,
                        "pins": [
                            {"net_id": p.net_id, "index": p.index}
                            for p in c.pins
                        ],
                    }
                )
            nets_data: list[dict[str, Any]] = []
            for n in model.nets.values():
                nets_data.append(
                    {
                        "id": n.id,
                        "is_gnd": n.is_gnd,
                        "pins": [
                            {"comp_name": p.comp_name, "index": p.index} for p in n.pins
                        ],
                    }
                )
            data = {"components": comps_data, "nets": nets_data}
            bom: dict[str, int] = {}
            for comp in model.components:
                bom[comp.type] = bom.get(comp.type, 0) + 1
            metadata = {
                "generator": "ESIMulador",
                "datetime": datetime.now().isoformat(),
                "component_count": len(model.components),
                "net_count": len(model.nets),
            }
            layout: list[dict[str, Any]] = []
            for comp in getattr(self, "componentes", []):
                try:
                    pos = comp.scenePos() if hasattr(comp, "scenePos") else comp.pos()
                    layout.append(
                        {
                            "name": getattr(comp, "nombre", ""),
                            "type": getattr(comp, "tipo", ""),
                            "x": float(pos.x()),
                            "y": float(pos.y()),
                            "vertical": bool(getattr(comp, "vertical", False)),
                            "invertir": bool(getattr(comp, "invertir", False)),
                            "scale": float(comp.scale()) if hasattr(comp, "scale") else 1.0,
                        }
                    )
                except Exception:
                    continue
            layers: list[dict[str, Any]] = []
            for comp in getattr(self, "componentes", []):
                try:
                    layers.append(
                        {
                            "name": getattr(comp, "nombre", ""),
                            "layer": float(comp.zValue()) if hasattr(comp, "zValue") else 0.0,
                        }
                    )
                except Exception:
                    continue
            netlist_lines: list[str] = ["* ESIMulador netlist"]
            try:
                gnd_id = self._obtener_nodo_gnd()
            except Exception:
                gnd_id = None
            nodemap: dict[Any, str] = {}
            next_num = 1
            def id_spice(node_id):
                nonlocal next_num
                if node_id is None:
                    return "0"
                if node_id in nodemap:
                    return nodemap[node_id]
                nodemap[node_id] = str(next_num)
                next_num += 1
                return nodemap[node_id]
            if gnd_id is not None:
                nodemap[gnd_id] = "0"
            for comp in getattr(self, "componentes", []):
                if not getattr(comp, "terminales", []):
                    continue
                ctype = getattr(comp, "tipo", "")
                if ctype == "GND":
                    continue
                if ctype in ["Resistencia", "Amperímetro", "Voltímetro", "Ohmímetro"] or "(eq." in getattr(comp, "nombre", ""):
                    pref = "R"
                elif ctype == "Capacitor":
                    pref = "C"
                elif ctype == "Bobina":
                    pref = "L"
                elif ctype in ["Fuente Voltaje", "Batería", "Generador de funciones"]:
                    pref = "V"
                elif ctype == "Fuente Corriente":
                    pref = "I"
                else:
                    pref = "R"
                nombre_spice = getattr(comp, "nombre", "").replace(" ", "_")
                nodos: list[str] = []
                for term in getattr(comp, "terminales", []):
                    try:
                        nodos.append(id_spice(getattr(term, "node_group", None)))
                    except Exception:
                        nodos.append("0")
                    if len(nodos) == 2:
                        break
                if len(nodos) < 2:
                    continue
                valor = "1"
                try:
                    params = getattr(comp, "parametros", [])
                    if params:
                        raw = params[0]
                        if ":" in raw:
                            val = raw.split(":", 1)[1].strip()
                        else:
                            val = raw.strip()
                        try:
                            valor_num = interpretar_valor_prefijo(val)
                            valor = f"{valor_num}"
                        except Exception:
                            valor = val
                except Exception:
                    valor = "1"
                netlist_lines.append(f"{pref}{nombre_spice} {nodos[0]} {nodos[1]} {valor}")
            netlist_str = "\n".join(netlist_lines) + "\n"
            sim_data: dict[str, Any] = {}
            try:
                sim_data["resultados"] = [
                    {"tipo": tipo, "mensaje": mensaje}
                    for tipo, mensaje in getattr(self, "resultados_simulaciones", [])
                ]
            except Exception:
                sim_data["resultados"] = []
            try:
                if getattr(self, "ultima_grafica", None):
                    volt_labels, volt_values, current_labels, current_values = self.ultima_grafica
                    sim_data["ultima_grafica"] = {
                        "volt_labels": volt_labels,
                        "volt_values": volt_values,
                        "current_labels": current_labels,
                        "current_values": current_values,
                    }
            except Exception:
                pass
            with zipfile.ZipFile(ruta, 'w', compression=zipfile.ZIP_DEFLATED) as zf:
                zf.writestr('circuit.json', json.dumps(data, indent=4, ensure_ascii=False))
                zf.writestr('bom.json', json.dumps(bom, indent=4, ensure_ascii=False))
                zf.writestr('metadata.json', json.dumps(metadata, indent=4, ensure_ascii=False))
                zf.writestr('layout.json', json.dumps(layout, indent=4, ensure_ascii=False))
                zf.writestr('layers.json', json.dumps(layers, indent=4, ensure_ascii=False))
                zf.writestr('netlist.cir', netlist_str)
                zf.writestr('simulation.json', json.dumps(sim_data, indent=4, ensure_ascii=False))
            QMessageBox.information(self, _("Información"), _("Proyecto exportado correctamente"))
            try:
                import os as _os
                self.last_folder = _os.path.dirname(ruta)
                preferences['last_folder'] = self.last_folder
                self.registrar_reciente(str(ruta))
                guardar_preferencias()
                self.statusBar().showMessage("Proyecto extendido exportado", 3000)
            except Exception:
                pass
            try:
                llamar_hooks("project_exported_extended", nombre_archivo=ruta, metadata=metadata)
            except Exception:
                pass
        except Exception as exc:
            try:
                logger.exception("Error al exportar proyecto extendido: %s", exc)
            except Exception:
                pass
            QMessageBox.critical(self, _("Error"), _("La exportación del proyecto ha fallado"))

    def distribuir_componentes_automaticamente(self) -> None:
        try:
            comps = list(self.componentes)
        except Exception:
            return
        n = len(comps)
        if n == 0:
            return
        import math
        cols = int(math.ceil(math.sqrt(n)))
        spacing_x = 120
        spacing_y = 120
        for idx, comp in enumerate(comps):
            row = idx // cols
            col = idx % cols
            x = col * spacing_x
            y = row * spacing_y
            try:
                pos = QPointF(x, y)
                comp.setPos(pos)
            except Exception:
                try:
                    comp.setPos(x, y)
                except Exception:
                    pass
        try:
            for comp in comps:
                for terminal in getattr(comp, 'terminales', []):
                    for conn in getattr(terminal, 'conexiones', []):
                        try:
                            conn.actualizar_posicion()
                        except Exception:
                            pass
        except Exception:
            pass

    def extraer_topologia(self):
        class UnionFind:
            def __init__(self):
                self.parent = {}

            def encontrar(self, x):
                if self.parent.get(x, x) != x:
                    self.parent[x] = self.encontrar(self.parent[x])
                return self.parent.get(x, x)

            def unir(self, x, y):
                self.parent[self.encontrar(x)] = self.encontrar(y)

        uf = UnionFind()
        for conexion in self.conexiones:
            uf.unir(conexion.t1, conexion.t2)

        nodo_ids = {}
        next_id = 0
        for comp in self.componentes:
            for terminal in comp.terminales:
                grupo = uf.encontrar(terminal)
                if grupo not in nodo_ids:
                    nodo_ids[grupo] = next_id
                    next_id += 1

        for comp in self.componentes:
            for terminal in comp.terminales:
                grupo = uf.encontrar(terminal)
                idx = nodo_ids.get(grupo, None)
                try:
                    terminal.node_group = idx 
                    terminal.node_idx = idx
                except Exception:
                    pass

        elementos = []
        k_C = 1e5  
        k_L = 1e-3  

        for comp in self.componentes:
            tipo = comp.tipo
            parametros = comp.parametros
            terminal_nodos = [nodo_ids[uf.encontrar(t)] for t in comp.terminales]
            try:
                valor_str = parametros[0].split(":")[1].strip()
                valor = interpretar_valor_prefijo(valor_str)
            except Exception:
                valor = 1.0

            if tipo in ["Resistencia", "Fuente Voltaje", "Fuente Corriente", "Batería", "Generador de funciones"]:
                tipo_fuente = "Fuente Voltaje" if tipo == "Generador de funciones" else tipo
                elementos.append({
                    "tipo": tipo_fuente,
                    "valor": valor,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })

            elif tipo in ["Amperímetro", "Voltímetro", "Ohmímetro"]:
                try:
                    if tipo == "Amperímetro":
                        r_interna = config.R_INT_A
                    elif tipo == "Voltímetro":
                        r_interna = config.R_INT_V
                    else:
                        r_interna = config.R_TEST_OHM
                except Exception:
                    if tipo == "Amperímetro":
                        r_interna = 1e-9
                    elif tipo == "Voltímetro":
                        r_interna = 1e12
                    else:
                        r_interna = 1e3
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": r_interna,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre + " (eq. instrumento)"
                })

            elif tipo == "Capacitor":
                elementos.append({
                    "tipo": "Capacitor",
                    "valor": valor,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })
            elif tipo == "Bobina":
                elementos.append({
                    "tipo": "Bobina",
                    "valor": valor,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })

            elif tipo == "Motor":
                resistencia_equivalente = valor if abs(valor) > 0 else 1.0
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": resistencia_equivalente,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre + " (eq. Motor)"
                })
                elementos.append({
                    "tipo": "Motor",
                    "valor": valor,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })
            elif tipo == "Transformador":
                V_in = V_out = I_in = I_out = None
                for param in parametros:
                    try:
                        clave, valor_str = param.split(":", 1)
                    except Exception:
                        continue
                        
                    clave_l = clave.strip().lower()
                    valor_limpo = valor_str.strip()
                    try:
                        num = interpretar_valor_prefijo(valor_limpo)
                    except Exception:
                        try:
                            parte_num = valor_limpo.split()[0]
                            num = float(parte_num.replace(",", "."))
                        except Exception:
                            num = None
                    if num is None:
                        continue
                    if "voltaje entrada" in clave_l:
                        V_in = num
                    elif "voltaje salida" in clave_l:
                        V_out = num
                    elif "corriente entrada" in clave_l:
                        I_in = num
                    elif "corriente salida" in clave_l:
                        I_out = num
                R_in = 1e12
                R_out = 1e12
                if V_in is not None and I_in is not None:
                    try:
                        if abs(I_in) > 0:
                            R_in = abs(V_in / I_in)
                        else:
                            R_in = 1e12
                    except Exception:
                        R_in = 1e12
                if V_out is not None and I_out is not None:
                    try:
                        if abs(I_out) > 0:
                            R_out = abs(V_out / I_out)
                        else:
                            R_out = 1e12
                    except Exception:
                        R_out = 1e12
                if len(terminal_nodos) >= 2:
                    nodos_prim = terminal_nodos[:2]
                    elementos.append({
                        "tipo": "Resistencia",
                        "valor": R_in,
                        "nodos": nodos_prim,
                        "nombre": comp.nombre + " (Primario)"
                    })
                if len(terminal_nodos) >= 4:
                    nodos_sec = terminal_nodos[2:4]
                    elementos.append({
                        "tipo": "Resistencia",
                        "valor": R_out,
                        "nodos": nodos_sec,
                        "nombre": comp.nombre + " (Secundario)"
                    })
            elif tipo == "OpAmp":
                elementos.append({
                    "tipo": "OpAmp",
                    "valor": valor,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })
            elif tipo == "Mosfet":
                mos_tipo = "N"
                vth_val = 2.0
                rds_val = 1.0
                for param in parametros:
                    try:
                        clave, valor_str = param.split(":", 1)
                    except Exception:
                        continue
                    clave_l = clave.strip().lower()
                    valor_s = valor_str.strip()
                    if "tipo" in clave_l:
                        if 'p' in valor_s.lower():
                            mos_tipo = "P"
                        else:
                            mos_tipo = "N"
                    elif "vth" in clave_l:
                        valor_clean = valor_s.replace("V", "").replace("v", "").strip()
                        try:
                            vth_val = interpretar_valor_prefijo(valor_clean)
                        except Exception:
                            try:
                                vth_val = float(valor_clean.replace(",", "."))
                            except Exception:
                                pass
                    elif "rds" in clave_l:
                        valor_clean = valor_s.replace("Ω", "").replace("ohm", "").replace("Ω", "").strip()
                        try:
                            rds_val = interpretar_valor_prefijo(valor_clean)
                        except Exception:
                            try:
                                rds_val = float(valor_clean.replace(",", "."))
                            except Exception:
                                pass
                elementos.append({
                    "tipo": "Mosfet",
                    "valor": {"tipo": mos_tipo, "Vth": vth_val, "Rds": rds_val},
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre
                })
            elif tipo == "GND":
                elementos.append({
                    "tipo": "GND",
                    "nodos": terminal_nodos
                })
            elif tipo == "Push Button":
                resistencia_equivalente = 1e12
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": resistencia_equivalente,
                    "nodos": terminal_nodos,
                    "nombre": f"{comp.nombre} (Abierto)"
                })
            elif tipo == "Push Button Cerrado":
                resistencia_equivalente = 1e-3
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": resistencia_equivalente,
                    "nodos": terminal_nodos,
                    "nombre": f"{comp.nombre} (Cerrado)"
                })

            elif tipo == "LED":
                
                Vf = 2.0
                If = 0.02
                
                for param in parametros:
                    try:
                        clave, valor_str = param.split(":")
                        clave = clave.strip()
                        valor_str = valor_str.strip()
                        valor_str_sin_unidad = valor_str.replace("V", "").replace("A", "").strip()
                        if clave.lower().startswith("vf"):
                            Vf = interpretar_valor_prefijo(valor_str_sin_unidad)
                        elif clave.lower().startswith("if"):
                            If = interpretar_valor_prefijo(valor_str_sin_unidad)
                    except Exception:
                        pass
                
                if If != 0:
                    resistencia_equivalente = Vf / If
                else:
                    resistencia_equivalente = 1e12
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": resistencia_equivalente,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre + " (eq. LED)"
                })

            elif tipo == "Diodo Zener":
                Vz = 5.0
                Iz = 0.05
                for param in parametros:
                    try:
                        clave, valor_str = param.split(":")
                        clave = clave.strip()
                        valor_str = valor_str.strip()
                        valor_str_sin_unidad = valor_str.replace("V", "").replace("A", "").strip()
                        if clave.lower().startswith("vz"):
                            Vz = interpretar_valor_prefijo(valor_str_sin_unidad)
                        elif "iz" in clave.lower():
                            Iz = interpretar_valor_prefijo(valor_str_sin_unidad)
                    except Exception:
                        pass
                if Iz != 0:
                    resistencia_equivalente = Vz / Iz
                else:
                    resistencia_equivalente = 1e12
                elementos.append({
                    "tipo": "Resistencia",
                    "valor": resistencia_equivalente,
                    "nodos": terminal_nodos,
                    "nombre": comp.nombre + " (eq. Zener)"
                })

            elif tipo == "Potenciómetro":
                            try:
                                valor_total_str = parametros[0].split(":")[1].strip()
                                cursor_str = parametros[1].split(":")[1].strip().replace('%', '')
                                R_total = interpretar_valor_prefijo(valor_total_str)
                                porcentaje = float(cursor_str) / 100.0
                                R1 = R_total * porcentaje
                                R2 = R_total * (1 - porcentaje)
                            except Exception:
                                R1 = R2 = 1.0

                            n1, n2, n3 = terminal_nodos
                            elementos.append({
                                "tipo": "Resistencia",
                                "valor": R1,
                                "nodos": [n1, n2],
                                "nombre": comp.nombre + " (R1)"
                            })
                            elementos.append({
                                "tipo": "Resistencia",
                                "valor": R2,
                                "nodos": [n2, n3],
                                "nombre": comp.nombre + " (R2)"
                            })
        return elementos

    def mostrar_resultados_graficos(self, nodo_idx, voltajes, nodo_gnd):
        from PyQt5.QtWidgets import QGraphicsTextItem
        from PyQt5.QtCore import QPointF, Qt

        for item in self.scene.items():
            if isinstance(item, QGraphicsTextItem) and getattr(item, "_sim_result", False):
                self.scene.removeItem(item)

        volt_dict = {nodo_gnd: 0.0}
        for nodo, idx in nodo_idx.items():
            volt_dict[nodo] = voltajes[idx]

        for comp in self.componentes:
            if not comp.terminales:
                continue
            nodo_0 = None
            if len(comp.terminales) >= 1:
                nodo_0 = comp.terminales[0]
            punto = comp.scenePos()
            texto = f"{comp.nombre}\n"
            for i, t in enumerate(comp.terminales):
                grupo = t
                for key, val in volt_dict.items():
                    if key == grupo:
                        texto += f"V{i}: {val:.2f} V\n"
                        break
            etiqueta = QGraphicsTextItem(texto.strip())
            etiqueta.setPos(punto + QPointF(10, -10))
            etiqueta.setDefaultTextColor(Qt.darkGreen)
            etiqueta._sim_result = True
            self.scene.addItem(etiqueta)

    def resolver_matriz_nodos(self):
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        import numpy as np
        try:
            self.statusBar().showMessage(traducir("Simulando voltajes..."), 0)
        except Exception:
            pass

        issues = self.ejecutar_comprobaciones_previas()
        has_errors = any(i.level == "error" for i in issues)
        has_warnings = any(i.level == "warning" for i in issues)
        
        if has_errors:
            from PyQt5.QtCore import Qt
            html = problemas_a_html(issues)
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Critical)
            msg_pre.setText(html)
            msg_pre.exec_()
            return
        if has_warnings:
            from PyQt5.QtCore import Qt
            html = problemas_a_html([i for i in issues if i.level == "warning"])
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Warning)
            msg_pre.setText(html)
            msg_pre.exec_()
        else:
            try:
                logger.info("Chequeo eléctrico: sin problemas.")
            except Exception:
                pass
        
        elementos = self.extraer_topologia()
        lista_fuentes_voltaje = [e for e in elementos if e['tipo'] in ['Fuente Voltaje', 'Batería']]
        M = len(lista_fuentes_voltaje)

        conexiones_efectivas = 0
        conteo_por_nodo = {}

        for e in elementos:
            for nodo in e["nodos"]:
                conteo_por_nodo[nodo] = conteo_por_nodo.get(nodo, 0) + 1
           
        nodos_conectados = [n for n, c in conteo_por_nodo.items() if c > 1] 
        if len(nodos_conectados) < 2:
           QMessageBox.warning(self, traducir("Simulación inválida"), traducir("No hay suficientes componentes conectados para simular."))
           return

        nodos = set()
        nodo_gnd = None
        for e in elementos:
            for n in e["nodos"]:
                nodos.add(n)
            if e["tipo"] == "GND":
                nodo_gnd = e["nodos"][0]

        if nodo_gnd is None:
           QMessageBox.critical(self, traducir("Error de simulación"), traducir("No se encontró GND.\n\nAgrega un nodo de tierra (GND) para poder simular el circuito."))
           return

        nodos = sorted(n for n in nodos if n != nodo_gnd)
        nodo_idx = {n: i for i, n in enumerate(nodos)}
        try:
            voltajes, corrientes, Req = self.analizar_circuito(elementos, nodo_gnd)
        except Exception as e:
            mensaje_err = str(e)
            if isinstance(mensaje_err, str) and ("singular matrix" in mensaje_err.lower()):
                mensaje_err = "No se puede resolver el circuito: el sistema es singular. Verifica las conexiones."
            QMessageBox.critical(self, traducir("Error de simulación"), traducir(mensaje_err))
            return

        volt_dict = {nodo_gnd: 0.0}
        for nodo, idx in nodo_idx.items():
            volt_dict[nodo] = voltajes[idx]

        def volt_nodo(grupo):
            return volt_dict.get(grupo, 0.0)

        mensaje = " <b>" + traducir("Voltajes por componente:") + "</b><br>"
        for e in elementos:
            tipo_e = e.get("tipo")
            if tipo_e in ["Resistencia", "Fuente Voltaje", "Batería"]:
                n1, n2 = e["nodos"]
                v1 = volt_nodo(n1)
                v2 = volt_nodo(n2)
                v_comp = v1 - v2
                mensaje += f"<b>{e['nombre']}</b>: {formatear_con_prefijo(abs(v_comp), 'V')}<br>"
            elif tipo_e == "Motor":
                n1, n2 = e["nodos"]
                v1 = volt_nodo(n1)
                v2 = volt_nodo(n2)
                v_comp = v1 - v2
                mensaje += f"<b>{e['nombre']}</b>: {formatear_con_prefijo(abs(v_comp), 'V')}<br>"
            elif tipo_e == "OpAmp":
                nodos_comp = e.get("nodos", [])
                if len(nodos_comp) == 3:
                    n_plus, n_minus, n_out = nodos_comp
                    v_plus = volt_nodo(n_plus)
                    v_minus = volt_nodo(n_minus)
                    v_out = volt_nodo(n_out)
                    mensaje += (
                        f"<b>{e['nombre']}</b>: V+: {formatear_con_prefijo(abs(v_plus), 'V')} "
                        f"V-: {formatear_con_prefijo(abs(v_minus), 'V')} "
                        f"Vout: {formatear_con_prefijo(abs(v_out), 'V')}<br>"
                    )
        for comp in self.componentes:
            if comp.tipo not in ["Transistor NPN", "Transistor PNP"]:
                continue
            term = comp.terminales
            if len(term) != 3:
                continue
            t_base, t_col, t_emit = term[0], term[1], term[2]
            n_base = getattr(t_base, 'node_group', None)
            n_col = getattr(t_col, 'node_group', None)
            n_emit = getattr(t_emit, 'node_group', None)
            if n_base is None or n_col is None or n_emit is None:
                continue
            Vb = volt_nodo(n_base)
            Vc = volt_nodo(n_col)
            Ve = volt_nodo(n_emit)
            if comp.tipo == "Transistor NPN":
                Vbe = Vb - Ve
                Vce = Vc - Ve
            else:
                Vbe = Ve - Vb
                Vce = Ve - Vc
            mensaje += f"<b>{comp.nombre}</b>: Vbe: {formatear_con_prefijo(abs(Vbe), 'V')}    Vce: {formatear_con_prefijo(abs(Vce), 'V')}<br>"
        tiene_fuente_corriente = any(e.get("tipo") == "Fuente Corriente" for e in elementos)
        if not tiene_fuente_corriente and Req is not None:
            mensaje += f"<br><b>Req:</b> {formatear_con_prefijo(abs(Req), 'Ω')}"
        
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        from PyQt5.QtCore import Qt
        msgbox = QMessageBox(self)
        msgbox.setWindowTitle(traducir("Resultados de Simulación"))
        msgbox.setTextFormat(Qt.RichText)
        msgbox.setText(mensaje)
        msgbox.setIcon(QMessageBox.Information)
        msgbox.exec_()
        try:
            self.agregar_resultado("voltajes", mensaje)
        except Exception:
            pass
        try:
            llamar_hooks("simulation_completed", voltajes=voltajes, corrientes=corrientes, Req=Req, mensaje=mensaje)
        except Exception:
            pass
        try:
            self.statusBar().showMessage(traducir("Simulación de voltajes completada"), 5000)
        except Exception:
            pass
            
    def __init__(self, vertical=False, invertir=False, nuevo_tab_callback=None):
        super().__init__()
        self.setWindowIcon(QIcon(ICON_PATH))
        self.nuevo_tab_callback = nuevo_tab_callback
        self.setWindowIcon(QIcon(ICON_PATH))

        self.setWindowTitle("ESIMulador")
        self.terminales = []
        self.vertical = vertical
        self.invertir = invertir

        toolbar = self.addToolBar("Herramientas")
        self.accion_tema = None
        
        from PyQt5.QtWidgets import QHBoxLayout
        from PyQt5.QtGui import QPixmap
        
        self.setGeometry(100, 100, 1100, 600)

        try:
            self.last_folder = preferences.get('last_folder', str(Path(__file__).resolve().parent))
        except Exception:
            self.last_folder = str(Path(__file__).resolve().parent)
        try:
            self.modo_oscuro = preferences.get('modo_oscuro', False)
        except Exception:
            self.modo_oscuro = False

        menu_bar = self.menuBar()

        self.menu_archivo = menu_bar.addMenu("&Archivo")
        self.menu_simulacion = menu_bar.addMenu("&Simulación")
        try:
            self.menu_edicion = QMenu("&Edición", self)
            menu_bar.insertMenu(self.menu_simulacion.menuAction(), self.menu_edicion)
        except Exception:
            try:
                self.menu_edicion = menu_bar.addMenu("&Edición")
            except Exception:
                self.menu_edicion = None

        if self.menu_edicion is not None:
            try:
                base_path = Path(__file__).resolve().parent
            except Exception:
                base_path = Path('.')
            icon_rest = self.obtener_icono("restablecer.png")
            self.accion_restablecer_menu = QAction(icon_rest, "Reestablecer", self)
            self.accion_restablecer_menu.setShortcut("Ctrl+Z")
            self.accion_restablecer_menu.triggered.connect(self.restablecer_accion)
            self.menu_edicion.addAction(self.accion_restablecer_menu)
            icon_redo = self.obtener_icono("rehacer.png")
            self.accion_rehacer_menu = QAction(icon_redo, "Rehacer", self)
            self.accion_rehacer_menu.setShortcut("Ctrl+Y")
            self.accion_rehacer_menu.triggered.connect(self.rehacer_accion)
            self.menu_edicion.addAction(self.accion_rehacer_menu)
            self.menu_edicion.addSeparator()
            icon_cut = self.obtener_icono("Cortar.png")
            self.accion_cortar = QAction(icon_cut, "Cortar", self)
            self.accion_cortar.setShortcut("Ctrl+X")
            self.accion_cortar.triggered.connect(self.cortar_componentes)
            self.menu_edicion.addAction(self.accion_cortar)
            icon_copy = self.obtener_icono("Copiar.png")
            self.accion_copiar = QAction(icon_copy, "Copiar", self)
            self.accion_copiar.setShortcut("Ctrl+C")
            self.accion_copiar.triggered.connect(self.copiar_componentes)
            self.menu_edicion.addAction(self.accion_copiar)
            icon_paste = self.obtener_icono("Pegar.png")
            self.accion_pegar = QAction(icon_paste, "Pegar", self)
            self.accion_pegar.setShortcut("Ctrl+V")
            self.accion_pegar.triggered.connect(self.pegar_componentes)
            self.menu_edicion.addAction(self.accion_pegar)
            icon_delete = self.obtener_icono("eliminar.png")
            self.accion_eliminar_menu = QAction(icon_delete, "Eliminar", self)
            self.accion_eliminar_menu.setShortcut("Delete")
            self.accion_eliminar_menu.triggered.connect(self.eliminar_seleccionado)
            self.menu_edicion.addAction(self.accion_eliminar_menu)
            self.menu_edicion.addSeparator()
            icon_select = self.obtener_icono("Seleccionar.png")
            self.accion_seleccionar = QAction(icon_select, "Seleccionar", self)
            self.accion_seleccionar.triggered.connect(self.seleccionar_componentes)
            self.menu_edicion.addAction(self.accion_seleccionar)
            icon_select_all = self.obtener_icono("Seleccionartodo.png")
            self.accion_select_all = QAction(icon_select_all, "Seleccionar todo", self)
            self.accion_select_all.setShortcut("Ctrl+A")
            self.accion_select_all.triggered.connect(self.seleccionar_todo)
            self.menu_edicion.addAction(self.accion_select_all)
            self.menu_edicion.addSeparator()
            icon_rotate = self.obtener_icono("Rotar.png")
            self.accion_rotar = QAction(icon_rotate, "Rotar", self)
            self.accion_rotar.setShortcut("Ctrl+R")
            self.accion_rotar.triggered.connect(self.rotar_componentes)
            self.menu_edicion.addAction(self.accion_rotar)

            try:
                icon_dup = self.obtener_icono("Copiar.png")
            except Exception:
                icon_dup = QIcon()
            self.accion_duplicar = QAction(icon_dup, traducir('Duplicar'), self)
            self.accion_duplicar.setShortcut('Ctrl+D')
            try:
                self.accion_duplicar.triggered.connect(self.duplicar_componentes)
            except Exception:
                pass
            self.menu_edicion.addAction(self.accion_duplicar)

            self.menu_alinear_distribuir = QMenu(traducir('Alinear/Distribuir'), self)
            self.accion_alinear_izquierda = QAction(traducir('Alinear izquierda'), self)
            self.accion_alinear_izquierda.triggered.connect(self.alinear_izquierda)
            self.menu_alinear_distribuir.addAction(self.accion_alinear_izquierda)
            self.accion_alinear_derecha = QAction(traducir('Alinear derecha'), self)
            self.accion_alinear_derecha.triggered.connect(self.alinear_derecha)
            self.menu_alinear_distribuir.addAction(self.accion_alinear_derecha)
            self.accion_alinear_centro = QAction(traducir('Alinear centro'), self)
            self.accion_alinear_centro.triggered.connect(self.alinear_centro)
            self.menu_alinear_distribuir.addAction(self.accion_alinear_centro)
            self.accion_distribuir_horizontal = QAction(traducir('Distribuir horizontal'), self)
            self.accion_distribuir_horizontal.triggered.connect(self.distribuir_horizontal)
            self.menu_alinear_distribuir.addAction(self.accion_distribuir_horizontal)
            self.accion_distribuir_vertical = QAction(traducir('Distribuir vertical'), self)
            self.accion_distribuir_vertical.triggered.connect(self.distribuir_vertical)
            self.menu_alinear_distribuir.addAction(self.accion_distribuir_vertical)
            self.menu_edicion.addMenu(self.menu_alinear_distribuir)

            try:
                if hasattr(self.menu_alinear_distribuir, 'menuAction'):
                    self.menu_alinear_distribuir.menuAction().setVisible(False)
            except Exception:
                pass


            self.accion_bloquear = QAction(traducir('Bloquear elementos'), self)
            self.accion_bloquear.setCheckable(True)
            self.accion_bloquear.toggled.connect(self.bloquear_elementos)
            self.menu_edicion.addAction(self.accion_bloquear)

            try:
                self.accion_bloquear.setVisible(False)
            except Exception:
                pass

            self.accion_solo_cables = QAction(traducir('Solo seleccionar cables'), self)
            self.accion_solo_cables.triggered.connect(self.solo_seleccionar_cables)
            self.menu_edicion.addAction(self.accion_solo_cables)
            self.accion_solo_componentes = QAction(traducir('Solo seleccionar componentes'), self)
            self.accion_solo_componentes.triggered.connect(self.solo_seleccionar_componentes)
            self.menu_edicion.addAction(self.accion_solo_componentes)
            try:
                self.accion_solo_cables.setVisible(False)
            except Exception:
                pass
            try:
                self.accion_solo_componentes.setVisible(False)
            except Exception:
                pass

            self.menu_edicion.addSeparator()

        self.menu_herramientas = menu_bar.addMenu("&Herramientas")
        self.menu_herramientas.aboutToShow.connect(self._reconstruir_menu_herramientas) 

        self.accion_toggle_nodos = QAction(traducir('Mostrar nodos'), self)
        self.accion_toggle_nodos.setCheckable(True)
        self.accion_toggle_nodos.toggled.connect(self.alternar_nodos)
        self.menu_herramientas.addAction(self.accion_toggle_nodos)

        self.accion_resaltado_red = QAction(traducir('Resaltado de red'), self)
        self.accion_resaltado_red.triggered.connect(self.resaltar_red)
        self.menu_herramientas.addAction(self.accion_resaltado_red)

        self.accion_ruteo_inteligente = QAction(traducir('Ruteo inteligente'), self)
        self.accion_ruteo_inteligente.triggered.connect(self.ruteo_inteligente)
        self.menu_herramientas.addAction(self.accion_ruteo_inteligente)

        self.accion_puentes = QAction(traducir('Puentes'), self)
        self.accion_puentes.setCheckable(True)
        self.accion_puentes.toggled.connect(self.alternar_puentes)
        self.menu_herramientas.addAction(self.accion_puentes)

        self.menu_plantillas = menu_bar.addMenu("&Plantillas")
        try:
            self.menu_plantillas.aboutToShow.connect(self._reconstruir_menu_plantillas)
        except Exception:
            pass
        try:
            self._reconstruir_menu_plantillas()
        except Exception:
            pass

        try:
            self.menu_insertar = QMenu("&Insertar", self)
            self.menu_ver = QMenu("&Ver", self)
            try:
                base_path = Path(__file__).resolve().parent
            except Exception:
                base_path = Path('.')

            icon_tierra = self.obtener_icono("Tierra.png")
            self.accion_tierra = QAction(icon_tierra, "Tierra", self)
            self.accion_tierra.setToolTip("Agregar un nodo de tierra (GND) al circuito")
            self.accion_tierra.triggered.connect(lambda: self.agregar_componente("GND"))
            self.menu_insertar.addAction(self.accion_tierra)

            icon_etiqueta = self.obtener_icono("Etiqueta_Cable.png")
            self.accion_etiqueta = QAction(icon_etiqueta, "Etiqueta Cable", self)
            self.accion_etiqueta.setToolTip("Permite colocar una etiqueta de texto en el circuito")
            self.accion_etiqueta.triggered.connect(self.iniciar_modo_texto)
            self.menu_insertar.addAction(self.accion_etiqueta)

            icon_conexion = self.obtener_icono("conexion.png")
            self.accion_conexion = QAction(icon_conexion, "Conexión", self)
            self.accion_conexion.setToolTip("Iniciar una conexión entre terminales")
            self.accion_conexion.triggered.connect(lambda: self.boton_conectar.setChecked(True))
            self.menu_insertar.addAction(self.accion_conexion)

            try:
                icon_etiqueta_red = self.obtener_icono("Etiqueta_Cable.png")
            except Exception:
                icon_etiqueta_red = QIcon()
            self.accion_etiquetas_red = QAction(icon_etiqueta_red, traducir('Etiquetas de red'), self)
            try:
                self.accion_etiquetas_red.setToolTip(traducir_descripcion_emergente('Permite agregar etiquetas de red como VCC, OUT, IN'))
            except Exception:
                pass
            self.accion_etiquetas_red.triggered.connect(self.etiquetas_red)
            self.menu_insertar.addAction(self.accion_etiquetas_red)

            icon_ampliar = self.obtener_icono("Ampliar.png")
            self.accion_ampliar = QAction(icon_ampliar, "Ampliar", self)
            self.accion_ampliar.setToolTip("Acercar la vista (zoom in)")
            self.accion_ampliar.setShortcut("Ctrl+=")
            self.accion_ampliar.triggered.connect(self.ampliar)
            self.menu_ver.addAction(self.accion_ampliar)

            icon_reducir = self.obtener_icono("Reducir.png")
            self.accion_reducir = QAction(icon_reducir, "Reducir", self)
            self.accion_reducir.setToolTip("Alejar la vista (zoom out)")
            self.accion_reducir.setShortcut("Ctrl+-")
            self.accion_reducir.triggered.connect(self.reducir)
            self.menu_ver.addAction(self.accion_reducir)

            icon_ver_todo = self.obtener_icono("Ver_Todo.png")
            self.accion_ver_todo = QAction(icon_ver_todo, traducir('Ver todo'), self)
            self.accion_ver_todo.setToolTip(traducir_descripcion_emergente('Ajustar la vista para mostrar todos los elementos'))
            self.accion_ver_todo.setShortcut("Ctrl+Shift+V")
            self.accion_ver_todo.triggered.connect(self.ver_todo)
            self.menu_ver.addAction(self.accion_ver_todo)

            icon_vista_11 = self.obtener_icono("Vista1,1.png")
            self.accion_vista_11 = QAction(icon_vista_11, traducir('Establecer escala 1:1'), self)
            self.accion_vista_11.setToolTip(traducir_descripcion_emergente('Restablecer la vista a escala 1:1'))
            self.accion_vista_11.setShortcut("Ctrl+Shift+1")
            self.accion_vista_11.triggered.connect(self.vista_uno_a_uno)
            self.menu_ver.addAction(self.accion_vista_11)

            menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_ver)
            menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_insertar)
        except Exception:
            pass

        try:
            if not hasattr(self, 'menu_ver'):
                self.menu_ver = QMenu(traducir('Ver'), self)
                self.accion_ampliar = QAction(traducir('Ampliar'), self)
                self.accion_ampliar.setShortcut("Ctrl+=")
                self.accion_ampliar.triggered.connect(self.ampliar)
                self.menu_ver.addAction(self.accion_ampliar)
                self.accion_reducir = QAction(traducir('Reducir'), self)
                self.accion_reducir.setShortcut("Ctrl+-")
                self.accion_reducir.triggered.connect(self.reducir)
                self.menu_ver.addAction(self.accion_reducir)
                self.accion_ver_todo = QAction(traducir('Ver todo'), self)
                self.accion_ver_todo.setShortcut("Ctrl+Shift+V")
                self.accion_ver_todo.triggered.connect(self.ver_todo)
                self.menu_ver.addAction(self.accion_ver_todo)
                self.accion_vista_11 = QAction(traducir('Establecer escala 1:1'), self)
                self.accion_vista_11.setShortcut("Ctrl+Shift+1")
                self.accion_vista_11.triggered.connect(self.vista_uno_a_uno)
                self.menu_ver.addAction(self.accion_vista_11)
                try:
                    menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_ver)
                except Exception:
                    try:
                        menu_bar.addMenu(self.menu_ver)
                    except Exception:
                        pass
            if not hasattr(self, 'menu_insertar'):
                self.menu_insertar = QMenu(traducir('Insertar'), self)
                try:
                    menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_insertar)
                except Exception:
                    try:
                        menu_bar.addMenu(self.menu_insertar)
                    except Exception:
                        pass
        except Exception:
            pass

        self.action_bode_top = QAction("Gráfica Bode (AC)", self)
        try:
            self.action_bode_top.triggered.connect(self.iniciar_bode)
        except Exception:
            self.action_bode_top.triggered.connect(self.mostrar_graficas)
        menu_bar.addAction(self.action_bode_top)
        self.action_lgr_top = QAction("Gráfica LGR", self)
        try:
            self.action_lgr_top.triggered.connect(self.iniciar_lgr)
        except Exception:
            pass

        menu_bar.addAction(self.action_lgr_top)

        self.action_vlsm_top = QAction("Calculadora de subred VSLM", self)
        try:
            self.action_vlsm_top.triggered.connect(self.abrir_calculadora_subred_vslm)
        except Exception:
            pass
        menu_bar.addAction(self.action_vlsm_top)

        self.menu_ayuda = menu_bar.addMenu("&Ayuda")
        try:
            self.accion_ayuda_atajos = QAction("Atajos de Teclado", self)
            self.accion_ayuda_atajos.setToolTip("Mostrar los atajos de teclado disponibles")
            self.accion_ayuda_atajos.triggered.connect(self.mostrar_atajos_teclado)
            self.menu_ayuda.addAction(self.accion_ayuda_atajos)
        except Exception:
            pass

        try:
            self.menu_idiomas = QMenu("Idiomas", self)
            self.menu_ayuda.addMenu(self.menu_idiomas)
            self.action_espanol = QAction("Español", self)
            self.action_espanol.triggered.connect(lambda: self.cambiar_idioma('es'))
            self.menu_idiomas.addAction(self.action_espanol)
            self.action_ingles = QAction("Inglés", self)
            self.action_ingles.triggered.connect(lambda: self.cambiar_idioma('en'))
            self.menu_idiomas.addAction(self.action_ingles)
        except Exception:
            pass

        try:
            self.accion_calculadora = QAction("Calculadora de resistencias", self)
            self.accion_calculadora.triggered.connect(self.abrir_calculadora_resistencias)
            menu_bar.insertAction(self.action_bode_top, self.accion_calculadora)
            self.accion_vna = QAction("VNA View", self)
            self.accion_vna.setToolTip("Abrir módulo VNA para mediciones de red")
            self.accion_vna.triggered.connect(self.abrir_vna)
            menu_bar.insertAction(self.action_bode_top, self.accion_vna)
        except Exception:
            pass

        try:
            self.accion_export_ext = QAction("Exportar Proyecto Extendido", self)
            self.accion_export_ext.setToolTip("Exporta el proyecto con layout, capas y netlist SPICE/Qucs")
            self.accion_export_ext.triggered.connect(self.exportar_proyecto_extendido)
            self.menu_archivo.addAction(self.accion_export_ext)
        except Exception:
            pass
        self.modo_medicion = None        
        self.medicion_nodos = []         
        self.medicion_labels = []        

        self.current_measure_instrument = None

        icon_restablecer = self.obtener_icono("restablecer.png")
        icon_eliminar = self.obtener_icono("eliminar.png")
        icon_conectar = self.obtener_icono("conexion.png")
 
        icon_restablecer = self.obtener_icono("restablecer.png")
        icon_eliminar = self.obtener_icono("eliminar.png")
        icon_conectar = self.obtener_icono("conexion.png")
        icon_rehacer = self.obtener_icono("rehacer.png")
        icon_mostrar_conexiones = self.obtener_icono("Mostrar_Conexion.png")

        self.modo_oscuro = False
        self.resultados_simulaciones = []

        self.historial_mediciones: List[Tuple[str, str]] = []
        self.ultima_grafica: Optional[Tuple[List[str], List[float], List[str], List[float]]] = None

        self.guideline_items: List[Any] = []
        self.modo_texto: bool = False

        try:
            from PyQt5.QtCore import Qt
            from PyQt5.QtWidgets import QDockWidget, QListWidget

            self.dock_resultados = QDockWidget("Resultados y Mediciones", self)
            self.lista_resultados = QListWidget()
            self.dock_resultados.setWidget(self.lista_resultados)

            self.dock_resultados.setAllowedAreas(Qt.RightDockWidgetArea)

            self.dock_resultados.setFeatures(QDockWidget.DockWidgetClosable)

            self.removeDockWidget(self.dock_resultados)
            self.addDockWidget(Qt.RightDockWidgetArea, self.dock_resultados)

            self.dock_resultados.setMinimumWidth(260)
            try:
                self.lista_resultados.itemDoubleClicked.connect(self._mostrar_detalle_resultado)
            except Exception:
                pass
        except Exception:
            pass

        toolbar = self.addToolBar("Herramientas")

        from PyQt5.QtCore import Qt

        self.accion_exportar_resultados = QAction(self.obtener_icono("txt.png"), traducir("Exportar Resultados"), self)
        self.accion_exportar_resultados.setToolTip(traducir_descripcion_emergente("Exportar Resultados"))
        self.accion_exportar_resultados.triggered.connect(self.exportar_resultados)
        icon_guardar = self.obtener_icono("guardar.png")
        icon_cargar = self.obtener_icono("cargar.png")
        icon_imagen = self.obtener_icono("imagen.png")
        icon_pdf = self.obtener_icono("pdf.png")
        
        accion_guardar = QAction(icon_guardar, traducir("Guardar Circuito"), self)
        accion_guardar.setToolTip(traducir_descripcion_emergente("Guardar Circuito"))
        accion_guardar.triggered.connect(self.guardar_circuito)
        toolbar.addAction(accion_guardar)
     
        self.accion_guardar_toolbar = accion_guardar
        toolbar.addSeparator()
        toolbar.addSeparator()
        accion_restablecer = QAction(icon_restablecer, traducir("Reestablecer"), self)
        accion_restablecer.setToolTip(traducir_descripcion_emergente("Reestablecer"))
        accion_restablecer.setShortcut("Ctrl+Z")
        accion_restablecer.triggered.connect(self.restablecer_accion)
        toolbar.addAction(accion_restablecer)
        self.accion_restablecer_toolbar = accion_restablecer
        toolbar.addSeparator()
        accion_rehacer = QAction(icon_rehacer, traducir("Rehacer"), self)
        accion_rehacer.setToolTip(traducir_descripcion_emergente("Rehacer"))
        accion_rehacer.setShortcut("Ctrl+Y")
        accion_rehacer.triggered.connect(self.rehacer_accion)
        toolbar.addAction(accion_rehacer)
        self.accion_rehacer_toolbar = accion_rehacer
        toolbar.addSeparator()
        toolbar.addSeparator()
        icon_restablecer = self.obtener_icono("restablecer.png")
        icon_rehacer = self.obtener_icono("rehacer.png")
        icon_mostrar_conexiones = self.obtener_icono("Mostrar_Conexion.png")
        
        accion_cargar = QAction(self.obtener_icono("cargar.png"), traducir("Cargar Circuito"), self)
        accion_cargar.setToolTip(traducir_descripcion_emergente("Cargar Circuito"))
        accion_cargar.triggered.connect(self.cargar_circuito)
        accion_cargar2 = QAction(icon_cargar, traducir("Cargar"), self)
        accion_cargar2.setToolTip(traducir_descripcion_emergente("Cargar"))
        accion_cargar2.setShortcut("Ctrl+O")
        accion_cargar2.triggered.connect(self.cargar_circuito)
        toolbar.addAction(accion_cargar2)
        self.accion_cargar_toolbar = accion_cargar2
        toolbar.addSeparator()
        accion_exportar_img = QAction(self.obtener_icono("imagen.png"), traducir("Exportar Imagen"), self)
        accion_exportar_img.setToolTip(traducir_descripcion_emergente("Exportar Imagen"))
        accion_exportar_img.triggered.connect(self.exportar_imagen)
        accion_exportar_img2 = QAction(icon_imagen, traducir("Exportar Imagen"), self)
        accion_exportar_img2.setToolTip(traducir_descripcion_emergente("Exportar Imagen"))
        accion_exportar_img2.setShortcut("Ctrl+I")
        accion_exportar_img2.triggered.connect(self.exportar_imagen)
        toolbar.addAction(accion_exportar_img2)
        self.accion_exportar_imagen_toolbar = accion_exportar_img2
        toolbar.addSeparator()
        accion_exportar_pdf = QAction(icon_pdf, traducir("Exportar PDF"), self)
        accion_exportar_pdf.setToolTip(traducir_descripcion_emergente("Exportar PDF"))
        accion_exportar_pdf.setShortcut("Ctrl+E")
        accion_exportar_pdf.triggered.connect(self.exportar_pdf)
        toolbar.addAction(accion_exportar_pdf)
        self.accion_exportar_pdf_toolbar = accion_exportar_pdf

        try:
            icon_word_toolbar = self.obtener_icono("Word.png")
        except Exception:
            icon_word_toolbar = QIcon()
        accion_reporte_word_toolbar = QAction(icon_word_toolbar, traducir("Reporte Word completo"), self)
        accion_reporte_word_toolbar.setToolTip(traducir_descripcion_emergente("Reporte Word completo"))
        accion_reporte_word_toolbar.triggered.connect(self.exportar_reporte_word_completo)
        toolbar.addAction(accion_reporte_word_toolbar)
        toolbar.addAction(self.accion_exportar_resultados)
        self.accion_reporte_word_toolbar = accion_reporte_word_toolbar
        toolbar.addSeparator()
        accion_eliminar = QAction(icon_eliminar, traducir("Eliminar"), self)
        accion_eliminar.setToolTip(traducir_descripcion_emergente("Eliminar"))
        accion_eliminar.setShortcut("Delete")
        accion_eliminar.triggered.connect(self.eliminar_seleccionado)
        toolbar.addAction(accion_eliminar)
        self.accion_eliminar_toolbar = accion_eliminar
        icon_limpiar = QIcon("icons/limpiar_todo.png")
        accion_limpiar = QAction(icon_limpiar, traducir("Limpiar Todo"), self)
        accion_limpiar.setToolTip(traducir_descripcion_emergente("Limpiar Todo"))
        accion_limpiar.setShortcut("Ctrl+L")
        accion_limpiar.triggered.connect(self.limpiar_todo)
        toolbar.addAction(accion_limpiar)
        self.accion_limpiar_toolbar = accion_limpiar
        toolbar.addSeparator()

        try:
            icon_texto = self.obtener_icono("Texto.png")
        except Exception:
            icon_texto = QIcon()
        accion_texto = QAction(icon_texto, traducir("Texto"), self)
        accion_texto.setToolTip(traducir_descripcion_emergente("Agregar etiqueta de texto al circuito"))
        accion_texto.triggered.connect(self.iniciar_modo_texto)
        toolbar.addAction(accion_texto)
        self.accion_texto_toolbar = accion_texto
        toolbar.addSeparator()
        try:
            icon_csv = self.obtener_icono("Excel.png")
        except Exception:
            icon_csv = QIcon()
        self.accion_exportar_csv = QAction(icon_csv, traducir("Exportar CSV (Med./Gráf.)"), self)
        self.accion_exportar_csv.setToolTip(
            traducir_descripcion_emergente("Exporta datos de mediciones y gráficas a un archivo CSV con timestamp")
        )
        self.accion_exportar_csv.triggered.connect(self.exportar_mediciones_csv)
        toolbar.addAction(self.accion_exportar_csv)
        self.accion_exportar_csv_toolbar = self.accion_exportar_csv

        try:
            icono_bom_toolbar = QIcon(BOM_ICON_PATH)
        except Exception:
            icono_bom_toolbar = QIcon()
        accion_bom_toolbar = QAction(icono_bom_toolbar, traducir("Lista de Materiales (BOM)"), self)
        accion_bom_toolbar.setToolTip(
            traducir_descripcion_emergente("Mostrar y exportar la lista de materiales (BOM)")
        )
        accion_bom_toolbar.triggered.connect(self.mostrar_bom)
        toolbar.addAction(accion_bom_toolbar)
        self.accion_bom_toolbar = accion_bom_toolbar

        self.asegurar_menu_ver()

        toolbar.addSeparator()
        toolbar.addAction(self.accion_ampliar)
        toolbar.addAction(self.accion_reducir)
        toolbar.addAction(self.accion_ver_todo)
        toolbar.addAction(self.accion_vista_11)

        toolbar.addSeparator()
        self.accion_conectar_toolbar = QAction(icon_conectar, traducir('Conectar'), self)
        self.accion_conectar_toolbar.setShortcut("Ctrl+K")
        self.accion_conectar_toolbar.setToolTip(
            traducir_descripcion_emergente('Haz clic aquí para iniciar una conexión entre terminales')
        )
        self.accion_conectar_toolbar.triggered.connect(
            lambda: self.boton_conectar.setChecked(True)
        )
        toolbar.addAction(self.accion_conectar_toolbar)
        self.accion_mostrar_conexiones_toolbar = QAction(
            icon_mostrar_conexiones, traducir('Mostrar Conexiones'), self
        )
        self.accion_mostrar_conexiones_toolbar.setCheckable(True)
        self.accion_mostrar_conexiones_toolbar.setToolTip(traducir('Mostrar Conexiones'))
        self.accion_mostrar_conexiones_toolbar.triggered.connect(self.colorear_conexiones)
        toolbar.addAction(self.accion_mostrar_conexiones_toolbar)
        icon_guardar = QIcon("icons/guardar.png")
        icon_cargar = QIcon("icons/cargar.png")
        icon_imagen = QIcon("icons/imagen.png")
        icon_pdf = QIcon("icons/pdf.png")
        icon_eliminar = QIcon("icons/eliminar.png")
        icon_conectar = QIcon("icons/conexion.png")
        toolbar.setToolButtonStyle(Qt.ToolButtonIconOnly)

        icon_rehacer = QIcon("icons/rehacer.png")
        icon_mostrar_conexiones = QIcon("icons/Mostrar_Conexion.png")

        try:          
            self.menu_plugins = self.menuBar().addMenu("&Plugins")
            try:
                self.menu_plugins.menuAction().setVisible(False)
            except Exception:
                pass
        except Exception:
            self.menu_plugins = None

        try:
            try:
                icon_inspector = self.obtener_icono("Inspector.png")
            except Exception:
                icon_inspector = QIcon()
            accion_inspector = QAction(icon_inspector, traducir("Inspector de Errores"), self)
            accion_inspector.setToolTip(traducir_descripcion_emergente("Ejecuta la inspección de errores de conexión"))
            accion_inspector.triggered.connect(self.inspector_errores)
            toolbar.addAction(accion_inspector)
            self.accion_inspector_toolbar = accion_inspector
        except Exception:
            pass

        try:
            if self.menu_plugins is not None:
                try:
                    action_manage_plugins = QAction(_("Gestor de plugins"), self)
                    action_manage_plugins.triggered.connect(self._gestor_plugins)
                    self.menu_plugins.addAction(action_manage_plugins)
                    self.menu_plugins.menuAction().setVisible(True)
                except Exception:
                    pass
            for plugin in loaded_plugins:
                try:
                    if hasattr(plugin, "init_ui"):
                        plugin.init_ui(self)
                except Exception:
                    continue
        except Exception:
            pass

        self.imagenes = {
            "Resistencia": ("Resistencia.svg", 2, False, False, 0.16),
            "Potenciómetro": ("Potenciometro.svg", 3, False, False, 0.25),
            "Push Button": ("PushButton.svg", 2, False, False, 0.5),
            "Push Button Cerrado": ("PushButtonCerrado.svg", 2, False, False, 0.5),
            "Fuente Voltaje": ("FuenteVoltaje.svg", 2, False, False, 0.07),
            "Fuente Corriente": ("FuenteCorriente.svg", 2, False, False, 0.44),
            "Transformador": ("Transformador.svg", 4, False, False, 0.10),
            "Generador de funciones": ("GeneradorFunciones.svg", 2, False, False, 0.15),
            "Capacitor": ("Capacitor.svg", 2, False, False, 0.22),
            "Bobina": ("Bobina.svg", 2, False, False, 0.10),
            "Batería": ("Bateria.svg", 2, False, False, 0.06),        
            "Motor": ("Motor.svg", 2, False, False, 0.15),
            "GND": ("GND.svg", 1, False, False, 0.1),
            "LED": ("LED.svg", 2, False, False, 0.11),
            "Diodo Zener": ("DiodoZener.svg", 2, False, False, 0.08),
            "Mosfet": ("Mosfet.svg", 3, False, False, 0.2),
            "OpAmp": ("OpAmp.svg", 3, False, False, 0.39),
            "Transistor NPN": ("TransistorNPN.svg", 3, False, False, 0.08),
            "Transistor PNP": ("TransistorPNP.svg", 3, False, False, 0.08),
            "AND": ("AND.svg", 3, False, False, 0.30),
            "OR": ("OR.svg", 3, False, False, 0.30),
            "NAND": ("NAND.svg", 3, False, False, 0.30),
            "NOR": ("NOR.svg", 3, False, False, 0.30),
            "XOR": ("XOR.svg", 3, False, False, 0.13),
            "NOT": ("NOT.svg", 2, False, False, 0.15),
            "Amperímetro": ("Amperimetro.svg", 2, False, False, 0.15),
            "Voltímetro": ("Voltmetro.svg", 2, False, False, 0.15),
            "Ohmímetro": ("Ohmetro.svg", 2, False, False, 0.15),
        }
        self.imagenes = {
            tipo: (str(IMAGE_DIR / datos[0]), *datos[1:])
            for tipo, datos in self.imagenes.items()
        }

        self.buscador = QLineEdit(self)
        self.buscador.setPlaceholderText("Buscar componente...")
        self.buscador.textChanged.connect(self.filtrar_componentes)

        self.selector = QListWidget(self)
        self.selector.setFocusPolicy(Qt.NoFocus)
        self.selector.setViewMode(QListWidget.IconMode)
        self.selector.setIconSize(QSize(64, 64))
        self.selector.setResizeMode(QListWidget.Adjust)
        self.selector.setGridSize(QSize(120, 90))
        self.selector.setSpacing(10)
        self.selector.setMovement(QListWidget.Static)
        self.selector.setSelectionMode(QListWidget.SingleSelection)
        self.selector.setWrapping(True)
        self.selector.setVerticalScrollBarPolicy(Qt.ScrollBarAsNeeded)
        self.selector.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.selector.setMinimumWidth(250)
        self.selector.setStyleSheet("""
QListWidget {
    outline: none;
    border: none;
}
QListWidget::item {
    border: none;
    padding: 2px;
}
QListWidget::item:selected {
    background-color: #cce8ff;
    border: 2px solid #3399ff;
    border-radius: 6px;
}
QListWidget::item:focus,
QListWidget::item:selected:focus {
    outline: none;
}
""")

        self.items_selector = []
        self.item_widgets = {}
        self.item_widgets = {}
        for tipo in self.imagenes:
            icon = QIcon(self.imagenes[tipo][0])
            item = QListWidgetItem()
            item.setSizeHint(QSize(120, 100))
            item.setData(Qt.UserRole, tipo)

            nombre_mostrar = tipo
            if tipo == "Push Button":
                nombre_mostrar = "Push Button Abierto"

            widget = ItemComponente(icon, nombre_mostrar)

            self.items_selector.append((tipo, item))
            self.item_widgets[tipo] = widget
            self.selector.addItem(item)
            self.selector.setItemWidget(item, widget)

        self.boton_agregar = QPushButton("Agregar Componente", self)
        self.boton_agregar.setGeometry(20, 130, 250, 40)
        self.boton_agregar.clicked.connect(self.agregar_componente_desde_lista)
        
        self.scene = SceneConCuadricula()
        self.view = VistaConZoom(self.scene, self)
        self.view.setRenderHint(QPainter.Antialiasing)

        try:
            self.scene.selectionChanged.connect(self.actualizar_inspector)
        except Exception:
            pass

        try:
            from PyQt5.QtWidgets import QGraphicsView as _QGraphicsView
            self.view.setDragMode(_QGraphicsView.RubberBandDrag)
        except Exception:
            pass
        
        from PyQt5.QtWidgets import QVBoxLayout, QHBoxLayout
        
        self.boton_conectar = QPushButton("Conectar", self)
        self.boton_conectar.setCheckable(True)
        self.boton_conectar.setToolTip("Haz clic aquí para iniciar una conexión entre terminales")
        self.boton_conectar.clicked.connect(lambda: self.statusBar().showMessage(traducir('Modo conexión activado')))
        self.view.viewport().installEventFilter(self)
        self.dragging_connection = False
        self.drag_start_terminal = None
        self.pending_terminal = None
        self.preview_line = None
        self.highlighted_terminal = None
        self.drag_motion_detected = False
        self.nodes_visible = True
        try:
            self.scene.selectionChanged.connect(self._resaltar_red_al_seleccionar)
        except Exception:
            pass

        contenedor = QWidget()
        from PyQt5.QtWidgets import QSplitter
        splitter = QSplitter(Qt.Horizontal)

        panel_lateral = QVBoxLayout()
        panel_lateral.addWidget(self.buscador)
        panel_lateral.addWidget(self.selector)
        panel_lateral.addWidget(self.boton_agregar)
        panel_lateral.addWidget(self.boton_conectar)
        panel_lateral.addStretch()

        self.boton_simular = QPushButton("Simular Voltajes", self)
        self.boton_simular.setToolTip("Ejecuta la simulación de voltajes en el circuito")
        self.boton_simular.clicked.connect(self.resolver_matriz_nodos)
        panel_lateral.addWidget(self.boton_simular)

        self.boton_corrientes = QPushButton("Simular Corrientes", self)
        self.boton_corrientes.setToolTip("Ejecuta la simulación de corrientes en el circuito")
        self.boton_corrientes.clicked.connect(self.simular_corrientes)
        panel_lateral.addWidget(self.boton_corrientes)

        self.boton_simular_todo = QPushButton("Simular todo", self)
        self.boton_simular_todo.setToolTip("Simula voltajes y corrientes en una sola operación")
        self.boton_simular_todo.clicked.connect(self.simular_todo)
        panel_lateral.addWidget(self.boton_simular_todo)

        self.boton_simular_logica = QPushButton(traducir('Simular lógica'), self)
        self.boton_simular_logica.setToolTip(traducir_descripcion_emergente('Simula lógica digital y muestra la tabla de verdad del circuito'))
        self.boton_simular_logica.clicked.connect(self.simular_logica)
        panel_lateral.addWidget(self.boton_simular_logica)

        self.inspector_title = QLabel(traducir('Inspector de propiedades'), self)
        self.inspector_title.setStyleSheet("font-weight: bold;")
        self.inspector_widget = QWidget(self)
        self.inspector_layout = QFormLayout(self.inspector_widget)
        panel_lateral.addWidget(self.inspector_title)
        panel_lateral.addWidget(self.inspector_widget)
        self.inspector_title.setVisible(False)
        self.inspector_widget.setVisible(False)

        panel_widget = QWidget()
        panel_widget.setLayout(panel_lateral)
        splitter.addWidget(panel_widget)
        panel_widget.setMaximumWidth(300)
        panel_widget.setMinimumWidth(250)
        splitter.addWidget(self.view)
        splitter.setHandleWidth(0)
        splitter.setChildrenCollapsible(False)
        splitter.setSizes([280, 1000])

        splitter.setSizes([360, 800])
        self.setCentralWidget(splitter)

        try:
            self.menu_archivo.clear()
        except Exception:
            pass
        accion_nueva_pestana_menu = QAction(QIcon(), "Nueva pestaña", self)
        accion_nueva_pestana_menu.setObjectName("Nueva pestaña")
        self.accion_nueva_pestana_menu = accion_nueva_pestana_menu
        
        def crear_nueva_pestana():
            if callable(getattr(self, "nuevo_tab_callback", None)):
                try:
                    self.nuevo_tab_callback()
                except Exception:
                    pass
        accion_nueva_pestana_menu.triggered.connect(crear_nueva_pestana)
        self.menu_archivo.addAction(accion_nueva_pestana_menu)
        
        accion_guardar_menu = QAction(QIcon("icons/guardar.png"), "Guardar Circuito", self)
        accion_guardar_menu.setObjectName("Guardar Circuito")
        self.accion_guardar_menu = accion_guardar_menu
        accion_guardar_menu.triggered.connect(self.guardar_circuito)
        self.menu_archivo.addAction(accion_guardar_menu)
        
        accion_cargar_menu = QAction(QIcon("icons/cargar.png"), "Cargar", self)
        accion_cargar_menu.setObjectName("Cargar")
        self.accion_cargar_menu = accion_cargar_menu
        accion_cargar_menu.triggered.connect(self.cargar_circuito)
        self.menu_archivo.addAction(accion_cargar_menu)
        
        accion_pdf_menu = QAction(QIcon("icons/pdf.png"), "PDF", self)
        accion_pdf_menu.setObjectName("PDF")
        self.accion_pdf_menu = accion_pdf_menu
        accion_pdf_menu.triggered.connect(self.exportar_pdf)
        self.menu_archivo.addAction(accion_pdf_menu)

        accion_txt_menu = QAction(QIcon("icons/txt.png"), "Txt", self)
        accion_txt_menu.setObjectName("Txt")
        self.accion_txt_menu = accion_txt_menu
        accion_txt_menu.triggered.connect(self.exportar_resultados)
        self.menu_archivo.addAction(accion_txt_menu)

        accion_png_menu = QAction(QIcon("icons/imagen.png"), "PNG", self)
        accion_png_menu.setObjectName("PNG")
        self.accion_png_menu = accion_png_menu
        accion_png_menu.triggered.connect(self.exportar_imagen)
        self.menu_archivo.addAction(accion_png_menu)

        try:
            icon_bom = QIcon(BOM_ICON_PATH)
        except Exception:
            icon_bom = QIcon()
        accion_bom_menu = QAction(icon_bom, "Lista de Materiales (BOM)", self)
        accion_bom_menu.setObjectName("Lista de Materiales (BOM)")
        self.accion_bom_menu = accion_bom_menu
        accion_bom_menu.triggered.connect(self.mostrar_bom)
        self.menu_archivo.addAction(accion_bom_menu)

        try:
            icon_excel_menu = self.obtener_icono("Excel.png")
        except Exception:
            icon_excel_menu = QIcon()
        accion_exportar_datos_menu = QAction(icon_excel_menu, "Exportar datos y mediciones", self)
        accion_exportar_datos_menu.setObjectName("Exportar datos y mediciones")
        self.accion_exportar_datos_menu = accion_exportar_datos_menu
        accion_exportar_datos_menu.triggered.connect(self.exportar_mediciones_csv)
        self.menu_archivo.addAction(accion_exportar_datos_menu)

        try:
            icon_pdf_reporte = QIcon("icons/pdf.png")
        except Exception:
            icon_pdf_reporte = QIcon()
        accion_reporte_completo = QAction(icon_pdf_reporte, "Reporte PDF completo", self)
        accion_reporte_completo.setObjectName("Reporte PDF completo")
        self.accion_reporte_completo = accion_reporte_completo
        accion_reporte_completo.triggered.connect(self.exportar_reporte_pdf_completo)
        self.menu_archivo.addAction(accion_reporte_completo)

        try:
            icon_word_reporte = self.obtener_icono("Word.png")
        except Exception:
            icon_word_reporte = QIcon()
        accion_reporte_word = QAction(icon_word_reporte, "Reporte Word completo", self)
        accion_reporte_word.setObjectName("Reporte Word completo")
        self.accion_reporte_word = accion_reporte_word
        accion_reporte_word.triggered.connect(self.exportar_reporte_word_completo)
        self.menu_archivo.addAction(accion_reporte_word)
        try:
            self.menu_exportar_logica = QMenu(traducir('Exportar lógica'), self)
        except Exception:
            self.menu_exportar_logica = QMenu('Exportar lógica', self)
        try:
            icon_logica_pdf = QIcon("icons/pdf.png")
        except Exception:
            icon_logica_pdf = QIcon()
        accion_exportar_logica_pdf_menu = QAction(icon_logica_pdf, traducir('Exportar lógica a PDF'), self)
        accion_exportar_logica_pdf_menu.setObjectName('Exportar lógica a PDF')
        accion_exportar_logica_pdf_menu.triggered.connect(self.exportar_logica_pdf)
        self.accion_exportar_logica_pdf_menu = accion_exportar_logica_pdf_menu
        self.menu_exportar_logica.addAction(accion_exportar_logica_pdf_menu)

        try:
            icon_logica_word = self.obtener_icono("Word.png")
        except Exception:
            icon_logica_word = QIcon()
        accion_exportar_logica_word_menu = QAction(icon_logica_word, traducir('Exportar lógica a Word'), self)
        accion_exportar_logica_word_menu.setObjectName('Exportar lógica a Word')
        accion_exportar_logica_word_menu.triggered.connect(self.exportar_logica_word)
        self.accion_exportar_logica_word_menu = accion_exportar_logica_word_menu
        self.menu_exportar_logica.addAction(accion_exportar_logica_word_menu)
        try:
            icon_logica_txt = QIcon("icons/txt.png")
        except Exception:
            icon_logica_txt = QIcon()
        accion_exportar_logica_txt_menu = QAction(icon_logica_txt, traducir('Exportar lógica a TXT'), self)
        accion_exportar_logica_txt_menu.setObjectName('Exportar lógica a TXT')
        accion_exportar_logica_txt_menu.triggered.connect(self.exportar_logica_txt)
        self.accion_exportar_logica_txt_menu = accion_exportar_logica_txt_menu
        self.menu_exportar_logica.addAction(accion_exportar_logica_txt_menu)
        try:
            icon_logica_excel = self.obtener_icono("Excel.png")
        except Exception:
            icon_logica_excel = QIcon()
        accion_exportar_logica_excel_menu = QAction(icon_logica_excel, traducir('Exportar lógica a Excel'), self)
        accion_exportar_logica_excel_menu.setObjectName('Exportar lógica a Excel')
        accion_exportar_logica_excel_menu.triggered.connect(self.exportar_logica_excel)
        self.accion_exportar_logica_excel_menu = accion_exportar_logica_excel_menu
        self.menu_exportar_logica.addAction(accion_exportar_logica_excel_menu)
        self.menu_archivo.addMenu(self.menu_exportar_logica)

        accion_export_ext = QAction(QIcon("icons/zip.png"), "Exportar Proyecto Extendido", self)
        accion_export_ext.setObjectName("Exportar Proyecto Extendido")
        self.accion_export_ext = accion_export_ext
        accion_export_ext.setToolTip("Exporta layout, capas, parámetros de simulación y netlist SPICE/Qucs")
        accion_export_ext.triggered.connect(self.exportar_proyecto_extendido)
        self.menu_archivo.addAction(accion_export_ext)

        try:
            accion_abrir_ext = QAction(QIcon("icons/zip.png"), "Abrir Proyecto Extendido", self)
            accion_abrir_ext.setObjectName("Abrir Proyecto Extendido")
            self.accion_abrir_ext = accion_abrir_ext
            accion_abrir_ext.setToolTip("Abre un proyecto extendido existente")
            accion_abrir_ext.triggered.connect(self.abrir_proyecto_extendido)
        except Exception:
            pass

        try:
            self.menu_simulacion.clear()
        except Exception:
            pass
        accion_simular_voltaje = QAction("Simular Voltaje", self)
        accion_simular_voltaje.setObjectName("Simular Voltaje")
        self.accion_simular_voltaje = accion_simular_voltaje
        accion_simular_voltaje.setToolTip("Simula la distribución de voltajes en el circuito")
        accion_simular_voltaje.triggered.connect(self.resolver_matriz_nodos)
        self.menu_simulacion.addAction(accion_simular_voltaje)
        accion_simular_corrientes = QAction("Simular Corriente", self)
        accion_simular_corrientes.setObjectName("Simular Corriente")
        self.accion_simular_corriente = accion_simular_corrientes
        accion_simular_corrientes.setToolTip("Simula la distribución de corrientes en el circuito")
        accion_simular_corrientes.triggered.connect(self.simular_corrientes)
        self.menu_simulacion.addAction(accion_simular_corrientes)
        accion_simular_todo = QAction("Simular todo", self)
        accion_simular_todo.setObjectName("Simular todo")
        self.accion_simular_todo = accion_simular_todo
        accion_simular_todo.setToolTip("Simula tanto voltajes como corrientes en una sola operación")
        accion_simular_todo.triggered.connect(self.simular_todo)
        self.menu_simulacion.addAction(accion_simular_todo)

        accion_simular_logica = QAction(traducir('Simular lógica'), self)
        accion_simular_logica.setObjectName("Simular lógica")
        accion_simular_logica.setToolTip(traducir_descripcion_emergente('Simula lógica digital y muestra la tabla de verdad del circuito'))
        accion_simular_logica.triggered.connect(self.simular_logica)
        self.menu_simulacion.addAction(accion_simular_logica)
        accion_monte_carlo = QAction("Barrido Monte Carlo", self)
        accion_monte_carlo.setObjectName("Barrido Monte Carlo")
        self.accion_monte_carlo = accion_monte_carlo
        accion_monte_carlo.setToolTip("Realiza un análisis Monte Carlo de las tolerancias de los componentes")
        accion_monte_carlo.triggered.connect(self.iniciar_monte_carlo)
        self.menu_simulacion.addAction(accion_monte_carlo)

        try:
            accion_export_netlist = QAction("Exportar netlist SPICE/Qucs", self)
            accion_export_netlist.setObjectName("Exportar netlist SPICE/Qucs")
            self.accion_export_netlist = accion_export_netlist
            accion_export_netlist.setToolTip("Genera un netlist con prefijos/modelos seleccionables")
            accion_export_netlist.triggered.connect(self.exportar_netlist_modelos)
            accion_ejecutar_spice = QAction("Ejecutar SPICE/Qucs", self)
            accion_ejecutar_spice.setObjectName("Ejecutar SPICE/Qucs")
            self.accion_ejecutar_spice = accion_ejecutar_spice
            accion_ejecutar_spice.setToolTip("Ejecuta el netlist en un motor SPICE/Qucs externo")
            accion_ejecutar_spice.triggered.connect(self.ejecutar_netlist_spice)
        except Exception:
            pass

        self.componentes = []
        self.conexiones = []
        self.logic_conexiones = []
        self.contadores = {tipo: 0 for tipo in self.imagenes}
        self.terminal_seleccion = []
        self.historial = []
        self.rehacer_stack = []

        self._clipboard_components = None

        try:
            self.aplicar_tema()
        except Exception:
            pass

        try:
            self.menu_recientes = None
            self.actualizar_menu_recientes()
        except Exception:
            pass

    def registrar_reciente(self, ruta: str) -> None:
        try:
            import os as _os  
            recents = preferences.get('recent_files', []) or []
            ruta_norm = _os.path.abspath(ruta)
            if ruta_norm in recents:
                recents.remove(ruta_norm)
            recents.insert(0, ruta_norm)
            preferences['recent_files'] = recents[:10]
            guardar_preferencias()
        except Exception:
            pass
        try:
            self.actualizar_menu_recientes()
        except Exception:
            pass

    def actualizar_menu_recientes(self) -> None:
        try:
            from PyQt5.QtWidgets import QMenu, QAction
            import os as _os
            if self.menu_recientes is None:
                self.menu_recientes = QMenu("Recientes", self)
                try:
                    self.menu_archivo.addSeparator()
                    self.menu_archivo.addMenu(self.menu_recientes)
                except Exception:
                    pass
            self.menu_recientes.clear()
            recents = preferences.get('recent_files', []) or []
            if not recents:
                self.menu_recientes.setDisabled(True)
                return
            self.menu_recientes.setDisabled(False)
            for ruta in recents:
                nombre = _os.path.basename(ruta)
                accion = QAction(nombre, self)
                accion.setToolTip(ruta)
                accion.triggered.connect(lambda checked=False, p=ruta: self.cargar_circuito_archivo(p))
                self.menu_recientes.addAction(accion)
        except Exception:
            pass

    def filtrar_componentes(self, texto):
        texto = (texto or "").lower()
        for tipo, item in self.items_selector:
            widget = self.item_widgets.get(tipo)
            try:
                nombre = widget.etiqueta.text().lower()
            except Exception:
                nombre = ""
            search_terms = [nombre, str(tipo).lower()]
            for key, synonyms in COMPONENT_SYNONYMS.items():
                if key in nombre or key in str(tipo).lower():
                    search_terms.extend(synonyms)
            search_str = " ".join(search_terms)
            item.setHidden(texto not in search_str)

    def agregar_componente_desde_lista(self):
        item = self.selector.currentItem()
        if not item:
            return
        tipo = item.data(Qt.UserRole)
        self.agregar_componente(tipo_manual=tipo)

    def eventFilter(self, source, event):
        from PyQt5.QtCore import Qt
        try:
            if event.type() == event.MouseButtonDblClick:
                pos = self.view.mapToScene(event.pos()) if hasattr(self, 'view') else None
                try:
                    from PyQt5.QtWidgets import QGraphicsTextItem
                    from PyQt5.QtGui import QFont
                    from PyQt5.QtWidgets import QDialog as _QDialog
                except Exception:
                    QGraphicsTextItem = None
                    QFont = None
                    _QDialog = None
                item = self.scene.itemAt(pos, self.view.transform()) if hasattr(self, 'scene') and pos is not None else None
                if QGraphicsTextItem is not None and isinstance(item, QGraphicsTextItem) and getattr(item, '_is_label', False):
                    try:
                        dlg = TextFormatDialog(self)
                    except Exception:
                        dlg = None
                    if dlg is not None:
                        try:
                            dlg.text_input.setText(item.toPlainText())
                        except Exception:
                            pass
                        try:
                            current_font = item.font()
                            dlg.font_combo.setCurrentFont(current_font)
                            dlg.size_spin.setValue(current_font.pointSize())
                            dlg.chk_bold.setChecked(current_font.bold())
                            dlg.chk_italic.setChecked(current_font.italic())
                            dlg.chk_underline.setChecked(current_font.underline())
                        except Exception:
                            pass
                        try:
                            current_color = item.defaultTextColor()
                            dlg._selected_color = current_color
                            dlg._actualizar_boton_color()
                        except Exception:
                            pass
                        try:
                            result = dlg.exec_()
                        except Exception:
                            result = 0
                        accepted_const = getattr(_QDialog, 'Accepted', 1) if _QDialog is not None else 1
                        if result == accepted_const:
                            try:
                                values = dlg.obtener_valores()
                            except Exception:
                                values = None
                            if values and values.get('text'):
                                try:
                                    item.setPlainText(values['text'])
                                except Exception:
                                    pass
                                if QFont is not None:
                                    try:
                                        new_font = QFont(values['font_family'], values['font_size'])
                                        new_font.setBold(values.get('bold', False))
                                        new_font.setItalic(values.get('italic', False))
                                        new_font.setUnderline(values.get('underline', False))
                                        item.setFont(new_font)
                                    except Exception:
                                        pass
                                try:
                                    item.setDefaultTextColor(values['color'])
                                except Exception:
                                    pass
                                try:
                                    if hasattr(self, 'registrar_historial'):
                                        self.registrar_historial()
                                except Exception:
                                    pass
                            return True
        except Exception:
            pass

        if getattr(self, 'modo_texto', False) and event.type() == event.MouseButtonPress and event.button() == Qt.LeftButton:
            pos = self.view.mapToScene(event.pos())
            try:
                from PyQt5.QtWidgets import QDialog, QGraphicsTextItem
                from PyQt5.QtGui import QFont
            except Exception:
                QDialog = None
                QGraphicsTextItem = None
                QFont = None
            dlg = TextFormatDialog(self)
            result = dlg.exec_() if dlg else 0
            if dlg and result == getattr(QDialog, 'Accepted', 1):
                try:
                    values = dlg.obtener_valores()
                except Exception:
                    values = None
                if values and values.get('text'):
                    try:
                        etiqueta = QGraphicsTextItem(values['text'])
                        if QFont is not None:
                            fnt = QFont(values['font_family'], values['font_size'])
                            fnt.setBold(values.get('bold', False))
                            fnt.setItalic(values.get('italic', False))
                            fnt.setUnderline(values.get('underline', False))
                            etiqueta.setFont(fnt)
                        try:
                            etiqueta.setDefaultTextColor(values['color'])
                        except Exception:
                            pass
                        try:
                            from PyQt5.QtCore import Qt as _Qt
                            etiqueta.setTextInteractionFlags(_Qt.TextEditorInteraction)
                        except Exception:
                            pass
                        etiqueta.setFlag(QGraphicsItem.ItemIsMovable, True)
                        etiqueta.setFlag(QGraphicsItem.ItemIsSelectable, True)
                        setattr(etiqueta, "_is_label", True)
                        self.scene.addItem(etiqueta)
                        etiqueta.setPos(pos)
                    except Exception:
                        pass
                    try:
                        self.registrar_historial()
                    except Exception:
                        pass
            self.modo_texto = False
            return True

        if self.modo_medicion and event.type() == event.MouseButtonPress and event.button() == Qt.LeftButton:
            pos = self.view.mapToScene(event.pos())
            
            item = self.scene.itemAt(pos, self.view.transform())
            if isinstance(item, TerminalItem):
                try:
                    self.extraer_topologia()
                except Exception:
                    pass
                self.medicion_nodos.append(item)
                if len(self.medicion_nodos) == 1:
                    self.statusBar().showMessage(f"Seleccione el segundo terminal para medir {self.modo_medicion.lower()}")
                elif len(self.medicion_nodos) == 2:
                    t1, t2 = self.medicion_nodos[0], self.medicion_nodos[1]
                    modo = self.modo_medicion
                    self.modo_medicion = None
                    self.medicion_nodos = []
                    self.realizar_medicion(modo, t1, t2)
                return True
            else:
                self.statusBar().showMessage("Seleccione terminales válidos para medir", 2000)
                return True
        try:
            from PyQt5.QtWidgets import QGraphicsLineItem as _QGraphicsLineItem
        except Exception:
            _QGraphicsLineItem = None
        if event.type() == event.MouseButtonPress and event.button() == Qt.RightButton:
            if _QGraphicsLineItem is not None:
                pos = self.view.mapToScene(event.pos())
                itm = self.scene.itemAt(pos, self.view.transform())
                if isinstance(itm, _QGraphicsLineItem):
                    for _conn in list(getattr(self, 'conexiones', [])):
                        if itm in getattr(_conn, 'lines', []):
                            try:
                                _conn.eliminar()
                            except Exception:
                                pass
                            try:
                                self.conexiones.remove(_conn)
                            except ValueError:
                                pass
                            if hasattr(self, 'logic_conexiones') and _conn in self.logic_conexiones:
                                try:
                                    self.logic_conexiones.remove(_conn)
                                except ValueError:
                                    pass
                            try:
                                self.dragging_connection = False
                                self.drag_motion_detected = False
                                self.pending_terminal = None
                                self.drag_start_terminal = None
                                if hasattr(self, 'highlighted_terminal'):
                                    try:
                                        self.highlighted_terminal.setBrush(Qt.black)
                                    except Exception:
                                        pass
                                    self.highlighted_terminal = None
                                if hasattr(self, 'terminal_seleccion'):
                                    self.terminal_seleccion.clear()
                            except Exception:
                                pass
                            return True
        if self.boton_conectar.isChecked():
            if event.type() == event.MouseButtonPress and event.button() == Qt.LeftButton:
                pos = self.view.mapToScene(event.pos())
                itm = self.scene.itemAt(pos, self.view.transform())
                if not isinstance(itm, TerminalItem) and itm is not None:
                    try:
                        parent = itm.parentItem()
                    except Exception:
                        parent = None
                    if parent is not None and isinstance(parent, ComponenteInteractivo) and hasattr(parent, 'terminales'):
                        if hasattr(itm, 'logic_input_index'):
                            idx = getattr(itm, 'logic_input_index')
                            try:
                                idx_int = int(idx)
                            except Exception:
                                idx_int = -1
                            if 0 <= idx_int < len(parent.terminales):
                                itm = parent.terminales[idx_int]
                        elif getattr(itm, 'output_indicator_flag', False) or getattr(itm, 'output_label_flag', False):
                            if parent.terminales:
                                itm = parent.terminales[-1]
                if isinstance(itm, TerminalItem):
                    self.pending_terminal = itm
                    self.drag_start_terminal = itm
                    self._press_view_pos = event.pos()
                    self.drag_motion_detected = False
                    return True
                return False
            if event.type() == event.MouseMove:
                if getattr(self, 'pending_terminal', None) is None:
                    return False
                if not getattr(self, 'drag_motion_detected', False):
                    try:
                        start_pt = self._press_view_pos
                    except Exception:
                        start_pt = None
                    if start_pt is not None:
                        dx = event.pos().x() - start_pt.x()
                        dy = event.pos().y() - start_pt.y()
                        if (dx * dx + dy * dy) >= 25:
                            self.drag_motion_detected = True
                            self.dragging_connection = True
                            try:
                                from PyQt5.QtGui import QPen as _QPen
                            except Exception:
                                _QPen = None
                            if _QPen is not None and getattr(self, 'preview_line', None) is None:
                                try:
                                    from PyQt5.QtCore import Qt as _Qt
                                except Exception:
                                    _Qt = None
                                pen = _QPen(_Qt.darkGray, 1)
                                try:
                                    pen.setStyle(_Qt.DashLine)
                                except Exception:
                                    pass
                                start_scene = self.drag_start_terminal.scenePos()
                                self.preview_line = self.scene.addLine(start_scene.x(), start_scene.y(), start_scene.x(), start_scene.y(), pen)
                if getattr(self, 'dragging_connection', False):
                    current_scene = self.view.mapToScene(event.pos())
                    if getattr(self, 'preview_line', None) is not None:
                        ln = self.preview_line.line()
                        ln.setP2(current_scene)
                        self.preview_line.setLine(ln)
                    try:
                        from PyQt5.QtCore import QRectF as _QRectF
                    except Exception:
                        _QRectF = None
                    nearest = None
                    radius = 15
                    if _QRectF is not None:
                        rect = _QRectF(current_scene.x() - radius, current_scene.y() - radius, radius * 2, radius * 2)
                        candidates = self.scene.items(rect)
                    else:
                        candidates = []
                    for cand in candidates:
                        if isinstance(cand, TerminalItem) and cand is not self.drag_start_terminal:
                            try:
                                dist = abs(cand.scenePos().x() - current_scene.x()) + abs(cand.scenePos().y() - current_scene.y())
                            except Exception:
                                dist = radius + 1
                            if dist <= radius:
                                nearest = cand
                                break
                    prev_hlt = getattr(self, 'highlighted_terminal', None)
                    if nearest != prev_hlt:
                        if prev_hlt is not None:
                            try:
                                prev_hlt.setBrush(Qt.black)
                            except Exception:
                                pass
                            self.highlighted_terminal = None
                        if nearest is not None:
                            self.highlighted_terminal = nearest
                            try:
                                nearest.setBrush(Qt.red)
                            except Exception:
                                pass
                    return True
            if event.type() == event.MouseButtonRelease and event.button() == Qt.LeftButton:
                if getattr(self, 'dragging_connection', False):
                    if getattr(self, 'preview_line', None) is not None:
                        try:
                            self.scene.removeItem(self.preview_line)
                        except Exception:
                            pass
                        self.preview_line = None
                    tgt = getattr(self, 'highlighted_terminal', None)
                    if tgt is not None and tgt is not getattr(self, 'drag_start_terminal', None):
                        try:
                            self._conectar_terminales(self.drag_start_terminal, tgt)
                        except Exception:
                            pass
                    if getattr(self, 'highlighted_terminal', None) is not None:
                        try:
                            self.highlighted_terminal.setBrush(Qt.black)
                        except Exception:
                            pass
                        self.highlighted_terminal = None
                    self.dragging_connection = False
                    self.drag_motion_detected = False
                    self.pending_terminal = None
                    self.drag_start_terminal = None
                    return True
                else:
                    if getattr(self, 'pending_terminal', None) is not None:
                        itm = self.pending_terminal
                        self.pending_terminal = None
                        self.drag_start_terminal = None
                        self.terminal_seleccion.append(itm)
                        if len(self.terminal_seleccion) == 2:
                            t1, t2 = self.terminal_seleccion[0], self.terminal_seleccion[1]
                            try:
                                self._conectar_terminales(t1, t2)
                            except Exception:
                                pass
                            self.terminal_seleccion.clear()
                        return True
                return False
            return super().eventFilter(source, event)
        return super().eventFilter(source, event)

    def iniciar_medicion(self, tipo):
        for label in getattr(self, 'medicion_labels', []):
            try:
                self.scene.removeItem(label)
            except Exception:
                pass
        self.medicion_labels = []
        self.medicion_nodos = []
        self.modo_medicion = tipo
        try:
            self.boton_conectar.setChecked(False)
        except Exception:
            pass
        self.statusBar().showMessage(f"Modo medición {tipo.lower()}. Seleccione dos terminales para medir.")

    def iniciar_medicion_instrumento(self, tipo_inst):
        try:
            for label in getattr(self, 'medicion_labels', []):
                try:
                    self.scene.removeItem(label)
                except Exception:
                    pass
        except Exception:
            pass
        self.medicion_labels = []
        self.medicion_nodos = []
        try:
            self.agregar_componente(tipo_inst)
            if self.componentes:
                self.current_measure_instrument = self.componentes[-1]
            else:
                self.current_measure_instrument = None
        except Exception:
            self.current_measure_instrument = None
        self.modo_medicion = tipo_inst
        try:
            self.boton_conectar.setChecked(False)
        except Exception:
            pass
        self.statusBar().showMessage(
            f"Modo medición {tipo_inst.lower()}. Seleccione dos terminales para medir.", 5000
        )

    def iniciar_modo_texto(self):
        try:
            self.boton_conectar.setChecked(False)
        except Exception:
            pass
        self.modo_medicion = None
        self.modo_texto = True
        try:
            self.statusBar().showMessage(
                "Modo texto activado: haz clic en la escena para colocar texto.",
                5000
            )
        except Exception:
            pass

    def realizar_medicion(self, tipo, terminal1, terminal2):
        from PyQt5.QtWidgets import QMessageBox, QGraphicsTextItem
        from PyQt5.QtCore import QPointF, Qt
        try:
            elementos = self.extraer_topologia()
            nodo_gnd = None
            for comp in self.componentes:
                if getattr(comp, 'tipo', None) == "GND" and comp.terminales:
                    nodo_gnd = comp.terminales[0].node_group
                    break
            if nodo_gnd is None:
                QMessageBox.warning(self, "Medición imposible", "Debe existir un nodo GND para realizar mediciones.")
                return
            n1 = getattr(terminal1, 'node_group', None)
            n2 = getattr(terminal2, 'node_group', None)
            if n1 is None or n2 is None:
                QMessageBox.warning(self, "Medición no válida", "Los terminales seleccionados no están conectados a nodos válidos.")
                return
            valor_texto = ""
            if tipo == "Voltímetro":
                voltajes, corrientes, _ = self.analizar_circuito(elementos, nodo_gnd)
                nodos_set = set()
                for e in elementos:
                    for n in e["nodos"]:
                        nodos_set.add(n)
                nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
                nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
                v1 = 0.0 if n1 == nodo_gnd else voltajes[nodo_idx.get(n1, 0)]
                v2 = 0.0 if n2 == nodo_gnd else voltajes[nodo_idx.get(n2, 0)]
                dv = v1 - v2
                dv = abs(dv)
                try:
                    valor_texto = formatear_con_prefijo(dv, "V")
                except Exception:
                    valor_texto = f"{dv:.4g} V"
            elif tipo == "Amperímetro":
                try:
                    test_r = config.R_INT_A
                except Exception:
                    test_r = 1e-9
                nombre_temp = "A_med"
                test_element = {"tipo": "Resistencia", "valor": test_r, "nodos": [n1, n2], "nombre": nombre_temp}
                tmp_elementos = list(elementos) + [test_element]
                voltajes_tmp, corrientes_tmp, _ = self.analizar_circuito(tmp_elementos, nodo_gnd)
                nodos_set = set()
                for e in tmp_elementos:
                    for n in e["nodos"]:
                        nodos_set.add(n)
                nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
                nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
                v1 = 0.0 if n1 == nodo_gnd else voltajes_tmp[nodo_idx.get(n1, 0)]
                v2 = 0.0 if n2 == nodo_gnd else voltajes_tmp[nodo_idx.get(n2, 0)]
                dv = v1 - v2
                i_val = abs(dv / test_r)
                try:
                    valor_texto = formatear_con_prefijo(i_val, "A")
                except Exception:
                    valor_texto = f"{i_val:.4g} A"
            elif tipo == "Ohmímetro":
                nombre_temp = "V_med"
                test_element = {"tipo": "Fuente Voltaje", "valor": 1.0, "nodos": [n1, n2], "nombre": nombre_temp}
                tmp_elementos = list(elementos) + [test_element]
                voltajes_tmp, corrientes_tmp, _ = self.analizar_circuito(tmp_elementos, nodo_gnd)
                I = None
                for k, val in corrientes_tmp.items():
                    if k == nombre_temp:
                        I = val
                        break
                if I is None or abs(I) < 1e-12:
                    res_val = float('inf')
                else:
                    res_val = 1.0 / abs(I)
                try:
                    valor_texto = formatear_con_prefijo(abs(res_val), "Ω")
                except Exception:
                    valor_texto = f"{abs(res_val):.4g} Ω"
            else:
                return
            
            p1 = terminal1.scenePos()
            p2 = terminal2.scenePos()
            mid = QPointF((p1.x() + p2.x()) / 2.0, (p1.y() + p2.y()) / 2.0)

            inst = getattr(self, 'current_measure_instrument', None)
            if inst is not None and getattr(inst, 'tipo', None) == tipo:
                try:
                    inst.setPos(mid)
                    conn1 = Conexion(self.scene, terminal1, inst.terminales[0])
                    self.conexiones.append(conn1)
                    conn2 = Conexion(self.scene, terminal2, inst.terminales[1])
                    self.conexiones.append(conn2)
                    self._etiqueta_resultado(f"{inst.nombre}: {valor_texto}", inst)
                    self.current_measure_instrument = None
                    self.statusBar().showMessage(f"{tipo} = {valor_texto}", 5000)
                except Exception:
                    etiqueta = QGraphicsTextItem(f"{tipo}: {valor_texto}")
                    etiqueta.setDefaultTextColor(Qt.red)
                    etiqueta.setZValue(1000)
                    etiqueta.setPos(mid + QPointF(10, -10))
                    self.scene.addItem(etiqueta)
                    self.medicion_labels.append(etiqueta)
                    self.statusBar().showMessage(f"{tipo} = {valor_texto}", 5000)
            else:
                etiqueta = QGraphicsTextItem(f"{tipo}: {valor_texto}")
                etiqueta.setDefaultTextColor(Qt.red)
                etiqueta.setZValue(1000)
                etiqueta.setPos(mid + QPointF(10, -10))
                self.scene.addItem(etiqueta)
                self.medicion_labels.append(etiqueta)
                self.statusBar().showMessage(f"{tipo} = {valor_texto}", 5000)
            try:
                self.historial_mediciones.append((tipo, valor_texto))
            except Exception:
                pass
        except Exception as e: 
            mensaje = str(e)
            if isinstance(mensaje, str) and ("Singular matrix" in mensaje or "singular matrix" in mensaje):
                mensaje = "Error de medición o Error de conexión"
            QMessageBox.critical(self, "Error de medición", mensaje)
    
def _reconstruir_menu_herramientas(self):
    self.menu_herramientas.clear()

    m_volt = self.menu_herramientas.addMenu(traducir('Medir Voltaje'))
    m_volt.setObjectName('Medir Voltaje')
    
    for comp in [c for c in self.componentes if getattr(c, "tipo", "") == "Voltímetro"]:
        m_volt.addAction(f"{comp.nombre}", lambda c=comp: self._medir_con_voltimetro(c))

    m_amp = self.menu_herramientas.addMenu(traducir('Medir Corriente'))
    m_amp.setObjectName('Medir Corriente')
    for comp in [c for c in self.componentes if getattr(c, "tipo", "") == "Amperímetro"]:
        m_amp.addAction(f"{comp.nombre}", lambda c=comp: self._medir_con_amperimetro(c))

    m_ohm = self.menu_herramientas.addMenu(traducir('Medir Resistencia'))
    m_ohm.setObjectName('Medir Resistencia')
    for comp in [c for c in self.componentes if getattr(c, "tipo", "") == "Ohmímetro"]:
        m_ohm.addAction(f"{comp.nombre}", lambda c=comp: self._medir_con_ohmimetro(c))

    from PyQt5.QtWidgets import QMessageBox, QAction

    self.menu_herramientas.addSeparator()
    accion_mostrar_graficas = QAction(traducir('Mostrar Gráficas'), self)
    accion_mostrar_graficas.setObjectName('Mostrar Gráficas')
    self.accion_mostrar_graficas = accion_mostrar_graficas
    accion_mostrar_graficas.triggered.connect(self.mostrar_graficas)
    self.menu_herramientas.addAction(accion_mostrar_graficas)

    accion_osc = QAction(traducir('Osciloscopio'), self)
    accion_osc.setObjectName('Osciloscopio')
    self.accion_osc = accion_osc
    accion_osc.triggered.connect(self.mostrar_osciloscopio)
    self.menu_herramientas.addAction(accion_osc)

    try:
        accion_med = QAction(traducir('Mediciones activas'), self)
        accion_med.setObjectName('Mediciones activas')
        self.accion_mediciones_activas = accion_med
        accion_med.triggered.connect(self.mostrar_mediciones_activas)
        self.menu_herramientas.addAction(accion_med)
    except Exception:
        pass

    try:
        accion_script_console = QAction(traducir('Consola de scripts'), self)
        accion_script_console.setObjectName('Consola de scripts')
        self.accion_script_console = accion_script_console
        current_tip = accion_script_console.toolTip() or "Abrir consola de scripting interna (Python)"
        accion_script_console.setToolTip(traducir_descripcion_emergente(current_tip))
        accion_script_console.triggered.connect(self.abrir_consola_scripts)
        self.menu_herramientas.addAction(accion_script_console)
    except Exception:
        pass

def _nodo_gnd_actual(self):
    for comp in self.componentes:
        if getattr(comp, 'tipo', None) == "GND" and comp.terminales:
            return comp.terminales[0].node_group
    return None

def _etiqueta_resultado(self, texto, comp):
    from PyQt5.QtWidgets import QGraphicsTextItem
    from PyQt5.QtCore import QPointF, Qt
    try:
        for lbl in getattr(self, "medicion_labels", []):
            self.scene.removeItem(lbl)
    except Exception:
        pass
    if not hasattr(self, "medicion_labels"):
        self.medicion_labels = []
    try:
        pos = comp.scenePos()
    except Exception:
        try:
            pos = comp.pos()
        except Exception:
            pos = QPointF(0, 0)
    lbl = QGraphicsTextItem(texto)
    lbl.setDefaultTextColor(Qt.red)
    lbl.setZValue(1000)
    lbl.setPos(pos + QPointF(0, -35))
    try:
        self.scene.addItem(lbl)
    except Exception:
        pass
    self.medicion_labels = [lbl]

def _medir_con_voltimetro(self, comp):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        elementos = self.extraer_topologia()
        gnd = self._nodo_gnd_actual()
        if gnd is None:
            QMessageBox.warning(self, "Medición imposible", "Agrega un GND al circuito.")
            return
        if len(comp.terminales) < 2:
            QMessageBox.warning(self, "Instrumento inválido", "El voltímetro no tiene dos terminales.")
            return
        t1, t2 = comp.terminales[:2]
        n1, n2 = getattr(t1, 'node_group', None), getattr(t2, 'node_group', None)
        if n1 is None or n2 is None:
            QMessageBox.warning(self, "Medición no válida", "Conecta el instrumento a nodos válidos.")
            return
        voltajes, corrientes_dict, _ = self.analizar_circuito(elementos, gnd)
        nodos = sorted({n for e in elementos for n in e["nodos"] if n != gnd})
        idx = {n: i for i, n in enumerate(nodos)}
        nombre_eq = comp.nombre + " (eq. instrumento)"
        dv = None
        if nombre_eq in corrientes_dict:
            try:
                r_interna = None
                for e in elementos:
                    if e.get("nombre") == nombre_eq:
                        r_interna = e.get("valor")
                        break
                if r_interna is not None:
                    i_val = corrientes_dict[nombre_eq]
                    dv = abs(i_val * r_interna)
            except Exception:
                dv = None
        if dv is None:
            v1 = 0.0 if n1 == gnd else voltajes[idx.get(n1, 0)]
            v2 = 0.0 if n2 == gnd else voltajes[idx.get(n2, 0)]
            dv = abs(v1 - v2)
        try:
            texto = f"{comp.nombre}: {formatear_con_prefijo(dv, 'V')}"
        except Exception:
            texto = f"{comp.nombre}: {dv:.4g} V"
        self._etiqueta_resultado(texto, comp)
        self.statusBar().showMessage(texto, 5000)
        try:
            self.statusBar().showMessage("Sonda no intrusiva: medición sin alterar el circuito.", 5000)
        except Exception:
            pass
        try:
            if ':' in texto:
                _, val = texto.split(':', 1)
                self.historial_mediciones.append(("Voltímetro", val.strip()))
            else:
                self.historial_mediciones.append(("Voltímetro", texto.strip()))
        except Exception:
            pass
    except Exception as e:
        mensaje = str(e)
        if isinstance(mensaje, str) and ("Singular matrix" in mensaje or "singular matrix" in mensaje):
            mensaje = "Error de medición o Error de conexión"
        QMessageBox.critical(self, "Error de medición", mensaje)

def _medir_con_amperimetro(self, comp):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        elementos = self.extraer_topologia()
        gnd = self._nodo_gnd_actual()
        if gnd is None:
            QMessageBox.warning(self, "Medición imposible", "Agrega un GND al circuito.")
            return
        nombre_eq = comp.nombre + " (eq. instrumento)"
        voltajes, corrientes, _ = self.analizar_circuito(elementos, gnd)
        if nombre_eq in corrientes:
            i_val = corrientes[nombre_eq]
        else:
            t1, t2 = comp.terminales[:2]
            n1, n2 = getattr(t1, 'node_group', None), getattr(t2, 'node_group', None)
            if n1 is None or n2 is None:
                QMessageBox.warning(self, "Medición no válida", "Conecta el instrumento a nodos válidos.")
                return
            try:
                test_r = config.R_INT_A
            except Exception:
                test_r = 1e-9
            tmp = list(elementos) + [{"tipo":"Resistencia","valor":test_r,"nodos":[n1,n2],"nombre":"A_tmp"}]
            volt, _, _ = self.analizar_circuito(tmp, gnd)
            nodos = sorted({n for e in tmp for n in e["nodos"] if n != gnd})
            idx = {n:i for i,n in enumerate(nodos)}
            v1 = 0.0 if n1 == gnd else volt[idx.get(n1, 0)]
            v2 = 0.0 if n2 == gnd else volt[idx.get(n2, 0)]
            i_val = (v1 - v2) / test_r
        i_val = abs(i_val)
        try:
            texto = f"{comp.nombre}: {formatear_con_prefijo(i_val, 'A')}"
        except Exception:
            texto = f"{comp.nombre}: {i_val:.4g} A"
        self._etiqueta_resultado(texto, comp)
        self.statusBar().showMessage(texto, 5000)
        try:
            self.statusBar().showMessage("Sonda no intrusiva: medición sin alterar el circuito.", 5000)
        except Exception:
            pass
        try:
            if ':' in texto:
                _, val = texto.split(':', 1)
                self.historial_mediciones.append(("Amperímetro", val.strip()))
            else:
                self.historial_mediciones.append(("Amperímetro", texto.strip()))
        except Exception:
            pass
    except Exception as e:
        mensaje = str(e)
        if isinstance(mensaje, str) and ("Singular matrix" in mensaje or "singular matrix" in mensaje):
            mensaje = "Error de medición o Error de conexión"
        QMessageBox.critical(self, "Error de medición", mensaje)

def _medir_con_ohmimetro(self, comp):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        elementos = self.extraer_topologia()
        gnd = self._nodo_gnd_actual()
        if gnd is None:
            QMessageBox.warning(self, "Medición imposible", "Agrega un GND al circuito.")
            return
        if len(comp.terminales) < 2:
            QMessageBox.warning(self, "Instrumento inválido", "El ohmímetro no tiene dos terminales.")
            return
        t1, t2 = comp.terminales[:2]
        n1, n2 = getattr(t1, 'node_group', None), getattr(t2, 'node_group', None)
        if n1 is None or n2 is None:
            QMessageBox.warning(self, "Medición no válida", "Conecta el instrumento a nodos válidos.")
            return
        
        nombre_temp = "V_med_ohm"
        tmp = list(elementos) + [{"tipo":"Fuente Voltaje","valor":1.0,"nodos":[n1,n2],"nombre":nombre_temp}]
        _, corr, _ = self.analizar_circuito(tmp, gnd)
        I = corr.get(nombre_temp, 0.0)
        R = float('inf') if abs(I) < 1e-12 else 1.0/abs(I)
        try:
            texto = f"{comp.nombre}: {formatear_con_prefijo(abs(R), 'Ω')}"
        except Exception:
            texto = f"{comp.nombre}: {abs(R):.4g} Ω"
        self._etiqueta_resultado(texto, comp)
        self.statusBar().showMessage(texto, 5000)
        try:
            if ':' in texto:
                _, val = texto.split(':', 1)
                self.historial_mediciones.append(("Ohmímetro", val.strip()))
            else:
                self.historial_mediciones.append(("Ohmímetro", texto.strip()))
        except Exception:
            pass
    except Exception as e:
        mensaje = str(e)
        if isinstance(mensaje, str) and ("Singular matrix" in mensaje or "singular matrix" in mensaje):
            mensaje = "Error de medición o Error de conexión"
        QMessageBox.critical(self, "Error de medición", mensaje)

def mostrar_graficas(self):
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        try:
            import matplotlib.pyplot as plt
        except Exception:
            QMessageBox.warning(self, "Dependencia faltante", "matplotlib no está disponible para mostrar gráficas.")
            return
        try:
            elementos = self.extraer_topologia()
            nodo_gnd = None
            for comp in self.componentes:
                if getattr(comp, 'tipo', None) == "GND" and comp.terminales:
                    nodo_gnd = comp.terminales[0].node_group
                    break
            if nodo_gnd is None:
                QMessageBox.warning(self, "Gráficas no disponibles", "Debes incluir un nodo GND para simular y graficar.")
                return
            
            voltajes, corrientes_dict, _ = self.analizar_circuito(elementos, nodo_gnd)
            
            nodos_set = set()
            for e in elementos:
                for n in e["nodos"]:
                    nodos_set.add(n)
            nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
            nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
            
            volt_labels = []
            volt_values = []
            for elem in elementos:
                try:
                    n1, n2 = elem.get("nodos", [None, None])
                    v1 = 0.0 if n1 == nodo_gnd else voltajes[nodo_idx.get(n1, 0)]
                    v2 = 0.0 if n2 == nodo_gnd else voltajes[nodo_idx.get(n2, 0)]
                    dv = abs(v1 - v2)
                    volt_labels.append(elem.get("nombre", ""))
                    volt_values.append(dv)
                except Exception:
                    continue

            current_labels = []
            current_values = []
            for name, val in corrientes_dict.items():
                current_labels.append(name)
                current_values.append(val)

            try:
                self.ultima_grafica = (volt_labels, volt_values, current_labels, current_values)
            except Exception:
                self.ultima_grafica = None

            fig, axes = plt.subplots(2, 1, figsize=(8, 6))

            x_v = list(range(len(volt_labels)))
            y_v = [abs(v) for v in volt_values]
            axes[0].plot(x_v, y_v, marker='o', linestyle='-', color='tab:blue')
            axes[0].set_xticks(x_v)
            axes[0].set_xticklabels(volt_labels, rotation=45, ha='right', fontsize=8)
            axes[0].set_ylabel('Voltaje (V)')
            axes[0].set_title('Voltajes por elemento')
            axes[0].grid(True, linestyle='--', alpha=0.5)
            if y_v:
                y_min = min(y_v)
                y_max = max(y_v)
                if y_min == y_max:
                    margen = 0.1 * y_max if y_max != 0 else 1.0
                    axes[0].set_ylim(y_min - margen, y_max + margen)
                else:
                    margen = 0.1 * (y_max - y_min)
                    axes[0].set_ylim(y_min - margen, y_max + margen)

            x_i = list(range(len(current_labels)))
            y_i = [abs(i) for i in current_values]
            axes[1].plot(x_i, y_i, marker='o', linestyle='-', color='tab:green')
            axes[1].set_xticks(x_i)
            axes[1].set_xticklabels(current_labels, rotation=45, ha='right', fontsize=8)
            axes[1].set_ylabel('Corriente (A)')
            axes[1].set_title('Corrientes por elemento')
            axes[1].grid(True, linestyle='--', alpha=0.5)
            if y_i:
                yi_min = min(y_i)
                yi_max = max(y_i)
                if yi_min == yi_max:
                    margen = 0.1 * yi_max if yi_max != 0 else 1.0
                    axes[1].set_ylim(yi_min - margen, yi_max + margen)
                else:
                    margen = 0.1 * (yi_max - yi_min)
                    axes[1].set_ylim(yi_min - margen, yi_max + margen)

            try:
                axes[0].relim()
                axes[0].autoscale_view()
                axes[1].relim()
                axes[1].autoscale_view()
            except Exception:
                pass
            plt.tight_layout()
            plt.show()

        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron generar las gráficas: {e}")

def mostrar_osciloscopio(self):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        import numpy as np
        import matplotlib.pyplot as plt
    except Exception:
        QMessageBox.warning(self, "Dependencia faltante", "matplotlib o numpy no están disponibles para mostrar el osciloscopio.")
        return

    def base_tiempo_automatica(max_freq_hz: float, periods: int = 10, t_min: float = 0.002, t_max: float = 1.0) -> float:
        if max_freq_hz <= 0:
            return 1.0
        t_end = periods / max_freq_hz
        return max(t_min, min(t_end, t_max))

    def onda_triangular(ampl: float, freq: float, t: np.ndarray) -> np.ndarray:
        return (2.0 * ampl / np.pi) * np.arcsin(np.sin(2.0 * np.pi * freq * t))

    def onda_diente_sierra(ampl: float, freq: float, t: np.ndarray) -> np.ndarray:
        return 2.0 * ampl * ((t * freq) - np.floor((t * freq) + 0.5))

    signals = []
    max_amp = 0.0
    max_freq = 0.0
    for comp in self.componentes:
        try:
            tipo = getattr(comp, "tipo", "")
            params = getattr(comp, "parametros", [])
            nombre = getattr(comp, "nombre", "")
            if tipo == "Generador de funciones":
                val_v_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("voltaje")), None)
                val_f_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("frecuencia")), None)
                forma = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("forma")), "Senoidal")
                off_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("offset")), None)
                fase_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("fase")), None)
                duty_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("duty")), None)
                if val_v_str is None or val_f_str is None:
                    continue
                try:
                    amp = interpretar_valor_prefijo(val_v_str)
                    freq = interpretar_valor_prefijo(val_f_str)
                except Exception:
                    continue
                try:
                    offset_val = interpretar_valor_prefijo(off_str) if off_str else 0.0
                except Exception:
                    offset_val = 0.0
                try:
                    if fase_str:
                        fase_clean = fase_str.replace("°", "").replace("deg", "").replace("grados", "").strip()
                        fase_num = float(fase_clean.replace(",", "."))
                        phase_rad = fase_num * np.pi / 180.0
                    else:
                        phase_rad = 0.0
                except Exception:
                    phase_rad = 0.0
                try:
                    if duty_str:
                        duty_clean = duty_str.replace("%", "").strip()
                        duty_num = float(duty_clean.replace(",", ".")) / 100.0
                        duty_frac = max(0.0, min(duty_num, 1.0))
                    else:
                        duty_frac = 0.5
                except Exception:
                    duty_frac = 0.5
                max_amp = max(max_amp, abs(offset_val) + abs(amp))
                max_freq = max(max_freq, freq)
                forma_lower = forma.lower()
                if "cuadrada" in forma_lower:
                    def construir_y(t, A=amp, f=freq, off=offset_val, ph=phase_rad, duty=duty_frac):
                        if f <= 0:
                            base = 0.0
                        else:
                            base = (f * t + ph / (2.0 * np.pi))
                        frac = base % 1.0
                        s = np.where(frac < duty, 1.0, -1.0)
                        return off + A * s
                elif "triangular" in forma_lower:
                    def construir_y(t, A=amp, f=freq, off=offset_val, ph=phase_rad):
                        return off + onda_triangular(A, f, t + (ph / (2.0 * np.pi * f)) if f > 0 else t)
                elif "diente" in forma_lower or "sierra" in forma_lower:
                    def construir_y(t, A=amp, f=freq, off=offset_val, ph=phase_rad):
                        return off + onda_diente_sierra(A, f, t + (ph / (2.0 * np.pi * f)) if f > 0 else t)
                else:
                    def construir_y(t, A=amp, f=freq, off=offset_val, ph=phase_rad):
                        return off + A * np.sin(2.0 * np.pi * f * t + ph)
                signals.append((nombre, construir_y, amp, freq))
            elif tipo in ["Fuente Voltaje", "Batería"]:
                val_str = next((p.split(":", 1)[1].strip() for p in params if p.lower().startswith("tensión") or p.lower().startswith("voltaje")), None)
                if val_str is None:
                    continue
                try:
                    amp = interpretar_valor_prefijo(val_str)
                except Exception:
                    continue
                max_amp = max(max_amp, abs(amp))
                def construir_y_dc(t, A=amp):
                    return np.full_like(t, A, dtype=float)
                signals.append((nombre, construir_y_dc, amp, 0.0))
        except Exception:
            continue

    if not signals:
        QMessageBox.information(self, "Osciloscopio", "No se encontraron señales para mostrar. Agrega generadores de funciones o fuentes de voltaje/baterías.")
        return
    from PyQt5.QtWidgets import QInputDialog
    try:
        coupling_mode, ok = QInputDialog.getItem(self, "Osciloscopio", "Modo de acoplamiento:", ["AC", "DC"], 0, False)
        ac_coupling = (ok and coupling_mode == "AC")
    except Exception:
        ac_coupling = True
    try:
        trigger_str, ok_trig = QInputDialog.getText(self, "Osciloscopio", "Nivel de trigger (deja vacío para sin trigger):", text="0")
        if ok_trig and trigger_str.strip():
            trigger_enabled = True
            try:
                trigger_level = float(trigger_str)
            except Exception:
                trigger_level = 0.0
            edge_choice, ok_edge = QInputDialog.getItem(self, "Osciloscopio", "Tipo de trigger:", ["subida", "bajada"], 0, False)
            trigger_edge = (edge_choice if ok_edge else "subida")
        else:
            trigger_enabled = False
            trigger_level = 0.0
            trigger_edge = "subida"
    except Exception:
        trigger_enabled = False
        trigger_level = 0.0
        trigger_edge = "subida"
    try:
        export_csv = (ok_export and export_choice == "Sí")
    except Exception:
        export_csv = False

    if max_freq <= 0:
        max_freq = 1.0
    t_end = base_tiempo_automatica(max_freq, periods=10, t_min=0.002, t_max=1.0)
    samples_per_period = 500
    fs_min = 20000.0  
    fs = max(fs_min, max_freq * samples_per_period)
    num_points = int(max(2000, fs * t_end))
    t = np.linspace(0.0, t_end, num_points, endpoint=False)

    plt.style.use("default")
    fig, ax = plt.subplots(figsize=(10, 5), dpi=100)
    ax.set_facecolor("#000000")
    fig.patch.set_facecolor("#000000")

    processed_signals = []
    max_amp_ac = 0.0
    for (label, builder, A, f) in signals:
        try:
            y = builder(t)
            if ac_coupling:
                y = y - np.mean(y)
            processed_signals.append((label, y, A, f))
            max_amp_ac = max(max_amp_ac, np.max(np.abs(y)))
        except Exception:
            continue

    start_idx = 0
    if trigger_enabled and processed_signals:
        ref_y = processed_signals[0][1]
        try:
            for i in range(len(ref_y) - 1):
                if trigger_edge.startswith("sub"):
                    if ref_y[i] < trigger_level <= ref_y[i + 1]:
                        start_idx = i
                        break
                else:
                    if ref_y[i] > trigger_level >= ref_y[i + 1]:
                        start_idx = i
                        break
        except Exception:
            start_idx = 0
    if start_idx < len(t):
        t_plot = t[start_idx:] - t[start_idx]
    else:
        t_plot = t - t[0]

    for (label, y, A, f) in processed_signals:
        y_plot = y[start_idx:]
        try:
            formatted_amp = formatear_con_prefijo(A, "V")
        except Exception:
            formatted_amp = f"{A:.4g} V"
        if f == 0 or abs(f) < 1e-30:
            formatted_freq = "DC"
        else:
            try:
                formatted_freq = formatear_con_prefijo(f, "Hz")
            except Exception:
                formatted_freq = f"{f:.4g} Hz"
        plot_label = f"{label} (A={formatted_amp}, f={formatted_freq})"
        ax.plot(t_plot, y_plot, linewidth=1.2, label=plot_label, alpha=0.9)

    y_span = max(1.0, 1.2 * (max_amp_ac if max_amp_ac > 0 else 1.0))
    if len(t_plot) > 0:
        ax.set_xlim(0.0, t_plot[-1])
    else:
        ax.set_xlim(0.0, t_end)
    ax.set_ylim(-y_span, y_span)

    if export_csv and processed_signals:
        try:
            import os, datetime
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"oscilloscope_export_{timestamp}.csv"
            filepath = os.path.join(str(Path(__file__).resolve().parent), filename)
            cols = [t_plot]
            for (_, y_vals, _, _) in processed_signals:
                cols.append(y_vals[start_idx:])
            data = np.column_stack(cols)
            header = "t," + ",".join([ps[0] for ps in processed_signals])
            np.savetxt(filepath, data, delimiter=",", header=header, comments="")
            QMessageBox.information(self, "Exportación CSV", f"Datos exportados a {filepath}")
        except Exception as ex:
            try:
                mostrar_error(self, "exportar CSV de osciloscopio", ex)
            except Exception:
                pass
            QMessageBox.warning(self, "Exportación CSV", f"No se pudo exportar a CSV: {ex}")
    ax.set_title("Osciloscopio", color="white")
    ax.set_xlabel("Tiempo (s)", color="white")
    ax.set_ylabel("Amplitud (V)", color="white")
    ax.grid(True, which="both", linestyle="--", alpha=0.35, color="white")
    ax.tick_params(colors="white")
    legend = ax.legend(loc="upper right")
    if legend:
        for text in legend.get_texts():
            text.set_color("white")
   
    cursor_hist = []  
    vline = ax.axvline(x=0, color="yellow", linestyle="--", alpha=0.8)
    hline = ax.axhline(y=0, color="yellow", linestyle="--", alpha=0.8)
    annotation = ax.text(0.01, 0.99, "", transform=ax.transAxes, color="yellow",
                         fontsize=8, verticalalignment="top",
                         bbox=dict(facecolor="black", alpha=0.6, pad=3.0))

    def al_hacer_clic(event):
        if event.inaxes != ax:
            return
        x = event.xdata
        y = event.ydata
        vline.set_xdata(x)
        hline.set_ydata(y)
        cursor_hist.append((x, y))
        lineas = [f"t = {x:.5g} s", f"V = {y:.5g} V"]
        if len(cursor_hist) >= 2:
            dx = cursor_hist[-1][0] - cursor_hist[-2][0]
            dy = cursor_hist[-1][1] - cursor_hist[-2][1]
            dt = abs(dx)
            dv = abs(dy)
            lineas.append(f"Δt = {dt:.5g} s")
            if dt > 0:
                fval = 1.0 / dt
                lineas.append(f"f = {fval:.5g} Hz")
            lineas.append(f"ΔV = {dv:.5g} V")
        annotation.set_text("\n".join(lineas))
        fig.canvas.draw_idle()

    fig.canvas.mpl_connect("button_press_event", al_hacer_clic)

    plt.tight_layout()
    plt.show()

def actualizar_inspector(self):
    try:
        seleccion = [item for item in self.scene.selectedItems() if isinstance(item, ComponenteInteractivo)]
    except Exception:
        seleccion = []
    if len(seleccion) != 1:
        try:
            self.inspector_title.setVisible(False)
            self.inspector_widget.setVisible(False)
        except Exception:
            pass
        try:
            while self.inspector_layout.count():
                item = self.inspector_layout.takeAt(0)
                w = item.widget()
                if w:
                    w.deleteLater()
        except Exception:
            pass
        return
    comp = seleccion[0]
    try:
        if CURRENT_LANGUAGE == 'en':
            base = traducir('Inspector de propiedades')
            self.inspector_title.setText(f"{base} - {comp.nombre}")
        else:
            self.inspector_title.setText(f"Inspector de {comp.nombre}")
    except Exception:
        self.inspector_title.setText(traducir('Inspector de propiedades'))
    try:
        self.inspector_title.setVisible(True)
    except Exception:
        pass
    try:
        while self.inspector_layout.count():
            item = self.inspector_layout.takeAt(0)
            w = item.widget()
            if w:
                w.deleteLater()
    except Exception:
        pass
    try:
        self.inspector_widget.setVisible(True)
    except Exception:
        pass
    param_list = []
    existing_keys = set()
    for param in comp.parametros:
        if isinstance(param, str) and ':' in param:
            key, value = param.split(':', 1)
            key = key.strip()
            value = value.strip()
        else:
            key = str(param).strip()
            value = ""
        param_list.append((key, value))
        existing_keys.add(key.lower())
    if comp.tipo == "Resistencia":
        if 'tolerancia' not in existing_keys:
            param_list.append(("Tolerancia", ""))
        if 'potencia' not in existing_keys:
            param_list.append(("Potencia", ""))
    elif comp.tipo in ("Capacitor", "Bobina"):
        if 'tolerancia' not in existing_keys:
            param_list.append(("Tolerancia", ""))
    for key, value in param_list:
        label = QLabel(traducir(key))
        line_edit = QLineEdit(value)
        line_edit.setObjectName(f"inspector_{key}")

        def confirmar_cambio(ed=line_edit, k=key, componente=comp, old_value=value):
            new_val = ed.text().strip()
            try:
                if '-' in new_val:
                    QMessageBox.warning(self, "Valor inválido", "No se permiten valores negativos ni caracteres no válidos.")
                    ed.setText(old_value)
                    return
                allowed_letters = set('pnmuµμkKMGΩAUVHFWw%')
                for ch in new_val:
                    if ch.isalpha() and ch not in allowed_letters:
                        QMessageBox.warning(self, "Valor inválido", "No se permiten valores negativos ni caracteres no válidos.")
                        ed.setText(old_value)
                        return
            except Exception:
                try:
                    ed.setText(old_value)
                except Exception:
                    pass
                return
            nuevos_params = []
            found_key = False
            for p in componente.parametros:
                if isinstance(p, str) and ':' in p:
                    pk, _pv = p.split(':', 1)
                    pk = pk.strip()
                    if pk.lower() == k.lower():
                        nuevos_params.append(f"{pk}: {new_val}")
                        found_key = True
                        continue
                elif isinstance(p, str) and p.strip().lower() == k.lower():
                    nuevos_params.append(f"{p}: {new_val}")
                    found_key = True
                    continue
                nuevos_params.append(p)
            if not found_key:
                nuevos_params.append(f"{k}: {new_val}")
            componente.parametros = nuevos_params
            try:
                if len(componente.parametros) > 1:
                    tooltip = f"<b>{componente.nombre}:</b><br>" + "<br>".join(componente.parametros)
                else:
                    tooltip = f"<b>{componente.nombre}:</b> " + " ".join(componente.parametros)
                componente.svg_item.setToolTip(tooltip)
            except Exception:
                pass
            try:
                preferences.setdefault('default_values', {})[componente.tipo] = list(componente.parametros)
                guardar_preferencias()
            except Exception:
                pass
            try:
                self.registrar_historial()
            except Exception:
                pass
        line_edit.editingFinished.connect(confirmar_cambio)
        try:
            self.inspector_layout.addRow(label, line_edit)
        except Exception:
            pass

    try:
        nombre_label = QLabel(traducir('Nombre'))
        nombre_edit = QLineEdit(comp.nombre)

        def confirmar_nombre(ed=nombre_edit, componente=comp):
            nuevo = ed.text().strip()
            if not nuevo:
                return
            old = componente.nombre
            if nuevo == old:
                return
            componente.nombre = nuevo
            try:
                componente.nombre_label.setPlainText(nuevo)
            except Exception:
                pass
            try:
                if CURRENT_LANGUAGE == 'en':
                    base = traducir('Inspector de propiedades')
                    self.inspector_title.setText(f"{base} - {nuevo}")
                else:
                    self.inspector_title.setText(f"Inspector de {nuevo}")
            except Exception:
                pass
            try:
                if len(componente.parametros) > 1:
                    tooltip = f"<b>{componente.nombre}:</b><br>" + "<br>".join(componente.parametros)
                else:
                    tooltip = f"<b>{componente.nombre}:</b> " + " ".join(componente.parametros)
                componente.svg_item.setToolTip(tooltip)
            except Exception:
                pass
            try:
                self.registrar_historial()
            except Exception:
                pass
        nombre_edit.editingFinished.connect(confirmar_nombre)
        self.inspector_layout.addRow(nombre_label, nombre_edit)
    except Exception:
        pass

    try:
        color_label = QLabel(traducir('Color etiqueta'))
        color_combo = QComboBox()
        colores = {
            "Negro": Qt.black,
            "Rojo": Qt.red,
            "Azul": Qt.blue,
            "Verde": Qt.darkGreen,
            "Magenta": Qt.magenta,
            "Cyan": Qt.darkCyan,
            "Amarillo": Qt.darkYellow,
        }
        color_combo.addItems(list(colores.keys()))
        try:
            current_name = getattr(comp, 'custom_color_name', None)
            if current_name in colores:
                idx = list(colores.keys()).index(current_name)
                color_combo.setCurrentIndex(idx)
        except Exception:
            pass

        def confirmar_color(index=0, componente=comp, combo=color_combo):
            try:
                nombre_color = combo.currentText()
                qt_color = colores.get(nombre_color, Qt.black)
                componente.nombre_label.setDefaultTextColor(qt_color)
                componente.custom_color_name = nombre_color
            except Exception:
                pass
            try:
                self.registrar_historial()
            except Exception:
                pass
        color_combo.currentIndexChanged.connect(confirmar_color)
        self.inspector_layout.addRow(color_label, color_combo)
    except Exception:
        pass

def eliminar_seleccionado(self):
        from PyQt5.QtWidgets import QGraphicsTextItem

        self.registrar_historial()
        seleccionados = list(self.scene.selectedItems())
        for item in seleccionados:
            if hasattr(item, 'terminales'):
                for term in getattr(item, 'terminales', []):
                    for conn in list(self.conexiones):
                        if conn.t1 == term or conn.t2 == term:
                            try:
                                conn.eliminar()
                            except Exception:
                                pass
                            try:
                                self.conexiones.remove(conn)
                            except ValueError:
                                pass
        for item in seleccionados:
            try:
                self.scene.removeItem(item)
            except Exception:
                pass
            if hasattr(item, 'nombre') and item in self.componentes:
                try:
                    self.componentes.remove(item)
                except ValueError:
                    pass
        for it in list(self.scene.items()):
            if isinstance(it, QGraphicsTextItem) and getattr(it, "_sim_result", False):
                try:
                    self.scene.removeItem(it)
                except Exception:
                    pass
        if hasattr(self, 'medicion_labels'):
            for lbl in getattr(self, 'medicion_labels', []):
                try:
                    self.scene.removeItem(lbl)
                except Exception:
                    pass
            self.medicion_labels = []

        try:
            self.scene.update()
        except Exception:
            pass
        try:
            if hasattr(self, 'view'):
                self.view.viewport().update()
        except Exception:
            pass

        try:
            renumerar_componentes(self)
        except Exception:
            pass


def duplicar_seleccionado(self) -> None:
    from PyQt5.QtCore import QPointF
    seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    if not seleccion:
        return
    for item in seleccion:
        try:
            tipo = getattr(item, 'tipo', None)
            if not tipo:
                continue
            params = list(getattr(item, 'parametros', []))
            datos = self.imagenes.get(tipo)
            if not datos:
                continue
            archivo, terminales, vertical, invertir, *resto = datos
            escala = resto[0] if resto else 0.15
            self.contadores[tipo] += 1
            prefijos = {
                "Resistencia": "R", "Capacitor": "C", "Bobina": "L",
                "Transistor NPN": "Q", "Transistor PNP": "Q", "Mosfet": "Q",
                "Diodo Zener": "D", "LED": "D", "OpAmp": "U", "Motor": "M",
                "Transformador": "T", "Potenciómetro": "VR", "Fuente Voltaje": "V",
                "Fuente Corriente": "I", "Batería": "B", "Push Button": "SW", "Push Button Cerrado": "SW",
                "GND": "GND",
                "Amperímetro": "A",
                "Voltímetro": "VM",
                "Ohmímetro": "OM",
                "Generador de funciones": "GF"
            }
            prefijo = prefijos.get(tipo, "X")
            numero = self.contadores[tipo]
            nombre = f"{prefijo}{numero}" if tipo != "GND" else "GND"
            if tipo == "Transistor NPN":
                nombre += " (NPN)"
            elif tipo == "Transistor PNP":
                nombre += " (PNP)"
            elif tipo == "Mosfet":
                nombre += " (MOS)"
            comp = ComponenteInteractivo(archivo, params, nombre, tipo, terminales=terminales, vertical=vertical, invertir=invertir)
            comp.svg_item.setScale(escala)
            try:
                comp.setRotation(getattr(item, 'rotation')())
            except Exception:
                try:
                    comp.setRotation(item.rotation())
                except Exception:
                    pass
            comp.reposicionar_terminales()
            try:
                pos_orig = item.pos()
            except Exception:
                pos_orig = QPointF(0, 0)
            comp.setPos(pos_orig + QPointF(20, 20))
            self.scene.addItem(comp)
            self.componentes.append(comp)
            self.registrar_historial()
        except Exception:
            continue

    def copiar_componentes(self) -> None:
        try:
            seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
        except Exception:
            seleccion = []
        if not seleccion:
            self._clipboard_components = None
            return
        self._clipboard_components = list(seleccion)

    def pegar_componentes(self) -> None:
        if not self._clipboard_components:
            return
        try:
            from PyQt5.QtCore import QPointF
        except Exception:
            return
        mapping = {}
        offset = QPointF(20, 20)
        for item in self._clipboard_components:
            try:
                tipo = getattr(item, 'tipo', None)
            except Exception:
                continue
            if not tipo:
                continue
            try:
                params = list(getattr(item, 'parametros', []))
            except Exception:
                params = []
            datos = self.imagenes.get(tipo)
            if not datos:
                continue
            archivo, terminales, vertical, invertir, *resto = datos
            escala_base = resto[0] if resto else 0.15
            self.contadores[tipo] += 1
            prefijos = {
                "Resistencia": "R", "Capacitor": "C", "Bobina": "L",
                "Transistor NPN": "Q", "Transistor PNP": "Q", "Mosfet": "Q",
                "Diodo Zener": "D", "LED": "D", "OpAmp": "U", "Motor": "M",
                "Transformador": "T", "Potenciómetro": "VR", "Fuente Voltaje": "V",
                "Fuente Corriente": "I", "Batería": "B", "Push Button": "SW",
                "Push Button Cerrado": "SW", "GND": "GND", "Amperímetro": "A",
                "Voltímetro": "VM", "Ohmímetro": "OM", "Generador de funciones": "GF"
            }
            prefijo = prefijos.get(tipo, "X")
            numero = self.contadores[tipo]
            nombre = f"{prefijo}{numero}" if tipo != "GND" else "GND"
            if tipo == "Transistor NPN":
                nombre += " (NPN)"
            elif tipo == "Transistor PNP":
                nombre += " (PNP)"
            elif tipo == "Mosfet":
                nombre += " (MOS)"
            nuevo = ComponenteInteractivo(archivo, params, nombre, tipo,
                                          terminales=terminales, vertical=vertical,
                                          invertir=invertir)
            try:
                nuevo.setScale(item.scale())
            except Exception:
                pass
            try:
                nuevo.svg_item.setScale(item.svg_item.scale())
            except Exception:
                nuevo.svg_item.setScale(escala_base)
            try:
                nuevo.setRotation(item.rotation())
            except Exception:
                pass
            try:
                nuevo.reposicionar_terminales()
            except Exception:
                pass
            try:
                pos_orig = item.pos()
            except Exception:
                pos_orig = QPointF(0, 0)
            nuevo.setPos(pos_orig + offset)
            try:
                self.scene.addItem(nuevo)
                self.componentes.append(nuevo)
            except Exception:
                pass
            mapping[item] = nuevo
        try:
            for conn in list(self.conexiones):
                try:
                    c1 = conn.t1.parentItem()
                    c2 = conn.t2.parentItem()
                except Exception:
                    continue
                if c1 in mapping and c2 in mapping:
                    try:
                        idx1 = c1.terminales.index(conn.t1)
                        idx2 = c2.terminales.index(conn.t2)
                    except Exception:
                        continue
                    nuevo1 = mapping[c1]
                    nuevo2 = mapping[c2]
                    try:
                        t1_new = nuevo1.terminales[idx1]
                        t2_new = nuevo2.terminales[idx2]
                    except Exception:
                        continue
                    try:
                        nueva_conn = Conexion(self.scene, t1_new, t2_new)
                        self.conexiones.append(nueva_conn)
                    except Exception:
                        pass
        except Exception:
            pass
        try:
            self.registrar_historial()
        except Exception:
            pass

        try:
            self.scene.update()
        except Exception:
            pass
        try:
            if hasattr(self, 'view'):
                self.view.viewport().update()
        except Exception:
            pass

def agregar_componente(self, tipo_manual=None):
        tipo = tipo_manual or self.selector.currentText()
        entradas = self.pedir_valores_componente(tipo)

        archivo, terminales, vertical, invertir, *resto = self.imagenes[tipo]
        escala = resto[0] if resto else 0.15
        if entradas is None:
            return

        self.contadores[tipo] += 1
        prefijos = {
            "Resistencia": "R", "Capacitor": "C", "Bobina": "L",
            "Transistor NPN": "Q", "Transistor PNP": "Q", "Mosfet": "Q",
            "Diodo Zener": "D", "LED": "D", "OpAmp": "U", "Motor": "M",
            "Transformador": "T", "Potenciómetro": "VR", "Fuente Voltaje": "V",
            "Fuente Corriente": "I", "Batería": "B", "Push Button": "SW","Push Button Cerrado": "SW",
            "GND": "GND",
            
            "Amperímetro": "A",
            "Voltímetro": "VM",
            "Ohmímetro": "OM",
            "Generador de funciones": "GF"
        }
        prefijo = prefijos.get(tipo, "X")
        numero = self.contadores[tipo]
        nombre = f"{prefijo}{numero}" if tipo != "GND" else "GND"
        if tipo == "Transistor NPN":
            nombre += " (NPN)"
        elif tipo == "Transistor PNP":
            nombre += " (PNP)"
        elif tipo == "Mosfet":
            nombre += " (MOS)"

        datos = self.imagenes[tipo]
        if len(datos) == 5:
           archivo, self.tipo_terminales, vertical, invertir, escala = datos
        else:
           archivo, self.tipo_terminales, vertical, invertir = datos
           escala = 0.15  

        archivo_svg = archivo

        comp = ComponenteInteractivo(archivo_svg, entradas, nombre, tipo, terminales=terminales, vertical=vertical, invertir=invertir)
        comp.svg_item.setScale(escala)
        comp.reposicionar_terminales()
        comp.setPos(100 + len(self.componentes) * 120, 100)
        self.scene.addItem(comp)
        self.componentes.append(comp)
        self.registrar_historial()
            
def pedir_valor_numerico(self, titulo, etiqueta):
        while True:
            title_translated = traducir(titulo) if isinstance(titulo, str) else titulo
            label_translated = traducir(etiqueta) if isinstance(etiqueta, str) else etiqueta
            val, ok = QInputDialog.getText(self, title_translated, label_translated)
            if not ok:
                return None
            try:
                num = float(val.replace(",", "."))
            except ValueError:
                mostrar_advertencia(self, "Valor inválido", "Por favor ingresa un número válido (ej. 4.7, 100, 0.01).")
                continue
            if num < 0:
                mostrar_advertencia(self, "Valor inválido", "El valor no puede ser negativo.")
                continue
            return val
    
def pedir_valores_componente(self, tipo):
        if tipo == "GND":  
            return ["GND"]

        if tipo in ["Amperímetro", "Voltímetro", "Ohmímetro"]:
            return [tipo]

        campos = {
            "Resistencia": [("Resistencia", "Ω")],
            "Capacitor": [("Capacitancia", "F")],
            "Bobina": [("Inductancia", "H")],
            "Fuente Voltaje": [("Voltaje", "V")],
            "Fuente Corriente": [("Corriente", "A")],
            "Motor": [("Voltaje nominal", "V")],
            "Transformador": [
                ("Voltaje entrada", "V"), ("Voltaje salida", "V"),
                ("Corriente entrada", "A"), ("Corriente salida", "A")
            ]
        }

        if tipo in campos:
            try:
                defaults = preferences.get('default_values', {}).get(tipo)
            except Exception:
                defaults = None
            valores = []
            for idx, (campo, unidad_base) in enumerate(campos[tipo]):
                num_default = ""
                unidad_default = unidad_base
                if defaults and idx < len(defaults):
                    try:
                        _, dv = defaults[idx].split(":", 1)
                        dv = dv.strip()
                        if " " in dv:
                            num_default, unidad_default = dv.split(None, 1)
                        else:
                            num_default = dv
                    except Exception:
                        pass
                while True:
                    try:
                        title_param = traducir('Entrada de parámetro')
                        prompt_param = f"{traducir(campo)} ({traducir('solo números')}):"
                        val, ok1 = QInputDialog.getText(self, title_param, prompt_param, text=str(num_default))
                    except Exception:
                        title_param = traducir('Entrada de parámetro')
                        prompt_param = f"{traducir(campo)} ({traducir('solo números')}):"
                        val, ok1 = QInputDialog.getText(self, title_param, prompt_param)
                    if not ok1:
                        return None
                    try:
                        num = float(val.replace(",", "."))
                    except ValueError:
                        mostrar_advertencia(self, "Valor inválido", "Por favor ingresa un número válido (ej. 4.7, 100, 0.01).")
                        continue
                    if num < 0:
                        mostrar_advertencia(self, "Valor inválido", "El valor no puede ser negativo.")
                        continue
                    break
                if tipo == "Capacitor":
                   unidades = ["F", "pF", "nF", "μF", "mF", "MF"]
                elif tipo == "Resistencia":
                   unidades = ["Ω", "kΩ", "MΩ"]
                elif tipo == "Bobina":
                   unidades = ["μH", "mH", "H"]
                elif tipo == "Fuente Voltaje" or (tipo == "Transformador" and "Voltaje" in campo) or tipo == "Motor":
                   unidades = ["μV", "mV", "V", "kV"]
                elif tipo == "Fuente Corriente" or (tipo == "Transformador" and "Corriente" in campo):
                   unidades = ["μA", "mA", "A"]
                else:
                   unidades = [unidad_base]
                try:
                    default_index = unidades.index(unidad_default) if unidad_default in unidades else 0
                except Exception:
                    default_index = 0
                unidad_title = traducir('Unidad')
                unidad_prompt = f"{traducir('Selecciona unidad para')} {campo}:"
                unidad, ok2 = QInputDialog.getItem(self, unidad_title, unidad_prompt, unidades, default_index, False)
                if not ok2:
                    return None
                valores.append(f"{campo}: {val} {unidad}")

            if tipo == "Resistencia":
                tol_default = "5%"
                pot_default_idx = 1  
                try:
                    defaults_tipo = preferences.get('default_values', {}).get(tipo, [])
                    for p in defaults_tipo:
                        if ":" in p:
                            pk, pv = p.split(":", 1)
                            pk = pk.strip().lower()
                            pv = pv.strip()
                            if pk.startswith("tolerancia"):
                                tol_default = pv
                            elif pk.startswith("potencia"):
                                potencias_std = ["0.125 W", "0.25 W", "0.5 W", "1 W", "2 W"]
                                if pv in potencias_std:
                                    pot_default_idx = potencias_std.index(pv)
                except Exception:
                    pass
                tol_opciones = ["20%", "10%", "5%", "2%", "1%"]
                try:
                    def_tol_index = tol_opciones.index(tol_default) if tol_default in tol_opciones else 2
                except Exception:
                    def_tol_index = 2
                tol_title = traducir('Tolerancia')
                tol_prompt = traducir('Selecciona la tolerancia (%):')
                tol, ok_tol = QInputDialog.getItem(self, tol_title, tol_prompt, tol_opciones, def_tol_index, False)
                if not ok_tol:
                    return None
                valores.append(f"{traducir('Tolerancia')}: {tol}")
                potencias_std = ["0.125 W", "0.25 W", "0.5 W", "1 W", "2 W"]
                try:
                    default_pot_index = pot_default_idx if 0 <= pot_default_idx < len(potencias_std) else 1
                except Exception:
                    default_pot_index = 1
                pot_title = traducir('Potencia nominal')
                pot_prompt = traducir('Selecciona la potencia nominal:')
                potencia_sel, ok_pot = QInputDialog.getItem(self, pot_title, pot_prompt, potencias_std, default_pot_index, False)
                if not ok_pot:
                    return None
                valores.append(f"{traducir('Potencia')}: {potencia_sel}")
            try:
                preferences.setdefault('default_values', {})[tipo] = list(valores)
                guardar_preferencias()
            except Exception:
                pass
            return valores

        elif tipo == "Batería":
            opciones = ["5 V", "9 V", "12 V"]
            default_index = 1  
            try:
                defaults = preferences.get('default_values', {}).get('Batería')
                if defaults:
                    dv = defaults[0].split(":", 1)[1].strip()
                    if dv in opciones:
                        default_index = opciones.index(dv)
            except Exception:
                pass
            title_bat = traducir('Tensión de la batería')
            prompt_bat = traducir('Selecciona tensión:')
            val, ok = QInputDialog.getItem(self, title_bat, prompt_bat, opciones, default_index, False)
            if ok:
                resultado = [f"{traducir('Tensión')}: {val}"] if TRANSLATIONS.get('Tensión') else [f"Tensión: {val}"]
                try:
                    preferences.setdefault('default_values', {})['Batería'] = resultado
                    guardar_preferencias()
                except Exception:
                    pass
                return resultado
            else:
                return None

        elif tipo == "Potenciómetro":
            valores = ["1kΩ", "5kΩ", "10kΩ", "50kΩ", "100kΩ", "500kΩ", "1MΩ"]
            default_val = 2  
            default_cursor = "50"  
            try:
                defaults = preferences.get('default_values', {}).get('Potenciómetro')
                if defaults and len(defaults) >= 2:
                    try:
                        _, val_str = defaults[0].split(":", 1)
                        val_str = val_str.strip()
                        if val_str in valores:
                            default_val = valores.index(val_str)
                    except Exception:
                        pass
                    try:
                        _, cur_str = defaults[1].split(":", 1)
                        cur_str = cur_str.strip()
                        if cur_str.endswith("%"):
                            default_cursor = cur_str[:-1].strip()
                    except Exception:
                        pass
            except Exception:
                pass
            title_pot = traducir('Valor del Potenciómetro')
            prompt_pot = traducir('Selecciona valor comercial:')
            resistencia, ok1 = QInputDialog.getItem(self, title_pot, prompt_pot, valores, default_val, False)
            if not ok1:
                return None
            while True:
                cursor_title = traducir('Porcentaje del cursor')
                cursor_prompt = traducir('Ingresa el % del cursor (0-100):')
                try:
                    porcentaje, ok2 = QInputDialog.getText(self, cursor_title, cursor_prompt, text=str(default_cursor))
                except Exception:
                    porcentaje, ok2 = QInputDialog.getText(self, cursor_title, cursor_prompt)
                if not ok2:
                    return None
                try:
                    valor_porcentaje = float(porcentaje.replace(",", "."))
                    if 0 <= valor_porcentaje <= 100:
                        break
                    else:
                        mostrar_advertencia(self, traducir('Valor fuera de rango'), traducir('Por favor ingresa un valor entre 0 y 100.'))
                except ValueError:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido (ej. 4.7, 100, 0.01).'))
            resultado = [f"{traducir('Valor')}: {resistencia}", f"{traducir('Cursor')}: {valor_porcentaje}%"]
            try:
                preferences.setdefault('default_values', {})['Potenciómetro'] = list(resultado)
                guardar_preferencias()
            except Exception:
                pass
            return resultado

        elif tipo == "LED":
            colores = {
                "Rojo": ("1.8", "20"), "Verde": ("2.0", "20"),
                "Azul": ("3.0", "20"), "Amarillo": ("2.1", "20"),
                "Blanco": ("3.2", "20")
            }
            color_title = traducir('Color del LED')
            color_prompt = traducir('Selecciona color:')
            color, ok = QInputDialog.getItem(self, color_title, color_prompt, list(colores.keys()), 0, False)
            if not ok:
                return None
            voltaje, corriente = colores[color]
            return [f"{traducir('Color')}: {color}", f"Vf: {voltaje} V", f"If: {corriente} mA"]
        
        elif tipo == "OpAmp":
            modelos = {
                "LM741": ["Ganancia: 200k", "Slew Rate: 0.5 V/μs", "Offset: 6 mV"],
                "TL081": ["Ganancia: 200k", "Slew Rate: 13 V/μs", "Offset: 3 mV"],
                "TL082": ["Ganancia: 100k", "Slew Rate: 8 V/μs", "Offset: 3 mV"],
                "LM324": ["Ganancia: 100k", "Slew Rate: 0.5 V/μs", "Offset: 2 mV"],
                "OP07": ["Ganancia: 200k", "Slew Rate: 0.3 V/μs", "Offset: 75 μV"],
                "NE5532": ["Ganancia: 100k", "Slew Rate: 9 V/μs", "Offset: 0.5 mV"],
                "TL074": ["Ganancia: 100k", "Slew Rate: 13 V/μs", "Offset: 3 mV"],
                "LF356": ["Ganancia: 100k", "Slew Rate: 12 V/μs", "Offset: 2 mV"],
                "CA3140": ["Ganancia: 150k", "Slew Rate: 9 V/μs", "Offset: 4 mV"],
                "Personalizado": None
            }
            op_title = traducir('Modelo Op-Amp')
            op_prompt = traducir('Selecciona modelo:')
            modelo, ok = QInputDialog.getItem(self, op_title, op_prompt, list(modelos.keys()), 0, False)
            if not ok:
                return None
            if modelo == "Personalizado":
                g = self.pedir_valor_numerico("Ganancia", "Ganancia:")
                if g is None:
                    return None
                sr = self.pedir_valor_numerico("Slew Rate", "Slew Rate (V/μs):")
                if sr is None:
                    return None
                off = self.pedir_valor_numerico("Offset", "Offset (mV):")
                if off is None:
                    return None
                return [f"Modelo: Personalizado", f"Ganancia: {g}", f"Slew Rate: {sr} V/μs", f"Offset: {off} mV"]
            return [f"Modelo: {modelo}"] + modelos[modelo]
        elif tipo == "Diodo Zener":
            modelos = {
                "1N4728": ["Vz: 3.3 V", "Iz max: 76 mA"],
                "1N4729": ["Vz: 3.6 V", "Iz max: 69 mA"],
                "1N4730": ["Vz: 3.9 V", "Iz max: 64 mA"],
                "1N4731": ["Vz: 4.3 V", "Iz max: 58 mA"],
                "1N4732": ["Vz: 4.7 V", "Iz max: 53 mA"],
                "1N4733": ["Vz: 5.1 V", "Iz max: 49 mA"],
                "1N4734": ["Vz: 5.6 V", "Iz max: 45 mA"],
                "1N4735": ["Vz: 6.2 V", "Iz max: 41 mA"],
                "1N4736": ["Vz: 6.8 V", "Iz max: 37 mA"],
                "Personalizado": None
            }
            dz_title = traducir('Diodo Zener')
            dz_prompt = traducir('Selecciona modelo:')
            modelo, ok = QInputDialog.getItem(self, dz_title, dz_prompt, list(modelos.keys()), 0, False)
            if not ok: return None
            if modelo == "Personalizado":
                vz = self.pedir_valor_numerico("Vz", "Voltaje Zener (V):")
                if vz is None: return None
                iz = self.pedir_valor_numerico("Iz", "Corriente máxima (mA):")
                if iz is None: return None                
                return [f"Modelo: Personalizado", f"Vz: {vz} V", f"Iz max: {iz} mA"]
            else:
                return[f"Modelo: {modelo}"] + modelos[modelo]

        elif tipo in ["Transistor NPN", "Transistor PNP"]:
            modelos = {
                "2N2222": ["Beta: 100", "Vbe: 0.7 V", "hfe: 200"],
                "BC547": ["Beta: 110", "Vbe: 0.6 V", "hfe: 300"],
                "2N3904": ["Beta: 150", "Vbe: 0.7 V", "hfe: 200"],
                "TIP120": ["Beta: 1000", "Vbe: 2.0 V", "hfe: 1000"],
                "2N5401": ["Beta: 100", "Vbe: 0.7 V", "hfe: 150"],
                "BC557": ["Beta: 125", "Vbe: 0.6 V", "hfe: 200"],
                "2N2907": ["Beta: 100", "Vbe: 0.7 V", "hfe: 100"],
                "MPSA42": ["Beta: 50", "Vbe: 0.7 V", "hfe: 30"],
                "S8050": ["Beta: 120", "Vbe: 0.6 V", "hfe: 200"],
                "Personalizado": None
            }
            tr_title = traducir('Modelo Transistor')
            tr_prompt = traducir('Selecciona modelo:')
            modelo, ok = QInputDialog.getItem(self, tr_title, tr_prompt, list(modelos.keys()), 0, False)
            if not ok:
                return None
            if modelo == "Personalizado":
                beta = self.pedir_valor_numerico("Beta", "Ganancia de corriente (β):")
                if beta is None: return None
                vbe = self.pedir_valor_numerico("Vbe", "Voltaje base-emisor (V):")
                if vbe is None: return None
                hfe = self.pedir_valor_numerico("hfe", "hfe (factor de amplificación):")
                if hfe is None: return None               
                return [f"Modelo: Personalizado", f"Beta: {beta}", f"Vbe: {vbe} V", f"hfe: {hfe}"]
            else:
                return[f"Modelo: {modelo}"] + modelos[modelo]

        elif tipo == "Generador de funciones":
            while True:
                v_title = traducir('Voltaje del generador')
                v_prompt = traducir('Voltaje (solo números):')
                val_v, ok_v = QInputDialog.getText(self, v_title, v_prompt)
                if not ok_v:
                    return None
                try:
                    num_v = float(val_v.replace(",", "."))
                except ValueError:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido para el voltaje.'))
                    continue
                if num_v < 0:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('El voltaje no puede ser negativo.'))
                    continue
                break
            unidades_v = ["pV", "nV", "µV", "mV", "V", "kV", "MV", "GV"]
            unidad_v_title = traducir('Unidad de voltaje')
            unidad_v_prompt = traducir('Selecciona unidad para el voltaje:')
            unidad_v, ok_uv = QInputDialog.getItem(self, unidad_v_title, unidad_v_prompt, unidades_v, 4, False)
            if not ok_uv:
                return None
            while True:
                f_title = traducir('Frecuencia del generador')
                f_prompt = traducir('Frecuencia (solo números):')
                val_f, ok_f = QInputDialog.getText(self, f_title, f_prompt)
                if not ok_f:
                    return None
                try:
                    num_f = float(val_f.replace(",", "."))
                except ValueError:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido para la frecuencia.'))
                    continue
                if num_f < 0:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('La frecuencia no puede ser negativa.'))
                    continue
                break
            unidades_f = ["pHz", "nHz", "µHz", "mHz", "Hz", "kHz", "MHz", "GHz"]
            unidad_f_title = traducir('Unidad de frecuencia')
            unidad_f_prompt = traducir('Selecciona unidad para la frecuencia:')
            unidad_f, ok_uf = QInputDialog.getItem(self, unidad_f_title, unidad_f_prompt, unidades_f, 4, False)
            if not ok_uf:
                return None
            while True:
                off_title = traducir('Offset DC')
                off_prompt = traducir('Offset DC (solo números):')
                val_off, ok_off = QInputDialog.getText(self, off_title, off_prompt, text="0")
                if not ok_off:
                    return None
                try:
                    num_off = float(val_off.replace(",", "."))
                except ValueError:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido para el offset.'))
                    continue
                break
            unidad_off_title = traducir('Unidad de offset')
            unidad_off_prompt = traducir('Selecciona unidad para el offset:')
            default_off_index = unidades_v.index(unidad_v) if unidad_v in unidades_v else 4
            unidad_off, ok_uo = QInputDialog.getItem(self, unidad_off_title, unidad_off_prompt, unidades_v, default_off_index, False)
            if not ok_uo:
                return None
            while True:
                ph_title = traducir('Fase')
                ph_prompt = traducir('Fase (grados, solo números):')
                val_phase, ok_phase = QInputDialog.getText(self, ph_title, ph_prompt, text="0")
                if not ok_phase:
                    return None
                try:
                    num_phase = float(val_phase.replace(",", "."))
                except ValueError:
                    mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido para la fase.'))
                    continue
                break
            formas = [traducir('Senoidal'), traducir('Cuadrada'), traducir('Triangular'), traducir('Diente de sierra')]
            formas_map = dict(zip(formas, ["Senoidal", "Cuadrada", "Triangular", "Diente de sierra"]))
            forma_title = traducir('Forma de onda')
            forma_prompt = traducir('Selecciona la forma de onda:')
            forma_trans, ok_forma = QInputDialog.getItem(self, forma_title, forma_prompt, formas, 0, False)
            if not ok_forma:
                return None
            forma = formas_map.get(forma_trans, forma_trans)
            entradas = [
                f"{traducir('Voltaje')}: {val_v} {unidad_v}",
                f"{traducir('Frecuencia')}: {val_f} {unidad_f}",
                f"{traducir('Offset')}: {val_off} {unidad_off}",
                f"{traducir('Fase')}: {val_phase}°",
                f"{traducir('Forma')}: {traducir(forma)}"
            ]
            if "Cuadrada" in forma:
                while True:
                    duty_title = traducir('Ciclo de trabajo')
                    duty_prompt = traducir('Ciclo de trabajo (%) (0-100):')
                    val_duty, ok_duty = QInputDialog.getText(self, duty_title, duty_prompt, text="50")
                    if not ok_duty:
                        return None
                    try:
                        num_duty = float(val_duty.replace(",", "."))
                    except ValueError:
                        mostrar_advertencia(self, traducir('Valor inválido'), traducir('Por favor ingresa un número válido para el ciclo de trabajo.'))
                        continue
                    if num_duty < 0 or num_duty > 100:
                        mostrar_advertencia(self, traducir('Valor fuera de rango'), traducir('El ciclo de trabajo debe estar entre 0 y 100.'))
                        continue
                    break
                entradas.append(f"{traducir('Ciclo de trabajo')}: {val_duty} %")
            return entradas

        elif tipo == "Mosfet":
            modelos = {
                "IRF540N": ["Tipo: Canal N", "Vth: 2 V", "Rds(on): 0.077 Ω"],
                "IRF9540N": ["Tipo: Canal P", "Vth: -2 V", "Rds(on): 0.117 Ω"],
                "IRLZ44N": ["Tipo: Canal N", "Vth: 1 V", "Rds(on): 0.022 Ω"],
                "IRF3205": ["Tipo: Canal N", "Vth: 2 V", "Rds(on): 0.008 Ω"],
                "IRF630": ["Tipo: Canal N", "Vth: 2.0 V", "Rds(on): 0.4 Ω"],
                "IRF840": ["Tipo: Canal N", "Vth: 2.0 V", "Rds(on): 0.85 Ω"],
                "IRF3708": ["Tipo: Canal N", "Vth: 1.5 V", "Rds(on): 0.0028 Ω"],
                "IRF4905": ["Tipo: Canal P", "Vth: -2.0 V", "Rds(on): 0.02 Ω"],
                "IRF540": ["Tipo: Canal N", "Vth: 2.0 V", "Rds(on): 0.077 Ω"],
                "Personalizado": None
            }
            mosfet_title = traducir('Modelo MOSFET') if TRANSLATIONS.get('Modelo MOSFET') else 'Modelo MOSFET'
            mosfet_prompt = traducir('Selecciona modelo:')
            modelo, ok = QInputDialog.getItem(self, mosfet_title, mosfet_prompt, list(modelos.keys()), 0, False)
            if not ok:
                return None
            if modelo == "Personalizado":
                tipo_title = traducir('Tipo') if TRANSLATIONS.get('Tipo') else 'Tipo'
                canal_prompt = traducir('Canal:') if TRANSLATIONS.get('Canal:') else 'Canal:'
                tipo_m, ok1 = QInputDialog.getItem(self, tipo_title, canal_prompt, [traducir('Canal N') if TRANSLATIONS.get('Canal N') else 'Canal N', traducir('Canal P') if TRANSLATIONS.get('Canal P') else 'Canal P'], 0, False)
                vth = self.pedir_valor_numerico("Vth", traducir('Tensión umbral (V):') if TRANSLATIONS.get('Tensión umbral (V):') else 'Tensión umbral (V):')
                if vth is None:
                    return None
                rds = self.pedir_valor_numerico("Rds(on)", traducir('Resistencia de conducción (Ω):') if TRANSLATIONS.get('Resistencia de conducción (Ω):') else 'Resistencia de conducción (Ω):')
                if rds is None:
                    return None
                return [f"Modelo: Personalizado", f"Tipo: {tipo_m}", f"Vth: {vth} V", f"Rds(on): {rds} Ω"]
            else:
                return [f"Modelo: {modelo}"] + modelos[modelo]

        return ["Sin parámetros"]

def guardar_circuito(self):
        import json
        from PyQt5.QtWidgets import QFileDialog
        datos = {
            "componentes": [],
            "conexiones": []
        }

        for comp in self.componentes:
            datos["componentes"].append({
                "tipo": comp.tipo,
                "nombre": comp.nombre,
                "parametros": comp.parametros,
                "pos": [comp.pos().x(), comp.pos().y()],
                "rotacion": comp.rotation(),
                "escala": comp.svg_item.scale()
            })

        for conexion in self.conexiones:
            t1_index = self.componentes.index(conexion.t1.parentItem())
            t2_index = self.componentes.index(conexion.t2.parentItem())
            i1 = conexion.t1.parentItem().terminales.index(conexion.t1)
            i2 = conexion.t2.parentItem().terminales.index(conexion.t2)
            datos["conexiones"].append({
                "c1": [t1_index, i1],
                "c2": [t2_index, i2]
            })

        try:
            inicio = getattr(self, 'last_folder', preferences.get('last_folder', ''))
        except Exception:
            inicio = preferences.get('last_folder', '')

        filtros = "Circuito JSON (*.json);;Circuito YAML (*.yaml);;Circuito TOML (*.toml)"
        archivo, filtro = QFileDialog.getSaveFileName(self, "Guardar circuito", inicio, filtros)
        if archivo:
            path = Path(archivo)
            ext = path.suffix.lower()
            if ext not in {'.json', '.yaml', '.toml'}:
                if filtro and 'yaml' in filtro.lower():
                    ext = '.yaml'
                elif filtro and 'toml' in filtro.lower():
                    ext = '.toml'
                else:
                    ext = '.json'
                archivo = f"{archivo}{ext}"
                path = Path(archivo)
            try:
                with open(path, "w", encoding='utf-8') as f:
                    if ext == '.yaml':
                        if yaml is None:
                            json.dump(datos, f, indent=4, ensure_ascii=False)
                        else:
                            yaml.safe_dump(datos, f, allow_unicode=True)
                    elif ext == '.toml':
                        if toml is None:
                            json.dump(datos, f, indent=4, ensure_ascii=False)
                        else:
                            toml.dump(datos, f)
                    else:
                        json.dump(datos, f, indent=4, ensure_ascii=False)
                try:
                    self.last_folder = os.path.dirname(archivo)
                    preferences['last_folder'] = self.last_folder
                    guardar_preferencias()
                except Exception:
                    pass
                try:
                    self.registrar_reciente(str(path))
                except Exception:
                    pass
                try:
                    self.statusBar().showMessage(traducir("Circuito guardado"), 3000)
                except Exception:
                    pass
            except Exception:
                try:
                    logger.exception("Error al guardar el circuito en %s", archivo)
                except Exception:
                    pass

def cargar_circuito(self):
        import json
        from PyQt5.QtWidgets import QFileDialog
        try:
            inicio = getattr(self, 'last_folder', preferences.get('last_folder', ''))
        except Exception:
            inicio = preferences.get('last_folder', '')
        filtros = "Circuito (*.json *.yaml *.toml);;Circuito JSON (*.json);;Circuito YAML (*.yaml);;Circuito TOML (*.toml)"
        archivo, _ = QFileDialog.getOpenFileName(self, "Cargar circuito", inicio, filtros)
        if not archivo:
            return

        try:
            self.last_folder = os.path.dirname(archivo)
            preferences['last_folder'] = self.last_folder
            guardar_preferencias()
        except Exception:
            pass

        path = Path(archivo)
        ext = path.suffix.lower()
        try:
            with open(path, "r", encoding='utf-8') as f:
                if ext in {'.yaml', '.yml'}:
                    if yaml is None:
                        datos = json.load(f)
                    else:
                        datos = yaml.safe_load(f)
                elif ext == '.toml':
                    if toml is None:
                        datos = json.load(f)
                    else:
                        datos = toml.load(f)
                else:
                    datos = json.load(f)
        except Exception as exc:
            try:
                mostrar_error(self, "cargar circuito", exc)
            except Exception:
                pass
            QMessageBox.critical(self, "Error", f"No se pudo cargar el archivo seleccionado: {exc}")
            return

        for c in self.conexiones[:]:
            c.eliminar()
            self.conexiones.remove(c)
        for comp in self.componentes:
            self.scene.removeItem(comp)
        self.componentes.clear()

        for comp_data in datos["componentes"]:
            tipo = comp_data["tipo"]
            nombre = comp_data["nombre"]
            parametros = comp_data["parametros"]
            x, y = comp_data["pos"]
            rotacion = comp_data.get("rotacion", 0)
            escala = comp_data.get("escala", 0.15)

            archivo, terminales, vertical, invertir, *_ = self.imagenes[tipo]

            comp = ComponenteInteractivo(archivo, parametros, nombre, tipo, terminales, vertical, invertir)
            comp.svg_item.setScale(escala)
            comp.reposicionar_terminales()
            comp.setPos(x, y)
            comp.setRotation(rotacion)

            self.scene.addItem(comp)
            self.componentes.append(comp)

        for conn in datos["conexiones"]:
            c1_idx, t1_idx = conn["c1"]
            c2_idx, t2_idx = conn["c2"]
            t1 = self.componentes[c1_idx].terminales[t1_idx]
            t2 = self.componentes[c2_idx].terminales[t2_idx]
            conexion = Conexion(self.scene, t1, t2)
            self.conexiones.append(conexion)
        self.statusBar().showMessage(traducir('Acción restablecida (Ctrl+Z)'), 2000)
        try:
            self.registrar_reciente(str(path))
        except Exception:
            pass
        try:
            self.statusBar().showMessage(traducir("Circuito cargado"), 3000)
        except Exception:
            pass

def cargar_circuito_archivo(self, archivo_json):
        import json
        from pathlib import Path

        path = Path(archivo_json)
        if not path.exists():
            QMessageBox.critical(self, "Error", f"El archivo {path} no existe.")
            return
        ext = path.suffix.lower()
        try:
            with open(path, "r", encoding="utf-8") as f:
                if ext in {'.yaml', '.yml'}:
                    if yaml is None:
                        datos = json.load(f)
                    else:
                        datos = yaml.safe_load(f)
                elif ext == '.toml':
                    if toml is None:
                        datos = json.load(f)
                    else:
                        datos = toml.load(f)
                else:
                    datos = json.load(f)
        except Exception as exc:
            QMessageBox.critical(self, "Error", f"No se pudo cargar el archivo seleccionado: {exc}")
            return

        for c in self.conexiones[:]:
            try:
                c.eliminar()
            except Exception:
                pass
            try:
                self.conexiones.remove(c)
            except ValueError:
                pass
        for comp in list(self.componentes):
            try:
                self.scene.removeItem(comp)
            except Exception:
                pass
        self.componentes.clear()

        try:
            for comp_data in datos.get("componentes", []):
                tipo = comp_data.get("tipo")
                nombre = comp_data.get("nombre")
                parametros = comp_data.get("parametros", [])
                x, y = comp_data.get("pos", (0, 0))
                rotacion = comp_data.get("rotacion", 0)
                escala = comp_data.get("escala", 0.15)

                archivo, terminales, vertical, invertir, *_ = self.imagenes[tipo]
                comp = ComponenteInteractivo(archivo, parametros, nombre, tipo, terminales, vertical, invertir)
                comp.svg_item.setScale(escala)
                comp.reposicionar_terminales()
                comp.setPos(x, y)
                comp.setRotation(rotacion)
                self.scene.addItem(comp)
                self.componentes.append(comp)
            for conn in datos.get("conexiones", []):
                c1_idx, t1_idx = conn.get("c1", (None, None))
                c2_idx, t2_idx = conn.get("c2", (None, None))
                if c1_idx is None or t1_idx is None or c2_idx is None or t2_idx is None:
                    continue
                try:
                    t1 = self.componentes[c1_idx].terminales[t1_idx]
                    t2 = self.componentes[c2_idx].terminales[t2_idx]
                except Exception:
                    continue
                conexion = Conexion(self.scene, t1, t2)
                self.conexiones.append(conexion)
            try:
                self.statusBar().showMessage(f"Circuito '{path.name}' cargado", 3000)
            except Exception:
                pass
            try:
                self.last_folder = str(path.parent)
                preferences['last_folder'] = self.last_folder
                self.registrar_reciente(str(path))
                guardar_preferencias()
            except Exception:
                pass
        except Exception as exc:
            QMessageBox.critical(self, "Error", f"Error al reconstruir el circuito: {exc}")

def _reconstruir_menu_plantillas(self):
        from pathlib import Path
        self.menu_plantillas.clear()
        try:
            accion_gestor = QAction("Gestor de plantillas", self)
            accion_gestor.triggered.connect(self.gestor_plantillas)
            self.menu_plantillas.addAction(accion_gestor)
            self.menu_plantillas.addSeparator()
        except Exception:
            pass
        try:
            base_dir = Path(__file__).resolve().parent
        except Exception:
            base_dir = Path('.')
        archivos = []
        try:
            archivos.extend(base_dir.glob('*.json'))
            archivos.extend(base_dir.glob('*.yaml'))
            archivos.extend(base_dir.glob('*.toml'))
        except Exception:
            pass
        if not archivos:
            action = QAction("No hay plantillas disponibles", self)
            action.setEnabled(False)
            self.menu_plantillas.addAction(action)
            return

        for fpath in archivos:
            nombre = fpath.stem
            action = QAction(nombre, self)
            action.triggered.connect(lambda _, p=fpath: self.cargar_circuito_archivo(p))
            self.menu_plantillas.addAction(action)

CircuitSimulator.cargar_circuito_archivo = cargar_circuito_archivo
CircuitSimulator._reconstruir_menu_plantillas = _reconstruir_menu_plantillas

def exportar_imagen(self):
        from PyQt5.QtWidgets import QFileDialog, QInputDialog, QMessageBox
        from PyQt5.QtGui import QImage
        opciones = ["Con cuadriculas", "Sin cuadriculas"]
        opcion, ok = QInputDialog.getItem(
            self,
            "Exportar imagen",
            "Seleccione cómo exportar la imagen:",
            opciones,
            0,
            False
        )
        if not ok:
            return
        incluir_grid = opcion == "Con cuadriculas"
        rect = self.scene.itemsBoundingRect()
        imagen = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        imagen.fill(Qt.white)
        original_pen = None
        if not incluir_grid and hasattr(self.scene, 'grid_pen'):
            try:
                original_pen = QPen(self.scene.grid_pen)  
                transparent_pen = QPen(Qt.transparent)
                self.scene.grid_pen = transparent_pen
            except Exception:
                original_pen = None
        painter = QPainter(imagen)
        self.scene.render(painter, target=QRectF(imagen.rect()), source=rect)
        painter.end()
        if not incluir_grid and original_pen is not None:
            try:
                self.scene.grid_pen = original_pen
            except Exception:
                pass
        archivo, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar como imagen",
            "",
            "Imagen PNG (*.png);;Imagen JPEG (*.jpg)"
        )
        if archivo:
            imagen.save(archivo)
            try:
                QMessageBox.information(self, "Guardar imagen", f"Imagen guardada en:\n{archivo}")
            except Exception:
                try:
                    self.statusBar().showMessage(f"Imagen guardada en: {archivo}", 5000)
                except Exception:
                    pass

def exportar_pdf(self):
        from PyQt5.QtWidgets import QFileDialog, QInputDialog
        from PyQt5.QtGui import QImage, QPen
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import Table, TableStyle
        from reportlab.lib import colors
        import tempfile

        opciones = ["Con cuadriculas", "Sin cuadriculas"]
        opcion, ok = QInputDialog.getItem(
            self,
            "Exportar imagen",
            "Seleccione cómo exportar la imagen:",
            opciones,
            0,
            False
        )
        if not ok:
            return
        incluir_grid = opcion == "Con cuadriculas"
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        original_pen_pdf = None
        if not incluir_grid and hasattr(self.scene, 'grid_pen'):
            try:
                original_pen_pdf = QPen(self.scene.grid_pen)
                self.scene.grid_pen = QPen(Qt.transparent)
            except Exception:
                original_pen_pdf = None
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        if not incluir_grid and original_pen_pdf is not None:
            try:
                self.scene.grid_pen = original_pen_pdf
            except Exception:
                pass

        archivo, _ = QFileDialog.getSaveFileName(self, "Guardar como PDF", "", "PDF (*.pdf)")
        if not archivo:
            return

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img_path = tmp.name
            img.save(img_path)

        c = canvas.Canvas(archivo, pagesize=letter)
        c.setFont("Helvetica", 12)
        c.drawString(50, 750, "ESIMulador - Exportación PDF")
        c.drawImage(ImageReader(img_path), 50, 450, width=500, preserveAspectRatio=True)

        data = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            data.append([
                comp.nombre,
                comp.tipo,
                ", ".join(comp.parametros)
            ])

        table = Table(data, colWidths=[100, 120, 280])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
            ('FONTSIZE', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))

        table.wrapOn(c, 480, 300)
        table.drawOn(c, 50, 100)

        c.save()
        try:
            QMessageBox.information(self, "Guardar PDF", f"PDF guardado en:\n{archivo}")
        except Exception:
            try:
                self.statusBar().showMessage(f"PDF guardado en: {archivo}", 5000)
            except Exception:
                pass

def _calcular_bom(self):
    logic_types = {
        'AND', 'OR', 'NAND', 'NOR', 'XOR', 'NOT'
    }
    tipos_comp = [(getattr(comp, 'tipo', '') or '').upper() for comp in getattr(self, 'componentes', [])]
    has_logic = any(t in logic_types for t in tipos_comp)
    analog_types = {
        'RESISTENCIA', 'CAPACITOR', 'BOBINA', 'FUENTE CORRIENTE',
        'FUENTE VOLTAJE', 'BATERÍA', 'TRANSFORMADOR', 'POTENCIÓMETRO',
        'LED', 'DIODO ZENER', 'MOSFET', 'OPAMP', 'TRANSISTOR NPN',
        'TRANSISTOR PNP', 'GENERADOR DE FUNCIONES'
    }
    has_analog = any(t in analog_types for t in tipos_comp)
    if not has_logic or has_analog:
        issues = self.ejecutar_comprobaciones_previas()
        if any(i.level == "error" for i in issues):
            return None, "El circuito no pasa el chequeo eléctrico. Añade GND y corrige problemas antes de generar la BOM."
    try:
        elementos = self.extraer_topologia()
    except Exception as e:
        return None, f"Error al extraer la topología: {e}"
    if has_logic and not has_analog:
        bom_rows = []
        for comp in getattr(self, 'componentes', []):
            nombre = getattr(comp, "nombre", "")
            tipo = getattr(comp, "tipo", "")
            try:
                params_str = ", ".join(comp.parametros)
            except Exception:
                params_str = ""
            elem = next((e for e in elementos if e.get("nombre") == nombre), None)
            nodos_str = ""
            try:
                if elem is not None:
                    nodos_str = ",".join(str(n) for n in elem.get("nodos", []))
            except Exception:
                nodos_str = ""
            valor_nominal = ""
            try:
                if comp.parametros:
                    primera = comp.parametros[0]
                    if ":" in primera:
                        _, val = primera.split(":", 1)
                        valor_nominal = val.strip()
                    else:
                        valor_nominal = primera.strip()
            except Exception:
                valor_nominal = ""
            bom_rows.append([
                nombre,
                tipo,
                params_str,
                valor_nominal,
                nodos_str,
                "",  
                ""   
            ])
        bom_rows.append(["TOTAL", "", "", "", "", "", ""])
        header = ["Nombre", "Tipo", "Parámetros", "Valor", "Nodos", "Potencia", "Alerta"]
        return [header] + bom_rows, None
    try:
        nodo_gnd = self._nodo_gnd_actual()
    except Exception:
        nodo_gnd = None
    if nodo_gnd is None:
        return None, "Debes colocar un nodo de tierra (GND) para poder calcular la BOM."
    try:
        voltajes, corrientes_fuente, Req = self.analizar_circuito(elementos, nodo_gnd)
    except Exception as exc:
        return None, f"No se pudo resolver el circuito: {exc}"
    nodos_set = set()
    for e in elementos:
        for n in e["nodos"]:
            nodos_set.add(n)
    nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
    nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
    volt_dict = {nodo_gnd: 0.0}
    for nodo, idx in nodo_idx.items():
        volt_dict[nodo] = voltajes[idx]
    lista_fuentes = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
    bom_rows = []
    total_p_dis = 0.0
    for comp in self.componentes:
        nombre = getattr(comp, "nombre", "")
        tipo = getattr(comp, "tipo", "")
        try:
            params_str = ", ".join(comp.parametros)
        except Exception:
            params_str = ""
        valor_nominal = ""
        try:
            if comp.parametros:
                primera = comp.parametros[0]
                if ":" in primera:
                    _, val = primera.split(":", 1)
                    val = val.strip()
                    try:
                        num_val = interpretar_valor_prefijo(val)
                        valor_nominal = f"{num_val}"
                    except Exception:
                        valor_nominal = val
        except Exception:
            valor_nominal = ""
        tolerancia = ""
        try:
            for p in comp.parametros:
                key = p.split(":", 1)[0].strip().lower() if ":" in p else ""
                val = p.split(":", 1)[1].strip() if ":" in p else p
                if "tol" in key or "tolerancia" in key or ("%" in val or "±" in val):
                    tolerancia = val
                    break
        except Exception:
            tolerancia = ""
        p_nom_val = None
        p_nom_str = ""
        try:
            for p in comp.parametros:
                key = p.split(":", 1)[0].strip().lower() if ":" in p else ""
                val = p.split(":", 1)[1].strip() if ":" in p else p
                if "w" in val.lower() or "potencia" in key:
                    try:
                        p_nom_val = interpretar_valor_prefijo(val)
                    except Exception:
                        try:
                            p_nom_val = float(val.replace(",", "."))
                        except Exception:
                            p_nom_val = None
                    break
        except Exception:
            p_nom_val = None
        if p_nom_val is None and tipo == "Resistencia":
            p_nom_val = 0.25  
        p_dis_val = None
        elem = next((e for e in elementos if e.get("nombre") == nombre), None)
        if elem is not None:
            tipo_e = elem.get("tipo")
            nodos = elem.get("nodos", [])
            if len(nodos) >= 2:
                n1, n2 = nodos[0], nodos[1]
                v1 = volt_dict.get(n1, 0.0)
                v2 = volt_dict.get(n2, 0.0)
                v_comp = v1 - v2
            else:
                v_comp = 0.0
            if tipo_e == "Resistencia":
                try:
                    r_val = float(elem.get("valor"))
                    i_val = (v_comp / r_val) if r_val != 0 else None
                    p_dis_val = (v_comp ** 2) / r_val if r_val != 0 else None
                except Exception:
                    i_val = None
                    p_dis_val = None
            elif tipo_e in ["Fuente Voltaje", "Batería"]:
                try:
                    idx_fuente = lista_fuentes.index(elem)
                    i_src = corrientes_fuente[idx_fuente]
                    i_val = i_src
                    p_dis_val = v_comp * i_src
                except Exception:
                    i_val = None
                    p_dis_val = None
            elif tipo_e == "Fuente Corriente":
                try:
                    i_val = -elem.get("valor", 0.0)
                    p_dis_val = v_comp * i_val
                except Exception:
                    i_val = None
                    p_dis_val = None
            else:
                i_val = None
                p_dis_val = None
        if p_dis_val is not None:
            total_p_dis += abs(p_dis_val)
        nodos_str = ""
        if elem is not None:
            try:
                nodos_str = ",".join(str(n) for n in elem.get("nodos", []))
            except Exception:
                nodos_str = ""
        p_nom_str = f"{p_nom_val:.6g}" if p_nom_val is not None else ""
        p_dis_str = f"{p_dis_val:.6g}" if p_dis_val is not None else ""
        alerta = ""
        if p_nom_val is not None and p_dis_val is not None:
            try:
                if abs(p_dis_val) > p_nom_val:
                    alerta = "⚠"
            except Exception:
                alerta = ""
        bom_rows.append([
            nombre,
            tipo,
            params_str,
            valor_nominal,
            nodos_str,
            p_dis_str,
            alerta
        ])
    bom_rows.append(["TOTAL", "", "", "", "", f"{total_p_dis:.6g}", ""])
    header = ["Nombre", "Tipo", "Parámetros", "Valor", "Nodos", "Potencia", "Alerta"]
    return [header] + bom_rows, None

def mostrar_bom(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    try:
        data, err = self._calcular_bom()
    except Exception as exc:
        data, err = None, str(exc)
    if err:
        QMessageBox.warning(self, "Lista de materiales", err)
        return
    if not data:
        QMessageBox.information(self, "Lista de materiales", "No se encontraron componentes para listar.")
        return
    try:
        from PyQt5.QtWidgets import QInputDialog
        header = data[0]
        rows = data[1:]
        idx_tipo = header.index("Tipo") if "Tipo" in header else 1
        tipos = sorted(set(r[idx_tipo] for r in rows if r[0] != "TOTAL" and r[idx_tipo]))
        opciones = [_("Todos")] + tipos
        tipo_sel, ok = QInputDialog.getItem(self, _("Filtrar por tipo"), _("Elige tipo de componente:"), opciones, 0, False)
        if not ok:
            return
        if tipo_sel != _("Todos"):
            idx_pot = header.index("Potencia") if "Potencia" in header else 5
            filtered = [r for r in rows if (r[0] != "TOTAL" and r[idx_tipo] == tipo_sel)]
            total = 0.0
            for r in filtered:
                try:
                    if r[idx_pot]:
                        total += float(r[idx_pot])
                except Exception:
                    continue
            filtered.append(["TOTAL", "", "", "", "", f"{total:.6g}", ""])
            data_to_use = [header] + filtered
        else:
            data_to_use = data
    except Exception:
        data_to_use = data
    criticos = []
    try:
        for r in data_to_use[1:-1]:
            try:
                if len(r) > 6 and r[-1] == "⚠":
                    criticos.append(str(r[0]))
            except Exception:
                continue
    except Exception:
        pass
    preview_lines = []
    for row in data_to_use:
        try:
            preview_lines.append(" | ".join(str(x) for x in row))
        except Exception:
            continue
    preview_text = "\n".join(preview_lines[:15])
    extra_info = ""
    if criticos:
        extra_info = "\n\n" + _("Componentes críticos") + ": " + ", ".join(criticos)
    QMessageBox.information(self, _("Lista de materiales (BOM) - Vista previa"),
                            _("Se generó la lista de materiales.") + "\n\n" + preview_text + extra_info + "\n\n" + _("Se exportará a PDF."))
    from pathlib import Path
    archivo, selected_filter = QFileDialog.getSaveFileName(self, _("Guardar lista de materiales"), "", "PDF (*.pdf);;CSV (*.csv)")
    if not archivo:
        return
    try:
        ext = Path(archivo).suffix.lower()
        export_data = data_to_use
        if ext == ".csv" or (selected_filter and "csv" in selected_filter.lower()):
            if not archivo.lower().endswith(".csv"):
                archivo = f"{archivo}.csv"
            self._exportar_bom_csv(export_data, archivo)
        else:
            if not archivo.lower().endswith(".pdf"):
                archivo = f"{archivo}.pdf"
            self._exportar_bom_pdf(export_data, archivo)
    except Exception as exc:
        QMessageBox.critical(self, _("Error al exportar BOM"),
                             _("No se pudo exportar la lista de materiales:") + f" {exc}")

def _exportar_bom_pdf(self, data, filename):
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    c = canvas.Canvas(filename, pagesize=letter)
    c.setFont("Helvetica", 12)
    c.drawString(50, 750, "ESIMulador - Lista de materiales")
    if data:
        cols = len(data[0])
        if cols == 6:
            col_widths = [90, 80, 200, 60, 60, 80]
        elif cols == 7:
            col_widths = [80, 70, 180, 50, 60, 60, 40]
        elif cols == 9:
            col_widths = [80, 70, 180, 50, 60, 50, 60, 60, 40]
        else:
            col_widths = [500.0 / max(cols, 1)] * cols
    else:
        col_widths = []
    table = Table(data, colWidths=col_widths)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
    ]))
    w, h = table.wrapOn(c, 480, 600)
    table.drawOn(c, 50, 720 - h)
    try:
        criticos = []
        for row in data[1:]:
            try:
                if row and isinstance(row, (list, tuple)) and len(row) >= 7:
                    if str(row[-1]).strip() == "⚠":
                        criticos.append(str(row[0]))
            except Exception:
                continue
        if criticos:
            try:
                summary_label = _("Componentes críticos")
            except Exception:
                summary_label = "Componentes críticos"
            summary_text = summary_label + ": " + ", ".join(criticos)
            y_pos = 720 - h - 15
            c.setFont("Helvetica", 9)
            c.drawString(50, y_pos, summary_text)
    except Exception:
        pass
    c.save()

def _exportar_bom_csv(self, data, filename):
    import csv
    try:
        with open(filename, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            for row in data:
                writer.writerow([str(x) if x is not None else "" for x in row])
    except Exception as exc:
        raise exc

def iniciar_monte_carlo(self):
    from PyQt5.QtWidgets import QInputDialog, QMessageBox, QDialog, QVBoxLayout, QLabel
    import random
    try:
        import numpy as np
        import matplotlib.pyplot as plt
    except Exception:
        QMessageBox.warning(self, "Dependencia faltante",
                            "Necesitas 'numpy' y 'matplotlib' para ejecutar el análisis de Monte Carlo.")
        return
    n_iter, ok = QInputDialog.getInt(self, "Monte Carlo", "Número de iteraciones:", 50, 1, 1000, 1)
    if not ok:
        return
    tol_perc, ok = QInputDialog.getDouble(self, "Monte Carlo", "Tolerancia (%):", 5.0, 0.0, 100.0, 1)
    if not ok:
        return
    tolerance = tol_perc / 100.0
    issues = self.ejecutar_comprobaciones_previas()
    if any(i.level == "error" for i in issues):
        msg = problemas_a_html([i for i in issues if i.level == "error"])
        QMessageBox.critical(self, "Monte Carlo", f"No se puede simular:\n{msg}")
        return
    try:
        elementos_base = self.extraer_topologia()
    except Exception as exc:
        QMessageBox.critical(self, "Monte Carlo", f"Error al preparar la topología: {exc}")
        return
    nodo_gnd = self._nodo_gnd_actual()
    if nodo_gnd is None:
        QMessageBox.critical(self, "Monte Carlo", "Debes colocar un nodo GND para poder simular.")
        return
    res_indices = []
    for idx, e in enumerate(elementos_base):
        if e.get("tipo") == "Resistencia":
            nombre = e.get("nombre", "")
            if "(eq." not in nombre:
                res_indices.append(idx)
    if not res_indices:
        QMessageBox.information(self, "Monte Carlo", "No se encontraron resistencias para variar.")
        return
    results = []
    for _ in range(n_iter):
        elems = []
        for e in elementos_base:
            elems.append(dict(e))
        for idx in res_indices:
            orig = elementos_base[idx]
            val = orig.get("valor", 0.0)
            try:
                val_f = float(val)
            except Exception:
                val_f = 0.0
            if val_f != 0.0:
                delta = random.uniform(-tolerance, tolerance)
                new_val = val_f * (1.0 + delta)
                elems[idx]["valor"] = new_val
        try:
            _, _, Req = self.analizar_circuito(elems, nodo_gnd)
            if Req is not None:
                results.append(float(Req))
        except Exception:
            continue
    if not results:
        QMessageBox.information(self, "Monte Carlo", "No se pudieron obtener resultados de las iteraciones.")
        return
    try:
        fig, ax = plt.subplots()
        ax.hist(results, bins=20, color='#4472C4', edgecolor='black')
        ax.set_title("Histograma de Req (Monte Carlo)")
        ax.set_xlabel("Resistencia equivalente (Ohmios)")
        ax.set_ylabel("Frecuencia")
    except Exception as exc:
        QMessageBox.critical(self, "Monte Carlo", f"No se pudo generar el histograma: {exc}")
        return
    dlg = QDialog(self)
    dlg.setWindowTitle("Resultado Monte Carlo")
    layout = QVBoxLayout(dlg)
    from io import BytesIO
    buf = BytesIO()
    fig.savefig(buf, format='png')
    plt.close(fig)
    buf.seek(0)
    from PyQt5.QtGui import QPixmap
    pix = QPixmap()
    pix.loadFromData(buf.getvalue(), "PNG")
    lbl_img = QLabel()
    lbl_img.setPixmap(pix)
    layout.addWidget(lbl_img)
    dlg.resize(pix.width() + 50, pix.height() + 50)
    dlg.exec_()

def keyPressEvent(self, event):
        scale_factor = 1.1
        if event.key() == Qt.Key_Plus or event.key() == Qt.Key_Equal:
            for item in self.scene.selectedItems():
                self.registrar_historial()
                item.setScale(item.scale() * scale_factor)
            return
        elif event.key() == Qt.Key_Minus:
            for item in self.scene.selectedItems():
                self.registrar_historial()
                item.setScale(item.scale() / scale_factor)
            return
        
        elif event.key() == Qt.Key_R and not (event.modifiers() & Qt.ControlModifier):
            for item in self.scene.selectedItems():
                if hasattr(item, 'setRotation'):
                    self.registrar_historial()
                    try:
                        item.setRotation((item.rotation() + 90) % 360)
                    except Exception:
                        pass
            return
        
        elif event.key() == Qt.Key_D and (event.modifiers() & Qt.ControlModifier):
            try:
                duplicar_seleccionado(self)
            except Exception:
                pass
            return
        elif event.key() == Qt.Key_C and (event.modifiers() & Qt.ControlModifier):
            try:
                self.copiar_componentes()
            except Exception:
                pass
            return
        elif event.key() == Qt.Key_V and (event.modifiers() & Qt.ControlModifier):
            try:
                self.pegar_componentes()
            except Exception:
                pass
            return
        elif event.key() == Qt.Key_Delete:
            removed_any = False
            from PyQt5.QtWidgets import QGraphicsLineItem
            for item in list(self.scene.selectedItems()):
                if isinstance(item, QGraphicsLineItem):
                    for conn in list(getattr(self, 'conexiones', [])):
                        if item in getattr(conn, 'lines', []):
                            try:
                                conn.eliminar()
                            except Exception:
                                pass
                            try:
                                self.conexiones.remove(conn)
                            except ValueError:
                                pass
                            removed_any = True
                            break
            if removed_any:
                try:
                    if hasattr(self, 'extraer_topologia'):
                        self.extraer_topologia()
                except Exception:
                    pass
                try:
                    for lbl in getattr(self, '_node_labels', []):
                        self.scene.removeItem(lbl)
                except Exception:
                    pass
                try:
                    self._node_labels = []
                    self._node_labels_active = False
                except Exception:
                    pass
                try:
                    self.scene.update()
                except Exception:
                    pass
                try:
                    if hasattr(self, 'view'):
                        self.view.viewport().update()
                except Exception:
                    pass
                try:
                    self.dragging_connection = False
                    self.drag_motion_detected = False
                    self.pending_terminal = None
                    self.drag_start_terminal = None
                    if hasattr(self, 'highlighted_terminal') and self.highlighted_terminal:
                        try:
                            self.highlighted_terminal.setBrush(Qt.black)
                        except Exception:
                            pass
                        self.highlighted_terminal = None
                    if hasattr(self, 'terminal_seleccion'):
                        self.terminal_seleccion.clear()
                except Exception:
                    pass
                return
            try:
                self.eliminar_seleccionado()
            except Exception:
                pass
            return
        elif event.key() == Qt.Key_N:
            try:
                self.nodes_visible = not getattr(self, 'nodes_visible', True)
                for comp in getattr(self, 'componentes', []):
                    for term in getattr(comp, 'terminales', []):
                        try:
                            term.setVisible(self.nodes_visible)
                        except Exception:
                            pass
            except Exception:
                pass
            return
        QMainWindow.keyPressEvent(self, event)

def agregar_conexion(self, comp1, comp2):
        if not comp1.terminales or not comp2.terminales:
            return
        conexion = Conexion(self.scene, comp1.terminales[0], comp2.terminales[0])
        self.conexiones.append(conexion)
        self.statusBar().showMessage(traducir('Acción restablecida (Ctrl+Z)'), 2000)

def registrar_historial(self):
        import copy
        estado = {
            "componentes": [],
            "conexiones": []
        }
        for comp in self.componentes:
            estado["componentes"].append({
                "tipo": comp.tipo,
                "nombre": comp.nombre,
                "parametros": copy.deepcopy(comp.parametros),
                "pos": [comp.pos().x(), comp.pos().y()],
                "rotacion": comp.rotation(),
                "escala": comp.svg_item.scale()
            })
        for conexion in self.conexiones:
            t1_idx = self.componentes.index(conexion.t1.parentItem())
            t2_idx = self.componentes.index(conexion.t2.parentItem())
            i1 = conexion.t1.parentItem().terminales.index(conexion.t1)
            i2 = conexion.t2.parentItem().terminales.index(conexion.t2)
            estado["conexiones"].append({
                "c1": [t1_idx, i1],
                "c2": [t2_idx, i2]
            })
        self.historial.append(estado)
        self.rehacer_stack.clear()

def restablecer_accion(self):
        if not self.historial:
            return
        estado = self.historial.pop()
        self.rehacer_stack.append(estado)
        for c in self.conexiones[:]:
            c.eliminar()
            self.conexiones.remove(c)
        for comp in self.componentes:
            self.scene.removeItem(comp)
        self.componentes.clear()
        for comp_data in estado["componentes"]:
            tipo = comp_data["tipo"]
            nombre = comp_data["nombre"]
            parametros = comp_data["parametros"]
            x, y = comp_data["pos"]
            rotacion = comp_data.get("rotacion", 0)
            escala = comp_data.get("escala", 0.15)
            archivo, terminales, vertical, invertir, *_ = self.imagenes[tipo]
            comp = ComponenteInteractivo(archivo, parametros, nombre, tipo, terminales, vertical, invertir)
            comp.svg_item.setScale(escala)
            comp.reposicionar_terminales()
            comp.setPos(x, y)
            comp.setRotation(rotacion)
            self.scene.addItem(comp)
            self.componentes.append(comp)
        for conn in estado["conexiones"]:
            c1_idx, t1_idx = conn["c1"]
            c2_idx, t2_idx = conn["c2"]
            t1 = self.componentes[c1_idx].terminales[t1_idx]
            t2 = self.componentes[c2_idx].terminales[t2_idx]
            conexion = Conexion(self.scene, t1, t2)
            self.conexiones.append(conexion)
        self.statusBar().showMessage(traducir('Acción restablecida (Ctrl+Z)'), 2000)

def rehacer_accion(self):
        if not self.rehacer_stack:
            return
        estado = self.rehacer_stack.pop()
        self.historial.append(estado)  

        for c in self.conexiones[:]:
            c.eliminar()
            self.conexiones.remove(c)
        for comp in self.componentes:
            self.scene.removeItem(comp)
        self.componentes.clear()

        for comp_data in estado["componentes"]:
            tipo = comp_data["tipo"]
            nombre = comp_data["nombre"]
            parametros = comp_data["parametros"]
            x, y = comp_data["pos"]
            rotacion = comp_data.get("rotacion", 0)
            escala = comp_data.get("escala", 0.15)
            archivo, terminales, vertical, invertir, *_ = self.imagenes[tipo]
            comp = ComponenteInteractivo(archivo, parametros, nombre, tipo, terminales, vertical, invertir)
            comp.svg_item.setScale(escala)
            comp.reposicionar_terminales()
            comp.setPos(x, y)
            comp.setRotation(rotacion)
            self.scene.addItem(comp)
            self.componentes.append(comp)

        for conn in estado["conexiones"]:
            c1_idx, t1_idx = conn["c1"]
            c2_idx, t2_idx = conn["c2"]
            t1 = self.componentes[c1_idx].terminales[t1_idx]
            t2 = self.componentes[c2_idx].terminales[t2_idx]
            conexion = Conexion(self.scene, t1, t2)
            self.conexiones.append(conexion)

def limpiar_todo(self):
        for conexion in self.conexiones[:]:
            conexion.eliminar()
            self.conexiones.remove(conexion)

        for comp in self.componentes:
            self.scene.removeItem(comp)
        self.componentes.clear()
        self.historial.clear()
        self.rehacer_stack.clear()

        for item in self.scene.items():
            if isinstance(item, QGraphicsTextItem):
                self.scene.removeItem(item)

        self.statusBar().showMessage(traducir("Circuito limpiado completamente."), 3000)

def simular_corrientes(self):
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        try:
            self.statusBar().showMessage(traducir("Simulando corrientes..."), 0)
        except Exception:
            pass
        issues = self.ejecutar_comprobaciones_previas()
        has_errors = any(i.level == "error" for i in issues)
        has_warnings = any(i.level == "warning" for i in issues)
    
        if has_errors:
            from PyQt5.QtCore import Qt
            html = problemas_a_html(issues)
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Critical)
            msg_pre.setText(html)
            msg_pre.exec_()
            return
        
        if has_warnings:
            from PyQt5.QtCore import Qt
            html = problemas_a_html([i for i in issues if i.level == "warning"])
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Warning)
            msg_pre.setText(html)
            msg_pre.exec_()
        else:
            try:
                logger.info("Chequeo eléctrico: sin problemas.")
            except Exception:
                pass
        
        elementos = self.extraer_topologia()
        nodo_gnd = None
        for e in elementos:
            if e["tipo"] == "GND":
                nodo_gnd = e["nodos"][0]

        if nodo_gnd is None:
            QMessageBox.critical(self, traducir("Error de simulación"), traducir("No se encontró GND.\n\nAgrega un nodo de tierra (GND) para poder simular el circuito."))
            return
        try:
            voltajes, corrientes, Req = self.analizar_circuito(elementos, nodo_gnd)
        except Exception as e:
            mensaje = str(e)
            if isinstance(mensaje, str) and ("singular matrix" in mensaje.lower()):
                mensaje = "No se puede resolver el circuito: el sistema es singular. Verifica las conexiones."
            QMessageBox.critical(self, traducir("Error de simulación"), traducir(mensaje))
            return
        mensaje = "<b>" + traducir("Corrientes por componente:") + "</b><br>"
        for k, i in corrientes.items():
            valor = formatear_con_prefijo(i, unidad="A")
            mensaje += f"{k}: {valor} <br>"

        tiene_fuente_corriente = any(e["tipo"] == "Fuente Corriente" for e in elementos)
        if not tiene_fuente_corriente and Req is not None:
            mensaje += f"<br><b>Req:</b> {formatear_con_prefijo(Req, 'Ω')}"
        QMessageBox.information(self, traducir("Resultados de Corrientes"), mensaje)
        try:
            self.agregar_resultado("corrientes", mensaje)
        except Exception:
            pass
        try:
            self.statusBar().showMessage(traducir("Simulación de corrientes completada"), 5000)
        except Exception:
            pass

def simular_todo(self):
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        try:
            self.statusBar().showMessage(traducir("Simulando circuito..."), 0)
        except Exception:
            pass
        issues = self.ejecutar_comprobaciones_previas()
        has_errors = any(i.level == "error" for i in issues)
        has_warnings = any(i.level == "warning" for i in issues)
        
        if has_errors:
            from PyQt5.QtCore import Qt
            html = problemas_a_html(issues)
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Critical)
            msg_pre.setText(html)
            msg_pre.exec_()
            return
        
        if has_warnings:
            from PyQt5.QtCore import Qt
            html = problemas_a_html([i for i in issues if i.level == "warning"])
            msg_pre = QMessageBox(self)
            msg_pre.setWindowTitle(traducir("Chequeo eléctrico"))
            msg_pre.setTextFormat(Qt.RichText)
            msg_pre.setIcon(QMessageBox.Warning)
            msg_pre.setText(html)
            msg_pre.exec_()
        else:
            try:
                logger.info("Chequeo eléctrico: sin problemas.")
            except Exception:
                pass
    
        elementos = self.extraer_topologia()
        nodo_gnd = None
        for e in elementos:
            if e.get("tipo") == "GND":
                nodo_gnd = e["nodos"][0]
        if nodo_gnd is None:
            QMessageBox.critical(self, traducir("Error de simulación"), traducir("No se encontró GND.\n\nAgrega un nodo de tierra (GND) para poder simular el circuito."))
            return
        
        try:
            voltajes, corrientes, Req = self.analizar_circuito(elementos, nodo_gnd)
        except Exception as e:
            mensaje_err = str(e)
            if isinstance(mensaje_err, str) and ("singular matrix" in mensaje_err.lower()):
                mensaje_err = "No se puede resolver el circuito: el sistema es singular. Verifica las conexiones."
            QMessageBox.critical(self, traducir("Error de simulación"), traducir(mensaje_err))
            return
        mensaje = "<b>" + traducir("Voltajes por componente:") + "</b><br>"
        nodos = set()
        for e in elementos:
            for n in e["nodos"]:
                nodos.add(n)
        nodos = sorted(n for n in nodos if n != nodo_gnd)
        nodo_idx = {n: i for i, n in enumerate(nodos)}
        def volt_nodo(grupo):
            if grupo is None or grupo == nodo_gnd:
                return 0.0
            idx = nodo_idx.get(grupo, None)
            return voltajes[idx] if idx is not None else 0.0
        for e in elementos:
            tipo_e = e.get("tipo")
            if tipo_e in ["Resistencia", "Fuente Voltaje", "Batería", "Motor", "Capacitor", "Condensador", "Inductor", "Bobina"]:
                nodos_e = e.get("nodos", [])
                if len(nodos_e) >= 2:
                    n1, n2 = nodos_e[0], nodos_e[1]
                    v1 = volt_nodo(n1)
                    v2 = volt_nodo(n2)
                    v_comp = v1 - v2
                    mensaje += f"<b>{e['nombre']}</b>: {formatear_con_prefijo(abs(v_comp), 'V')}<br>"
                continue
            elif tipo_e == "OpAmp":
                nodos_comp = e.get("nodos", [])
                if len(nodos_comp) == 3:
                    n_plus, n_minus, n_out = nodos_comp
                    v_plus = volt_nodo(n_plus)
                    v_minus = volt_nodo(n_minus)
                    v_out = volt_nodo(n_out)
                    mensaje += (
                        f"<b>{e['nombre']}</b>: V+: {formatear_con_prefijo(abs(v_plus), 'V')} "
                        f"V-: {formatear_con_prefijo(abs(v_minus), 'V')} "
                        f"Vout: {formatear_con_prefijo(abs(v_out), 'V')}<br>"
                    )
            elif tipo_e == "Mosfet":
                nodos_mos = e.get("nodos", [])
                if len(nodos_mos) == 3:
                    n_gate, n_drain, n_source = nodos_mos
                    v_gate = volt_nodo(n_gate)
                    v_drain = volt_nodo(n_drain)
                    v_source = volt_nodo(n_source)
                    vgs = v_gate - v_source
                    vds = v_drain - v_source
                    mensaje += (
                        f"<b>{e['nombre']}</b>: Vgs: {formatear_con_prefijo(abs(vgs), 'V')} "
                        f"Vds: {formatear_con_prefijo(abs(vds), 'V')}<br>"
                    )
        
        for comp in self.componentes:
            if comp.tipo not in ["Transistor NPN", "Transistor PNP"]:
                continue
            term = comp.terminales
            if len(term) != 3:
                continue
            t_base, t_col, t_emit = term[0], term[1], term[2]
            n_base = getattr(t_base, 'node_group', None)
            n_col = getattr(t_col, 'node_group', None)
            n_emit = getattr(t_emit, 'node_group', None)
            if n_base is None or n_col is None or n_emit is None:
                continue
            Vb = volt_nodo(n_base)
            Vc = volt_nodo(n_col)
            Ve = volt_nodo(n_emit)
            if comp.tipo == "Transistor NPN":
                Vbe = Vb - Ve
                Vce = Vc - Ve
            else:
                Vbe = Ve - Vb
                Vce = Ve - Vc
            mensaje += f"<b>{comp.nombre}</b>: Vbe: {formatear_con_prefijo(abs(Vbe), 'V')}    Vce: {formatear_con_prefijo(abs(Vce), 'V')}<br>"
        tiene_fuente_corriente = any(e.get("tipo") == "Fuente Corriente" for e in elementos)
        if not tiene_fuente_corriente and Req is not None:
            mensaje += f"<br><b>Req:</b> {formatear_con_prefijo(abs(Req), 'Ω')}"
        mensaje += "<br><br><b>" + traducir("Corrientes por componente:") + "</b><br>"
        for nombre, corriente in corrientes.items():
            valor_c = formatear_con_prefijo(abs(corriente), unidad="A")
            mensaje += f"{nombre}: {valor_c}<br>"
    
        msgbox = QMessageBox(self)
        from PyQt5.QtCore import Qt
        msgbox.setWindowTitle(traducir("Resultados de Simulación"))
        msgbox.setTextFormat(Qt.RichText)
        msgbox.setText(mensaje)
        msgbox.setIcon(QMessageBox.Information)
        msgbox.exec_()
        try:
            self.agregar_resultado("todo", mensaje)
        except Exception:
            pass
        try:
            llamar_hooks('simulation_completed', voltajes=voltajes, corrientes=corrientes, Req=Req, elementos=elementos)
        except Exception:
            pass
        try:
            self.statusBar().showMessage(traducir("Simulación completa finalizada"), 5000)
        except Exception:
            pass

def analizar_circuito(self, elementos, nodo_gnd):
        import numpy as np

        nodos = set()
        for e in elementos:
            for n in e["nodos"]:
                nodos.add(n)
        nodos = sorted(n for n in nodos if n != nodo_gnd)
        nodo_idx = {n: i for i, n in enumerate(nodos)}

        N = len(nodos)
        fuentes_v = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
        op_amps = [e for e in elementos if e.get("tipo") == "OpAmp"]
        M = len(fuentes_v)
        O = len(op_amps)

        B_total = np.zeros((N, M + O))
        E_total = np.zeros(M + O)
        idx_fuente = 0
        idx_opamp = 0
        for e in elementos:
            tipo = e.get("tipo")
            if tipo in ["Fuente Voltaje", "Batería"]:
                n1, n2 = e["nodos"]
                v_val = e["valor"]
                if n1 != nodo_gnd:
                    B_total[nodo_idx[n1], idx_fuente] = 1.0
                if n2 != nodo_gnd:
                    B_total[nodo_idx[n2], idx_fuente] = -1.0
                E_total[idx_fuente] = v_val
                idx_fuente += 1
        for op in op_amps:
            nodes = op.get("nodos", [])
            if len(nodes) >= 3:
                n_plus, n_minus, n_out = nodes[0], nodes[1], nodes[2]
            else:
                continue
            if n_out != nodo_gnd:
                B_total[nodo_idx[n_out], M + idx_opamp] = 1.0
            idx_opamp += 1

        def resolver_mna(G_mat, I_vec):
            dim = N + M + O
            A_mat = np.zeros((dim, dim))
            A_mat[:N, :N] = G_mat
        
            if M + O > 0:
                A_mat[:N, N:] = B_total
                A_mat[N:, :N] = B_total.T
            z_vec = np.zeros(dim)
            z_vec[:N] = I_vec
            if M > 0:
                z_vec[N:N + M] = E_total[:M]
            mu = 1.0e5
            for j, op in enumerate(op_amps):
                row = N + M + j
                A_mat[row, :] = 0.0
                nodes = op.get("nodos", [])
                if len(nodes) >= 3:
                    n_plus, n_minus, n_out = nodes[0], nodes[1], nodes[2]
                else:
                    continue
                if n_out != nodo_gnd and n_out in nodo_idx:
                    A_mat[row, nodo_idx[n_out]] = 1.0
                if n_plus != nodo_gnd and n_plus in nodo_idx:
                    A_mat[row, nodo_idx[n_plus]] -= mu
                if n_minus != nodo_gnd and n_minus in nodo_idx:
                    A_mat[row, nodo_idx[n_minus]] += mu
            try:
                x_vec = np.linalg.solve(A_mat, z_vec)
            except np.linalg.LinAlgError as e:
                raise e
            return x_vec[:N], x_vec[N:]

        G_base = np.zeros((N, N))
        I_base = np.zeros(N)
        for e in elementos:
            tipo = e.get("tipo")
            if tipo == "Resistencia":
                n1, n2 = e["nodos"]
                g_val = 0.0
                try:
                    g_val = 1.0 / float(e["valor"])
                except Exception:
                    g_val = 0.0
                if n1 != nodo_gnd:
                    G_base[nodo_idx[n1], nodo_idx[n1]] += g_val
                if n2 != nodo_gnd:
                    G_base[nodo_idx[n2], nodo_idx[n2]] += g_val
                if n1 != nodo_gnd and n2 != nodo_gnd:
                    G_base[nodo_idx[n1], nodo_idx[n2]] -= g_val
                    G_base[nodo_idx[n2], nodo_idx[n1]] -= g_val
            elif tipo == "Fuente Corriente":
                n1, n2 = e["nodos"]
                i_val = e.get("valor", 0.0)
                try:
                    i_val = float(i_val)
                except Exception:
                    i_val = 0.0
                if n1 != nodo_gnd:
                    I_base[nodo_idx[n1]] -= i_val
                if n2 != nodo_gnd:
                    I_base[nodo_idx[n2]] += i_val
            elif tipo == "Capacitor":
                continue
            elif tipo in ("Inductor", "Bobina"):
                n1, n2 = e["nodos"]
                try:
                    Lval = float(e.get("valor", 0.0))
                except Exception:
                    Lval = 0.0
                if Lval and abs(Lval) > 0.0:
                    R_equiv = 1e-3 * float(Lval)  
                else:
                    R_equiv = 1e-6
                try:
                    g_val = 1.0 / R_equiv
                except Exception:
                    g_val = 0.0
                if n1 != nodo_gnd:
                    G_base[nodo_idx[n1], nodo_idx[n1]] += g_val
                if n2 != nodo_gnd:
                    G_base[nodo_idx[n2], nodo_idx[n2]] += g_val
                if n1 != nodo_gnd and n2 != nodo_gnd:
                    G_base[nodo_idx[n1], nodo_idx[n2]] -= g_val
                    G_base[nodo_idx[n2], nodo_idx[n1]] -= g_val

        voltajes0, corrientes_fuente0 = resolver_mna(G_base, I_base)

        conduction_list = []

        def _voltaje_nodo(node):
            try:
                from PyQt5.QtWidgets import QGraphicsItem
                if isinstance(node, QGraphicsItem) or hasattr(node, 'node_group'):
                    node = getattr(node, 'node_group', None)
            except Exception:
                pass
            if node is None or node == nodo_gnd:
                return 0.0
            idx = nodo_idx.get(node, None)
            if idx is None:
                return 0.0
            return voltajes0[idx]

        for comp in self.componentes:
            
            ctype = getattr(comp, 'tipo', None)
            if ctype not in ["Transistor NPN", "Transistor PNP"]:
                continue
            term = getattr(comp, 'terminales', [])
            if len(term) != 3:
                continue
            t_base, t_col, t_emit = term[0], term[1], term[2]
            n_base = getattr(t_base, 'node_group', None)
            n_col = getattr(t_col, 'node_group', None)
            n_emit = getattr(t_emit, 'node_group', None)
            
            if n_base is None or n_col is None or n_emit is None:
                continue
            beta = 100.0
            vbe_th = 0.7
            for p in getattr(comp, 'parametros', []):
                if not isinstance(p, str):
                    continue
                partes = p.split(":", 1)
                if len(partes) != 2:
                    continue
                clave = partes[0].strip().lower()
                valor = partes[1].strip()

                def _extraer_numero_con_prefijo(text):
                    cleaned = text.replace("V", "").replace("v", "").replace("A", "").replace("a", "").strip()
                    try:
                        return interpretar_valor_prefijo(cleaned)
                    except Exception:
                        try:
                            return float(cleaned.replace(",", "."))
                        except Exception:
                            return None
                if clave.startswith("beta") or clave.startswith("hfe"):
                    val = _extraer_numero_con_prefijo(valor)
                    if val is not None:
                        beta = val
                elif clave.startswith("vbe"):
                    val = _extraer_numero_con_prefijo(valor)
                    if val is not None:
                        vbe_th = val
            Vb = _voltaje_nodo(n_base)
            Vc = _voltaje_nodo(n_col)
            Ve = _voltaje_nodo(n_emit)
            if ctype == "Transistor NPN":
                if Vb - Ve >= vbe_th:
                    R_ce = 1.0
                    conduction_list.append((n_col, n_emit, R_ce))
                    R_be = R_ce * beta if beta > 0 else R_ce * 100.0
                    conduction_list.append((n_base, n_emit, R_be))
            elif ctype == "Transistor PNP":
                if Ve - Vb >= vbe_th:
                    R_ce = 1.0
                    conduction_list.append((n_emit, n_col, R_ce))
                    R_be = R_ce * beta if beta > 0 else R_ce * 100.0
                    conduction_list.append((n_emit, n_base, R_be))

        for comp in self.componentes:
            ctype = getattr(comp, 'tipo', None)
            if ctype != "Mosfet":
                continue
            term = getattr(comp, 'terminales', [])
            if len(term) != 3:
                continue
            
            t_gate, t_drain, t_source = term[0], term[1], term[2]
            n_gate = getattr(t_gate, 'node_group', None)
            n_drain = getattr(t_drain, 'node_group', None)
            n_source = getattr(t_source, 'node_group', None)
            if n_gate is None or n_drain is None or n_source is None:
                continue
            mos_params = None
            try:
                for e in elementos:
                    if e.get("tipo") == "Mosfet" and e.get("nombre") == comp.nombre:
                        mos_params = e.get("valor", None)
                        break
            except Exception:
                mos_params = None
            mos_tipo = 'N'
            Vth_val = 2.0
            Rds_val = 1.0
            if mos_params:
                mos_tipo = mos_params.get('tipo', 'N')
                Vth_val = mos_params.get('Vth', 2.0)
                Rds_val = mos_params.get('Rds', 1.0)
            Vg = _voltaje_nodo(n_gate)
            Vd = _voltaje_nodo(n_drain)
            Vs = _voltaje_nodo(n_source)
            if mos_tipo.upper().startswith('N'):
                Vgs = Vg - Vs
                if Vgs >= Vth_val:
                    R_on = Rds_val if abs(Rds_val) > 1e-12 else 1e-6
                    conduction_list.append((n_drain, n_source, R_on))
            else:
                Vsg = Vs - Vg
                if Vsg >= abs(Vth_val):
                    R_on = Rds_val if abs(Rds_val) > 1e-12 else 1e-6
                    conduction_list.append((n_source, n_drain, R_on))

        if conduction_list:
            G_second = G_base.copy()
            for (n1, n2, R_on) in conduction_list:
                if n1 == nodo_gnd and n2 == nodo_gnd:
                    continue
                try:
                    g_val = 1.0 / R_on
                except Exception:
                    g_val = 0.0
                if n1 != nodo_gnd:
                    idx1 = nodo_idx[n1]
                    G_second[idx1, idx1] += g_val
                if n2 != nodo_gnd:
                    idx2 = nodo_idx[n2]
                    G_second[idx2, idx2] += g_val
                if n1 != nodo_gnd and n2 != nodo_gnd:
                    idx1 = nodo_idx[n1]
                    idx2 = nodo_idx[n2]
                    G_second[idx1, idx2] -= g_val
                    G_second[idx2, idx1] -= g_val
            voltajes, corrientes_fuente = resolver_mna(G_second, I_base)
        else:
            voltajes, corrientes_fuente = voltajes0, corrientes_fuente0

        corrientes = {}
        for e in elementos:
            tipo_e = e.get("tipo")
            if tipo_e == "Resistencia":
                n1, n2 = e["nodos"]
                v1 = voltajes[nodo_idx[n1]] if n1 != nodo_gnd else 0.0
                v2 = voltajes[nodo_idx[n2]] if n2 != nodo_gnd else 0.0
                v = v1 - v2
                try:
                    corrientes[e["nombre"]] = v / float(e["valor"])
                except Exception:
                    corrientes[e["nombre"]] = 0.0
            elif tipo_e in ("Capacitor", "Condensador"):
                corrientes[e["nombre"]] = 0.0
            elif tipo_e in ("Inductor", "Bobina"):
                n1, n2 = e["nodos"]
                v1 = voltajes[nodo_idx[n1]] if n1 != nodo_gnd else 0.0
                v2 = voltajes[nodo_idx[n2]] if n2 != nodo_gnd else 0.0
                v = v1 - v2
                try:
                    Lval = float(e.get("valor", 0.0))
                except Exception:
                    Lval = 0.0
                if Lval and abs(Lval) > 0.0:
                    R_equiv = 1e-3 * Lval
                else:
                    R_equiv = 1e-6
                try:
                    corrientes[e["nombre"]] = v / R_equiv
                except Exception:
                    corrientes[e["nombre"]] = 0.0
        for idx, fv in enumerate(fuentes_v):
            corrientes[fv["nombre"]] = corrientes_fuente[idx]

        V_total = sum(abs(e["valor"]) for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"])
        I_total = sum(abs(e["valor"]) for e in elementos if e.get("tipo") == "Fuente Corriente")
        if any(e["tipo"] == "Fuente Corriente" for e in elementos):
            Req = None
        elif I_total > 0:
            Req = V_total / I_total
        elif len(corrientes_fuente) > 0:
            Req = V_total / abs(corrientes_fuente[0])
        else:
            Req = None

        def obtener_beta_vbe(parametros):
            beta = None
            hfe = None
            vbe = 0.7
            for p in parametros:
                if not isinstance(p, str):
                    continue
                partes = p.split(":", 1)
                if len(partes) != 2:
                    continue
                clave = partes[0].strip().lower()
                valor = partes[1].strip()
                
                def extraer_numero(cadena):
                    c = cadena.replace("V", "").replace("v", "").replace("A", "").replace("a", "").strip()
                    try:
                        return interpretar_valor_prefijo(c)
                    except Exception:
                        try:
                            return float(c)
                        except Exception:
                            return None
                if clave.startswith("beta"):
                    val = extraer_numero(valor)
                    if val is not None:
                        beta = val
                elif clave.startswith("hfe"):
                    val = extraer_numero(valor)
                    if val is not None:
                        hfe = val
                elif clave.startswith("vbe"):
                    val = extraer_numero(valor)
                    if val is not None:
                        vbe = val
            if beta is None:
                beta = hfe if hfe is not None else 100.0
            return beta, vbe

        def voltaje_nodo(nodo):
            try:
                from PyQt5.QtWidgets import QGraphicsItem
                if isinstance(nodo, QGraphicsItem) or hasattr(nodo, 'node_group'):
                    nodo = getattr(nodo, 'node_group', None)
            except Exception:
                pass
            if nodo is None or nodo == nodo_gnd:
                return 0.0
            idx = nodo_idx.get(nodo, None)
            if idx is None:
                return 0.0
            return voltajes[idx]

        fuente_corr_dict = {}
        for idx, fv in enumerate(fuentes_v):
            n1, n2 = fv["nodos"]
            curr = corrientes_fuente[idx]
            fuente_corr_dict.setdefault(n1, []).append((n1, n2, curr))
            fuente_corr_dict.setdefault(n2, []).append((n1, n2, curr))

        for comp in self.componentes:
            if comp.tipo not in ["Transistor NPN", "Transistor PNP"]:
                continue
           
            term = comp.terminales
            if len(term) != 3:
                continue
            t_base = term[0]
            t_col = term[1]
            t_emit = term[2]
            try:
                n_base = getattr(t_base, 'node_group', None)
                n_col = getattr(t_col, 'node_group', None)
                n_emit = getattr(t_emit, 'node_group', None)
            except Exception:
                continue
            if n_base is None or n_col is None or n_emit is None:
                continue
            beta, vbe_th = obtener_beta_vbe(comp.parametros)
            Vb = voltaje_nodo(n_base)
            Vc = voltaje_nodo(n_col)
            Ve = voltaje_nodo(n_emit)
            I_b_in = 0.0  
            I_b_out = 0.0  
            for e in elementos:
                if e["tipo"] != "Resistencia":
                    continue
                n1, n2 = e["nodos"]
                Rv = e["valor"]
                if n_base in (n1, n2):
                    other = n2 if n_base == n1 else n1
                    V_other = voltaje_nodo(other)
                    Ibase_to_other = (Vb - V_other) / Rv
                    Iother_to_base = -Ibase_to_other
                    I_b_in += Iother_to_base
                    I_b_out += Ibase_to_other
            for e in elementos:
                if e["tipo"] == "Fuente Corriente":
                    n1, n2 = e["nodos"]
                    i_val = e["valor"]
                    if n_base == n2:
                        I_b_in += i_val
                        I_b_out -= i_val
                    elif n_base == n1:
                        I_b_out += i_val
                        I_b_in -= i_val
            
            for idx_f, fv in enumerate(fuentes_v):
                n1, n2 = fv["nodos"]
                cur = corrientes_fuente[idx_f]
                if n_base == n2:
                    I_b_in += cur
                    I_b_out -= cur
                elif n_base == n1:
                    I_b_out += cur
                    I_b_in -= cur

            if comp.tipo == "Transistor NPN":
                I_b = abs(I_b_in)
                vbe_diff = Vb - Ve
                I_c_av = 0.0
                for e in elementos:
                    if e["tipo"] != "Resistencia":
                        continue
                    n1, n2 = e["nodos"]
                    Rv = e["valor"]
                    if n_col in (n1, n2):
                        other = n2 if n_col == n1 else n1
                        V_other = voltaje_nodo(other)
                        Ic_other_to_col = (V_other - Vc) / Rv
                        I_c_av += Ic_other_to_col
                for e in elementos:
                    if e["tipo"] == "Fuente Corriente":
                        n1, n2 = e["nodos"]
                        i_val = e["valor"]
                        if n_col == n2:
                            I_c_av += i_val
                        elif n_col == n1:
                            I_c_av -= i_val
                for idx_f, fv in enumerate(fuentes_v):
                    n1, n2 = fv["nodos"]
                    cur = corrientes_fuente[idx_f]
                    if n_col == n2:
                        I_c_av += cur
                    elif n_col == n1:
                        I_c_av -= cur
                I_c_av = max(0.0, I_c_av)
                if I_b <= 0 or vbe_diff < vbe_th:
                    I_b = 0.0
                    I_c = 0.0
                    I_e = 0.0
                else:
                    I_c_desired = beta * I_b
                    if I_c_av <= 0:
                        I_c = 0.0
                    else:
                        I_c = I_c_desired if I_c_desired < I_c_av else I_c_av
                    I_e = I_b + I_c
                corrientes[f"{comp.nombre} Ib"] = I_b
                corrientes[f"{comp.nombre} Ic"] = I_c
                corrientes[f"{comp.nombre} Ie"] = I_e
            else:
                I_b = abs(I_b_out)
                vbe_diff = Ve - Vb
                I_c_av = 0.0
                for e in elementos:
                    if e["tipo"] != "Resistencia":
                        continue
                    n1, n2 = e["nodos"]
                    Rv = e["valor"]
                    if n_col in (n1, n2):
                        other = n2 if n_col == n1 else n1
                        V_other = voltaje_nodo(other)
                        Ic_col_to_other = (Vc - V_other) / Rv
                        I_c_av += Ic_col_to_other
                for e in elementos:
                    if e["tipo"] == "Fuente Corriente":
                        n1, n2 = e["nodos"]
                        i_val = e["valor"]
                        if n_col == n1:
                            I_c_av += i_val
                        elif n_col == n2:
                            I_c_av -= i_val
                for idx_f, fv in enumerate(fuentes_v):
                    n1, n2 = fv["nodos"]
                    cur = corrientes_fuente[idx_f]
                    if n_col == n1:
                        I_c_av += cur
                    elif n_col == n2:
                        I_c_av -= cur
                I_c_av = max(0.0, I_c_av)
                if I_b <= 0 or vbe_diff < vbe_th:
                    I_b = 0.0
                    I_c = 0.0
                    I_e = 0.0
                else:
                    I_c_desired = beta * I_b
                    if I_c_av <= 0:
                        I_c = 0.0
                    else:
                        I_c = I_c_desired if I_c_desired < I_c_av else I_c_av
                    I_e = I_b + I_c
                corrientes[f"{comp.nombre} Ib"] = I_b
                corrientes[f"{comp.nombre} Ic"] = I_c
                corrientes[f"{comp.nombre} Ie"] = I_e

        mosfet_params = {e['nombre']: e['valor'] for e in elementos if e.get('tipo') == 'Mosfet'}
        
        def voltaje_mosfet(nodo):
            try:
                from PyQt5.QtWidgets import QGraphicsItem
                if isinstance(nodo, QGraphicsItem) or hasattr(nodo, 'node_group'):
                    nodo = getattr(nodo, 'node_group', None)
            except Exception:
                pass
            if nodo is None or nodo == nodo_gnd:
                return 0.0
            idx = nodo_idx.get(nodo, None)
            if idx is None:
                return 0.0
            return voltajes[idx]
        for comp in self.componentes:
            if getattr(comp, 'tipo', None) != 'Mosfet':
                continue
            term = getattr(comp, 'terminales', [])
            if len(term) != 3:
                continue
            t_gate, t_drain, t_source = term[0], term[1], term[2]
            n_gate = getattr(t_gate, 'node_group', None)
            n_drain = getattr(t_drain, 'node_group', None)
            n_source = getattr(t_source, 'node_group', None)
            Vg = voltaje_mosfet(n_gate)
            Vd = voltaje_mosfet(n_drain)
            Vs = voltaje_mosfet(n_source)
            params = mosfet_params.get(comp.nombre, None)
            if params:
                mos_tipo = params.get('tipo', 'N')
                Vth_val = params.get('Vth', 2.0)
                Rds_val = params.get('Rds', 1.0)
            else:
                mos_tipo = 'N'
                Vth_val = 2.0
                Rds_val = 1.0
            if mos_tipo.upper().startswith('N'):
                Vgs = Vg - Vs
                Vds = Vd - Vs
                if Vgs >= Vth_val and abs(Rds_val) > 1e-12:
                    I_ds = Vds / Rds_val
                else:
                    I_ds = 0.0
            else:
                Vsg = Vs - Vg
                Vsd = Vs - Vd
                th = abs(Vth_val)
                if Vsg >= th and abs(Rds_val) > 1e-12:
                    I_ds = Vsd / Rds_val
                else:
                    I_ds = 0.0
            corrientes[f"{comp.nombre} Ids"] = I_ds
        return voltajes, corrientes, Req

def alternar_modo_oscuro(self):
        accion_tema = getattr(self, "accion_tema", None)
        if accion_tema is None:
            return
    
        self.modo_oscuro = not getattr(self, "modo_oscuro", False)
        try:
            preferences['modo_oscuro'] = self.modo_oscuro
            guardar_preferencias()
        except Exception:
            pass
        if self.modo_oscuro:
            accion_tema.setIcon(QIcon("icons/sol_oscuro.png"))
            accion_tema.setText(traducir('Modo Claro'))
        else:
            accion_tema.setIcon(QIcon("icons/sol_claro.png"))
            accion_tema.setText(traducir('Modo Oscuro'))
        
        self.aplicar_tema()

def aplicar_tema(self):
        app = QApplication.instance()
        if app is None:
            app = QApplication(sys.argv)
        try:
            app.setWindowIcon(QIcon(ICON_PATH))
        except Exception:
            pass

        if self.modo_oscuro:
            stylesheet = """
            QMainWindow, QWidget, QMenuBar, QMenu, QToolBar, QListWidget, QLineEdit {
                background-color: #2b2b2b;
                color: #f0f0f0;
            }
            /* Resaltar elementos seleccionados con un borde sutil */
            QListWidget::item:selected {
                background-color: #3e3e3e;
                border: 2px solid #5e5e5e;
                border-radius: 6px;
            }
            QSplitter::handle {
                background: #3b3b3b;
            }
            QTabBar::tab:selected {
                background: #3b3b3b;
            }
            /* Extender el tema oscuro a editores de texto y tablas */
            QTextEdit, QPlainTextEdit, QTableWidget, QTableWidget::item {
                background-color: #2b2b2b;
                color: #f0f0f0;
                selection-background-color: #3e3e3e;
            }
            QHeaderView::section {
                background-color: #3b3b3b;
                color: #f0f0f0;
            }
            QToolTip {
                background-color: #353535;
                color: #f0f0f0;
                border: 1px solid #5e5e5e;
            }
            """         
            try:
                self.scene.grid_pen.setColor(Qt.darkGray)
            except Exception:
                pass
        else:
            stylesheet = ""
            try:
                self.scene.grid_pen.setColor(Qt.lightGray)
            except Exception:
                pass
        if app:
            app.setStyleSheet(stylesheet)
        try:
            from PyQt5.QtCore import Qt as QT
            for item in self.scene.items():
                try:
                    if isinstance(item, TerminalItem):
                        item.setBrush(QT.white if self.modo_oscuro else QT.black)
                except Exception:
                    continue
            for conexion in getattr(self, "conexiones", []):
                try:
                    for line in conexion.lines:
                        pen = line.pen()
                        if getattr(self, "coloreado", False):
                            continue
                        if self.modo_oscuro:
                            pen.setColor(QT.white)
                        else:
                            pen.setColor(QT.black)
                        line.setPen(pen)
                except Exception:
                    continue
        except Exception:
            pass
        try:
            self.scene.update()
        except Exception:
            pass

def agregar_resultado(self, tipo, mensaje):
        try:
            self.resultados_simulaciones.append((tipo, mensaje))
        except Exception:
            self.resultados_simulaciones = [(tipo, mensaje)]
        try:
            if hasattr(self, "lista_resultados") and self.lista_resultados is not None:
                try:
                    texto = mensaje.replace("<br>", " ").replace("<br/>", " ").replace("<br />", " ")
                    import re
                    limpio = re.sub(r"<[^>]+>", "", texto)
                    resumen = limpio[:60] + ("…" if len(limpio) > 60 else "")
                except Exception:
                    resumen = str(mensaje)[:60]
                from PyQt5.QtWidgets import QListWidgetItem
                item = QListWidgetItem(f"{tipo}: {resumen}")
                item.setData(Qt.UserRole, (tipo, mensaje))
                self.lista_resultados.addItem(item)
        except Exception:
            pass
        try:
            llamar_hooks("result_added", tipo=tipo, mensaje=mensaje)
        except Exception:
            pass

def exportar_resultados(self):
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        if not self.resultados_simulaciones:
            QMessageBox.information(self, "Exportar resultados", "No hay resultados para exportar.")
            return
        opciones = "Archivo de texto (*.txt);;Archivo CSV (*.csv)"
        ruta, filtro = QFileDialog.getSaveFileName(self, "Guardar resultados", "", opciones)
        if not ruta:
            return
        try:
            es_csv = ruta.lower().endswith(".csv") or "csv" in filtro.lower()
            from PyQt5.QtCore import Qt as _Qt
            results_iter: list[tuple[str, str]] = []
            try:
                if hasattr(self, 'lista_resultados') and self.lista_resultados is not None:
                    selected_items = self.lista_resultados.selectedItems()
                    if selected_items:
                        for _it in selected_items:
                            data = _it.data(_Qt.UserRole)
                            if data:
                                results_iter.append(data)
                    elif self.resultados_simulaciones:
                        results_iter.append(self.resultados_simulaciones[-1])
                elif self.resultados_simulaciones:
                    results_iter.append(self.resultados_simulaciones[-1])
            except Exception:
                results_iter = list(self.resultados_simulaciones)
            if not results_iter:
                results_iter = list(self.resultados_simulaciones)
            with open(ruta, "w", encoding="utf-8") as f:
                for tipo, mensaje in results_iter:
                    if es_csv:
                        import re
                        texto = re.sub(r'<br\\s*/?>', '\n', mensaje, flags=re.IGNORECASE)
                        texto = re.sub(r'<[^>]+>', '', texto)  
                        lineas = [l.strip() for l in texto.splitlines() if l.strip()]
                        for linea in lineas:
                            f.write(f"{tipo},{linea}\n")
                    else:
                        f.write(f"=== {tipo.upper()} ===\n")
                        texto = mensaje.replace("<br>", "\n").replace("<br/>", "\n").replace("<br />", "\n")
                        import re
                        texto = re.sub(r'<[^>]+>', '', texto)
                        f.write(texto)
                        f.write("\n\n")
            QMessageBox.information(self, "Exportar resultados", f"Resultados exportados en: {ruta}")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudieron exportar los resultados:\n{e}")

def exportar_logica_pdf(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.platypus import Table, TableStyle
            from reportlab.lib import colors
            from reportlab.lib.utils import ImageReader
        except Exception:
            title = traducir('Exportar lógica a PDF')
            msg = (
                "La librería reportlab no está instalada. Instálala para generar archivos PDF."
                if CURRENT_LANGUAGE == 'es' else
                "The reportlab library is not installed. Install it to generate PDF files."
            )
            QMessageBox.critical(self, title, msg)
            return
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a PDF'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a PDF'),
            "",
            "PDF (*.pdf)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".pdf"):
            ruta = f"{ruta}.pdf"
        try:
            c = canvas.Canvas(ruta, pagesize=letter)
            try:
                from PyQt5.QtGui import QImage, QPainter
                from PyQt5.QtCore import QRectF, Qt as _Qt
                import tempfile, os
                rect = self.scene.itemsBoundingRect()
                img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
                img.fill(_Qt.white)
                painter = QPainter(img)
                self.scene.render(painter, target=QRectF(img.rect()), source=rect)
                painter.end()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    img_path = tmp_img.name
                    img.save(img_path)
                c.setFont("Helvetica-Bold", 14)
                circ_label = "Circuito" if CURRENT_LANGUAGE == 'es' else "Circuit"
                c.drawString(50, 750, circ_label)
                try:
                    c.drawImage(ImageReader(img_path), 50, 450, width=500, preserveAspectRatio=True)
                except Exception:
                    pass
                c.showPage()
                try:
                    os.unlink(img_path)
                except Exception:
                    pass
            except Exception:
                c.showPage()
            c.setFont("Helvetica", 12)
            title_text = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
            c.drawString(50, 750, title_text)
            data = [columns]
            for r in rows:
                data.append([r.get(col, 0) for col in columns])
            col_width = 500 // len(columns) if columns else 500
            table = Table(data, colWidths=[col_width] * len(columns))
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#d9d9d9')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.25, colors.grey),
            ]))
            w, h = table.wrapOn(c, 480, 600)
            table.drawOn(c, 50, 750 - 30 - h)
            y_position = 750 - 30 - h - 20
            try:
                functions_info = api.expresiones_booleanas_digitales()
            except Exception:
                functions_info = {}
            if functions_info:
                label_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                c.drawString(50, y_position, label_funcs)
                y_position -= 15
                for out_name, exprs in functions_info.items():
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                    canonical_expr = exprs.get('canonical', '')
                    simplified_expr = exprs.get('simplified', '')
                    line_text = f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}"
                    c.drawString(60, y_position, line_text)
                    y_position -= 12
            c.save()
            success_msg = (
                f"PDF guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"PDF saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a PDF'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a PDF\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to PDF\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a PDF'), error_msg)

def exportar_logica_word(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            from docx import Document
            from docx.shared import Inches
        except Exception:
            title = traducir('Exportar lógica a Word')
            msg = (
                "La librería python‑docx no está instalada. Instálala para generar archivos DOCX."
                if CURRENT_LANGUAGE == 'es' else
                "The python‑docx library is not installed. Install it to generate DOCX files."
            )
            QMessageBox.critical(self, title, msg)
            return
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a Word'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a Word'),
            "",
            "Word Document (*.docx)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".docx"):
            ruta = f"{ruta}.docx"
        try:
            doc = Document()
            heading_text = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
            doc.add_heading(heading_text, level=1)
            try:
                from PyQt5.QtGui import QImage, QPainter
                from PyQt5.QtCore import QRectF, Qt as _Qt
                import tempfile, os
                rect = self.scene.itemsBoundingRect()
                img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
                img.fill(_Qt.white)
                painter = QPainter(img)
                self.scene.render(painter, target=QRectF(img.rect()), source=rect)
                painter.end()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
                    circ_path = tmp_img.name
                    img.save(circ_path)
                circ_label = "Circuito" if CURRENT_LANGUAGE == 'es' else "Circuit"
                doc.add_heading(circ_label, level=2)
                try:
                    doc.add_picture(circ_path, width=Inches(5.5))
                except Exception:
                    doc.add_paragraph("No se pudo insertar la imagen del circuito.")
                try:
                    os.unlink(circ_path)
                except Exception:
                    pass
                doc.add_page_break()
            except Exception:
                pass
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            table = doc.add_table(rows=len(rows_sorted) + 1, cols=len(columns))
            table.style = "Table Grid"
            for j, col in enumerate(columns):
                table.cell(0, j).text = str(col)
            for i, row in enumerate(rows_sorted):
                for j, col in enumerate(columns):
                    val = row.get(col, 0)
                    table.cell(i + 1, j).text = str(val)
            try:
                functions_info = api.expresiones_booleanas_digitales()
            except Exception:
                functions_info = {}
            if functions_info:
                heading_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                doc.add_heading(heading_funcs, level=2)
                for out_name, exprs in functions_info.items():
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                    canonical_expr = exprs.get('canonical', '')
                    simplified_expr = exprs.get('simplified', '')
                    doc.add_paragraph(
                        f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}"
                    )
            doc.save(ruta)
            success_msg = (
                f"Word guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"Word saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a Word'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a Word\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to Word\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a Word'), error_msg)

def exportar_logica_txt(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a TXT'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a TXT'),
            "",
            "Archivo de texto (*.txt)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".txt"):
            ruta = f"{ruta}.txt"
        try:
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            with open(ruta, 'w', encoding='utf-8') as f:
                header = "Tabla de verdad" if CURRENT_LANGUAGE == 'es' else "Truth table"
                f.write(header + "\n")
                f.write('\t'.join(columns) + '\n')
                for row in rows_sorted:
                    f.write('\t'.join(str(row.get(col, 0)) for col in columns) + '\n')
                try:
                    functions_info = api.expresiones_booleanas_digitales()
                except Exception:
                    functions_info = {}
                if functions_info:
                    label_funcs = "Funciones booleanas" if CURRENT_LANGUAGE == 'es' else "Boolean functions"
                    f.write('\n' + label_funcs + '\n')
                    for out_name, exprs in functions_info.items():
                        if CURRENT_LANGUAGE == 'en':
                            canonical_label = 'Canonical'
                            simplified_label = 'Simplified'
                        else:
                            canonical_label = 'Canónica'
                            simplified_label = 'Simplificada'
                        canonical_expr = exprs.get('canonical', '')
                        simplified_expr = exprs.get('simplified', '')
                        f.write(f"{out_name}: {canonical_label}: {canonical_expr}; {simplified_label}: {simplified_expr}\n")
            success_msg = (
                f"TXT guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"TXT saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a TXT'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a TXT\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to TXT\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a TXT'), error_msg)

def exportar_logica_excel(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        try:
            api = SimulationAPI(self.componentes)
            rows, columns = api.tabla_verdad_digital_extendida()
        except Exception:
            rows, columns = [], []
        if not rows or not columns:
            QMessageBox.information(
                self,
                traducir('Exportar lógica a Excel'),
                traducir('No se encontraron compuertas lógicas en el circuito.')
            )
            return
        ruta, _ = QFileDialog.getSaveFileName(
            self,
            traducir('Exportar lógica a Excel'),
            "",
            "Excel (*.xlsx)"
        )
        if not ruta:
            return
        if not ruta.lower().endswith(".xlsx"):
            ruta = f"{ruta}.xlsx"
        try:
            try:
                import pandas as pd
            except Exception:
                title = traducir('Exportar lógica a Excel')
                msg = (
                    "La librería pandas no está instalada. Instálala para generar archivos Excel."
                    if CURRENT_LANGUAGE == 'es' else
                    "The pandas library is not installed. Install it to generate Excel files."
                )
                QMessageBox.critical(self, title, msg)
                return
            try:
                rows_sorted = sorted(rows, key=lambda row: [row.get(col, 0) for col in columns])
            except Exception:
                rows_sorted = rows
            df = pd.DataFrame(rows_sorted)
            try:
                df = df.reindex(columns=columns)
            except Exception:
                pass
            sheet_truth = 'TablaDeVerdad' if CURRENT_LANGUAGE == 'es' else 'TruthTable'
            sheet_funcs = 'FuncionesBooleanas' if CURRENT_LANGUAGE == 'es' else 'BooleanFunctions'
            engine_name = None
            try:
                import xlsxwriter  
                engine_name = 'xlsxwriter'
            except Exception:
                try:
                    import openpyxl  
                    engine_name = 'openpyxl'
                except Exception:
                    engine_name = None
            if engine_name is None:
                title = traducir('Exportar lógica a Excel')
                msg = (
                    "No se encontró un motor de Excel (necesitas instalar 'xlsxwriter' o 'openpyxl') para generar archivos .xlsx."
                    if CURRENT_LANGUAGE == 'es' else
                    "No Excel engine found (please install 'xlsxwriter' or 'openpyxl') to generate .xlsx files."
                )
                QMessageBox.critical(self, title, msg)
                return
            writer_obj = pd.ExcelWriter(ruta, engine=engine_name)
            with writer_obj as writer:
                df.to_excel(writer, index=False, sheet_name=sheet_truth)
                try:
                    functions_info = api.expresiones_booleanas_digitales()
                except Exception:
                    functions_info = {}
                if functions_info:
                    funcs_rows: list[dict[str, str]] = []
                    if CURRENT_LANGUAGE == 'en':
                        canonical_label = 'Canonical'
                        simplified_label = 'Simplified'
                        output_label = 'Output'
                    else:
                        canonical_label = 'Canónica'
                        simplified_label = 'Simplificada'
                        output_label = 'Salida'
                    for out_name, exprs in functions_info.items():
                        funcs_rows.append({
                            output_label: out_name,
                            canonical_label: exprs.get('canonical', ''),
                            simplified_label: exprs.get('simplified', '')
                        })
                    df2 = pd.DataFrame(funcs_rows)
                    df2.to_excel(writer, index=False, sheet_name=sheet_funcs)
            success_msg = (
                f"Excel guardado en:\n{ruta}" if CURRENT_LANGUAGE == 'es' else f"Excel saved to:\n{ruta}"
            )
            QMessageBox.information(self, traducir('Exportar lógica a Excel'), success_msg)
        except Exception as exc:
            error_msg = (
                f"No se pudo exportar la lógica a Excel\n{exc}" if CURRENT_LANGUAGE == 'es' else
                f"Could not export logic to Excel\n{exc}"
            )
            QMessageBox.critical(self, traducir('Exportar lógica a Excel'), error_msg)

def mostrar_atajos_teclado(self):
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        from PyQt5.QtCore import Qt
        atajos = [
            ("Ctrl+X", traducir("Cortar componentes")),
            ("Ctrl+C", traducir("Copiar componentes")),
            ("Ctrl+V", traducir("Pegar componentes")),
            ("Delete", traducir("Eliminar selección")),
            ("Ctrl+Z", traducir("Deshacer")),
            ("Ctrl+Y", traducir("Rehacer")),
            ("Ctrl+A", traducir("Seleccionar todo")),
            ("Ctrl+=", traducir("Acercar vista (zoom in)")),
            ("Ctrl+-", traducir("Alejar vista (zoom out)")),
            ("Ctrl+Shift+V", traducir("Ver todo")),
            ("Ctrl+Shift+1", traducir("Restablecer escala 1:1")),
            ("Ctrl+O", traducir("Cargar circuito")),
        ]
        texto = f"<b>{traducir('Atajos de teclado disponibles:')}</b><br><ul>"
        for combo, desc in atajos:
            texto += f"<li><b>{combo}</b>: {desc}</li>"
        texto += "</ul>"
        msg = QMessageBox(self)
        msg.setWindowTitle(traducir("Atajos de Teclado"))
        msg.setTextFormat(Qt.RichText)
        msg.setIcon(QMessageBox.Information)
        msg.setText(texto)
        msg.exec_()



def _obtener_nodo_gnd(self):
    for comp in self.componentes:
        if getattr(comp, "tipo", "") == "GND" and getattr(comp, "terminales", []):
            try:
                return getattr(comp.terminales[0], "node_group", None)
            except Exception:
                return None
    return None

def exportar_mediciones_csv(self) -> None:
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    import datetime

    try:
        no_med = not getattr(self, 'historial_mediciones', None)
        no_graf = not getattr(self, 'ultima_grafica', None)
    except Exception:
        no_med = True
        no_graf = True

    if no_med and no_graf:
        try:
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is not None:
                voltajes, corrientes_dict, _ = self.analizar_circuito(elementos, nodo_gnd)
                nodos_set = set()
                for e in elementos:
                    for n in e["nodos"]:
                        nodos_set.add(n)
                nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
                nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
                volt_labels = []
                volt_values = []
                for elem in elementos:
                    try:
                        n1, n2 = elem.get("nodos", [None, None])
                        v1 = 0.0 if n1 == nodo_gnd else voltajes[nodo_idx.get(n1, 0)]
                        v2 = 0.0 if n2 == nodo_gnd else voltajes[nodo_idx.get(n2, 0)]
                        dv = abs(v1 - v2)
                        volt_labels.append(elem.get("nombre", ""))
                        volt_values.append(dv)
                    except Exception:
                        continue
                current_labels = []
                current_values = []
                for name, val in corrientes_dict.items():
                    current_labels.append(name)
                    current_values.append(val)
                try:
                    self.ultima_grafica = (volt_labels, volt_values, current_labels, current_values)
                except Exception:
                    self.ultima_grafica = None
        except Exception:
            pass
    has_hist = bool(getattr(self, 'historial_mediciones', None))
    has_graph = False
    try:
        if getattr(self, 'ultima_grafica', None):
            vlabels, vvals, clabels, cvals = self.ultima_grafica
            has_graph = bool(vlabels) or bool(clabels)
    except Exception:
        has_graph = False
    if not has_hist and not has_graph:
        QMessageBox.information(self, "Exportar CSV", "No hay datos de mediciones ni gráficas para exportar.")
        return

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    sugerido = f"mediciones_{timestamp}.csv"
    ruta, _ = QFileDialog.getSaveFileName(
        self,
        "Guardar datos de mediciones y gráficas",
        sugerido,
        "CSV (*.csv)"
    )
    if not ruta:
        return
    try:
        import csv
        with open(ruta, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            if has_hist:
                writer.writerow(["Tipo de medición", "Valor"])
                try:
                    for tipo, valor in self.historial_mediciones:
                        writer.writerow([tipo, valor])
                except Exception:
                    pass
                writer.writerow([])
            try:
                if getattr(self, 'ultima_grafica', None):
                    volt_labels, volt_values, current_labels, current_values = self.ultima_grafica
                    if volt_labels:
                        writer.writerow(["Elemento", "Voltaje (V)"])
                        for lbl, val in zip(volt_labels, volt_values):
                            writer.writerow([lbl, f"{val}"])
                        writer.writerow([])
                    if current_labels:
                        writer.writerow(["Elemento", "Corriente (A)"])
                        for lbl, val in zip(current_labels, current_values):
                            writer.writerow([lbl, f"{val}"])
                        writer.writerow([])
            except Exception:
                pass
        try:
            import os
            preferences['last_folder'] = os.path.dirname(ruta)
            guardar_preferencias()
        except Exception:
            pass
        mostrar_informacion(self, "Exportar CSV", f"Datos exportados exitosamente a:\n{ruta}")
        try:
            self.agregar_resultado("exportar mediciones", f"Mediciones exportadas a: {ruta}")
        except Exception:
            pass
        try:
            llamar_hooks("measurement_exported", ruta=ruta)
        except Exception:
            pass
    except Exception as exc:
        try:
            mostrar_error(self, "exportar mediciones CSV", exc)
        except Exception:
            pass
        mostrar_error_detallado(self, "Exportar CSV", f"No se pudo exportar a CSV:\n{exc}")

def mostrar_mediciones_activas(self):
    from PyQt5.QtWidgets import QDialog, QTableWidget, QTableWidgetItem, QHeaderView, QVBoxLayout, QHBoxLayout, QCheckBox, QPushButton, QMessageBox
    from PyQt5.QtCore import Qt, QPointF
    from PyQt5.QtGui import QFont
    def calcular_mediciones():
        filas = []
        try:
            issues = self.ejecutar_comprobaciones_previas()
        except Exception as exc:
            return [], str(exc)
        if any(i.level == "error" for i in issues):
            return [], "El circuito no pasa el chequeo eléctrico. Añade GND y corrige problemas antes de medir."
        try:
            elementos = self.extraer_topologia()
        except Exception as e:
            return [], f"Error al extraer la topología: {e}"
        try:
            nodo_gnd = self._nodo_gnd_actual()
        except Exception:
            nodo_gnd = None
        if nodo_gnd is None:
            return [], "Debes colocar un nodo de tierra (GND) para poder medir."
        try:
            voltajes, corrientes_fuente, _ = self.analizar_circuito(elementos, nodo_gnd)
        except Exception as exc:
            return [], f"No se pudo resolver el circuito: {exc}"
        nodos_set = set()
        for e in elementos:
            for n in e["nodos"]:
                nodos_set.add(n)
        nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
        nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
        volt_dict = {nodo_gnd: 0.0}
        for n, idx in nodo_idx.items():
            try:
                volt_dict[n] = voltajes[idx]
            except Exception:
                volt_dict[n] = 0.0
        lista_fuentes = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
        for comp in self.componentes:
            nombre = getattr(comp, "nombre", "")
            tipo = getattr(comp, "tipo", "")
            if tipo in ["Amperímetro", "Voltímetro", "Ohmímetro"]:
                continue
            elem = next((e for e in elementos if e.get("nombre") == nombre), None)
            v_diff = 0.0
            i_val = None
            p_val = None
            if elem is not None:
                nodos = elem.get("nodos", [])
                if len(nodos) >= 2:
                    n1, n2 = nodos[0], nodos[1]
                    v1 = volt_dict.get(n1, 0.0)
                    v2 = volt_dict.get(n2, 0.0)
                    v_diff = v1 - v2
                tipo_e = elem.get("tipo")
                if tipo_e == "Resistencia":
                    try:
                        r_val = float(elem.get("valor"))
                        i_val = (v_diff / r_val) if r_val != 0 else None
                        p_val = (v_diff ** 2) / r_val if r_val != 0 else None
                    except Exception:
                        i_val = None
                        p_val = None
                elif tipo_e in ["Fuente Voltaje", "Batería"]:
                    try:
                        idx_fuente = lista_fuentes.index(elem)
                        i_src = corrientes_fuente[idx_fuente]
                        i_val = i_src
                        p_val = v_diff * i_src
                    except Exception:
                        i_val = None
                        p_val = None
                elif tipo_e == "Fuente Corriente":
                    try:
                        i_val = -elem.get("valor", 0.0)
                        p_val = v_diff * i_val
                    except Exception:
                        i_val = None
                        p_val = None
            filas.append([nombre, tipo, v_diff, i_val, p_val])
        return filas, None
    filas, error = calcular_mediciones()
    if error:
        QMessageBox.warning(self, "Mediciones activas", error)
        return
    try:
        mensaje_med = "<b>Mediciones activas:</b><br>"
        for (nomb, tipo, v_val, i_val, p_val) in filas:
            partes = []
            try:
                partes.append(f"V={v_val:.3g}V")
            except Exception:
                pass
            if i_val is not None:
                try:
                    partes.append(f"I={i_val:.3g}A")
                except Exception:
                    pass
            if p_val is not None:
                try:
                    partes.append(f"P={p_val:.3g}W")
                except Exception:
                    pass
            mensaje_med += f"<b>{nomb}</b> ({tipo}): " + " ".join(partes) + "<br>"
        try:
            self.agregar_resultado("mediciones", mensaje_med)
        except Exception:
            pass
        try:
            llamar_hooks("measurement_completed", filas=filas)
        except Exception:
            pass
    except Exception:
        pass
    dlg = QDialog(self)
    dlg.setWindowTitle("Mediciones activas")
    layout = QVBoxLayout(dlg)
    tabla = QTableWidget(len(filas), 5, dlg)
    tabla.setHorizontalHeaderLabels(["Nombre", "Tipo", "V (V)", "I (A)", "P (W)"])
    tabla.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    font_header = tabla.horizontalHeader().font()
    font_header.setBold(True)
    tabla.horizontalHeader().setFont(font_header)
    def poblar_tabla(d):
        tabla.setRowCount(len(d))
        for i, (n, t, v, I, P) in enumerate(d):
            tabla.setItem(i, 0, QTableWidgetItem(str(n)))
            tabla.setItem(i, 1, QTableWidgetItem(str(t)))
            try:
                tabla.setItem(i, 2, QTableWidgetItem(f"{v:.6g}"))
            except Exception:
                tabla.setItem(i, 2, QTableWidgetItem(""))
            try:
                if I is not None:
                    tabla.setItem(i, 3, QTableWidgetItem(f"{I:.6g}"))
                else:
                    tabla.setItem(i, 3, QTableWidgetItem(""))
            except Exception:
                tabla.setItem(i, 3, QTableWidgetItem(""))
            try:
                if P is not None:
                    tabla.setItem(i, 4, QTableWidgetItem(f"{P:.6g}"))
                else:
                    tabla.setItem(i, 4, QTableWidgetItem(""))
            except Exception:
                tabla.setItem(i, 4, QTableWidgetItem(""))
        tabla.resizeRowsToContents()
        if check_mostrar.isChecked():
            actualizar_etiquetas(d)
    check_mostrar = QCheckBox("Mostrar en pantalla", dlg)
    def eliminar_etiquetas():
        try:
            for lbl in getattr(self, "mediciones_labels", []):
                try:
                    self.scene.removeItem(lbl)
                except Exception:
                    pass
        except Exception:
            pass
        self.mediciones_labels = []
    def actualizar_etiquetas(d):
        eliminar_etiquetas()
        if not hasattr(self, "mediciones_labels"):
            self.mediciones_labels = []
        from PyQt5.QtWidgets import QGraphicsTextItem
        for (nombre, tipo, v, I, P) in d:
            comp_obj = None
            for c in self.componentes:
                try:
                    if getattr(c, "nombre", "") == nombre:
                        comp_obj = c
                        break
                except Exception:
                    continue
            if comp_obj is None:
                continue
            try:
                pos = comp_obj.scenePos()
            except Exception:
                try:
                    pos = comp_obj.pos()
                except Exception:
                    pos = QPointF(0, 0)
            texto = []
            try:
                texto.append(f"V={v:.3g}V")
            except Exception:
                pass
            if I is not None:
                try:
                    texto.append(f"I={I:.3g}A")
                except Exception:
                    pass
            if P is not None:
                try:
                    texto.append(f"P={P:.3g}W")
                except Exception:
                    pass
            if not texto:
                continue
            label_text = "\n".join(texto)
            lbl = QGraphicsTextItem(label_text)
            lbl.setDefaultTextColor(Qt.darkMagenta)
            lbl.setPos(pos + QPointF(0, -35))
            lbl.setZValue(1000)
            try:
                self.scene.addItem(lbl)
            except Exception:
                pass
            self.mediciones_labels.append(lbl)
    def al_cambiar_marcacion(state):
        if state == Qt.Checked:
            actualizar_etiquetas(filas)
        else:
            eliminar_etiquetas()
    check_mostrar.stateChanged.connect(al_cambiar_marcacion)
    btn_actualizar = QPushButton("Actualizar", dlg)
    def al_actualizar():
        nonlocal filas, error
        filas, err = calcular_mediciones()
        if err:
            QMessageBox.warning(self, "Mediciones", err)
            return
        poblar_tabla(filas)
    btn_actualizar.clicked.connect(al_actualizar)
    layout.addWidget(tabla)
    hbox = QHBoxLayout()
    hbox.addWidget(check_mostrar)
    hbox.addStretch(1)
    hbox.addWidget(btn_actualizar)
    layout.addLayout(hbox)
    poblar_tabla(filas)
    def al_cerrar(result):
        eliminar_etiquetas()
    dlg.finished.connect(al_cerrar)
    dlg.setLayout(layout)
    dlg.exec_()

def gestor_plantillas(self):
    from PyQt5.QtWidgets import QDialog, QListWidget, QListWidgetItem, QLineEdit, QLabel, QVBoxLayout, QHBoxLayout, QPushButton, QMessageBox
    from PyQt5.QtGui import QIcon
    from PyQt5.QtCore import Qt
    from pathlib import Path
    try:
        base_dir = Path(__file__).resolve().parent
    except Exception:
        base_dir = Path('.')
    json_files = list(base_dir.glob('*.json'))
    dlg = QDialog(self)
    dlg.setWindowTitle("Gestor de plantillas")
    dlg.resize(480, 400)
    layout = QVBoxLayout(dlg)
    buscador = QLineEdit(dlg)
    buscador.setPlaceholderText("Buscar plantilla...")
    layout.addWidget(buscador)
    lista = QListWidget(dlg)
    layout.addWidget(lista)
    def poblar_lista():
        lista.clear()
        for fpath in json_files:
            nombre = fpath.stem
            icono = None
            try:
                png_path = fpath.with_suffix('.png')
                if png_path.exists():
                    icono = QIcon(str(png_path))
            except Exception:
                icono = None
            if not icono or icono.isNull():
                try:
                    icono = QIcon(ICON_PATH)
                except Exception:
                    icono = QIcon()
            item = QListWidgetItem(icono, nombre)
            item.setData(Qt.UserRole, str(fpath))
            lista.addItem(item)
    poblar_lista()
    def filtrar(text):
        text = text.lower()
        for i in range(lista.count()):
            item = lista.item(i)
            name = item.text().lower()
            item.setHidden(text not in name)
    buscador.textChanged.connect(filtrar)
    boton_layout = QHBoxLayout()
    btn_abrir = QPushButton("Abrir", dlg)
    btn_cerrar = QPushButton("Cerrar", dlg)
    boton_layout.addWidget(btn_abrir)
    boton_layout.addWidget(btn_cerrar)
    layout.addLayout(boton_layout)
    def abrir_plantilla():
        item = lista.currentItem()
        if not item:
            QMessageBox.information(dlg, "Seleccionar plantilla", "Selecciona una plantilla de la lista.")
            return
        path_str = item.data(Qt.UserRole)
        try:
            ruta = Path(path_str)
            self.cargar_circuito_archivo(ruta)
        except Exception as exc:
            QMessageBox.critical(dlg, "Error", f"No se pudo cargar la plantilla:\n{exc}")
        dlg.accept()
    btn_abrir.clicked.connect(abrir_plantilla)
    lista.itemDoubleClicked.connect(lambda _: abrir_plantilla())
    btn_cerrar.clicked.connect(dlg.reject)
    dlg.setLayout(layout)
    dlg.exec_()

def exportar_reporte_pdf_completo(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    from PyQt5.QtGui import QImage, QPainter
    from PyQt5.QtCore import QRectF, Qt
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.utils import ImageReader
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors
    import datetime
    import tempfile
    import os
    archivo, _ = QFileDialog.getSaveFileName(self, "Guardar reporte completo", "", "PDF (*.pdf)")
    if not archivo:
        return
    if not archivo.lower().endswith('.pdf'):
        archivo = f"{archivo}.pdf"
    try:
        c = canvas.Canvas(archivo, pagesize=letter)
        c.setFont("Helvetica-Bold", 20)
        c.drawCentredString(300, 700, "Reporte de Simulación")
        c.setFont("Helvetica", 12)
        fecha = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        c.drawCentredString(300, 670, f"Generado el {fecha}")
        try:
            c.drawImage(ImageReader(ICON_PATH), 200, 480, width=200, preserveAspectRatio=True)
        except Exception:
            pass
        c.showPage()
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_img:
            img_path = tmp_img.name
            img.save(img_path)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Circuito")
        try:
            c.drawImage(ImageReader(img_path), 50, 400, width=500, preserveAspectRatio=True)
        except Exception:
            pass
        c.showPage()
        datos_bom, err_bom = self._calcular_bom()
        if err_bom:
            raise Exception(err_bom)
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Lista de materiales (BOM)")
        if datos_bom:
            ncols = len(datos_bom[0])
            if ncols == 9:
                col_widths = [80, 70, 180, 50, 60, 50, 60, 60, 40]
            elif ncols == 6:
                col_widths = [90, 80, 180, 60, 60, 80]
            else:
                col_widths = [500.0 / ncols] * ncols
            table_bom = Table(datos_bom, colWidths=col_widths)
            table_bom.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            w, h = table_bom.wrapOn(c, 500, 600)
            table_bom.drawOn(c, 50, 720 - h)
        c.showPage()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Gráficas de voltaje y corriente")
        try:
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is None:
                raise Exception("No se encuentra nodo GND para graficar")
            voltajes, corrientes_fuente, _ = self.analizar_circuito(elementos, nodo_gnd)
            nodos_set_g = set()
            for e in elementos:
                for n in e["nodos"]:
                    nodos_set_g.add(n)
            nodos_sorted_g = sorted(n for n in nodos_set_g if n != nodo_gnd)
            nodo_idx_g = {n: i for i, n in enumerate(nodos_sorted_g)}
            volt_dict_g = {nodo_gnd: 0.0}
            for n, idx in nodo_idx_g.items():
                volt_dict_g[n] = voltajes[idx]
            volt_labels = []
            volt_vals = []
            curr_labels = []
            curr_vals = []
            lista_fuentes_g = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
            for elem in elementos:
                try:
                    nombre = elem.get("nombre", "")
                    tipo = elem.get("tipo", "")
                    n1, n2 = elem.get("nodos", [None, None])
                    v1 = volt_dict_g.get(n1, 0.0)
                    v2 = volt_dict_g.get(n2, 0.0)
                    dv = abs(v1 - v2)
                    volt_labels.append(nombre)
                    volt_vals.append(dv)
                    i_val = None
                    if tipo == "Resistencia":
                        try:
                            r_val = float(elem.get("valor"))
                            i_val = dv / r_val if r_val != 0 else None
                        except Exception:
                            i_val = None
                    elif tipo in ["Fuente Voltaje", "Batería"]:
                        try:
                            idxf = lista_fuentes_g.index(elem)
                            i_src = corrientes_fuente[idxf]
                            i_val = abs(i_src)
                        except Exception:
                            i_val = None
                    elif tipo == "Fuente Corriente":
                        try:
                            i_val = abs(elem.get("valor", 0.0))
                        except Exception:
                            i_val = None
                    if i_val is not None:
                        curr_labels.append(nombre)
                        curr_vals.append(i_val)
                except Exception:
                    continue
            if volt_labels or curr_labels:
                fig, axes = plt.subplots(2, 1, figsize=(8, 6))
                if volt_labels:
                    xv = list(range(len(volt_labels)))
                    axes[0].plot(xv, volt_vals, marker='o', linestyle='-', color='tab:blue')
                    axes[0].set_xticks(xv)
                    axes[0].set_xticklabels(volt_labels, rotation=45, ha='right', fontsize=6)
                    axes[0].set_ylabel('Voltaje (V)')
                    axes[0].set_title('Voltajes por elemento')
                    axes[0].grid(True, linestyle='--', alpha=0.5)
                if curr_labels:
                    xi = list(range(len(curr_labels)))
                    axes[1].plot(xi, curr_vals, marker='o', linestyle='-', color='tab:green')
                    axes[1].set_xticks(xi)
                    axes[1].set_xticklabels(curr_labels, rotation=45, ha='right', fontsize=6)
                    axes[1].set_ylabel('Corriente (A)')
                    axes[1].set_title('Corrientes por elemento')
                    axes[1].grid(True, linestyle='--', alpha=0.5)
                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_plot:
                    plot_path = tmp_plot.name
                fig.savefig(plot_path)
                plt.close(fig)
                try:
                    c.drawImage(ImageReader(plot_path), 50, 400, width=500, height=300, preserveAspectRatio=True)
                except Exception:
                    pass
            else:
                c.setFont("Helvetica", 12)
                c.drawString(50, 720, "No se pudieron generar gráficas debido a la falta de datos.")
        except Exception as gexc:
            c.setFont("Helvetica", 12)
            c.drawString(50, 720, f"No se pudieron generar gráficas: {gexc}")
        c.showPage()
        c.setFont("Helvetica-Bold", 14)
        c.drawString(50, 750, "Parámetros de simulación")
        datos_param = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            try:
                datos_param.append([
                    getattr(comp, "nombre", ""),
                    getattr(comp, "tipo", ""),
                    ", ".join(getattr(comp, "parametros", []))
                ])
            except Exception:
                continue
        if datos_param:
            ncols_p = len(datos_param[0])
            col_widths_p = [160, 100, 240] if ncols_p == 3 else [500.0 / ncols_p] * ncols_p
            table_params = Table(datos_param, colWidths=col_widths_p)
            table_params.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#d9d9d9")),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            w_p, h_p = table_params.wrapOn(c, 500, 600)
            table_params.drawOn(c, 50, 720 - h_p)
        c.save()
        try:
            if os.path.exists(img_path):
                os.unlink(img_path)
        except Exception:
            pass
        try:
            if 'plot_path' in locals() and os.path.exists(plot_path):
                os.unlink(plot_path)
        except Exception:
            pass
        QMessageBox.information(self, "Reporte PDF", f"Reporte guardado en:\n{archivo}")
    except Exception as exc:
        QMessageBox.critical(self, "Reporte PDF", f"No se pudo exportar el reporte completo:\n{exc}")

def exportar_reporte_word_completo(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    from PyQt5.QtGui import QImage, QPainter
    from PyQt5.QtCore import QRectF, Qt
    import datetime
    import tempfile
    import os

    archivo, _ = QFileDialog.getSaveFileName(
        self, "Guardar reporte completo", "", "Word Document (*.docx)"
    )
    if not archivo:
        return
    if not archivo.lower().endswith(".docx"):
        archivo = f"{archivo}.docx"

    try:
        try:
            from docx import Document
            from docx.shared import Inches, Cm
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except Exception:
            QMessageBox.critical(
                self,
                "Reporte Word",
                "No se pudo exportar el reporte completo:\n"
                "La librería python-docx no está instalada. Instálala para generar archivos DOCX.",
            )
            return

        doc = Document()
        titulo = doc.add_heading("Reporte de Simulación", level=1)
        titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        fecha = datetime.datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        par_fecha = doc.add_paragraph(f"Generado el {fecha}")
        par_fecha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        if ICON_PATH and os.path.isfile(ICON_PATH):
            try:
                doc.add_picture(ICON_PATH, width=Inches(3))
            except Exception:
                pass
        doc.add_page_break()

        doc.add_heading("Circuito", level=2)
        rect = self.scene.itemsBoundingRect()
        img = QImage(rect.size().toSize(), QImage.Format_ARGB32)
        img.fill(Qt.white)
        painter = QPainter(img)
        self.scene.render(painter, target=QRectF(img.rect()), source=rect)
        painter.end()
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_circ:
            circ_path = tmp_circ.name
            img.save(circ_path)
        try:
            doc.add_picture(circ_path, width=Inches(5.5))
        except Exception:
            doc.add_paragraph("No se pudo insertar la imagen del circuito.")
        try:
            if os.path.exists(circ_path):
                os.unlink(circ_path)
        except Exception:
            pass
        doc.add_page_break()

        doc.add_heading("Lista de materiales (BOM)", level=2)
        datos_bom, err_bom = self._calcular_bom()
        if err_bom:
            raise Exception(err_bom)
        if datos_bom:
            nrows = len(datos_bom)
            ncols = len(datos_bom[0])
            table = doc.add_table(rows=nrows, cols=ncols)
            table.style = "Table Grid"
            for i, row in enumerate(datos_bom):
                for j, val in enumerate(row):
                    cell = table.cell(i, j)
                    try:
                        cell.text = str(val)
                    except Exception:
                        cell.text = ""
                    for paragraph in cell.paragraphs:
                        paragraph.style = doc.styles['Normal']
            doc.add_page_break()
        else:
            doc.add_paragraph("No se pudo generar la lista de materiales.")

        doc.add_heading("Gráficas de voltaje y corriente", level=2)
        plot_temp_path = None
        try:
            import matplotlib
            matplotlib.use("Agg")
            import matplotlib.pyplot as plt
            elementos = self.extraer_topologia()
            nodo_gnd = self._nodo_gnd_actual()
            if nodo_gnd is None:
                raise Exception("No se encuentra nodo GND para graficar")
            voltajes, corrientes_fuente, _ = self.analizar_circuito(elementos, nodo_gnd)
            nodos_set_g = set()
            for e in elementos:
                for n in e["nodos"]:
                    nodos_set_g.add(n)
            nodos_sorted_g = sorted(n for n in nodos_set_g if n != nodo_gnd)
            nodo_idx_g = {n: i for i, n in enumerate(nodos_sorted_g)}
            volt_dict_g = {nodo_gnd: 0.0}
            for n, idx in nodo_idx_g.items():
                volt_dict_g[n] = voltajes[idx]
            volt_labels = []
            volt_vals = []
            curr_labels = []
            curr_vals = []
            lista_fuentes_g = [e for e in elementos if e.get("tipo") in ["Fuente Voltaje", "Batería"]]
            for elem in elementos:
                try:
                    nombre = elem.get("nombre", "")
                    tipo = elem.get("tipo", "")
                    n1, n2 = elem.get("nodos", [None, None])
                    v1 = volt_dict_g.get(n1, 0.0)
                    v2 = volt_dict_g.get(n2, 0.0)
                    dv = abs(v1 - v2)
                    volt_labels.append(nombre)
                    volt_vals.append(dv)
                    i_val = None
                    if tipo == "Resistencia":
                        try:
                            r_val = float(elem.get("valor"))
                            i_val = dv / r_val if r_val != 0 else None
                        except Exception:
                            i_val = None
                    elif tipo in ["Fuente Voltaje", "Batería"]:
                        try:
                            idxf = lista_fuentes_g.index(elem)
                            i_src = corrientes_fuente[idxf]
                            i_val = abs(i_src)
                        except Exception:
                            i_val = None
                    elif tipo == "Fuente Corriente":
                        try:
                            i_val = abs(elem.get("valor", 0.0))
                        except Exception:
                            i_val = None
                    if i_val is not None:
                        curr_labels.append(nombre)
                        curr_vals.append(i_val)
                except Exception:
                    continue
            if volt_labels or curr_labels:
                fig, axes = plt.subplots(2, 1, figsize=(8, 6))
                if volt_labels:
                    xv = list(range(len(volt_labels)))
                    axes[0].plot(
                        xv,
                        volt_vals,
                        marker="o",
                        linestyle="-", color="tab:blue"
                    )
                    axes[0].set_xticks(xv)
                    axes[0].set_xticklabels(
                        volt_labels, rotation=45, ha="right", fontsize=6
                    )
                    axes[0].set_ylabel("Voltaje (V)")
                    axes[0].set_title("Voltajes por elemento")
                    axes[0].grid(True, linestyle="--", alpha=0.5)
                if curr_labels:
                    xi = list(range(len(curr_labels)))
                    axes[1].plot(
                        xi,
                        curr_vals,
                        marker="o",
                        linestyle="-", color="tab:green"
                    )
                    axes[1].set_xticks(xi)
                    axes[1].set_xticklabels(
                        curr_labels, rotation=45, ha="right", fontsize=6
                    )
                    axes[1].set_ylabel("Corriente (A)")
                    axes[1].set_title("Corrientes por elemento")
                    axes[1].grid(True, linestyle="--", alpha=0.5)
                plt.tight_layout()
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_plot:
                    plot_temp_path = tmp_plot.name
                fig.savefig(plot_temp_path)
                plt.close(fig)
                doc.add_picture(plot_temp_path, width=Inches(5.5))
            else:
                doc.add_paragraph(
                    "No se pudieron generar gráficas debido a la falta de datos."
                )
        except Exception as gexc:
            doc.add_paragraph(f"No se pudieron generar gráficas: {gexc}")
        try:
            if plot_temp_path and os.path.exists(plot_temp_path):
                os.unlink(plot_temp_path)
        except Exception:
            pass
        doc.add_page_break()

        doc.add_heading("Parámetros de simulación", level=2)
        datos_param = [["Nombre", "Tipo", "Parámetros"]]
        for comp in self.componentes:
            try:
                datos_param.append(
                    [
                        getattr(comp, "nombre", ""),
                        getattr(comp, "tipo", ""),
                        ", ".join(getattr(comp, "parametros", [])),
                    ]
                )
            except Exception:
                continue
        if len(datos_param) > 1:
            nrows_p = len(datos_param)
            ncols_p = len(datos_param[0])
            table_p = doc.add_table(rows=nrows_p, cols=ncols_p)
            table_p.style = "Table Grid"
            for i, row in enumerate(datos_param):
                for j, val in enumerate(row):
                    cell = table_p.cell(i, j)
                    cell.text = str(val)
                    for paragraph in cell.paragraphs:
                        paragraph.style = doc.styles['Normal']
        else:
            doc.add_paragraph("No hay parámetros registrados.")

        doc.save(archivo)
        QMessageBox.information(self, "Reporte Word", f"Reporte guardado en:\n{archivo}")
    except Exception as exc:
        QMessageBox.critical(self, "Reporte Word", f"No se pudo exportar el reporte completo:\n{exc}")

def exportar_netlist_spice(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    gnd_id = self._obtener_nodo_gnd()
    if gnd_id is None:
        QMessageBox.warning(self, "Exportar SPICE", "Debes colocar un GND en el circuito para exportar.")
        return
    nodemap = {gnd_id: "0"}
    next_num = 1

    def id_spice(node_id):
        nonlocal next_num
        if node_id in nodemap:
            return nodemap[node_id]
        nodemap[node_id] = str(next_num)
        next_num += 1
        return nodemap[node_id]

    lines = ["* ESIMulador netlist"]
    for comp in self.componentes:
        if not getattr(comp, "terminales", []):
            continue
        ctype = getattr(comp, "tipo", "")
        if ctype == "GND":
            continue
        if ctype in ["Resistencia", "Amperímetro", "Voltímetro", "Ohmímetro"] or "(eq." in comp.nombre:
            pref = "R"
        elif ctype == "Capacitor":
            pref = "C"
        elif ctype == "Bobina":
            pref = "L"
        elif ctype in ["Fuente Voltaje", "Batería", "Generador de funciones"]:
            pref = "V"
        elif ctype == "Fuente Corriente":
            pref = "I"
        else:
            pref = "R"
        nombre_spice = comp.nombre.replace(" ", "_")
        nodos = []
        for term in comp.terminales:
            try:
                nodos.append(id_spice(getattr(term, "node_group", None)))
            except Exception:
                nodos.append("0")
            if len(nodos) == 2:
                break
        if len(nodos) < 2:
            continue
        valor = "1"
        try:
            if getattr(comp, "parametros", []):
                raw = comp.parametros[0]
                if ":" in raw:
                    val = raw.split(":", 1)[1].strip()
                else:
                    val = raw.strip()
                try:
                    valor_num = interpretar_valor_prefijo(val)
                    valor = f"{valor_num}"
                except Exception:
                    valor = val
        except Exception:
            valor = "1"
        lines.append(f"{pref}{nombre_spice} {nodos[0]} {nodos[1]} {valor}")
    ruta, _ = QFileDialog.getSaveFileName(self, "Guardar netlist SPICE", "circuito.cir", "Netlist (*.cir *.sp *.spice *.txt)")
    if not ruta:
        return
    try:
        with open(ruta, "w", encoding="utf-8") as f:
            f.write("\n".join(lines) + "\n")
        try:
            logger.info(f"Netlist SPICE exportado a {ruta}")
        except Exception:
            pass
    except Exception as e:
        QMessageBox.critical(self, "Exportar SPICE", f"No se pudo guardar el netlist:\n{e}")

def _poner_componente_simple(self, tipo, nombre, valor_txt, pos):
    from PyQt5.QtCore import QPointF
    
    datos = self.imagenes.get(tipo, None)
    if not datos:
        return None
    archivo, terminales, vertical, invertir, *resto = datos
    escala = resto[0] if resto else 0.15
    
    comp = ComponenteInteractivo(archivo, [f"Valor: {valor_txt}"], nombre, tipo, terminales=terminales, vertical=vertical, invertir=invertir)
    comp.setScale(escala)
    try:
        comp.setPos(pos)
    except Exception:
        comp.setPos(QPointF(pos.x(), pos.y()))
    self.scene.addItem(comp)
    self.componentes.append(comp)
    try:
        comp.reposicionar_terminales()
    except Exception:
        pass
    return comp

def importar_netlist_spice(self):
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    from PyQt5.QtCore import QPointF
    import re
    ruta, _ = QFileDialog.getOpenFileName(self, "Abrir netlist SPICE", "", "Netlist (*.cir *.sp *.spice *.txt)")
    if not ruta:
        return
    try:
        with open(ruta, "r", encoding="utf-8") as f:
            lines = [ln.strip() for ln in f if ln.strip() and not ln.strip().startswith("*")]
    except Exception as e:
        QMessageBox.critical(self, "Importar SPICE", f"No se pudo leer el archivo:\n{e}")
        return
    try:
        self.limpiar_todo()
    except Exception:
        pass
    comp_por_nodo = {}
    x, y = 100, 80
    dx = 180
    for ln in lines:
        m = re.match(r"^([RCLVI])([A-Za-z0-9_]+)\s+(\S+)\s+(\S+)\s+(\S+)$", ln, flags=re.IGNORECASE)
        if not m:
            
            continue
        pref, name, n1, n2, val = m.groups()
        pref = pref.upper()
        tipo = {
            "R": "Resistencia",
            "C": "Capacitor",
            "L": "Bobina",
            "V": "Fuente Voltaje",
            "I": "Fuente Corriente"
        }.get(pref, "Resistencia")
        nombre_comp = name.replace("_", " ")
        comp = self._poner_componente_simple(tipo, nombre_comp, val, QPointF(x, y))
        if comp is None:
            continue
        comp_por_nodo.setdefault(n1, []).append((comp, 0))
        comp_por_nodo.setdefault(n2, []).append((comp, 1))
        x += dx
    if "0" in comp_por_nodo:
        try:
            datos_gnd = self.imagenes.get("GND", None)
            if datos_gnd:
                archivo, terminales, vertical, invertir, *resto = datos_gnd
                escala = resto[0] if resto else 0.1
                gnd_comp = ComponenteInteractivo(archivo, [], "GND1", "GND", terminales=terminales, vertical=vertical, invertir=invertir)
                gnd_comp.setScale(escala)
                gnd_comp.setPos(QPointF(120, y + 120))
                self.scene.addItem(gnd_comp)
                self.componentes.append(gnd_comp)
                comp_por_nodo.setdefault("0", []).append((gnd_comp, 0))
        except Exception:
            pass
    
    for nodo, extremos in comp_por_nodo.items():
        if len(extremos) < 2:
            continue
        base_comp, base_idx = extremos[0]
        base_term = base_comp.terminales[base_idx]
        for comp, idx in extremos[1:]:
            try:
                conexion = Conexion(self.scene, base_term, comp.terminales[idx])
                self.conexiones.append(conexion)
            except Exception:
                pass
    try:
        logger.info("Importación SPICE terminada.")
    except Exception:
        pass

def colorear_conexiones(self):
        from PyQt5.QtGui import QColor
        if not hasattr(self, 'coloreado'):
            self.coloreado = False  

        colores = [Qt.red, Qt.green, Qt.blue, Qt.magenta, Qt.darkCyan, Qt.darkYellow]

        for i, conexion in enumerate(self.conexiones):
            color = QColor(0, 0, 0)  
            if not self.coloreado:
                color = colores[i % len(colores)]
            for line in conexion.lines:
                pen = line.pen()
                pen.setColor(color)
                pen.setWidth(3 if not self.coloreado else 1)
                line.setPen(pen)

        self.coloreado = not self.coloreado
        msg = "Conexiones coloreadas" if self.coloreado else "Conexiones restauradas"
        self.statusBar().showMessage(msg, 3000)

def autodetectar_bode(self):
    nodo_gnd, nodos_sorted, nodo_idx = self._mapear_nodos_actual()

    def es_fuente_v(c):
        t = (getattr(c, "tipo", "") or "").lower()
        return ("fuente" in t and "volt" in t) or ("voltaje" in t)

    fuentes = [c for c in self.componentes if es_fuente_v(c)]
    if len(fuentes) != 1:
        return None
    fv = fuentes[0]
    if len(getattr(fv, "terminales", [])) < 2:
        return None
    tinp = fv.terminales[0] 
    tinm = fv.terminales[1]  

    for comp in self.componentes:
        nom = (getattr(comp, "nombre", "") or "").upper()
        if nom in ("OUT", "SALIDA") and getattr(comp, "terminales", []):
            toutp = comp.terminales[0]

            class DummyTerm:
                def __init__(self, node): self.node_group = node
            toutm = DummyTerm(nodo_gnd)
            return (tinp, tinm, toutp, toutm)

    from collections import defaultdict, deque
    grafo = defaultdict(set)

    for comp in self.componentes:
        terms = getattr(comp, "terminales", [])
        if len(terms) >= 2:
            n1 = getattr(terms[0], "node_group", None)
            n2 = getattr(terms[1], "node_group", None)
            if None not in (n1, n2) and n1 != n2:
                grafo[n1].add(n2); grafo[n2].add(n1)

    origen = getattr(tinp, "node_group", None)
    if origen is None:
        return None

    dist = {origen: 0}
    q = deque([origen])
    while q:
        u = q.popleft()
        for v in grafo.get(u, []):
            if v not in dist:
                dist[v] = dist[u] + 1
                q.append(v)

    candidatos = [(d, n) for n, d in dist.items() if n != nodo_gnd]
    if not candidatos:
        return None
    candidatos.sort(reverse=True)
    n_out = candidatos[0][1]

    toutp = None
    for comp in self.componentes:
        for t in getattr(comp, "terminales", []):
            if getattr(t, "node_group", None) == n_out:
                toutp = t
                break
        if toutp:
            break
    if toutp is None:
        return None

    class DummyTerm:
        def __init__(self, node): self.node_group = node
    toutm = DummyTerm(nodo_gnd)

    return (tinp, tinm, toutp, toutm)

def dialogo_bode_combo(self):
    from PyQt5.QtWidgets import QDialog, QFormLayout, QDialogButtonBox, QComboBox

    try:
        nodo_gnd, nodos_sorted, nodo_idx = self._mapear_nodos_actual()
    except Exception as e:
        QMessageBox.warning(self, "Topología no válida", str(e))
        return None

    opciones = ["GND"] + [f"N{n}" for n in nodos_sorted]

    def crear_combo(idx0=0):
        cb = QComboBox(); cb.addItems(opciones); cb.setCurrentIndex(idx0); return cb

    dlg = QDialog(self); dlg.setWindowTitle("Configurar Bode (AC)")
    lay = QFormLayout(dlg)
    c_inp = crear_combo(); c_inm = crear_combo(0); c_outp = crear_combo(); c_outm = crear_combo(0)
    lay.addRow("IN +", c_inp); lay.addRow("IN −", c_inm)
    lay.addRow("OUT +", c_outp); lay.addRow("OUT −", c_outm)

    btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, parent=dlg)
    lay.addWidget(btns)
    btns.accepted.connect(dlg.accept); btns.rejected.connect(dlg.reject)

    if dlg.exec_() != QDialog.Accepted:
        return None

    class DummyTerm:
        def __init__(self, node): self.node_group = node

    def a_terminal(txt):
        if txt == "GND":
            return DummyTerm(nodo_gnd)
        node_val = int(txt[1:])  
        for comp in self.componentes:
            for t in getattr(comp, "terminales", []):
                if getattr(t, "node_group", None) == node_val:
                    return t
        
        dt = DummyTerm(node_val)
        return dt

    tinp = a_terminal(c_inp.currentText())
    tinm = a_terminal(c_inm.currentText())
    toutp = a_terminal(c_outp.currentText())
    toutm = a_terminal(c_outm.currentText())

    if None in (tinp, tinm, toutp, toutm):
        QMessageBox.warning(self, "Selección inválida",
                            "Alguno de los nodos seleccionados no existe en el esquema.")
        return None
    return (tinp, tinm, toutp, toutm)

def iniciar_bode(self):
    try:
        import numpy as np, matplotlib.pyplot as plt  
    except Exception:
        QMessageBox.warning(self, "Dependencia faltante",
                            "Necesitas 'numpy' y 'matplotlib' para graficar Bode.")
        return

    sel = None
    try:
        comps_util = [c for c in getattr(self,'componentes',[]) if getattr(c,'tipo','') not in ('GND',)]
        if len(comps_util) < 2:
            self._bode_por_transferencia(); return
    except Exception:
        pass
    try:
        sel = self.autodetectar_bode()
    except Exception:
        sel = None

    if sel is None:
        sel = self.dialogo_bode_combo()
        if sel is None:
            self._bode_por_transferencia()
            return

    tinp, tinm, toutp, toutm = sel
    self._graficar_bode(tinp, tinm, toutp, toutm) 

def _mapear_nodos_actual(self):
    try:
        self.extraer_topologia()
    except Exception:
        pass
    
    try:
        nodo_gnd = self._nodo_gnd_actual()
    except Exception:
        nodo_gnd = None
    if nodo_gnd is None:
        raise RuntimeError("Debes colocar un GND para poder simular.")
    
    nodos_set = set()
    for comp in getattr(self, "componentes", []):
        for t in getattr(comp, "terminales", []):
            grp = getattr(t, "node_group", None)
            if grp is not None:
                nodos_set.add(grp)
    
    nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
    nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
    return nodo_gnd, nodos_sorted, nodo_idx

def _graficar_bode(self, tinp, tinm, toutp, toutm):
    import numpy as np
    import matplotlib.pyplot as plt
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    
    try:
        f, mag_db, phase_deg = simulation_core.calcular_bode(
            getattr(self, 'componentes', []), tinp, tinm, toutp, toutm
        )
    except Exception as exc:
        QMessageBox.warning(self, "Simulación AC",
                            f"No se pudo preparar la simulación: {exc}")
        return
    import matplotlib.pyplot as plt
    try:
        if getattr(self, 'modo_oscuro', False):
            plt.style.use('dark_background')
        else:
            plt.style.use('default')
    except Exception:
        pass
    fig = plt.figure(figsize=(8, 6))
    ax1 = fig.add_subplot(211)
    ax1.semilogx(f, mag_db)
    ax1.set_ylabel('|H(jω)| (dB)')
    ax1.grid(True, which='both', ls='--', alpha=0.5)
    ax1.set_title('Bode: Magnitud y Fase')
    if len(mag_db) > 0:
        m_min = float(np.min(mag_db))
        m_max = float(np.max(mag_db))
        if m_min == m_max:
            margen = 1.0
            ax1.set_ylim(m_min - margen, m_max + margen)
        else:
            margen = 0.1 * (m_max - m_min)
            ax1.set_ylim(m_min - margen, m_max + margen)
    ax2 = fig.add_subplot(212)
    ax2.semilogx(f, phase_deg)
    ax2.set_xlabel('Frecuencia (Hz)')
    ax2.set_ylabel('Fase (°)')
    ax2.grid(True, which='both', ls='--', alpha=0.5)
    if len(phase_deg) > 0:
        p_min = float(np.min(phase_deg))
        p_max = float(np.max(phase_deg))
        if p_min == p_max:
            margen = 10.0
            ax2.set_ylim(p_min - margen, p_max + margen)
        else:
            margen = 0.1 * (p_max - p_min)
            ax2.set_ylim(p_min - margen, p_max + margen)
    try:
        ax1.relim()
        ax1.autoscale_view()
        ax2.relim()
        ax2.autoscale_view()
    except Exception:
        pass
    cursor_lines = {'ax1': None, 'ax2': None}
    annotation = fig.text(0.01, 0.95, '', transform=fig.transFigure, color='red', fontsize=9)

    def al_mover(event):
        if event.inaxes not in (ax1, ax2):
            return
        freq = event.xdata
        if freq is None:
            return
        idx = int(np.argmin(np.abs(f - freq)))
        idx = max(0, min(len(f) - 1, idx))
        freq_val = f[idx]
        mag_val = mag_db[idx]
        phase_val = phase_deg[idx]
        for key, ax in [('ax1', ax1), ('ax2', ax2)]:
            if cursor_lines[key] is not None:
                try:
                    cursor_lines[key].remove()
                except Exception:
                    pass
            cursor_lines[key] = ax.axvline(freq_val, color='red', linestyle='--', alpha=0.7)
        annotation.set_text(
            f"f = {freq_val:.4g} Hz | |H| = {mag_val:.2f} dB | ∠H = {phase_val:.1f}°"
        )
        fig.canvas.draw_idle()
    fig.canvas.mpl_connect('motion_notify_event', al_mover)

    plt.tight_layout()
    plt.show()

def _bode_por_transferencia(self):
    from PyQt5.QtWidgets import QInputDialog, QMessageBox
    import numpy as np
    import cmath, math

    expr, ok = QInputDialog.getText(
        self,
        traducir('Bode de función de transferencia'),
        traducir('Ingresa H(s) (usa s para la variable compleja):'),
        text="s(s-1)/((1+s+3s**2)(s+1))"
    )
    if not ok or not expr.strip():
        return
    
    expr_input = expr.strip()
    expr_eval = expr_input.replace('^', '**')
    import re
    expr_eval = re.sub(r'S', 's', expr_eval)
    expr_eval = re.sub(r'(\d|\))\s*(s)', r'\1*\2', expr_eval)  
    expr_eval = re.sub(r'(\d|\))\s*\(', r'\1*(', expr_eval)    
    expr_eval = re.sub(r'(s)\s*(\d|\()', r'\1*\2', expr_eval)  
    expr_eval = re.sub(r'(\))\s*(\()', r'\1*\2', expr_eval)    

    try:
        cfg = preferences.get('bode_config', None)
        if isinstance(cfg, (list, tuple)) and len(cfg) == 3:
            fmin, fmax, npts = float(cfg[0]), float(cfg[1]), int(cfg[2])
        else:
            fmin, fmax, npts = 1e-3, 1e4, 401
    except Exception:
        fmin, fmax, npts = 1e-3, 1e4, 401
    try:
        default_text = f"{fmin}, {fmax}, {npts}"
    except Exception:
        default_text = "0.001, 10000, 401"
    try:
        rng_str, ok2 = QInputDialog.getText(self, _("Rango de frecuencias"),
                                            _("fmin, fmax, puntos (Hz):"),
                                            text=default_text)
        if ok2 and rng_str and rng_str.strip():
            parts = [p.strip() for p in rng_str.split(',')]
            if len(parts) >= 2:
                fmin = float(parts[0]); fmax = float(parts[1])
            if len(parts) >= 3:
                npts = int(parts[2])
            try:
                preferences['bode_config'] = [fmin, fmax, npts]
                guardar_preferencias()
            except Exception:
                pass
    except Exception:
        pass
    if fmin <= 0 or fmax <= 0 or fmax <= fmin:
        QMessageBox.warning(self, "Parámetros inválidos", "Rango de frecuencias no válido.")
        return

    f = np.logspace(np.log10(fmin), np.log10(fmax), int(npts))
    w = 2*np.pi*f

    safe = {
        '__builtins__': {},
        's': 0j,
        'j': 1j,
        'pi': math.pi,
        'e': math.e,
        'sin': math.sin, 'cos': math.cos, 'tan': math.tan,
        'exp': math.exp, 'sqrt': math.sqrt, 'log': math.log, 'abs': abs
    }
    H = np.zeros_like(w, dtype=complex)
    for i, omega in enumerate(w):
        safe['s'] = 1j*omega
        try:
            val = eval(expr_eval, safe, {})
        except Exception as ex:
            QMessageBox.critical(self, "Error al evaluar H(s)", f"No se pudo evaluar la expresión:\n{ex}")
            return
        H[i] = complex(val)

    mag_db = 20*np.log10(np.maximum(np.abs(H), 1e-20))
    phase_deg = np.angle(H, deg=True)

    import matplotlib.pyplot as plt
    fig = plt.figure(figsize=(8,6))
    ax1 = fig.add_subplot(211)
    ax1.semilogx(f, mag_db)
    ax1.set_ylabel('|H(jω)| (dB)')
    ax1.grid(True, which='both', ls='--', alpha=0.5)
    try:
        ax1.set_title(f'Bode: {expr_input}')
    except Exception:
        ax1.set_title('Bode')
    if len(mag_db) > 0:
        m_min = float(np.min(mag_db))
        m_max = float(np.max(mag_db))
        if m_min == m_max:
            margen = 1.0
            ax1.set_ylim(m_min - margen, m_max + margen)
        else:
            margen = 0.1 * (m_max - m_min)
            ax1.set_ylim(m_min - margen, m_max + margen)
    ax2 = fig.add_subplot(212)
    ax2.semilogx(f, phase_deg)
    ax2.set_xlabel(traducir('Frecuencia (Hz)'))
    ax2.set_ylabel(traducir('Fase (°)'))
    ax2.grid(True, which='both', ls='--', alpha=0.5)
    if len(phase_deg) > 0:
        p_min = float(np.min(phase_deg))
        p_max = float(np.max(phase_deg))
        if p_min == p_max:
            margen = 10.0
            ax2.set_ylim(p_min - margen, p_max + margen)
        else:
            margen = 0.1 * (p_max - p_min)
            ax2.set_ylim(p_min - margen, p_max + margen)
    try:
        ax1.relim()
        ax1.autoscale_view()
        ax2.relim()
        ax2.autoscale_view()
    except Exception:
        pass
    plt.tight_layout(); plt.show()

def iniciar_lgr(self):
    from PyQt5.QtWidgets import QInputDialog, QMessageBox
    try:
        import numpy as np
        import matplotlib.pyplot as plt
        import sympy as sp
        try:
            if getattr(self, 'modo_oscuro', False):
                plt.style.use('dark_background')
            else:
                plt.style.use('default')
        except Exception:
            pass
    except Exception:
        QMessageBox.warning(self, "Dependencia faltante",
                            "Necesitas 'numpy', 'matplotlib' y 'sympy' para graficar el LGR.")
        return
    try:
        import control as ct
        s = sp.Symbol('s')

        default_lgr = preferences.get('lgr_function', "1/(s*(s+1)*(s**2+2*s+5))")
        func_str, ok = QInputDialog.getText(
            self,
            _("Gráfica LGR limpia"),
            _("Ingresa la función de transferencia G(s):"),
            text=str(default_lgr)
        )
        if not ok or not func_str:
            return
        try:
            preferences['lgr_function'] = func_str
            guardar_preferencias()
        except Exception:
            pass

        expr = sp.sympify(func_str.replace('^', '**'))
        num_expr, den_expr = expr.as_numer_denom()
        num = [float(c) for c in sp.Poly(num_expr, s).all_coeffs()]
        den = [float(c) for c in sp.Poly(den_expr, s).all_coeffs()]

        G = ct.tf(num, den)
        ct.root_locus(G, grid=True)
        try:
            fig = plt.gcf()
            fig.suptitle("")
        except Exception:
            pass
        plt.title(f"G(S): {func_str}")
        plt.show()
        return  
    except Exception:
        pass  
    default_lgr2 = preferences.get('lgr_function', "1/(s*(s+1)*(s**2+2*s+5))")
    expr, ok = QInputDialog.getText(
        self,
        _("LGR de función de transferencia"),
        _("Ingresa H(s) (usa s para la variable compleja):"),
        text=str(default_lgr2)
    )
    if not ok or not expr or not expr.strip():
        return
    try:
        preferences['lgr_function'] = expr.strip()
        guardar_preferencias()
    except Exception:
        pass
    expr_input = expr.strip()
    expr_eval = expr_input.replace('^', '**')
    import re
    expr_eval = re.sub(r'S', 's', expr_eval)
    expr_eval = re.sub(r'(\d|\))\s*(s)', r'\1*\2', expr_eval)
    expr_eval = re.sub(r'(\d|\))\s*\(', r'\1*(', expr_eval)
    expr_eval = re.sub(r'(s)\s*(\d|\()', r'\1*\2', expr_eval)
    expr_eval = re.sub(r'(\))\s*(\()', r'\1*\2', expr_eval)

    try:
        s = sp.symbols('s')
        expr_sym = sp.sympify(expr_eval, {'s': s})
    except Exception as ex:
        QMessageBox.critical(self, "Error al evaluar H(s)",
                             f"No se pudo interpretar la expresión:\n{ex}")
        return

    try:
        num, den = sp.fraction(sp.simplify(expr_sym))
        num_poly = sp.Poly(num, s)
        den_poly = sp.Poly(den, s)
        num_coeffs = [complex(coef) for coef in num_poly.all_coeffs()]
        den_coeffs = [complex(coef) for coef in den_poly.all_coeffs()]
    except Exception as ex:
        QMessageBox.critical(self, "Error al procesar H(s)",
                             f"No se pudo obtener numerador y denominador:\n{ex}")
        return
    from PyQt5.QtWidgets import QApplication
    from PyQt5.QtCore import Qt
    try:
        QApplication.setOverrideCursor(Qt.WaitCursor)
    except Exception:
        pass
    K_vals = np.logspace(-2, 3, 200)
    len_diff = len(den_coeffs) - len(num_coeffs)
    if len_diff > 0:
        num_coeffs_pad = [0] * len_diff + num_coeffs
        den_coeffs_pad = den_coeffs
    elif len_diff < 0:
        den_coeffs_pad = [0] * (-len_diff) + den_coeffs
        num_coeffs_pad = num_coeffs
    else:
        den_coeffs_pad = den_coeffs
        num_coeffs_pad = num_coeffs
    den_coeffs_pad = np.array(den_coeffs_pad, dtype=complex)
    num_coeffs_pad = np.array(num_coeffs_pad, dtype=complex)

    roots_by_k = []
    for K in K_vals:
        char_coeffs = den_coeffs_pad + K * num_coeffs_pad
        try:
            roots = np.roots(char_coeffs)
        except Exception:
            roots = []
        roots_by_k.append(roots)
    try:
        QApplication.restoreOverrideCursor()
        QApplication.processEvents()
    except Exception:
        pass

    if not roots_by_k or len(roots_by_k[0]) == 0:
        QMessageBox.warning(self, "LGR", "No se pudo calcular raíces para el rango de ganancia.")
        return

    nbranches = len(roots_by_k[0])
    initial_roots = sorted(roots_by_k[0], key=lambda r: r.real)
    branches = [[r] for r in initial_roots]
    for roots in roots_by_k[1:]:
        sorted_roots = sorted(roots, key=lambda r: r.real)
        for idx in range(min(nbranches, len(sorted_roots))):
            branches[idx].append(sorted_roots[idx])

    fig, ax = plt.subplots(figsize=(8, 6))
    colors = plt.cm.viridis(np.linspace(0, 1, nbranches))
    for branch, color in zip(branches, colors):
        b_arr = np.array(branch)
        ax.plot(b_arr.real, b_arr.imag, color=color)
    try:
        poles = np.roots(den_coeffs)
    except Exception:
        poles = []
    try:
        zeros = np.roots(num_coeffs)
    except Exception:
        zeros = []
    if len(poles) > 0:
        ax.plot(np.real(poles), np.imag(poles), 'x', markersize=10, color='red', label='Polos')
    if len(zeros) > 0:
        ax.plot(np.real(zeros), np.imag(zeros), 'o', markersize=8, mfc='none', color='blue', label='Ceros')

    ax.axhline(0, color='black', linewidth=0.5, linestyle='--')
    ax.axvline(0, color='black', linewidth=0.5, linestyle='--')
    ax.set_xlabel('Real Axis')
    ax.set_ylabel('Imaginary Axis')
    try:
        ax.set_title(f'Root Locus: {expr_input}')
    except Exception:
        ax.set_title('Root Locus')
    ax.grid(True, which='both', ls='--', alpha=0.5)
    if len(poles) + len(zeros) > 0:
        ax.legend()
    plt.tight_layout()
    plt.show()

def simular_transitorio_stub(self):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    QMessageBox.information(self, "Transitorio (próximo)",
                            "Módulo de Transitorio simple quedará activado en la siguiente iteración.")

def modelos_no_lineales(self):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        self.resolver_matriz_nodos()
        QMessageBox.information(
            self,
            "Modelos no lineales",
            "Simulación completada con modelos no lineales para diodos, BJT y MOSFET.\n"
            "Los valores mostrados consideran aproximaciones numéricas iterativas."
        )
    except Exception as exc:
        try:
            QMessageBox.warning(
                self,
                "Modelos no lineales",
                f"No se pudo completar la simulación no lineal: {exc}"
            )
        except Exception:
            pass

def iniciar_analisis_ac(self):
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        self.iniciar_bode()
    except Exception as exc:
        try:
            QMessageBox.warning(self, "Análisis AC", f"No se pudo iniciar el análisis AC: {exc}")
        except Exception:
            pass

def simular_transitorio(self):
    import numpy as np
    from PyQt5.QtWidgets import QInputDialog, QMessageBox
    try:
        nodo_gnd, nodos_sorted, nodo_idx = self._mapear_nodos_actual()
    except Exception as exc:
        QMessageBox.warning(self, "Simulación transitoria",
                            f"No se pudo preparar la simulación: {exc}")
        return
    try:
        params_str, ok = QInputDialog.getText(
            self, "Parámetros de simulación transitoria",
            "Ingresa dt, tmax (segundos):", text="0.001, 1.0")
        if not ok or not params_str.strip():
            return
        parts = [p.strip() for p in params_str.split(',')]
        if len(parts) >= 1:
            dt = float(parts[0])
        else:
            dt = 0.001
        if len(parts) >= 2:
            tmax = float(parts[1])
        else:
            tmax = 1.0
    except Exception:
        QMessageBox.warning(self, "Parámetros inválidos",
                            "No se pudo interpretar los parámetros ingresados.")
        return
    if dt <= 0 or tmax <= 0 or tmax <= dt:
        QMessageBox.warning(self, "Parámetros inválidos",
                            "El tiempo y el paso deben ser positivos y tmax > dt.")
        return
    try:
        sel = self.autodetectar_bode()
    except Exception:
        sel = None
    if sel is None:
        sel = self.dialogo_bode_combo()
        if sel is None:
            QMessageBox.information(self, "Simulación transitoria",
                                    "No se seleccionaron terminales; se cancela la simulación.")
            return
    tinp, tinm, toutp, toutm = sel
    def _id_nodo(term):
        return getattr(term, "node_group", None)
    n_inp, n_inm = _id_nodo(tinp), _id_nodo(tinm)
    n_outp, n_outm = _id_nodo(toutp), _id_nodo(toutm)
    elementos_res = []  
    elementos_cap = []  
    elementos_src_v = []  
    elementos_src_i = []  
    def obtener_valor(comp):
        try:
            params = getattr(comp, "parametros", [])
            val_str = params[0].split(":", 1)[1].strip()
            return interpretar_valor_prefijo(val_str)
        except Exception:
            return 1.0
    for comp in getattr(self, "componentes", []):
        tipo = getattr(comp, "tipo", "")
        nodos = []
        for t in getattr(comp, "terminales", []):
            nodos.append(getattr(t, "node_group", None))
        if len(nodos) < 2:
            continue
        val = obtener_valor(comp)
        nombre = getattr(comp, "nombre", tipo)
        if tipo in ("Resistencia", "Resistor", "Resistencia variable"):
            elementos_res.append((nodos[0], nodos[1], val, nombre))
        elif tipo in ("Capacitor", "Condensador"):
            elementos_cap.append((nodos[0], nodos[1], val, nombre))
        elif tipo in ("Fuente Voltaje", "Batería"):
            elementos_src_v.append((nodos[0], nodos[1], val, nombre))
        elif tipo in ("Fuente Corriente", ):
            elementos_src_i.append((nodos[0], nodos[1], val, nombre))
    N = len(nodos_sorted)
    M = len(elementos_src_v)
    B_total = np.zeros((N, M))
    E_total = np.zeros(M)
    for idx, (n1, n2, valor, _nom) in enumerate(elementos_src_v):
        if n1 != nodo_gnd:
            B_total[nodo_idx[n1], idx] = 1.0
        if n2 != nodo_gnd:
            B_total[nodo_idx[n2], idx] = -1.0
        E_total[idx] = valor
    G_res = np.zeros((N, N))
    for (n1, n2, Rval, _nom) in elementos_res:
        if abs(Rval) < 1e-30:
            continue
        g = 1.0 / Rval
        if n1 != nodo_gnd:
            i = nodo_idx[n1]
            G_res[i, i] += g
        if n2 != nodo_gnd:
            j = nodo_idx[n2]
            G_res[j, j] += g
        if n1 != nodo_gnd and n2 != nodo_gnd:
            i = nodo_idx[n1]; j = nodo_idx[n2]
            G_res[i, j] -= g
            G_res[j, i] -= g
    G_cap = np.zeros((N, N))
    caps_precalc = []
    for (n1, n2, Cval, _nom) in elementos_cap:
        if abs(Cval) < 1e-30:
            continue
        Gc = Cval / dt
        caps_precalc.append((n1, n2, Gc))
        if n1 != nodo_gnd:
            i = nodo_idx[n1]
            G_cap[i, i] += Gc
        if n2 != nodo_gnd:
            j = nodo_idx[n2]
            G_cap[j, j] += Gc
        if n1 != nodo_gnd and n2 != nodo_gnd:
            i = nodo_idx[n1]; j = nodo_idx[n2]
            G_cap[i, j] -= Gc
            G_cap[j, i] -= Gc
    G_total = G_res + G_cap
    I_const = np.zeros(N)
    for (n1, n2, Ival, _nom) in elementos_src_i:
        if n1 != nodo_gnd:
            I_const[nodo_idx[n1]] -= Ival
        if n2 != nodo_gnd:
            I_const[nodo_idx[n2]] += Ival
    V_prev = np.zeros(N)
    A_top = np.hstack((G_total, B_total)) if M > 0 else G_total
    if M > 0:
        A_bottom = np.hstack((B_total.T, np.zeros((M, M))))
        A_mat = np.vstack((A_top, A_bottom))
    else:
        A_mat = A_top
    n_steps = int(np.ceil(tmax / dt)) + 1
    t_vals = np.linspace(0.0, dt * (n_steps - 1), n_steps)
    v_out_vals = np.zeros(n_steps, dtype=float)
    z_vec = np.zeros(N + M)
    for k in range(n_steps):
        t = t_vals[k]
        I_vec = I_const.copy()
        for (n1, n2, Gc) in caps_precalc:
            v_diff_prev = 0.0
            if n1 != nodo_gnd:
                v1 = V_prev[nodo_idx[n1]]
            else:
                v1 = 0.0
            if n2 != nodo_gnd:
                v2 = V_prev[nodo_idx[n2]]
            else:
                v2 = 0.0
            v_diff_prev = v1 - v2
            I_hist = -Gc * v_diff_prev
            if n1 != nodo_gnd:
                I_vec[nodo_idx[n1]] += I_hist
            if n2 != nodo_gnd:
                I_vec[nodo_idx[n2]] -= I_hist
        if M > 0:
            z_vec[:N] = I_vec
            z_vec[N:] = E_total
        else:
            z_vec[:N] = I_vec
        try:
            x = np.linalg.solve(A_mat, z_vec)
        except Exception as ex:
            QMessageBox.critical(self, "Simulación transitoria",
                                 f"Error al resolver el sistema en t = {t:.4g}s: {ex}")
            return
        V_sol = x[:N] if M > 0 else x
        V_prev = V_sol.copy()
        def voltaje_nodo(n):
            if n is None or n == nodo_gnd:
                return 0.0
            return V_sol[nodo_idx[n]]
        v_out_vals[k] = voltaje_nodo(n_outp) - voltaje_nodo(n_outm)
    import matplotlib.pyplot as plt
    fig, ax = plt.subplots(figsize=(8, 4))
    ax.plot(t_vals, v_out_vals)
    ax.set_xlabel('Tiempo (s)')
    ax.set_ylabel('Tensión de salida (V)')
    ax.set_title('Simulación transitoria')
    ax.grid(True)
    cursor_line = {'line': None}
    annotation = fig.text(0.01, 0.95, '', transform=fig.transFigure, color='red', fontsize=9)
    def al_mover(event):
        if event.inaxes != ax:
            return
        t_cur = event.xdata
        if t_cur is None:
            return
        idx = int(np.argmin(np.abs(t_vals - t_cur)))
        idx = max(0, min(len(t_vals) - 1, idx))
        t_val = t_vals[idx]
        v_val = v_out_vals[idx]
        if cursor_line['line'] is not None:
            try:
                cursor_line['line'].remove()
            except Exception:
                pass
        cursor_line['line'] = ax.axvline(t_val, color='red', linestyle='--', alpha=0.7)
        annotation.set_text(f"t = {t_val:.5g} s | V = {v_val:.5g} V")
        fig.canvas.draw_idle()
    fig.canvas.mpl_connect('motion_notify_event', al_mover)
    plt.tight_layout()
    plt.show()

def carta_smith_interactiva(self):
    import numpy as np
    import matplotlib.pyplot as plt
    from PyQt5.QtWidgets import QInputDialog, QMessageBox
    try:
        nodo_gnd, nodos_sorted, nodo_idx = self._mapear_nodos_actual()
    except Exception as exc:
        QMessageBox.warning(self, "Carta Smith",
                            f"No se pudo preparar la simulación: {exc}")
        return
    try:
        sel = self.autodetectar_bode()
    except Exception:
        sel = None
    if sel is None:
        sel = self.dialogo_bode_combo()
        if sel is None:
            QMessageBox.information(self, "Carta Smith",
                                    "No se seleccionaron terminales; se cancela la carta Smith.")
            return
    tinp, tinm, _toutp, _toutm = sel
    def _id_nodo(term):
        return getattr(term, "node_group", None)
    n_inp, n_inm = _id_nodo(tinp), _id_nodo(tinm)
    if None in (n_inp, n_inm):
        QMessageBox.warning(self, "Carta Smith",
                            "Terminales seleccionadas no tienen nodos válidos.")
        return
    try:
        cfg = preferences.get('smith_config', None)
        if isinstance(cfg, (list, tuple)) and len(cfg) == 4:
            default_Z0, def_fmin, def_fmax, def_npts = cfg
            default_Z0 = float(default_Z0)
            def_fmin = float(def_fmin)
            def_fmax = float(def_fmax)
            def_npts = int(def_npts)
        else:
            default_Z0, def_fmin, def_fmax, def_npts = 50.0, 1e3, 1e6, 401
    except Exception:
        default_Z0, def_fmin, def_fmax, def_npts = 50.0, 1e3, 1e6, 401
    try:
        z0_str, ok = QInputDialog.getText(self, _("Impedancia característica"),
                                           _("Ingresa el valor de Z0 en ohms:"),
                                           text=str(default_Z0))
        if not ok or not z0_str or not z0_str.strip():
            return
        Z0 = float(z0_str.replace(',', '.'))
    except Exception:
        QMessageBox.warning(self, _("Carta Smith"), _("Valor de Z0 no válido."))
        return
    try:
        default_rng = f"{def_fmin}, {def_fmax}, {def_npts}"
    except Exception:
        default_rng = "1e3, 1e6, 401"
    try:
        rng_str, ok2 = QInputDialog.getText(self, _("Rango de frecuencias"),
                                            _("fmin, fmax, puntos (Hz):"),
                                            text=default_rng)
        if ok2 and rng_str and rng_str.strip():
            parts = [p.strip() for p in rng_str.split(',')]
            fmin = float(parts[0]) if len(parts) >= 1 else def_fmin
            fmax = float(parts[1]) if len(parts) >= 2 else def_fmax
            npts = int(parts[2]) if len(parts) >= 3 else def_npts
        else:
            fmin, fmax, npts = def_fmin, def_fmax, def_npts
        try:
            preferences['smith_config'] = [Z0, fmin, fmax, npts]
            guardar_preferencias()
        except Exception:
            pass
    except Exception:
        QMessageBox.warning(self, _("Carta Smith"), _("No se pudo interpretar el rango de frecuencias."))
        return
    if fmin <= 0 or fmax <= 0 or fmax <= fmin:
        QMessageBox.warning(self, _("Carta Smith"), _("Rango de frecuencias no válido."))
        return
    elementos = []
    def obtener_valor(comp):
        try:
            params = getattr(comp, "parametros", [])
            val_str = params[0].split(":", 1)[1].strip()
            return interpretar_valor_prefijo(val_str)
        except Exception:
            return 1.0
    for comp in getattr(self, "componentes", []):
        tipo = getattr(comp, "tipo", "")
        nodos = []
        for t in getattr(comp, "terminales", []):
            nodos.append(getattr(t, "node_group", None))
        if len(nodos) < 2:
            continue
        val = obtener_valor(comp)
        if tipo in ("Resistencia", "Resistor", "Resistencia variable"):
            elementos.append({"tipo": "Resistencia", "nodos": nodos[:2], "valor": val})
        elif tipo in ("Capacitor", "Condensador"):
            elementos.append({"tipo": "Capacitor", "nodos": nodos[:2], "valor": val})
        elif tipo in ("Bobina", "Inductor"):
            elementos.append({"tipo": "Bobina", "nodos": nodos[:2], "valor": val})
    f = np.logspace(np.log10(fmin), np.log10(fmax), int(npts))
    w = 2.0 * np.pi * f
    N = len(nodos_sorted)
    Z_in = np.zeros(len(w), dtype=complex)
    def construir_Y(omega):
        Y_mat = np.zeros((N, N), dtype=complex)
        def estampar_admitancia(n1, n2, y):
            if n1 != nodo_gnd:
                i = nodo_idx[n1]
                Y_mat[i, i] += y
            if n2 != nodo_gnd:
                j = nodo_idx[n2]
                Y_mat[j, j] += y
            if n1 != nodo_gnd and n2 != nodo_gnd:
                i = nodo_idx[n1]; j = nodo_idx[n2]
                Y_mat[i, j] -= y
                Y_mat[j, i] -= y
        for el in elementos:
            n1, n2 = el["nodos"][0], el["nodos"][1]
            val = el["valor"]
            if el["tipo"] == "Resistencia":
                if abs(val) < 1e-30:
                    continue
                y = 1.0 / val
            elif el["tipo"] == "Capacitor":
                y = 1j * omega * val
            elif el["tipo"] == "Bobina":
                if abs(val) < 1e-30:
                    continue
                y = 1.0 / (1j * omega * val)
            else:
                continue
            estampar_admitancia(n1, n2, y)
        return Y_mat
    for idx_w, omega in enumerate(w):
        Y_mat = construir_Y(omega)
        I_vec = np.zeros(N, dtype=complex)
        if n_inp != nodo_gnd:
            I_vec[nodo_idx[n_inp]] -= 1.0
        if n_inm != nodo_gnd:
            I_vec[nodo_idx[n_inm]] += 1.0
        try:
            V = np.linalg.solve(Y_mat, I_vec)
        except Exception:
            Z_in[idx_w] = complex('inf')
            continue
        def voltaje_nodo(n):
            if n is None or n == nodo_gnd:
                return 0.0
            return V[nodo_idx[n]]
        V_port = voltaje_nodo(n_inp) - voltaje_nodo(n_inm)
        Z_in[idx_w] = V_port
    Gamma = (Z_in - Z0) / (Z_in + Z0)
    fig, ax = plt.subplots(figsize=(6, 6))
    ax.set_title('Carta Smith')
    theta = np.linspace(0, 2 * np.pi, 400)
    ax.plot(np.cos(theta), np.sin(theta), color='black', linestyle='-')
    for r in [0, 0.5, 1, 2, 5, 10]:
        c = r / (1.0 + r)
        rad = 1.0 / (1.0 + r)
        x = c + rad * np.cos(theta)
        y = rad * np.sin(theta)
        ax.plot(x, y, color='lightgray', linestyle='--', linewidth=0.5)
    for x in [-5, -2, -1, -0.5, 0.5, 1, 2, 5]:
        xc = 1.0
        yc = 1.0 / x
        rad = 1.0 / abs(x)
        phi = np.linspace(0, 2 * np.pi, 400)
        xs = xc + rad * np.cos(phi)
        ys = yc + rad * np.sin(phi)
        mask = (xs**2 + ys**2) <= 1.0 + 1e-6
        ax.plot(xs[mask], ys[mask], color='lightgray', linestyle='--', linewidth=0.5)
    ax.plot(Gamma.real, Gamma.imag, color='blue', marker='o', markersize=3, linestyle='-', label='Γ(f)')
    ax.set_xlabel('Re{Γ}')
    ax.set_ylabel('Im{Γ}')
    ax.set_aspect('equal')
    ax.grid(False)
    ax.set_xlim(-1.1, 1.1)
    ax.set_ylim(-1.1, 1.1)
    annotation = fig.text(0.01, 0.95, '', transform=fig.transFigure, color='red', fontsize=9)
    marker_line = {'dot': None}
    def al_mover(event):
        if event.inaxes != ax:
            return
        xcur, ycur = event.xdata, event.ydata
        if xcur is None or ycur is None:
            return
        idx = int(np.argmin(np.abs(Gamma.real - xcur) + np.abs(Gamma.imag - ycur)))
        idx = max(0, min(len(Gamma) - 1, idx))
        f_val = f[idx]
        z_val = Z_in[idx]
        g_val = Gamma[idx]
        if marker_line['dot'] is not None:
            try:
                marker_line['dot'].remove()
            except Exception:
                pass
        marker_line['dot'] = ax.plot(g_val.real, g_val.imag, marker='o', markersize=6,
                                      color='red')[0]
        annotation.set_text(
            f"f = {f_val:.4g} Hz | Z_in = {z_val.real:.3g} + j{z_val.imag:.3g} Ω"
        )
        fig.canvas.draw_idle()
    fig.canvas.mpl_connect('motion_notify_event', al_mover)
    plt.tight_layout()
    plt.show()

for _name in [
    "_reconstruir_menu_herramientas",
    "_nodo_gnd_actual",
    "_etiqueta_resultado",
    "_medir_con_voltimetro",
    "_medir_con_amperimetro",
    "_medir_con_ohmimetro",
    "mostrar_graficas",
    "eliminar_seleccionado",
    "agregar_componente",
    "pedir_valor_numerico",
    "pedir_valores_componente",
    "guardar_circuito",
    "cargar_circuito",
    "exportar_imagen",
    "exportar_pdf",
    "keyPressEvent",
    "agregar_conexion",
    "registrar_historial",
    "restablecer_accion",
    "rehacer_accion",
    "limpiar_todo",
    "closeEvent",
    "simular_corrientes",
    "simular_todo",
    "analizar_circuito",
    "mostrar_osciloscopio",
    'alternar_modo_oscuro',
    'aplicar_tema',
    "agregar_resultado",
    "exportar_resultados",
    "colorear_conexiones",
    "_obtener_nodo_gnd",
    "exportar_netlist_spice",
    "_poner_componente_simple",
    "importar_netlist_spice",
    "iniciar_bode",
    "dialogo_bode_combo",
    "_bode_handle_click",
    "_mapear_nodos_actual",
    "_graficar_bode",
    "simular_transitorio_stub",
    "modelos_no_lineales",
    "_bode_por_transferencia",
    "_calcular_bom",
    "mostrar_bom",
    "_exportar_bom_pdf",
    "iniciar_monte_carlo",
    "actualizar_inspector",
    "iniciar_lgr",
]:
    if _name in globals():
        setattr(CircuitSimulator, _name, globals()[_name])

CircuitSimulator.exportar_mediciones_csv = exportar_mediciones_csv
CircuitSimulator._exportar_bom_csv = _exportar_bom_csv
CircuitSimulator.mostrar_mediciones_activas = mostrar_mediciones_activas
CircuitSimulator.gestor_plantillas = gestor_plantillas
CircuitSimulator.exportar_reporte_pdf_completo = exportar_reporte_pdf_completo
CircuitSimulator.exportar_reporte_word_completo = exportar_reporte_word_completo
CircuitSimulator.mostrar_atajos_teclado = mostrar_atajos_teclado

class ModoLaboratorioDialog(QDialog):
    def __init__(self, parent=None, simulacion: Optional[Tuple[List[str], List[float]]] = None):
        super().__init__(parent)
        self.setWindowTitle("Modo Laboratorio")
        from PyQt5.QtWidgets import QVBoxLayout, QPushButton
        if simulacion is not None:
            self.sim_labels, self.sim_values = simulacion
        else:
            self.sim_labels, self.sim_values = [], []
        self.med_labels: List[str] = []
        self.med_values: List[float] = []
        self.figure = Figure(figsize=(5, 4))
        self.canvas = FigureCanvas(self.figure)
        layout = QVBoxLayout(self)
        layout.addWidget(self.canvas)
        self.btn_cargar = QPushButton("Cargar medición")
        self.btn_cargar.clicked.connect(self.cargar_medicion)
        layout.addWidget(self.btn_cargar)
        self.graficar()

    def cargar_medicion(self) -> None:
        from PyQt5.QtWidgets import QFileDialog, QMessageBox
        fname, _ = QFileDialog.getOpenFileName(
            self,
            "Selecciona archivo de medición",
            "",
            "CSV (*.csv);;Todos los archivos (*.*)"
        )
        if not fname:
            return
        labels: List[str] = []
        values: List[float] = []
        import csv
        try:
            with open(fname, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                for row in reader:
                    if not row:
                        continue
                    try:
                        label = row[0]
                        val = float(row[1]) if len(row) > 1 else 0.0
                        labels.append(label)
                        values.append(val)
                    except Exception:
                        continue
            self.med_labels = labels
            self.med_values = values
            self.graficar()
        except Exception as exc:
            try:
                QMessageBox.critical(self, "Modo laboratorio", f"No se pudo cargar la medición:\n{exc}")
            except Exception:
                pass

    def graficar(self) -> None:
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        sim_plotted = False
        med_plotted = False
        if self.sim_values:
            try:
                import numpy as _np
                sim_vals = list(self.sim_values)
                if self.med_values and len(self.med_values) != len(self.sim_values):
                    x_sim = _np.linspace(0, len(sim_vals) - 1, len(sim_vals))
                    x_target = _np.linspace(0, len(sim_vals) - 1, len(self.med_values))
                    sim_vals_interp = _np.interp(x_target, x_sim, sim_vals)
                    x_sim_plot = list(range(len(self.med_values)))
                    ax.plot(x_sim_plot, sim_vals_interp, label="Simulación")
                else:
                    x_sim_plot = list(range(len(sim_vals)))
                    ax.plot(x_sim_plot, sim_vals, label="Simulación")
                sim_plotted = True
            except Exception:
                try:
                    x = list(range(len(self.sim_values)))
                    ax.plot(x, self.sim_values, label="Simulación")
                    sim_plotted = True
                except Exception:
                    pass
        if self.med_values:
            try:
                x_med = list(range(len(self.med_values)))
                ax.plot(x_med, self.med_values, label="Medición")
                med_plotted = True
            except Exception:
                pass
        try:
            if sim_plotted or med_plotted:
                ax.legend()
                ax.set_xlabel("Muestra")
                ax.set_ylabel("Magnitud")
                ax.set_title("Curva simulada vs medida")
                combined = []
                if sim_plotted:
                    combined.extend(self.sim_values)
                if med_plotted:
                    combined.extend(self.med_values)
                valid_vals = [v for v in combined if isinstance(v, (int, float))]
                if valid_vals:
                    y_min = min(valid_vals)
                    y_max = max(valid_vals)
                    if y_min == y_max:
                        margin = 0.1 * abs(y_max) if y_max != 0 else 1.0
                        ax.set_ylim(y_min - margin, y_max + margin)
                    else:
                        margin = 0.1 * (y_max - y_min)
                        ax.set_ylim(y_min - margin, y_max + margin)
                max_len = 0
                if sim_plotted:
                    max_len = max(max_len, len(self.sim_values))
                if med_plotted:
                    max_len = max(max_len, len(self.med_values))
                if max_len > 0:
                    ax.set_xlim(0, max_len - 1)
                ax.grid(True, which='both', linestyle='--', alpha=0.5)
        except Exception:
            pass
        try:
            self.canvas.draw()
        except Exception:
            pass


class TextFormatDialog(QDialog):
    def __init__(self, parent: Optional[Any] = None) -> None:
        super().__init__(parent)
        from PyQt5.QtWidgets import (
            QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QFontComboBox,
            QSpinBox, QCheckBox, QPushButton, QColorDialog, QWidget, QGridLayout
        )
        from PyQt5.QtGui import QColor
        try:
            self.setWindowTitle(traducir('Agregar etiqueta de texto al circuito'))
        except Exception:
            self.setWindowTitle('Agregar texto')
        layout = QVBoxLayout(self)
        form_layout = QGridLayout()
        row = 0
        self.text_input = QLineEdit()
        lbl_text = QLabel(traducir('Texto'))
        form_layout.addWidget(lbl_text, row, 0)
        form_layout.addWidget(self.text_input, row, 1, 1, 2)
        row += 1
        lbl_font = QLabel(traducir('Fuente') if 'Fuente' in TRANSLATIONS else 'Fuente')
        self.font_combo = QFontComboBox()
        try:
            from PyQt5.QtGui import QFontDatabase
            self.font_combo.setWritingSystem(QFontDatabase.Latin)
        except Exception:
            pass
        form_layout.addWidget(lbl_font, row, 0)
        form_layout.addWidget(self.font_combo, row, 1, 1, 2)
        row += 1
        lbl_size = QLabel(traducir('Tamaño') if 'Tamaño' in TRANSLATIONS else 'Tamaño')
        self.size_spin = QSpinBox()
        self.size_spin.setRange(6, 72)
        self.size_spin.setValue(12)
        form_layout.addWidget(lbl_size, row, 0)
        form_layout.addWidget(self.size_spin, row, 1)
        row += 1
        self.chk_bold = QCheckBox(traducir('Negrita') if 'Negrita' in TRANSLATIONS else 'Negrita')
        self.chk_italic = QCheckBox(traducir('Cursiva') if 'Cursiva' in TRANSLATIONS else 'Cursiva')
        self.chk_underline = QCheckBox(traducir('Subrayado') if 'Subrayado' in TRANSLATIONS else 'Subrayado')
        form_layout.addWidget(self.chk_bold, row, 0)
        form_layout.addWidget(self.chk_italic, row, 1)
        form_layout.addWidget(self.chk_underline, row, 2)
        row += 1
        lbl_color = QLabel(traducir('Color') if 'Color' in TRANSLATIONS else 'Color')
        self.color_button = QPushButton()
        self._selected_color = QColor('black')
        self._actualizar_boton_color()
        self.color_button.clicked.connect(self._elegir_color)
        form_layout.addWidget(lbl_color, row, 0)
        form_layout.addWidget(self.color_button, row, 1)
        row += 1
        layout.addLayout(form_layout)
        button_layout = QHBoxLayout()
        self.ok_button = QPushButton(traducir('Sí') if 'Sí' in TRANSLATIONS else 'OK')
        self.cancel_button = QPushButton(traducir('No') if 'No' in TRANSLATIONS else 'Cancel')
        self.ok_button.clicked.connect(self.accept)
        self.cancel_button.clicked.connect(self.reject)
        button_layout.addStretch()
        button_layout.addWidget(self.ok_button)
        button_layout.addWidget(self.cancel_button)
        layout.addLayout(button_layout)

    def _elegir_color(self) -> None:
        from PyQt5.QtWidgets import QColorDialog
        col = QColorDialog.getColor(self._selected_color, self)
        if col.isValid():
            self._selected_color = col
            self._actualizar_boton_color()

    def _actualizar_boton_color(self) -> None:
        col = self._selected_color.name()
        self.color_button.setStyleSheet(f"background-color: {col};")

    def obtener_valores(self) -> Dict[str, Any]:
        return {
            'text': self.text_input.text(),
            'font_family': self.font_combo.currentFont().family(),
            'font_size': self.size_spin.value(),
            'bold': self.chk_bold.isChecked(),
            'italic': self.chk_italic.isChecked(),
            'underline': self.chk_underline.isChecked(),
            'color': self._selected_color,
        }

class EditorComponentesDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Editor de Componentes")
        from PyQt5.QtWidgets import QFormLayout, QLineEdit, QPushButton, QMessageBox
        layout = QFormLayout(self)
        self.nombre_input = QLineEdit()
        self.pines_input = QLineEdit()
        self.modelo_input = QLineEdit()
        layout.addRow("Nombre del componente:", self.nombre_input)
        layout.addRow("Pines (separados por coma):", self.pines_input)
        layout.addRow("Modelo SPICE:", self.modelo_input)
        self.btn_guardar = QPushButton("Guardar")
        layout.addRow(self.btn_guardar)
        self.btn_guardar.clicked.connect(self.guardar)

    def guardar(self) -> None:
        from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
        nombre = self.nombre_input.text().strip()
        pines_text = self.pines_input.text().strip()
        modelo = self.modelo_input.text().strip()
        if not nombre or not pines_text or not modelo:
            QMessageBox.warning(self, "Editor de Componentes", "Debe completar todos los campos.")
            return
        pines = [p.strip() for p in pines_text.split(",") if p.strip()]
        if not pines:
            QMessageBox.warning(self, "Editor de Componentes", "Debe especificar al menos un pin.")
            return
        try:
            if CUSTOM_COMPONENTS_FILE.exists():
                try:
                    with open(CUSTOM_COMPONENTS_FILE, 'r', encoding='utf-8') as f:
                        data = json.load(f)
                except Exception:
                    data = {}
            else:
                data = {}
            data[nombre] = {"pins": pines, "model": modelo}
            with open(CUSTOM_COMPONENTS_FILE, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=4, ensure_ascii=False)
            try:
                COMPONENT_SYNONYMS.setdefault(nombre.lower(), [nombre.lower()])
            except Exception:
                pass
            QMessageBox.information(self, "Editor de Componentes", "Componente guardado correctamente.")
            self.accept()
        except Exception as exc:
            QMessageBox.critical(self, "Editor de Componentes", f"No se pudo guardar el componente:\n{exc}")

def abrir_modo_laboratorio(self) -> None:
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        sim_data = None
        if getattr(self, "ultima_grafica", None):
            try:
                volt_labels, volt_values, _cur_labels, _cur_values = self.ultima_grafica
                sim_data = (volt_labels, volt_values)
            except Exception:
                sim_data = None
        if sim_data is None:
            try:
                elementos = self.extraer_topologia()
                nodo_gnd = None
                for comp in self.componentes:
                    if getattr(comp, 'tipo', None) == "GND" and comp.terminales:
                        nodo_gnd = comp.terminales[0].node_group
                        break
                if nodo_gnd is None:
                    raise RuntimeError("Debes incluir un nodo GND para simular.")
                try:
                    voltajes, corrientes_dict, _ = self.analizar_circuito(elementos, nodo_gnd)
                except Exception as e:
                    raise RuntimeError(str(e))
                nodos_set = set()
                for e in elementos:
                    for n in e["nodos"]:
                        nodos_set.add(n)
                nodos_sorted = sorted(n for n in nodos_set if n != nodo_gnd)
                nodo_idx = {n: i for i, n in enumerate(nodos_sorted)}
                volt_labels = []
                volt_values = []
                for elem in elementos:
                    try:
                        n1, n2 = elem.get("nodos", [None, None])
                        v1 = 0.0 if n1 == nodo_gnd else voltajes[nodo_idx.get(n1, 0)]
                        v2 = 0.0 if n2 == nodo_gnd else voltajes[nodo_idx.get(n2, 0)]
                        dv = abs(v1 - v2)
                        volt_labels.append(elem.get("nombre", ""))
                        volt_values.append(dv)
                    except Exception:
                        continue
                try:
                    self.ultima_grafica = (volt_labels, volt_values, list(corrientes_dict.keys()), list(corrientes_dict.values()))
                except Exception:
                    self.ultima_grafica = None
                sim_data = (volt_labels, volt_values)
            except Exception as exc_sim:
                sim_data = None
                try:
                    QMessageBox.warning(self, traducir("Simulación inválida"), traducir(str(exc_sim)))
                except Exception:
                    pass
        dlg = ModoLaboratorioDialog(self, simulacion=sim_data)
        dlg.exec_()
    except Exception as exc:
        try:
            QMessageBox.critical(self, "Modo laboratorio", f"No se pudo abrir el modo laboratorio:\n{exc}")
        except Exception:
            pass

def guardar_medicion_a_db(self) -> None:
    from PyQt5.QtWidgets import QInputDialog, QMessageBox
    equipo, ok1 = QInputDialog.getText(self, "Guardar Medición", "Equipo de medición:")
    if not ok1:
        return
    circuito, ok2 = QInputDialog.getText(self, "Guardar Medición", "Circuito:")
    if not ok2:
        return
    parametros, ok3 = QInputDialog.getText(self, "Guardar Medición", "Parámetros de prueba:")
    if not ok3:
        return
    try:
        fecha = datetime.now().isoformat()
        inicializar_bd()
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        c.execute(
            "INSERT INTO mediciones (fecha, equipo, circuito, parametros) VALUES (?,?,?,?)",
            (fecha, equipo, circuito, parametros),
        )
        conn.commit()
        conn.close()
        QMessageBox.information(self, _("Información"), "Medición guardada en la base de datos.")
    except Exception as exc:
        try:
            logger.exception(f"Error guardando medición: {exc}")
        except Exception:
            pass
        QMessageBox.critical(self, _("Error"), f"No se pudo guardar la medición:\n{exc}")


def mostrar_base_de_datos(self) -> None:
    from PyQt5.QtWidgets import QDialog, QVBoxLayout, QTableWidget, QTableWidgetItem, QHeaderView, QMessageBox
    try:
        inicializar_bd()
        conn = sqlite3.connect(DB_PATH)
        c = conn.cursor()
        filas = c.execute(
            "SELECT id, fecha, equipo, circuito, parametros FROM mediciones ORDER BY id DESC"
        ).fetchall()
        conn.close()
    except Exception as exc:
        try:
            QMessageBox.critical(self, _("Error"), f"No se pudieron leer las mediciones:\n{exc}")
        except Exception:
            pass
        return
    dlg = QDialog(self)
    dlg.setWindowTitle("Historial de mediciones")
    layout = QVBoxLayout(dlg)
    table = QTableWidget(len(filas), 5)
    table.setHorizontalHeaderLabels(["ID", "Fecha", "Equipo", "Circuito", "Parámetros"])
    for i, fila in enumerate(filas):
        for j, col in enumerate(fila):
            item = QTableWidgetItem(str(col))
            table.setItem(i, j, item)
    try:
        table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
    except Exception:
        pass
    layout.addWidget(table)
    dlg.setLayout(layout)
    dlg.exec_()

def abrir_editor_componentes(self) -> None:
    dlg = EditorComponentesDialog(self)
    dlg.exec_()

def agregar_mejoras(self) -> None:
    try:
        accion_modo_lab = QAction(traducir('Modo Laboratorio'), self)
        accion_modo_lab.setToolTip(traducir_descripcion_emergente('Comparar curva simulada y medida en tiempo real'))
        accion_modo_lab.triggered.connect(lambda: self.abrir_modo_laboratorio())
        self.accion_modo_lab = accion_modo_lab

        accion_guardar = QAction(traducir('Guardar Medición'), self)
        accion_guardar.setToolTip(traducir_descripcion_emergente('Guardar medición actual en la base de datos'))
        accion_guardar.triggered.connect(lambda: self.guardar_medicion_a_db())
        self.accion_guardar_medicion = accion_guardar

        accion_mostrar_db = QAction(traducir('Ver mediciones'), self)
        accion_mostrar_db.setToolTip(traducir_descripcion_emergente('Mostrar historial de mediciones guardadas'))
        accion_mostrar_db.triggered.connect(lambda: self.mostrar_base_de_datos())
        self.accion_ver_mediciones = accion_mostrar_db

        if hasattr(self, 'menu_simulacion') and self.menu_simulacion is not None:
            self.menu_simulacion.addSeparator()
            self.menu_simulacion.addAction(accion_modo_lab)
            self.menu_simulacion.addAction(accion_guardar)
            self.menu_simulacion.addAction(accion_mostrar_db)
        else:
            self.menuBar().addAction(accion_modo_lab)
            self.menuBar().addAction(accion_guardar)
            self.menuBar().addAction(accion_mostrar_db)
    except Exception:
        pass
    try:
        accion_editor = QAction("Editor de Componentes", self)
        accion_editor.setToolTip("Crear o editar símbolos y componentes personalizados")
        accion_editor.triggered.connect(lambda: self.abrir_editor_componentes())
        if hasattr(self, 'menu_herramientas') and self.menu_herramientas is not None:
            self.menu_herramientas.addSeparator()
            self.menu_herramientas.addAction(accion_editor)
        else:
            self.menuBar().addAction(accion_editor)
    except Exception:
        pass

CircuitSimulator.abrir_modo_laboratorio = abrir_modo_laboratorio
CircuitSimulator.guardar_medicion_a_db = guardar_medicion_a_db
CircuitSimulator.mostrar_base_de_datos = mostrar_base_de_datos
CircuitSimulator.abrir_editor_componentes = abrir_editor_componentes
CircuitSimulator.agregar_mejoras = agregar_mejoras


def asegurar_menu_ver(self) -> None:
    try:
        menu_bar = self.menuBar()
        if hasattr(self, 'menu_ver') and self.menu_ver is not None:
            try:
                current_menus = [act.menu() for act in menu_bar.actions()
                                 if hasattr(act, 'menu') and act.menu() is not None]
                if self.menu_ver not in current_menus:
                    if hasattr(self, 'menu_plantillas') and self.menu_plantillas is not None:
                        try:
                            menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_ver)
                        except Exception:
                            menu_bar.addMenu(self.menu_ver)
                    else:
                        menu_bar.addMenu(self.menu_ver)
            except Exception:
                pass
            return
        from PyQt5.QtWidgets import QMenu, QAction
        self.menu_ver = QMenu(traducir('Ver'), self)
        try:
            icon_ampliar = self.obtener_icono("Ampliar.png")
        except Exception:
            icon_ampliar = None
        self.accion_ampliar = QAction(icon_ampliar, traducir('Ampliar'), self)
        self.accion_ampliar.setShortcut("Ctrl+=")
        self.accion_ampliar.setToolTip(traducir_descripcion_emergente('Acercar vista'))
        self.accion_ampliar.triggered.connect(self.ampliar)
        self.menu_ver.addAction(self.accion_ampliar)
        try:
            icon_reducir = self.obtener_icono("Reducir.png")
        except Exception:
            icon_reducir = None
        self.accion_reducir = QAction(icon_reducir, traducir('Reducir'), self)
        self.accion_reducir.setShortcut("Ctrl+-")
        self.accion_reducir.setToolTip(traducir_descripcion_emergente('Alejar vista'))
        self.accion_reducir.triggered.connect(self.reducir)
        self.menu_ver.addAction(self.accion_reducir)
        try:
            icon_ver_todo = self.obtener_icono("Ver_Todo.png")
        except Exception:
            icon_ver_todo = None
        self.accion_ver_todo = QAction(icon_ver_todo, traducir('Ver todo'), self)
        self.accion_ver_todo.setShortcut("Ctrl+Shift+V")
        self.accion_ver_todo.setToolTip(traducir_descripcion_emergente('Ajustar la vista para mostrar todos los elementos'))
        self.accion_ver_todo.triggered.connect(self.ver_todo)
        self.menu_ver.addAction(self.accion_ver_todo)
        try:
            icon_vista_11 = self.obtener_icono("Vista1,1.png")
        except Exception:
            icon_vista_11 = None
        self.accion_vista_11 = QAction(icon_vista_11, traducir('Establecer escala 1:1'), self)
        self.accion_vista_11.setShortcut("Ctrl+Shift+1")
        self.accion_vista_11.setToolTip(traducir_descripcion_emergente('Restablecer la vista a escala 1:1'))
        self.accion_vista_11.triggered.connect(self.vista_uno_a_uno)
        self.menu_ver.addAction(self.accion_vista_11)
        if hasattr(self, 'menu_plantillas') and self.menu_plantillas is not None:
            try:
                menu_bar.insertMenu(self.menu_plantillas.menuAction(), self.menu_ver)
            except Exception:
                menu_bar.addMenu(self.menu_ver)
        else:
            menu_bar.addMenu(self.menu_ver)
    except Exception:
        pass

CircuitSimulator.asegurar_menu_ver = asegurar_menu_ver

def _conectar_terminales(self, t1, t2):
    try:
        if t1 is None or t2 is None or t1 == t2:
            return
    except Exception:
        return
    if getattr(t1, 'is_logic_pin', False) != getattr(t2, 'is_logic_pin', False):
        try:
            self.statusBar().showMessage("No se pueden conectar terminales lógicos con analógicos", 4000)
        except Exception:
            pass
        return
    if getattr(t1, 'is_logic_pin', False) and getattr(t2, 'is_logic_pin', False):
        is_out1 = getattr(t1, 'is_logic_output', False)
        is_out2 = getattr(t2, 'is_logic_output', False)
        if is_out1 == is_out2:
            try:
                self.statusBar().showMessage("Las conexiones lógicas deben ser de salida a entrada", 4000)
            except Exception:
                pass
            return
        if not is_out1:
            for _c in getattr(self, 'conexiones', []):
                if getattr(_c, 'is_logic', False) and (_c.t2 == t1 or _c.t1 == t1):
                    try:
                        self.statusBar().showMessage("Una entrada lógica ya está conectada", 4000)
                    except Exception:
                        pass
                    return
        if not is_out2:
            for _c in getattr(self, 'conexiones', []):
                if getattr(_c, 'is_logic', False) and (_c.t2 == t2 or _c.t1 == t2):
                    try:
                        self.statusBar().showMessage("Una entrada lógica ya está conectada", 4000)
                    except Exception:
                        pass
                    return
    for _c in getattr(self, 'conexiones', []):
        try:
            if (_c.t1 == t1 and _c.t2 == t2) or (_c.t1 == t2 and _c.t2 == t1):
                self.statusBar().showMessage("Las terminales ya están conectadas", 2000)
                return
        except Exception:
            pass
    try:
        nueva = Conexion(self.scene, t1, t2)
    except Exception:
        return
    try:
        self.conexiones.append(nueva)
    except Exception:
        pass
    if getattr(nueva, 'is_logic', False):
        if not hasattr(self, 'logic_conexiones'):
            self.logic_conexiones = []
        try:
            self.logic_conexiones.append(nueva)
        except Exception:
            pass
        try:
            if getattr(t1, 'is_logic_output', False):
                comp = t1.parentItem()
            elif getattr(t2, 'is_logic_output', False):
                comp = t2.parentItem()
            else:
                comp = None
            if comp is not None and hasattr(comp, 'actualizar_estado_logico'):
                comp.actualizar_estado_logico()
        except Exception:
            pass
    try:
        self.statusBar().showMessage(traducir('Acción restablecida (Ctrl+Z)'), 2000)
    except Exception:
        pass
    try:
        self.registrar_historial()
    except Exception:
        pass


def _resaltar_red_al_seleccionar(self):
    if not hasattr(self, 'scene'):
        return
    try:
        from PyQt5.QtWidgets import QGraphicsLineItem as _QGraphicsLineItem
        from PyQt5.QtGui import QPen as _QPen
        from PyQt5.QtCore import Qt as _Qt
    except Exception:
        _QGraphicsLineItem = None
        _QPen = None
    for _conn in getattr(self, 'conexiones', []):
        for _ln in getattr(_conn, 'lines', []):
            if _QPen is not None:
                pen = _QPen(_Qt.blue, 2) if getattr(_conn, 'is_logic', False) else _QPen(_Qt.black, 2)
                _ln.setPen(pen)
    for comp in getattr(self, 'componentes', []):
        for term in getattr(comp, 'terminales', []):
            try:
                term.setBrush(_Qt.black)
            except Exception:
                pass
    selected = []
    try:
        selected = self.scene.selectedItems()
    except Exception:
        pass
    try:
        if selected and hasattr(self, 'selection_filter') and self.selection_filter:
            filt = self.selection_filter
            wire_type = _QGraphicsLineItem
            if filt == 'cables':
                for itm in selected:
                    if wire_type is None or not isinstance(itm, wire_type):
                        try:
                            itm.setSelected(False)
                        except Exception:
                            pass
            elif filt == 'componentes':
                for itm in selected:
                    if wire_type is not None and isinstance(itm, wire_type):
                        try:
                            itm.setSelected(False)
                        except Exception:
                            pass
            try:
                selected = self.scene.selectedItems()
            except Exception:
                selected = []
    except Exception:
        pass
    if not selected:
        return
    visited_terms: set = set()
    visited_conns: set = set()
    queue_terms: List[TerminalItem] = []
    for item in selected:
        if isinstance(item, TerminalItem):
            queue_terms.append(item)
        elif _QGraphicsLineItem is not None and isinstance(item, _QGraphicsLineItem):
            for _conn in getattr(self, 'conexiones', []):
                if item in getattr(_conn, 'lines', []):
                    if _conn not in visited_conns:
                        visited_conns.add(_conn)
                        queue_terms.append(_conn.t1)
                        queue_terms.append(_conn.t2)
                    break
    while queue_terms:
        t = queue_terms.pop()
        if t in visited_terms:
            continue
        visited_terms.add(t)
        try:
            t.setBrush(_Qt.red)
        except Exception:
            pass
        for _conn in getattr(t, 'conexiones', []):
            if _conn not in visited_conns:
                visited_conns.add(_conn)
                for _ln in getattr(_conn, 'lines', []):
                    if _QPen is not None:
                        pen = _QPen(_Qt.darkGreen, 3)
                        _ln.setPen(pen)
                try:
                    other = _conn.t1 if _conn.t2 == t else _conn.t2
                    queue_terms.append(other)
                except Exception:
                    pass

CircuitSimulator._conectar_terminales = _conectar_terminales
CircuitSimulator._resaltar_red_al_seleccionar = _resaltar_red_al_seleccionar

def ampliar(self):
    try:
        factor = 1.15
        if hasattr(self, 'view') and self.view is not None:
            self.view.scale(factor, factor)
    except Exception:
        pass

def reducir(self):
    try:
        factor = 1.15
        if hasattr(self, 'view') and self.view is not None:
            self.view.scale(1.0 / factor, 1.0 / factor)
    except Exception:
        pass

def ver_todo(self):
    try:
        from PyQt5.QtCore import Qt
        if hasattr(self, 'view') and hasattr(self, 'scene') and self.view is not None and self.scene is not None:
            bounding_rect = self.scene.itemsBoundingRect()
            if not bounding_rect.isNull():
                self.view.fitInView(bounding_rect, Qt.KeepAspectRatio)
    except Exception:
        pass

def vista_uno_a_uno(self):
    try:
        if hasattr(self, 'view') and self.view is not None:
            self.view.resetTransform()
    except Exception:
        pass

def abrir_calculadora_resistencias(self) -> None:
    try:
        dlg = ResistenciasCalculatorDialog(self)
        dlg.exec_()
    except Exception as exc:
        try:
            QMessageBox.critical(self, "Calculadora de resistencias",
                                 f"No se pudo abrir la calculadora: {exc}")
        except Exception:
            pass

def abrir_vna(self) -> None:
    from pathlib import Path
    import sys
    import subprocess
    try:
        script_path = Path(__file__).resolve().parent / 'VNA3.1.51.py'
        if not script_path.exists():
            raise FileNotFoundError(f"No se encontró {script_path.name} en {script_path.parent}")
        subprocess.Popen([sys.executable, str(script_path)], cwd=str(script_path.parent))
    except Exception as exc:
        try:
            QMessageBox.critical(self, "VNA View", f"No se pudo abrir el módulo VNA: {exc}")
        except Exception:
            pass

CircuitSimulator.iniciar_analisis_ac = iniciar_analisis_ac
CircuitSimulator.simular_transitorio = simular_transitorio
CircuitSimulator.carta_smith_interactiva = carta_smith_interactiva

def cortar_componentes(self):
    try:
        self.copiar_componentes()
    except Exception:
        pass
    try:
        self.eliminar_seleccionado()
    except Exception:
        pass

    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass

def seleccionar_componentes(self):
    try:
        if hasattr(self, 'boton_conectar') and self.boton_conectar:
            self.boton_conectar.setChecked(False)
    except Exception:
        pass
    try:
        self.modo_medicion = None
        self.modo_texto = False
    except Exception:
        pass

    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass

def seleccionar_todo(self):
    try:
        for item in self.scene.items():
            try:
                item.setSelected(True)
            except Exception:
                pass
    except Exception:
        pass

    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass

def rotar_componentes(self):
    try:
        seleccionados = list(self.scene.selectedItems())
    except Exception:
        seleccionados = []
    for item in seleccionados:
        try:
            current_rot = item.rotation()
        except Exception:
            current_rot = 0
        try:
            item.setRotation(current_rot + 90)
        except Exception:
            continue
        try:
            if hasattr(item, 'vertical'):
                try:
                    item.vertical = not item.vertical
                except Exception:
                    pass
            if hasattr(item, 'reposicionar_terminales'):
                item.reposicionar_terminales()
            try:
                for term in getattr(item, 'terminales', []):
                    for conn in getattr(term, 'conexiones', []):
                        conn.actualizar_posicion()
            except Exception:
                pass
        except Exception:
            pass
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass

    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass

CircuitSimulator.cortar_componentes = cortar_componentes
CircuitSimulator.seleccionar_componentes = seleccionar_componentes
CircuitSimulator.seleccionar_todo = seleccionar_todo
CircuitSimulator.rotar_componentes = rotar_componentes

def duplicar_componentes(self):
    try:
        if hasattr(self, 'copiar_componentes') and hasattr(self, 'pegar_componentes'):
            self.copiar_componentes()
            self.pegar_componentes()
        try:
            self.statusBar().showMessage(traducir('Duplicar'), 2000)
        except Exception:
            pass
    except Exception:
        try:
            self.statusBar().showMessage(traducir('Función en desarrollo'), 2000)
        except Exception:
            pass

def alinear_izquierda(self):
    try:
        if getattr(self, 'elementos_bloqueados', False):
            return
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    if not seleccion:
        try:
            self.statusBar().showMessage(traducir('Alinear izquierda'), 2000)
        except Exception:
            pass
        return
    min_x = None
    for item in seleccion:
        try:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            lx = rect.left()
            if min_x is None or lx < min_x:
                min_x = lx
        except Exception:
            continue
    if min_x is None:
        return
    for item in seleccion:
        try:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            dx = rect.left() - min_x
            pos = item.pos()
            item.setPos(pos.x() - dx, pos.y())
        except Exception:
            continue
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Alinear izquierda'), 2000)
    except Exception:
        pass

def alinear_derecha(self):
    if getattr(self, 'elementos_bloqueados', False):
        return
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    if not seleccion:
        try:
            self.statusBar().showMessage(traducir('Alinear derecha'), 2000)
        except Exception:
            pass
        return
    max_x = None
    for item in seleccion:
        try:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            rx = rect.right()
            if max_x is None or rx > max_x:
                max_x = rx
        except Exception:
            continue
    if max_x is None:
        return
    for item in seleccion:
        try:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            dx = rect.right() - max_x
            pos = item.pos()
            item.setPos(pos.x() - dx, pos.y())
        except Exception:
            continue
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Alinear derecha'), 2000)
    except Exception:
        pass

def alinear_centro(self):
    if getattr(self, 'elementos_bloqueados', False):
        return
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    if not seleccion:
        try:
            self.statusBar().showMessage(traducir('Alinear centro'), 2000)
        except Exception:
            pass
        return
    center_x = None
    try:
        union_rect = None
        for item in seleccion:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            if union_rect is None:
                union_rect = rect
            else:
                union_rect = union_rect.united(rect)
        if union_rect is not None:
            center_x = union_rect.center().x()
    except Exception:
        center_x = None
    if center_x is None:
        total = 0
        count = 0
        for item in seleccion:
            try:
                rect = item.mapToScene(item.boundingRect()).boundingRect()
                total += rect.center().x()
                count += 1
            except Exception:
                continue
        if count > 0:
            center_x = total / count
    if center_x is None:
        return
    for item in seleccion:
        try:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            dx = rect.center().x() - center_x
            pos = item.pos()
            item.setPos(pos.x() - dx, pos.y())
        except Exception:
            continue
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Alinear centro'), 2000)
    except Exception:
        pass

def distribuir_horizontal(self):
    if getattr(self, 'elementos_bloqueados', False):
        return
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    n = len(seleccion)
    if n <= 2:
        try:
            self.statusBar().showMessage(traducir('Distribuir horizontal'), 2000)
        except Exception:
            pass
        return
    try:
        items_with_left = []
        for item in seleccion:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            items_with_left.append((rect.left(), item, rect))
        items_with_left.sort(key=lambda x: x[0])
    except Exception:
        items_with_left = [(0, item, None) for item in seleccion]
        items_with_left.sort(key=lambda x: getattr(x[1], 'x', lambda: 0)())
    total_width = 0.0
    rects = []
    for _, item, rect in items_with_left:
        if rect is None:
            try:
                rect = item.mapToScene(item.boundingRect()).boundingRect()
            except Exception:
                rect = None
        rects.append(rect)
        if rect:
            total_width += rect.width()
    try:
        first_left = min(r.left() for r in rects if r)
        last_right = max(r.right() for r in rects if r)
    except Exception:
        return
    available = last_right - first_left
    try:
        spacing = (available - total_width) / (n - 1)
    except Exception:
        spacing = 0.0
    running_x = first_left
    for idx, (_, item, rect) in enumerate(items_with_left):
        if rect is None:
            try:
                rect = item.mapToScene(item.boundingRect()).boundingRect()
            except Exception:
                rect = None
        if rect:
            desired_left = running_x
            dx = rect.left() - desired_left
            try:
                pos = item.pos()
                item.setPos(pos.x() - dx, pos.y())
            except Exception:
                pass
            running_x += rect.width() + spacing
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Distribuir horizontal'), 2000)
    except Exception:
        pass

def distribuir_vertical(self):
    if getattr(self, 'elementos_bloqueados', False):
        return
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    n = len(seleccion)
    if n <= 2:
        try:
            self.statusBar().showMessage(traducir('Distribuir vertical'), 2000)
        except Exception:
            pass
        return
    try:
        items_with_top = []
        for item in seleccion:
            rect = item.mapToScene(item.boundingRect()).boundingRect()
            items_with_top.append((rect.top(), item, rect))
        items_with_top.sort(key=lambda x: x[0])
    except Exception:
        items_with_top = [(0, item, None) for item in seleccion]
        items_with_top.sort(key=lambda x: getattr(x[1], 'y', lambda: 0)())
    total_height = 0.0
    rects = []
    for _, item, rect in items_with_top:
        if rect is None:
            try:
                rect = item.mapToScene(item.boundingRect()).boundingRect()
            except Exception:
                rect = None
        rects.append(rect)
        if rect:
            total_height += rect.height()
    try:
        first_top = min(r.top() for r in rects if r)
        last_bottom = max(r.bottom() for r in rects if r)
    except Exception:
        return
    available = last_bottom - first_top
    try:
        spacing = (available - total_height) / (n - 1)
    except Exception:
        spacing = 0.0
    running_y = first_top
    for idx, (_, item, rect) in enumerate(items_with_top):
        if rect is None:
            try:
                rect = item.mapToScene(item.boundingRect()).boundingRect()
            except Exception:
                rect = None
        if rect:
            desired_top = running_y
            dy = rect.top() - desired_top
            try:
                pos = item.pos()
                item.setPos(pos.x(), pos.y() - dy)
            except Exception:
                pass
            running_y += rect.height() + spacing
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Distribuir vertical'), 2000)
    except Exception:
        pass

def bloquear_elementos(self, checked: bool):
    try:
        self.elementos_bloqueados = checked
        try:
            from PyQt5.QtWidgets import QGraphicsLineItem, QGraphicsItem
        except Exception:
            QGraphicsLineItem = None
            QGraphicsItem = None
        if hasattr(self, 'scene') and QGraphicsItem is not None:
            for itm in self.scene.items():
                try:
                    is_wire = QGraphicsLineItem is not None and isinstance(itm, QGraphicsLineItem)
                except Exception:
                    is_wire = False
                if isinstance(itm, TerminalItem) or is_wire:
                    continue
                try:
                    itm.setFlag(QGraphicsItem.ItemIsMovable, not checked)
                except Exception:
                    pass
        estado = 'ON' if checked else 'OFF'
        self.statusBar().showMessage(f"{traducir('Bloquear elementos')}: {estado}", 2000)
    except Exception:
        pass

def solo_seleccionar_cables(self):
    try:
        self.selection_filter = 'cables'
        self.statusBar().showMessage(traducir('Solo seleccionar cables'), 2000)
    except Exception:
        pass

def solo_seleccionar_componentes(self):
    try:
        self.selection_filter = 'componentes'
        self.statusBar().showMessage(traducir('Solo seleccionar componentes'), 2000)
    except Exception:
        pass

def alternar_nodos(self, checked: bool):
    try:
        self.mostrar_nodos = checked
        estado = 'ON' if checked else 'OFF'
        self.statusBar().showMessage(f"{traducir('Nodos explícitos')}: {estado}", 2000)
    except Exception:
        pass

def resaltar_red(self):
    try:
        from PyQt5.QtGui import QPen, QColor
        from PyQt5.QtCore import Qt
    except Exception:
        return
    if not hasattr(self, '_resaltado_activo'):
        self._resaltado_activo = False
        self._old_pen_info = {}
    nodos_a_resaltar = set()
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    for comp in seleccion:
        for term in getattr(comp, 'terminales', []):
            node = getattr(term, 'node_group', None)
            if node is not None:
                nodos_a_resaltar.add(node)
    if not nodos_a_resaltar:
        try:
            for comp in getattr(self, 'componentes', []):
                for term in getattr(comp, 'terminales', []):
                    node = getattr(term, 'node_group', None)
                    if node is not None:
                        nodos_a_resaltar.add(node)
        except Exception:
            pass
    if not self._resaltado_activo:
        self._old_pen_info = {}
        for conn in getattr(self, 'conexiones', []):
            try:
                t1_node = getattr(conn.t1, 'node_group', None)
                t2_node = getattr(conn.t2, 'node_group', None)
                highlight = (t1_node in nodos_a_resaltar) or (t2_node in nodos_a_resaltar)
            except Exception:
                highlight = False
            for line in conn.lines:
                try:
                    pen = line.pen()
                except Exception:
                    continue
                self._old_pen_info.setdefault(conn, []).append(QPen(pen))
                color = QColor(Qt.red) if highlight else QColor(Qt.gray)
                width = 4 if highlight else 1
                try:
                    pen.setColor(color)
                    pen.setWidth(width)
                    line.setPen(pen)
                except Exception:
                    pass
        self._resaltado_activo = True
    else:
        for conn, pens in self._old_pen_info.items():
            for line, saved_pen in zip(conn.lines, pens):
                try:
                    line.setPen(saved_pen)
                except Exception:
                    pass
        self._old_pen_info = {}
        self._resaltado_activo = False
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Resaltado de red'), 2000)
    except Exception:
        pass

def etiquetas_red(self):
    try:
        from PyQt5.QtWidgets import QGraphicsTextItem
        from PyQt5.QtGui import QColor
    except Exception:
        return
    try:
        if hasattr(self, 'extraer_topologia'):
            self.extraer_topologia()
    except Exception:
        pass
    if not hasattr(self, '_node_labels_active'):
        self._node_labels_active = False
        self._node_labels = []
    if self._node_labels_active:
        for lbl in getattr(self, '_node_labels', []):
            try:
                self.scene.removeItem(lbl)
            except Exception:
                pass
        self._node_labels = []
        self._node_labels_active = False
    else:
        mapping = {}
        try:
            for comp in getattr(self, 'componentes', []):
                for term in getattr(comp, 'terminales', []):
                    node = getattr(term, 'node_group', None)
                    if node is None:
                        continue
                    pos = term.scenePos()
                    mapping.setdefault(node, []).append(pos)
        except Exception:
            pass
        for node, positions in mapping.items():
            try:
                if not positions:
                    continue
                avg_x = sum(p.x() for p in positions) / len(positions)
                avg_y = sum(p.y() for p in positions) / len(positions)
                texto = f"N{node}"  
                lbl = QGraphicsTextItem(texto)
                lbl.setDefaultTextColor(QColor('darkMagenta'))
                lbl.setPos(avg_x + 5, avg_y + 5)
                lbl.setZValue(10)
                self.scene.addItem(lbl)
                self._node_labels.append(lbl)
            except Exception:
                continue
        self._node_labels_active = True
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Etiquetas de red'), 2000)
    except Exception:
        pass

def ruteo_inteligente(self):
    try:
        conexiones = getattr(self, 'conexiones', [])
    except Exception:
        conexiones = []
    if not conexiones:
        try:
            self.statusBar().showMessage(traducir('Ruteo inteligente'), 2000)
        except Exception:
            pass
        return
    for conn in conexiones:
        try:
            conn.actualizar_posicion()
        except Exception:
            continue
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass
    try:
        self.scene.update()
    except Exception:
        pass
    try:
        if hasattr(self, 'view'):
            self.view.viewport().update()
    except Exception:
        pass
    try:
        self.statusBar().showMessage(traducir('Ruteo inteligente'), 2000)
    except Exception:
        pass

def alternar_puentes(self, checked: bool):
    try:
        self.puentes_activos = checked
        estado = 'ON' if checked else 'OFF'
        self.statusBar().showMessage(f"{traducir('Puentes')}: {estado}", 2000)
    except Exception:
        pass

CircuitSimulator.duplicar_componentes = duplicar_componentes
CircuitSimulator.alinear_izquierda = alinear_izquierda
CircuitSimulator.alinear_derecha = alinear_derecha
CircuitSimulator.alinear_centro = alinear_centro
CircuitSimulator.distribuir_horizontal = distribuir_horizontal
CircuitSimulator.distribuir_vertical = distribuir_vertical
CircuitSimulator.bloquear_elementos = bloquear_elementos
CircuitSimulator.solo_seleccionar_cables = solo_seleccionar_cables
CircuitSimulator.solo_seleccionar_componentes = solo_seleccionar_componentes
CircuitSimulator.alternar_nodos = alternar_nodos
CircuitSimulator.resaltar_red = resaltar_red
CircuitSimulator.etiquetas_red = etiquetas_red
CircuitSimulator.ruteo_inteligente = ruteo_inteligente
CircuitSimulator.alternar_puentes = alternar_puentes

def copiar_componentes(self):
    try:
        seleccion = [item for item in self.scene.selectedItems() if hasattr(item, 'tipo')]
    except Exception:
        seleccion = []
    if not seleccion:
        try:
            self._clipboard_components = None
        except Exception:
            pass
        return
    try:
        self._clipboard_components = list(seleccion)
    except Exception:
        pass

def pegar_componentes(self):
    try:
        clipboard = getattr(self, '_clipboard_components', None)
    except Exception:
        clipboard = None
    if not clipboard:
        return
    try:
        from PyQt5.QtCore import QPointF
    except Exception:
        return
    mapping = {}
    offset = QPointF(20, 20)
    for item in clipboard:
        try:
            tipo = getattr(item, 'tipo', None)
        except Exception:
            continue
        if not tipo:
            continue
        try:
            params = list(getattr(item, 'parametros', []))
        except Exception:
            params = []
        datos = self.imagenes.get(tipo)
        if not datos:
            continue
        archivo, terminales, vertical, invertir, *resto = datos
        escala_base = resto[0] if resto else 0.15
        try:
            self.contadores[tipo] += 1
        except Exception:
            self.contadores[tipo] = 1
        prefijos = {
            "Resistencia": "R", "Capacitor": "C", "Bobina": "L",
            "Transistor NPN": "Q", "Transistor PNP": "Q", "Mosfet": "Q",
            "Diodo Zener": "D", "LED": "D", "OpAmp": "U", "Motor": "M",
            "Transformador": "T", "Potenciómetro": "VR", "Fuente Voltaje": "V",
            "Fuente Corriente": "I", "Batería": "B", "Push Button": "SW",
            "Push Button Cerrado": "SW", "GND": "GND", "Amperímetro": "A",
            "Voltímetro": "VM", "Ohmímetro": "OM", "Generador de funciones": "GF"
        }
        prefijo = prefijos.get(tipo, "X")
        try:
            numero = self.contadores[tipo]
        except Exception:
            numero = 1
        nombre = f"{prefijo}{numero}" if tipo != "GND" else "GND"
        if tipo == "Transistor NPN":
            nombre += " (NPN)"
        elif tipo == "Transistor PNP":
            nombre += " (PNP)"
        elif tipo == "Mosfet":
            nombre += " (MOS)"
        try:
            nuevo = ComponenteInteractivo(archivo, params, nombre, tipo,
                                          terminales=terminales, vertical=vertical,
                                          invertir=invertir)
        except Exception:
            continue
        try:
            nuevo.setScale(item.scale())
        except Exception:
            pass
        try:
            nuevo.svg_item.setScale(item.svg_item.scale())
        except Exception:
            nuevo.svg_item.setScale(escala_base)
        try:
            nuevo.setRotation(item.rotation())
        except Exception:
            pass
        try:
            nuevo.reposicionar_terminales()
        except Exception:
            pass
        try:
            pos_orig = item.pos()
        except Exception:
            pos_orig = QPointF(0, 0)
        nuevo.setPos(pos_orig + offset)
        try:
            self.scene.addItem(nuevo)
            self.componentes.append(nuevo)
        except Exception:
            pass
        mapping[item] = nuevo
    try:
        for conn in list(self.conexiones):
            try:
                c1 = conn.t1.parentItem()
                c2 = conn.t2.parentItem()
            except Exception:
                continue
            if c1 in mapping and c2 in mapping:
                try:
                    idx1 = c1.terminales.index(conn.t1)
                    idx2 = c2.terminales.index(conn.t2)
                except Exception:
                    continue
                nuevo1 = mapping[c1]
                nuevo2 = mapping[c2]
                try:
                    t1_new = nuevo1.terminales[idx1]
                    t2_new = nuevo2.terminales[idx2]
                except Exception:
                    continue
                try:
                    nueva_conn = Conexion(self.scene, t1_new, t2_new)
                    self.conexiones.append(nueva_conn)
                except Exception:
                    pass
    except Exception:
        pass
    try:
        if hasattr(self, 'registrar_historial'):
            self.registrar_historial()
    except Exception:
        pass

CircuitSimulator.copiar_componentes = copiar_componentes
CircuitSimulator.pegar_componentes = pegar_componentes

def _gestor_plugins(self):
    try:
        from PyQt5.QtWidgets import QDialog, QVBoxLayout, QCheckBox, QDialogButtonBox, QLabel, QMessageBox
    except Exception:
        return
    try:
        plugins_list = sorted(plugin_enabled.keys())
    except Exception:
        plugins_list = []
    if not plugins_list:
        QMessageBox.information(self, _("Información"), _("No hay plugins disponibles para gestionar."))
        return
    dlg = QDialog(self)
    dlg.setWindowTitle(_("Gestor de plugins"))
    layout = QVBoxLayout(dlg)
    eventos = [
        "project_exported", "project_exported_extended",
        "simulation_completed", "result_added",
        "measurement_exported", "measurement_completed",
        "project_opened_extended"
    ]
    try:
        label = QLabel(_("Eventos disponibles:") + "\n" + "\n".join(eventos))
        layout.addWidget(label)
    except Exception:
        pass
    checkboxes: Dict[str, QCheckBox] = {}
    for name in plugins_list:
        try:
            cb = QCheckBox(name)
            cb.setChecked(plugin_enabled.get(name, True))
            layout.addWidget(cb)
            checkboxes[name] = cb
        except Exception:
            continue
    btns = QDialogButtonBox(QDialogButtonBox.Ok | QDialogButtonBox.Cancel, parent=dlg)
    layout.addWidget(btns)
    def al_aceptar():
        changed = False
        for pname, cb in checkboxes.items():
            val = cb.isChecked()
            if plugin_enabled.get(pname, True) != val:
                plugin_enabled[pname] = val
                changed = True
        preferences['enabled_plugins'] = plugin_enabled
        try:
            guardar_preferencias()
        except Exception:
            pass
        dlg.accept()
        if changed:
            QMessageBox.information(self, _("Información"), _("Los cambios se aplicarán al reiniciar la aplicación."))
    btns.accepted.connect(al_aceptar)
    btns.rejected.connect(dlg.reject)
    dlg.exec_()

CircuitSimulator._gestor_plugins = _gestor_plugins

def renumerar_componentes(self) -> None:
    try:
        prefijos = {
            "Resistencia": "R", "Capacitor": "C", "Bobina": "L",
            "Transistor NPN": "Q", "Transistor PNP": "Q", "Mosfet": "Q",
            "Diodo Zener": "D", "LED": "D", "OpAmp": "U", "Motor": "M",
            "Transformador": "T", "Potenciómetro": "VR", "Fuente Voltaje": "V",
            "Fuente Corriente": "I", "Batería": "B", "Push Button": "SW",
            "Push Button Cerrado": "SW", "GND": "GND",
            "Amperímetro": "A", "Voltímetro": "VM", "Ohmímetro": "OM",
            "Generador de funciones": "GF"
        }
        comps_by_type = {}
        for comp in getattr(self, 'componentes', []):
            tipo = getattr(comp, 'tipo', None)
            if not tipo:
                continue
            comps_by_type.setdefault(tipo, []).append(comp)
        import re
        for tipo, comps in comps_by_type.items():
            if tipo == "GND":
                try:
                    self.contadores[tipo] = len(comps)
                except Exception:
                    pass
                continue
            def _clave_numerica(c):
                name = getattr(c, 'nombre', '')
                m = re.search(r'(\d+)', name)
                return int(m.group(1)) if m else 0
            comps_sorted = sorted(comps, key=_clave_numerica)
            try:
                self.contadores[tipo] = len(comps_sorted)
            except Exception:
                pass
            prefijo = prefijos.get(tipo, 'X')
            for idx, comp in enumerate(comps_sorted):
                num = idx + 1
                nuevo_nombre = f"{prefijo}{num}" if tipo != "GND" else "GND"
                if tipo == "Transistor NPN":
                    nuevo_nombre += " (NPN)"
                elif tipo == "Transistor PNP":
                    nuevo_nombre += " (PNP)"
                elif tipo == "Mosfet":
                    nuevo_nombre += " (MOS)"
                try:
                    comp.nombre = nuevo_nombre
                except Exception:
                    pass
                try:
                    if hasattr(comp, 'nombre_label'):
                        comp.nombre_label.setPlainText(nuevo_nombre)
                except Exception:
                    pass
    except Exception:
        pass

CircuitSimulator.renumerar_componentes = renumerar_componentes

class MultiTabWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowIcon(QIcon(ICON_PATH))
        self.setWindowTitle("ESIMulador")
        from PyQt5.QtWidgets import QTabWidget
        
        try:
            size = preferences.get('window_size', [1100, 600])
            pos = preferences.get('window_pos', [100, 100])
            self.resize(size[0], size[1])
            self.move(pos[0], pos[1])
        except Exception:
            pass

        self.tab_widget = QTabWidget()
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.tabCloseRequested.connect(self.cerrar_tab)
        self.setCentralWidget(self.tab_widget)
        self.nueva_pestana()

    def nueva_pestana(self):
        sim = CircuitSimulator(nuevo_tab_callback=self.nueva_pestana)
        try:
            if hasattr(sim, 'agregar_mejoras'):
                sim.agregar_mejoras()
        except Exception:
            pass
        try:
            autosave_path = Path(__file__).resolve().parent / 'autosave.json'
            if autosave_path.exists():
                try:
                    sim.cargar_circuito_archivo(str(autosave_path))
                    autosave_path.unlink()
                except Exception:
                    pass
        except Exception:
            pass
        index = self.tab_widget.addTab(sim, f"{traducir('Pestaña')} {self.tab_widget.count() + 1}")
        self.tab_widget.setCurrentIndex(index)

    def _simulador_actual(self):
        try:
            return self.tab_widget.currentWidget()
        except Exception:
            return None

    def cerrar_tab(self, index):
        widget = self.tab_widget.widget(index)
        if widget is not None:
            widget.deleteLater()
        self.tab_widget.removeTab(index)

    def closeEvent(self, event):
        msgbox = QMessageBox(self)
        msgbox.setWindowTitle(traducir('Salir'))
        msgbox.setText(traducir('¿Estás seguro que quieres salir?'))
        msgbox.setIcon(QMessageBox.Question)
        msgbox.setStandardButtons(QMessageBox.Yes | QMessageBox.No)
        msgbox.button(QMessageBox.Yes).setText(traducir('Sí'))
        msgbox.button(QMessageBox.No).setText(traducir('No'))

        if msgbox.exec_() == QMessageBox.Yes:
            try:
                preferences['window_size'] = [self.width(), self.height()]
                preferences['window_pos'] = [self.x(), self.y()]
                try:
                    current_tab = self.tab_widget.currentWidget()
                    if current_tab and hasattr(current_tab, 'modo_oscuro'):
                        preferences['modo_oscuro'] = current_tab.modo_oscuro
                except Exception:
                    pass
                guardar_preferencias()
            except Exception:
                pass
            event.accept()
        else:
            event.ignore()

def _autoguardar_circuito(self) -> None:
    try:
        autosave_path = Path(__file__).resolve().parent / 'autosave.json'
        datos: Dict[str, Any] = {"componentes": [], "conexiones": []}
        for comp in getattr(self, "componentes", []):
            try:
                pos = comp.pos() if hasattr(comp, "pos") else QPointF(0, 0)
                datos["componentes"].append({
                    "tipo": getattr(comp, "tipo", ""),
                    "nombre": getattr(comp, "nombre", ""),
                    "parametros": getattr(comp, "parametros", []),
                    "pos": [float(pos.x()), float(pos.y())],
                    "rotacion": float(comp.rotation()) if hasattr(comp, "rotation") else 0.0,
                    "escala": float(comp.svg_item.scale()) if hasattr(comp, "svg_item") else 1.0
                })
            except Exception:
                continue
        for conexion in getattr(self, "conexiones", []):
            try:
                t1_index = self.componentes.index(conexion.t1.parentItem())
                t2_index = self.componentes.index(conexion.t2.parentItem())
                i1 = conexion.t1.parentItem().terminales.index(conexion.t1)
                i2 = conexion.t2.parentItem().terminales.index(conexion.t2)
                datos["conexiones"].append({
                    "c1": [t1_index, i1],
                    "c2": [t2_index, i2]
                })
            except Exception:
                continue
        with open(autosave_path, 'w', encoding='utf-8') as f:
            json.dump(datos, f, ensure_ascii=False, indent=4)
    except Exception as exc:
        try:
            logger.debug(f"Error en autoguardado: {exc}")
        except Exception:
            pass


def guardar_perfil_simulacion(self) -> None:
    from PyQt5.QtWidgets import QFileDialog, QMessageBox, QListWidgetItem
    ruta, _ = QFileDialog.getSaveFileName(
        self,
        _("Guardar perfil de simulación"),
        "",
        "Perfil de simulación (*.json);;Todos los archivos (*)"
    )
    if not ruta:
        return
    perfil: Dict[str, Any] = {}
    perfil["resultados"] = [
        {"tipo": tipo, "mensaje": mensaje}
        for tipo, mensaje in getattr(self, "resultados_simulaciones", [])
    ]
    if getattr(self, "ultima_grafica", None):
        volt_labels, volt_values, current_labels, current_values = self.ultima_grafica
        perfil["ultima_grafica"] = {
            "volt_labels": volt_labels,
            "volt_values": volt_values,
            "current_labels": current_labels,
            "current_values": current_values,
        }
    try:
        with open(ruta, 'w', encoding='utf-8') as f:
            json.dump(perfil, f, ensure_ascii=False, indent=4)
        QMessageBox.information(self, _("Información"), _("Proyecto exportado correctamente"))
    except Exception as exc:
        QMessageBox.critical(self, _("Error"), f"No se pudo guardar el perfil: {exc}")


def cargar_perfil_simulacion(self) -> None:
    from PyQt5.QtWidgets import QFileDialog, QMessageBox, QListWidgetItem
    from PyQt5.QtCore import Qt
    ruta, _ = QFileDialog.getOpenFileName(
        self,
        _("Abrir perfil de simulación"),
        "",
        "Perfil de simulación (*.json);;Todos los archivos (*)"
    )
    if not ruta:
        return
    try:
        with open(ruta, 'r', encoding='utf-8') as f:
            perfil = json.load(f)
    except Exception as exc:
        QMessageBox.critical(self, _("Error"), f"No se pudo cargar el perfil: {exc}")
        return
    self.resultados_simulaciones = [
        (item.get("tipo", ""), item.get("mensaje", ""))
        for item in perfil.get("resultados", [])
    ]
    if hasattr(self, 'lista_resultados'):
        self.lista_resultados.clear()
        for tipo, mensaje in self.resultados_simulaciones:
            resumen = mensaje.replace('\n', ' ')
            if len(resumen) > 60:
                resumen = resumen[:57] + '...'
            item_widget = QListWidgetItem(f"{tipo}: {resumen}")
            item_widget.setData(Qt.UserRole, (tipo, mensaje))
            self.lista_resultados.addItem(item_widget)
    graf = perfil.get("ultima_grafica")
    if graf:
        self.ultima_grafica = [
            graf.get("volt_labels", []),
            graf.get("volt_values", []),
            graf.get("current_labels", []),
            graf.get("current_values", []),
        ]
        if hasattr(self, 'mostrar_graficas'):
            try:
                self.mostrar_graficas()
            except Exception:
                pass
    QMessageBox.information(self, _("Información"), _("Proyecto exportado correctamente"))


def cambiar_idioma(self, lang_code: str) -> None:
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    if lang_code not in LANGUAGES:
        QMessageBox.warning(self, _("Error"), f"Idioma no soportado: {lang_code}")
        return
    if preferences.get('language') == lang_code:
        return
    preferences['language'] = lang_code
    guardar_preferencias()
    if lang_code == 'es':
        msg = "El idioma ha sido cambiado a español. Reinicia la aplicación para aplicar los cambios."
    else:
        msg = "Language has been changed to English. Please restart the application for the changes to take effect."
    QMessageBox.information(self, _("Información"), msg)

CircuitSimulator._autoguardar_circuito = _autoguardar_circuito
CircuitSimulator.guardar_perfil_simulacion = guardar_perfil_simulacion
CircuitSimulator.cargar_perfil_simulacion = cargar_perfil_simulacion
CircuitSimulator.cambiar_idioma = cambiar_idioma

def abrir_proyecto_extendido(self) -> None:
    from PyQt5.QtWidgets import QFileDialog, QMessageBox
    import json
    ruta, _ = QFileDialog.getOpenFileName(
        self,
        "Abrir proyecto extendido",
        "",
        "Proyecto ESIM extendido (*.esimx *.zip);;Todos los archivos (*)",
    )
    if not ruta:
        return
    try:
        with zipfile.ZipFile(ruta, 'r') as zf:
            try:
                circuit_data = json.loads(zf.read('circuit.json').decode('utf-8'))
            except Exception:
                circuit_data = {}
            try:
                layout_data = json.loads(zf.read('layout.json').decode('utf-8'))
            except Exception:
                layout_data = []
            try:
                layers_data = json.loads(zf.read('layers.json').decode('utf-8'))
            except Exception:
                layers_data = []
            try:
                sim_data = json.loads(zf.read('simulation.json').decode('utf-8'))
            except Exception:
                sim_data = {}
    except Exception as exc:
        QMessageBox.critical(self, _("Error"), f"No se pudo abrir el proyecto: {exc}")
        return
    try:
        for c in list(getattr(self, 'conexiones', [])):
            try:
                c.eliminar()
            except Exception:
                pass
            try:
                self.conexiones.remove(c)
            except Exception:
                pass
    except Exception:
        pass
    try:
        for comp in list(getattr(self, 'componentes', [])):
            try:
                self.scene.removeItem(comp)
            except Exception:
                pass
    except Exception:
        pass
    self.componentes = []
    self.conexiones = []
    comps_map = {}
    for comp_entry in circuit_data.get('components', []):
        tipo = comp_entry.get('type') or comp_entry.get('tipo')
        nombre = comp_entry.get('name') or comp_entry.get('nombre')
        if not tipo or not nombre:
            continue
        params = comp_entry.get('params', {})
        param_list = []
        if isinstance(params, dict):
            for k, v in params.items():
                param_list.append(f"{k}:{v}")
        elif isinstance(params, list):
            param_list = list(params)
        else:
            param_list = []
        try:
            archivo, terminales_count, vertical, invertir, *_ = self.imagenes.get(tipo, (None, 2, False, False))
        except Exception:
            archivo, terminales_count, vertical, invertir = None, 2, False, False
        if not archivo:
            continue
        comp = ComponenteInteractivo(archivo, param_list, nombre, tipo, terminales_count, vertical, invertir)
        layout_entry = None
        for lay in layout_data:
            if lay.get('name') == nombre:
                layout_entry = lay
                break
        if layout_entry:
            try:
                escala_layout = float(layout_entry.get('scale', 1.0))
            except Exception:
                escala_layout = 1.0
            try:
                default_scale = self.imagenes.get(tipo, (None, 2, False, False, 1.0))[-1]
            except Exception:
                default_scale = 1.0
            ratio = default_scale / escala_layout if escala_layout else 1.0
            try:
                x_val = float(layout_entry.get('x', 0.0)) * ratio
                y_val = float(layout_entry.get('y', 0.0)) * ratio
                comp.setPos(x_val, y_val)
            except Exception:
                try:
                    comp.setPos(float(layout_entry.get('x', 0.0)), float(layout_entry.get('y', 0.0)))
                except Exception:
                    pass
            try:
                comp.vertical = bool(layout_entry.get('vertical', comp.vertical))
                comp.invertir = bool(layout_entry.get('invertir', comp.invertir))
            except Exception:
                pass
            try:
                if hasattr(comp, 'svg_item'):
                    comp.svg_item.setScale(default_scale)
            except Exception:
                pass
        try:
            comp.reposicionar_terminales()
        except Exception:
            pass
        try:
            self.scene.addItem(comp)
        except Exception:
            pass
        self.componentes.append(comp)
        comps_map[nombre] = comp
    for layer_entry in layers_data:
        name = layer_entry.get('name')
        try:
            layer_val = float(layer_entry.get('layer', 0.0))
        except Exception:
            layer_val = 0.0
        comp_obj = comps_map.get(name)
        if comp_obj is not None:
            try:
                comp_obj.setZValue(layer_val)
            except Exception:
                pass
    for net_entry in circuit_data.get('nets', []):
        pins = net_entry.get('pins', [])
        if not pins or len(pins) < 2:
            continue
        terms = []
        for p in pins:
            cname = p.get('comp_name') or p.get('comp') or p.get('nombre')
            idx = p.get('index', 0)
            comp_obj = comps_map.get(cname)
            try:
                if comp_obj is not None and 0 <= idx < len(comp_obj.terminales):
                    terms.append(comp_obj.terminales[idx])
            except Exception:
                continue
        if len(terms) >= 2:
            base_term = terms[0]
            for other_term in terms[1:]:
                try:
                    conn = Conexion(self.scene, base_term, other_term)
                    self.conexiones.append(conn)
                except Exception:
                    pass
    sim_results = sim_data.get('resultados', []) if isinstance(sim_data, dict) else []
    self.resultados_simulaciones = [(item.get('tipo', ''), item.get('mensaje', '')) for item in sim_results]
    try:
        if hasattr(self, 'lista_resultados'):
            self.lista_resultados.clear()
            from PyQt5.QtCore import Qt
            for tipo_res, mensaje in self.resultados_simulaciones:
                resumen = mensaje.replace('\n', ' ')
                if len(resumen) > 60:
                    resumen = resumen[:57] + '...'
                item_widget = QListWidgetItem(f"{tipo_res}: {resumen}")
                item_widget.setData(Qt.UserRole, (tipo_res, mensaje))
                self.lista_resultados.addItem(item_widget)
    except Exception:
        pass
    try:
        graf = sim_data.get('ultima_grafica') if isinstance(sim_data, dict) else None
        if graf:
            self.ultima_grafica = [
                graf.get('volt_labels', []),
                graf.get('volt_values', []),
                graf.get('current_labels', []),
                graf.get('current_values', []),
            ]
            try:
                if hasattr(self, 'mostrar_graficas'):
                    self.mostrar_graficas()
            except Exception:
                pass
    except Exception:
        pass
    try:
        self.last_folder = os.path.dirname(ruta)
        preferences['last_folder'] = self.last_folder
        try:
            self.registrar_reciente(str(ruta))
        except Exception:
            pass
        guardar_preferencias()
    except Exception:
        pass
    try:
        QMessageBox.information(self, _( "Información"), _( "Proyecto exportado correctamente"))
    except Exception:
        pass
    try:
        llamar_hooks("project_opened_extended", nombre_archivo=ruta, metadata=sim_data.get('metadata', {}))
    except Exception:
        pass

def exportar_netlist_modelos(self) -> None:
    from PyQt5.QtWidgets import QInputDialog, QFileDialog, QMessageBox
    try:
        tipos = sorted({getattr(c, 'tipo', '') for c in self.componentes if getattr(c, 'tipo', '') and getattr(c, 'tipo', '') != 'GND'})
    except Exception:
        tipos = []
    pref_def = {
        'Resistencia': 'R',
        'Resistor': 'R',
        'Capacitor': 'C',
        'Bobina': 'L',
        'Inductor': 'L',
        'Fuente Voltaje': 'V',
        'Batería': 'V',
        'Generador de funciones': 'V',
        'Fuente Corriente': 'I',
        'Amperímetro': 'I',
        'Voltímetro': 'V',
        'Ohmímetro': 'R',
    }
    modelo_map = {}
    for tipo_item in tipos:
        valor_def = pref_def.get(tipo_item, tipo_item[:1].upper())
        try:
            valor, ok = QInputDialog.getText(self, _( "Modelo SPICE"), f"Ingrese modelo o prefijo para {tipo_item}:", text=valor_def)
        except Exception:
            valor, ok = (valor_def, True)
        if not ok or not valor:
            valor = valor_def
        modelo_map[tipo_item] = valor
    ruta, _ = QFileDialog.getSaveFileName(
        self,
        _( "Guardar netlist SPICE/Qucs"),
        "netlist.cir",
        "Netlist (*.cir *.net);;Todos los archivos (*)",
    )
    if not ruta:
        return
    try:
        netlist_lines = ["* Netlist generado por ESIMulador"]
        try:
            gnd_id = self._obtener_nodo_gnd()
        except Exception:
            gnd_id = None
        nodemap = {}
        next_num = 1
        def id_spice(node_id):
            nonlocal next_num
            if node_id is None:
                return "0"
            if node_id in nodemap:
                return nodemap[node_id]
            nodemap[node_id] = str(next_num)
            next_num += 1
            return nodemap[node_id]
        if gnd_id is not None:
            nodemap[gnd_id] = "0"
        for comp in getattr(self, 'componentes', []):
            if not getattr(comp, 'terminales', []):
                continue
            tipo_comp = getattr(comp, 'tipo', '')
            if tipo_comp == 'GND':
                continue
            pref = modelo_map.get(tipo_comp, tipo_comp[:1].upper())
            nombre_comp = getattr(comp, 'nombre', '').replace(' ', '_')
            nodos = []
            for term in getattr(comp, 'terminales', []):
                try:
                    nodos.append(id_spice(getattr(term, 'node_group', None)))
                except Exception:
                    nodos.append('0')
                if len(nodos) == 2:
                    break
            if len(nodos) < 2:
                continue
            valor = '1'
            try:
                params_list = getattr(comp, 'parametros', [])
                if params_list:
                    raw = params_list[0]
                    if ':' in raw:
                        val = raw.split(':', 1)[1].strip()
                    else:
                        val = raw.strip()
                    try:
                        valor_num = interpretar_valor_prefijo(val)
                        valor = f"{valor_num}"
                    except Exception:
                        valor = val
            except Exception:
                valor = '1'
            netlist_lines.append(f"{pref}{nombre_comp} {nodos[0]} {nodos[1]} {valor}")
        with open(ruta, 'w', encoding='utf-8') as f:
            f.write('\n'.join(netlist_lines) + '\n')
        QMessageBox.information(self, _( "Información"), _( "Proyecto exportado correctamente"))
    except Exception:
        try:
            QMessageBox.critical(self, _( "Error"), _( "La exportación del proyecto ha fallado"))
        except Exception:
            pass

def ejecutar_netlist_spice(self) -> None:
    from PyQt5.QtWidgets import QInputDialog, QFileDialog, QMessageBox
    import subprocess
    motores = ["ngspice", "qucs"]
    try:
        motor, ok = QInputDialog.getItem(self, _( "Seleccionar motor SPICE"), _( "Motor:"), motores, 0, False)
    except Exception:
        motor, ok = (motores[0], True)
    if not ok or not motor:
        return
    ruta, _ = QFileDialog.getOpenFileName(
        self,
        _( "Seleccionar netlist SPICE/Qucs"),
        "",
        "Netlist (*.cir *.net);;Todos los archivos (*)",
    )
    if not ruta:
        return
    try:
        if motor == "ngspice":
            resultado = subprocess.run([motor, "-b", ruta], capture_output=True, text=True)
        else:
            resultado = subprocess.run([motor, ruta], capture_output=True, text=True)
        salida_total = (resultado.stdout or '') + (resultado.stderr or '')
        try:
            self.resultados_simulaciones = getattr(self, 'resultados_simulaciones', [])
        except Exception:
            self.resultados_simulaciones = []
        self.resultados_simulaciones.append(("SPICE", salida_total))
        try:
            if hasattr(self, 'lista_resultados'):
                from PyQt5.QtCore import Qt
                resumen_out = salida_total.replace('\n', ' ')
                if len(resumen_out) > 60:
                    resumen_out = resumen_out[:57] + '...'
                item_res = QListWidgetItem(f"SPICE: {resumen_out}")
                item_res.setData(Qt.UserRole, ("SPICE", salida_total))
                self.lista_resultados.addItem(item_res)
        except Exception:
            pass
        QMessageBox.information(self, _( "Información"), _( "Proyecto exportado correctamente"))
    except Exception as exc:
        try:
            QMessageBox.critical(self, _( "Error"), f"No se pudo ejecutar {motor}: {exc}")
        except Exception:
            pass

class _ScriptConsole(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        from PyQt5.QtWidgets import (QVBoxLayout, QHBoxLayout, QPushButton,
                                     QTextEdit, QFileDialog, QComboBox, QLabel)
        self.setWindowTitle("Consola de scripts (Python)")
        self.resize(700, 500)

        main_layout = QVBoxLayout(self)

        toolbar_layout = QHBoxLayout()

        self.example_combo = QComboBox(self)
        self.example_combo.addItem("Seleccionar ejemplo…")
        for name in SCRIPT_EXAMPLES.keys():
            self.example_combo.addItem(name)
        self.example_combo.currentIndexChanged.connect(self._cargar_ejemplo)
        toolbar_layout.addWidget(QLabel("Ejemplos:", self))
        toolbar_layout.addWidget(self.example_combo)

        load_btn = QPushButton("Cargar", self)
        load_btn.setToolTip("Cargar un script desde un archivo .py")
        load_btn.clicked.connect(self._cargar_script)
        toolbar_layout.addWidget(load_btn)

        save_btn = QPushButton("Guardar", self)
        save_btn.setToolTip("Guardar el script actual a un archivo .py")
        save_btn.clicked.connect(self._guardar_script)
        toolbar_layout.addWidget(save_btn)

        run_btn = QPushButton("Ejecutar", self)
        run_btn.setToolTip("Ejecutar el código Python en el contexto actual")
        run_btn.clicked.connect(self._ejecutar_codigo)
        toolbar_layout.addWidget(run_btn)

        toolbar_layout.addStretch()
        main_layout.addLayout(toolbar_layout)

        self.code_edit = QTextEdit(self)
        self.code_edit.setPlaceholderText(
            "Escribe tu código Python aquí\nPor ejemplo:\n"
            "for c in componentes:\n    print(c.nombre, c.tipo)"
        )
        main_layout.addWidget(self.code_edit, stretch=3)

        self.output_edit = QTextEdit(self)
        self.output_edit.setReadOnly(True)
        self.output_edit.setPlaceholderText("Salida del script…")
        main_layout.addWidget(self.output_edit, stretch=2)

    def _cargar_ejemplo(self, index: int) -> None:
        if index <= 0:
            return
        example_name = list(SCRIPT_EXAMPLES.keys())[index - 1]
        script_code = SCRIPT_EXAMPLES.get(example_name, "")
        self.code_edit.setPlainText(script_code)

    def _guardar_script(self) -> None:
        from PyQt5.QtWidgets import QFileDialog
        filename, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar script",
            str(Path.cwd() / "script.py"),
            "Archivos Python (*.py);;Todos los archivos (*)"
        )
        if filename:
            try:
                with open(filename, 'w', encoding='utf-8') as f:
                    f.write(self.code_edit.toPlainText())
            except Exception as exc:
                mostrar_error_detallado(self, "Guardar script", exc)

    def _cargar_script(self) -> None:
        from PyQt5.QtWidgets import QFileDialog
        filename, _ = QFileDialog.getOpenFileName(
            self,
            "Cargar script",
            str(Path.cwd()),
            "Archivos Python (*.py);;Todos los archivos (*)"
        )
        if filename:
            try:
                with open(filename, 'r', encoding='utf-8') as f:
                    self.code_edit.setPlainText(f.read())
            except Exception as exc:
                mostrar_error_detallado(self, "Cargar script", exc)

    def _ejecutar_codigo(self) -> None:
        import io
        import contextlib
        codigo = self.code_edit.toPlainText()
        salida_buff = io.StringIO()
        sim_ref = self.parent()
        env = {
            'sim': sim_ref,
            'componentes': getattr(sim_ref, 'componentes', []),
            'conexiones': getattr(sim_ref, 'conexiones', []),
            'preferences': preferences,
            'interpretar_valor_prefijo': interpretar_valor_prefijo,
            'formatear_valor_con_prefijo': formatear_valor_con_prefijo,
            'simulation_core': simulation_core,
            'llamar_hooks': llamar_hooks,
        }
        try:
            with contextlib.redirect_stdout(salida_buff), contextlib.redirect_stderr(salida_buff):
                exec(codigo, env)
        except Exception as exc:
            print(exc)
        self.output_edit.setPlainText(salida_buff.getvalue())

def abrir_consola_scripts(self) -> None:
    dlg = _ScriptConsole(self)
    dlg.exec_()

CircuitSimulator.abrir_proyecto_extendido = abrir_proyecto_extendido
CircuitSimulator.exportar_netlist_modelos = exportar_netlist_modelos
CircuitSimulator.ejecutar_netlist_spice = ejecutar_netlist_spice
CircuitSimulator.abrir_consola_scripts = abrir_consola_scripts

def simular_logica(self) -> None:
    from PyQt5.QtWidgets import QMessageBox, QMenu, QApplication
    try:
        api = SimulationAPI(self.componentes)
        rows, columns = api.tabla_verdad_digital_extendida()
        if not rows or not columns:
            QMessageBox.information(self, "Tabla de verdad", "No se encontraron compuertas lógicas en el circuito.")
            return
        net_name_map = getattr(api, 'last_net_name_map', {})
        if net_name_map:
            try:
                model = ModelBuilder.desde_escena(self.componentes)
                comp_by_name = {comp.name: comp for comp in model.components}
            except Exception:
                comp_by_name = {}
            for comp in self.componentes:
                if not hasattr(comp, 'logic_inputs'):
                    continue
                mcomp = comp_by_name.get(getattr(comp, 'nombre', None))
                if mcomp is None or not mcomp.pins:
                    continue
                try:
                    num_inputs = len(comp.logic_inputs)
                    output_pin = mcomp.pins[-1] if mcomp.pins else None
                    output_net = output_pin.net_id if output_pin else None
                    input_names: List[str] = []
                    for idx in range(num_inputs):
                        if idx < len(mcomp.pins):
                            pin_net = mcomp.pins[idx].net_id
                            name = net_name_map.get(pin_net)
                            if name is None:
                                name = chr(ord('A') + idx)
                            input_names.append(name)
                    out_name = net_name_map.get(output_net, None)
                    if out_name is None:
                        out_name = 'Y'
                    comp.establecer_nombres_logicos(input_names, out_name)
                except Exception:
                    continue
        dlg = DigitalTruthTableDialog(rows, columns, parent=self)
        dlg.exec_()
        try:
            functions_info = api.expresiones_booleanas_digitales()
        except Exception:
            functions_info = {}
        if functions_info:
            expr_lines: list[str] = []
            lang = CURRENT_LANGUAGE if 'CURRENT_LANGUAGE' in globals() else 'es'
            if lang == 'en':
                canonical_label = 'Canonical'
                simplified_label = 'Simplified'
            else:
                canonical_label = 'Canónica'
                simplified_label = 'Simplificada'
            for out_name, exprs in functions_info.items():
                canonical_expr = exprs.get('canonical', '')
                simplified_expr = exprs.get('simplified', '')
                line = f"{out_name}:\n  {canonical_label}: {canonical_expr}\n  {simplified_label}: {simplified_expr}"
                expr_lines.append(line)
            expr_text = "\n\n".join(expr_lines)
            try:
                QMessageBox.information(self, "Funciones booleanas", expr_text)
            except Exception:
                print(expr_text)
    except Exception as exc:
        try:
            QMessageBox.critical(self, "Error", f"No se pudo generar la tabla de verdad: {exc}")
        except Exception:
            pass

CircuitSimulator.simular_logica = simular_logica

try:
    CircuitSimulator.exportar_logica_pdf = exportar_logica_pdf
except Exception:
    pass
try:
    CircuitSimulator.exportar_logica_word = exportar_logica_word
except Exception:
    pass
try:
    CircuitSimulator.exportar_logica_txt = exportar_logica_txt
except Exception:
    pass
try:
    CircuitSimulator.exportar_logica_excel = exportar_logica_excel
except Exception:
    pass

if __name__ == '__main__':
    if '--cli' in sys.argv:
        argv = [arg for arg in sys.argv[1:] if arg != '--cli']
        exit_code = ejecutar_cli(argv)
        sys.exit(exit_code)
    elif '--run-tests' in sys.argv:
        ok = ejecutar_pruebas_internas(verbose=True)
        sys.exit(0 if ok else 1)
    else:
        app = QApplication(sys.argv)
        ventana = MultiTabWindow()
        ventana.show()
        sys.exit(app.exec_())
